#!/usr/bin/env python3
"""Build detailed colored dependence-plot interpretation document."""

from __future__ import annotations

from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path("/home/xulc/flash_drought/process/SEM_analysis0401/codex")
ORTHO = ROOT / "GLEAM/plots2/prepeak_shap_nomulticollinearity/orthogonal_decomposition"
OUT = ROOT / "GLEAM/writing4"

BIOMES = ["Forest", "Grassland", "Savanna", "Cropland", "Shrubland"]
METRICS = ["GPP", "RECO"]
BIOME_CN = {
    "Forest": "森林",
    "Grassland": "草地",
    "Savanna": "稀树草原",
    "Cropland": "农田",
    "Shrubland": "灌丛",
}
LABELS = {
    "SSRD_z": "SSRD",
    "Pre_z": "Pre",
    "Duration_z": "Duration",
    "Intensity_z": "Intensity",
    "Wind_z": "Wind",
    "STRD_resid_after_SSRD": "STRD_resid",
    "TMP_resid_after_SSRD_STRD": "TMP_resid",
    "VPD_resid_after_SSRD_TMP_Wind": "VPD_resid",
    "EVA_resid_after_SSRD_Pre_VPD": "EVA_resid",
    "SMrz_resid_after_Pre_EVA": "SMrz_resid",
}
ENERGY = {"SSRD_z", "STRD_resid_after_SSRD", "TMP_resid_after_SSRD_STRD"}
WATER = {"EVA_resid_after_SSRD_Pre_VPD", "SMrz_resid_after_Pre_EVA", "VPD_resid_after_SSRD_TMP_Wind", "Pre_z"}
EVENT = {"Duration_z", "Intensity_z"}


def cn_font(run, name: str = "SimSun", size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    cn_font(r, "SimHei", 18, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    cn_font(r, "SimSun", 10)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        cn_font(r, "SimHei", 14 if level == 1 else 12, True)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.line_spacing = 1.18
    r = p.add_run(text)
    cn_font(r, "SimSun", 10)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], fill: str = "D9EAF7") -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        cn_font(r, "SimHei", 8, True)
        shade_cell(cell, fill)
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            cn_font(r, "SimSun", 8)
    doc.add_paragraph()


def add_image(doc: Document, path: Path, caption: str, width: float = 5.8) -> None:
    if not path.exists():
        add_para(doc, f"图件缺失：{path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    cn_font(r, "SimSun", 9)


def new_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.82)
    sec.bottom_margin = Inches(0.82)
    sec.left_margin = Inches(0.82)
    sec.right_margin = Inches(0.82)
    doc.styles["Normal"].font.name = "SimSun"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    doc.styles["Normal"].font.size = Pt(10)
    return doc


def load_importance() -> pd.DataFrame:
    frames = []
    for metric in METRICS:
        for biome in BIOMES:
            p = ORTHO / metric / biome / "feature_importance.csv"
            df = pd.read_csv(p)
            df["metric"] = metric
            df["biome"] = biome
            frames.append(df)
    return pd.concat(frames, ignore_index=True)


def dependence_profile(metric: str, biome: str, feature: str) -> dict[str, object]:
    p = ORTHO / metric / biome / "dependence_plot_data.parquet"
    df = pd.read_parquet(p, columns=[f"feature__{feature}", f"shap__{feature}"])
    x = pd.to_numeric(df[f"feature__{feature}"], errors="coerce").to_numpy(dtype=float)
    y = pd.to_numeric(df[f"shap__{feature}"], errors="coerce").to_numpy(dtype=float)
    ok = np.isfinite(x) & np.isfinite(y)
    x, y = x[ok], y[ok]
    if len(x) < 80:
        return {"rho": np.nan, "direction": "样本不足", "threshold": "", "delta": np.nan}
    rho = float(pd.Series(x).corr(pd.Series(y), method="spearman"))
    edges = np.unique(np.nanquantile(x, np.linspace(0.02, 0.98, 31)))
    cx, cy = [], []
    for left, right in zip(edges[:-1], edges[1:]):
        mask = (x >= left) & (x <= right if right == edges[-1] else x < right)
        if mask.sum() >= 30:
            cx.append(float(np.nanmedian(x[mask])))
            cy.append(float(np.nanmedian(y[mask])))
    if len(cx) < 5:
        return {"rho": rho, "direction": "样本不足", "threshold": "", "delta": np.nan}
    cx = np.asarray(cx)
    cy = np.asarray(cy)
    yr = float(np.nanmax(cy) - np.nanmin(cy))
    low = float(np.nanmedian(cy[cx <= np.nanquantile(cx, 0.25)]))
    high = float(np.nanmedian(cy[cx >= np.nanquantile(cx, 0.75)]))
    delta = high - low
    if abs(delta) < max(0.12 * yr, 1e-9):
        direction = "非线性/弱方向"
    elif delta > 0:
        direction = "高值端延长恢复"
    else:
        direction = "高值端缩短恢复"
    threshold = ""
    signs = np.sign(cy)
    signs[np.abs(cy) < max(0.02 * yr, 1e-9)] = 0
    for i in range(len(cx) - 1):
        if signs[i] == 0:
            threshold = f"{cx[i]:.2f} 零交叉"
            break
        if signs[i] * signs[i + 1] < 0:
            denom = cy[i + 1] - cy[i]
            val = cx[i] if abs(denom) <= 1e-12 else cx[i] - cy[i] * (cx[i + 1] - cx[i]) / denom
            threshold = f"{val:.2f} 零交叉"
            break
    if not threshold:
        slope = np.abs(np.diff(cy) / np.maximum(np.diff(cx), 1e-9))
        idx = int(np.nanargmax(slope))
        threshold = f"{((cx[idx] + cx[idx + 1]) / 2):.2f} 最大斜率转折"
    return {"rho": rho, "direction": direction, "threshold": threshold, "delta": delta}


def preferred_color(feature: str) -> str:
    if feature in ENERGY:
        return "EVA_resid_after_SSRD_Pre_VPD" if feature == "SSRD_z" else "VPD_resid_after_SSRD_TMP_Wind"
    if feature in WATER:
        return "SSRD_z" if feature in {"EVA_resid_after_SSRD_Pre_VPD", "Pre_z"} else "TMP_resid_after_SSRD_STRD"
    if feature in EVENT:
        return "Pre_z"
    return "VPD_resid_after_SSRD_TMP_Wind"


REPRESENTATIVE_PLOTS = {
    "Forest": {
        "GPP": [
            ("SSRD_z", "EVA_resid_after_SSRD_Pre_VPD"),
            ("TMP_resid_after_SSRD_STRD", "VPD_resid_after_SSRD_TMP_Wind"),
            ("EVA_resid_after_SSRD_Pre_VPD", "SSRD_z"),
        ],
        "RECO": [
            ("TMP_resid_after_SSRD_STRD", "VPD_resid_after_SSRD_TMP_Wind"),
            ("VPD_resid_after_SSRD_TMP_Wind", "TMP_resid_after_SSRD_STRD"),
        ],
    },
    "Grassland": {
        "GPP": [
            ("TMP_resid_after_SSRD_STRD", "VPD_resid_after_SSRD_TMP_Wind"),
            ("VPD_resid_after_SSRD_TMP_Wind", "TMP_resid_after_SSRD_STRD"),
            ("Duration_z", "Pre_z"),
        ],
        "RECO": [
            ("TMP_resid_after_SSRD_STRD", "VPD_resid_after_SSRD_TMP_Wind"),
            ("Duration_z", "Pre_z"),
        ],
    },
    "Savanna": {
        "GPP": [
            ("SSRD_z", "EVA_resid_after_SSRD_Pre_VPD"),
            ("TMP_resid_after_SSRD_STRD", "VPD_resid_after_SSRD_TMP_Wind"),
        ],
        "RECO": [
            ("SSRD_z", "EVA_resid_after_SSRD_Pre_VPD"),
            ("STRD_resid_after_SSRD", "VPD_resid_after_SSRD_TMP_Wind"),
        ],
    },
    "Cropland": {
        "GPP": [
            ("EVA_resid_after_SSRD_Pre_VPD", "SSRD_z"),
            ("SMrz_resid_after_Pre_EVA", "TMP_resid_after_SSRD_STRD"),
            ("Pre_z", "SSRD_z"),
        ],
        "RECO": [
            ("EVA_resid_after_SSRD_Pre_VPD", "TMP_resid_after_SSRD_STRD"),
            ("TMP_resid_after_SSRD_STRD", "VPD_resid_after_SSRD_TMP_Wind"),
        ],
    },
    "Shrubland": {
        "GPP": [
            ("SSRD_z", "EVA_resid_after_SSRD_Pre_VPD"),
            ("VPD_resid_after_SSRD_TMP_Wind", "SMrz_resid_after_Pre_EVA"),
            ("Duration_z", "Pre_z"),
        ],
        "RECO": [
            ("SSRD_z", "VPD_resid_after_SSRD_TMP_Wind"),
            ("VPD_resid_after_SSRD_TMP_Wind", "TMP_resid_after_SSRD_STRD"),
            ("Duration_z", "Pre_z"),
        ],
    },
}


def colored_plot_path(metric: str, biome: str, feature: str, color_feature: str) -> Path:
    fl = LABELS.get(feature, feature)
    cl = LABELS.get(color_feature, color_feature)
    return ORTHO / metric / biome / "dependence_all_colored_cross" / f"{fl}_colored_by_{cl}.png"


def top_feature_rows(imp: pd.DataFrame, metric: str, biome: str, n: int = 5) -> list[list[str]]:
    rows = []
    sub = imp[(imp.metric == metric) & (imp.biome == biome)].head(n)
    for r in sub.itertuples(index=False):
        st = dependence_profile(metric, biome, r.feature)
        rows.append([
            r.display_label,
            f"{r.percent:.1f}%",
            f"{st['rho']:.2f}",
            st["direction"],
            st["threshold"],
        ])
    return rows


def metric_biome_summary(imp: pd.DataFrame, metric: str, biome: str) -> str:
    sub = imp[(imp.metric == metric) & (imp.biome == biome)].head(5)
    names = [f"{r.display_label}({r.percent:.1f}%)" for r in sub.itertuples(index=False)]
    energy = [r.display_label for r in sub.itertuples(index=False) if r.feature in ENERGY]
    water = [r.display_label for r in sub.itertuples(index=False) if r.feature in WATER]
    event = [r.display_label for r in sub.itertuples(index=False) if r.feature in EVENT]
    bits = [f"{metric}-{biome} 的前五贡献特征为：{'、'.join(names)}。"]
    if energy:
        bits.append(f"能量/热量分量包括 {'、'.join(energy)}，说明恢复过程仍受辐射和热背景控制。")
    if water:
        bits.append(f"水分/干旱分量包括 {'、'.join(water)}，说明蒸散、根区水分、降水或大气干旱仍有独立作用。")
    if event:
        bits.append(f"事件属性分量包括 {'、'.join(event)}，提示恢复时间存在事件记忆。")
    return "".join(bits)


def biome_interpretation(metric: str, biome: str) -> str:
    if biome == "Forest":
        return (
            f"{metric}-Forest 的彩色 dependence plot 应重点关注能量分量与蒸散/大气干旱颜色分层。森林冠层和深根系会缓冲短期水分波动，"
            "因此很多响应不是简单线性，而表现为低值区间快速变化、中值区间平台、高值区间再次调整。若 SSRD 或 TMP_resid 的高 SHAP 区域伴随高 EVA_resid/VPD_resid 颜色，"
            "说明森林恢复并非纯辐射或纯温度效应，而是由能量输入、蒸散需求和水分状态共同调节。"
        )
    if biome == "Grassland":
        return (
            f"{metric}-Grassland 的彩色图主要体现浅根系草地对温度、大气干旱和降水脉冲的敏感性。"
            "TMP_resid 若随 VPD_resid 着色出现高值分层，说明温度促进恢复的同时可能伴随大气干旱导致恢复拖尾；"
            "Pre 或 Duration 面板若出现明显颜色分层，则说明降水补给和事件持续时间共同决定恢复阈值。"
        )
    if biome == "Savanna":
        return (
            f"{metric}-Savanna 的 dependence 形态通常体现热干边缘生态系统的能量-水分耦合。"
            "SSRD、TMP_resid 和 STRD_resid 的主效应需要结合 EVA_resid、VPD_resid 或 Pre 的颜色分层解释：如果高能量背景同时对应高蒸散或高干旱颜色，"
            "则恢复时间变化更可能来自能量输入与水分限制的共同作用，而不是单一因子。"
        )
    if biome == "Cropland":
        return (
            f"{metric}-Cropland 的彩色 dependence plot 需要把作物生育期、灌溉/管理背景和蒸散重启一起考虑。"
            "EVA_resid 或 SMrz_resid 面板若受 SSRD/TMP_resid 明显着色分层，说明水分恢复必须在适宜能量和温度背景下才会转化为碳通量恢复；"
            "Pre 与 Duration 的颜色关系则提示同样的降水补给在不同事件持续时间下具有不同生态意义。"
        )
    return (
        f"{metric}-Shrubland 的彩色图最适合解释复杂分段响应。灌丛具有耐旱和恢复滞后特征，因此 SSRD、EVA_resid、VPD_resid 和 Duration 常出现非单调或阈值后平台。"
        "如果 SSRD 面板在高 EVA_resid 或高 VPD_resid 颜色下出现下降或平台，说明强辐射并不一定促进恢复，而可能通过蒸散耗水和大气干旱放大水分限制。"
    )


def build_doc() -> Path:
    imp = load_importance()
    doc = new_doc()
    add_title(
        doc,
        "13 正交分解 SHAP 彩色 Dependence Plot 与阈值机制详细分析",
        "基于 dependence_all_colored_cross 小图：能量特征以水分特征着色，水分特征以能量特征着色",
    )
    add_para(
        doc,
        "本文件参考原始 02_dependence_threshold_analysis 的写法，但分析对象换成正交分解后的 SHAP dependence plot。"
        "与原始图不同，正交分解后的横轴为标准化锚点或残差 z-score，因此本文不把转折位置解释为原始物理单位阈值，而解释为正交空间中的相对低值、中值和高值转换。"
    )
    add_para(
        doc,
        "新生成的 dependence_all_colored_cross 小图采用交叉颜色映射：能量/热量类特征使用水分/干旱特征着色，水分/干旱类特征使用能量/热量特征着色。"
        "这样做的目的，是检验单个正交主效应是否受到另一个机制维度调制。例如 SSRD_z colored by EVA_resid 可判断短波辐射效应是否随蒸散恢复而改变；"
        "VPD_resid colored by TMP_resid 可判断大气干旱效应是否在热背景较强时放大。"
    )
    add_heading(doc, "1. 变量分组与读图原则", 1)
    add_table(
        doc,
        ["变量组", "正交变量", "解释重点", "颜色映射原则"],
        [
            ["能量/热量", "SSRD_z, STRD_resid, TMP_resid", "辐射输入、长波热背景、独立温度效应", "用 EVA_resid, SMrz_resid, VPD_resid, Pre 着色"],
            ["水分/干旱", "EVA_resid, SMrz_resid, VPD_resid, Pre", "蒸散重启、根区水分、大气干旱、降水补给", "用 SSRD_z, STRD_resid, TMP_resid 着色"],
            ["事件属性", "Duration_z, Intensity_z", "闪旱持续时间与强度记忆", "用能量与水分变量共同着色"],
            ["大气动力", "Wind_z", "风速及其对蒸散需求的辅助调节", "用能量与水分变量共同着色"],
        ],
    )
    add_para(
        doc,
        "读图时应同时看三层信息：第一，黑色趋势线的方向和转折位置，判断该正交变量在何处由缩短恢复转为延长恢复，或由快速变化进入平台；"
        "第二，散点颜色在同一横轴区间内是否分层，判断该主效应是否受着色变量调节；第三，GPP 与 RECO 是否在同一变量上出现不同转折，"
        "从而区分光合恢复与呼吸恢复的机制差异。"
    )

    add_heading(doc, "2. 总体结果", 1)
    add_para(
        doc,
        "总体上，SSRD_z 仍是最稳定的主效应，说明短波辐射背景在降低共线性后依然是 GPP 与 RECO 恢复时间的重要约束。"
        "但彩色 dependence plot 显示，SSRD 的作用并非孤立存在：当 EVA_resid、VPD_resid 或 Pre 的颜色在高 SHAP 区域明显分层时，说明辐射效应需要放在蒸散需求、水分亏缺和补水过程下解释。"
    )
    add_para(
        doc,
        "TMP_resid、STRD_resid 和 VPD_resid 的彩色图进一步说明，热量和大气干旱通常以耦合方式影响恢复。RECO 对 TMP_resid 和 STRD_resid 的响应更容易呈连续上升或高值端增强，"
        "符合呼吸过程对温度和热背景敏感的机制；GPP 则更容易出现平台或下降段，说明光合恢复在能量足够后可能转而受气孔调节和水分约束限制。"
    )

    figure_no = 1
    add_heading(doc, "3. 分 biome 与指标的详细解释", 1)
    for biome in BIOMES:
        add_heading(doc, f"{BIOME_CN[biome]}（{biome}）", 2)
        for metric in METRICS:
            add_heading(doc, f"{metric} - {biome}", 3)
            add_para(doc, metric_biome_summary(imp, metric, biome))
            add_para(doc, biome_interpretation(metric, biome))
            add_table(doc, ["特征", "贡献占比", "Spearman", "方向", "近似转折/零交叉"], top_feature_rows(imp, metric, biome, 5), "EAF3DF")
            reps = REPRESENTATIVE_PLOTS[biome][metric]
            for feature, color_feature in reps:
                feature_label = LABELS.get(feature, feature)
                color_label = LABELS.get(color_feature, color_feature)
                path = colored_plot_path(metric, biome, feature, color_feature)
                add_image(
                    doc,
                    path,
                    f"图 {figure_no}. {metric}-{biome}: {feature_label} colored by {color_label}。该图用于判断 {feature_label} 的主效应是否受 {color_label} 调制。",
                    5.6,
                )
                figure_no += 1
            if biome == "Forest":
                add_para(
                    doc,
                    "森林中应重点关注 EVA_resid、TMP_resid 与 Pre 的分段响应。GPP 常出现低值快速上升、中值平台和高值再次调整，反映冠层恢复和水分调节的缓冲；"
                    "RECO 更容易随水热条件改善而连续增强，说明呼吸恢复对温度、蒸散和底物活化更敏感。"
                )
            elif biome == "Grassland":
                add_para(
                    doc,
                    "草地中 TMP_resid 和 Duration_z 的 GPP-RECO 差异尤其重要。GPP 在适宜温度附近容易进入平台，RECO 则更接近温度激活型响应；"
                    "Duration 对 GPP 的阈值更早，对 RECO 的阈值更晚，说明光合恢复更早受到事件记忆影响，而呼吸恢复需要更长持续胁迫才明显改变。"
                )
            elif biome == "Savanna":
                add_para(
                    doc,
                    "稀树草原中 GPP 与 RECO 的 dependence 形态整体更一致，说明两类碳通量共同受热干背景控制。彩色图的重点在于判断能量增强是否伴随蒸散需求和大气干旱增加。"
                )
            elif biome == "Cropland":
                add_para(
                    doc,
                    "农田中 EVA_resid 和 SMrz_resid 的彩色图应结合管理和作物生育期解释。水分恢复对 GPP 更直接，而 RECO 的响应更容易受到温度、底物和土壤湿润共同控制。"
                )
            elif biome == "Shrubland":
                add_para(
                    doc,
                    "灌丛中最常见的是非单调和阈值后平台。SSRD、EVA_resid、VPD_resid 和 Duration 的颜色分层可用于判断强辐射、高蒸散和高干旱是否共同导致恢复滞后。"
                )

    add_heading(doc, "4. 与原始 dependence plot 文档的关系", 1)
    add_para(
        doc,
        "原始 writing3 的 02 文档负责解释真实物理单位下的阈值，例如实际 SSRD、TMP、PRE 或 |EVA| 的生态意义。"
        "本文件则回答另一个问题：当这些变量的共线部分被剥离后，主效应和条件分层是否仍然存在。若原始图和正交彩色图在方向上相同，说明该机制较稳健；"
        "若原始图有明确阈值但正交图中的转折位置移动，说明原始阈值包含了多个相关变量的共同信息，论文中应谨慎表述为原始变量总效应下的阈值。"
    )
    add_para(
        doc,
        "因此，推荐在结果部分使用原始 dependence plot 讨论物理单位阈值，在稳健性或补充分析中使用本文件的正交彩色 dependence plot，说明结论不完全依赖共线变量。"
        "彩色小图的最大价值，是指出阈值并不是固定的单变量边界，而会随蒸散、土壤水分、大气干旱、辐射和事件持续时间共同移动。"
    )

    add_heading(doc, "5. 结论", 1)
    add_para(
        doc,
        "正交分解后的彩色 dependence plot 表明，GPP 与 RECO 恢复时间的关键非线性响应主要由能量输入、热背景、水分供给、大气干旱和事件记忆共同塑造。"
        "SSRD_z 的稳健性说明辐射背景是共同约束；TMP_resid 和 STRD_resid 在 RECO 中更突出，说明呼吸恢复更依赖温度和热背景；"
        "EVA_resid、SMrz_resid、VPD_resid 和 Pre 的颜色分层说明 GPP 的光合恢复更容易受到蒸散需求、根区水分和降水补给调节。"
        "这些结果支持将原始 SHAP 阈值解释为物理单位下的总效应，将正交彩色 dependence 解释为降低共线性后的机制稳健性和条件依赖证据。"
    )

    out = OUT / "13_orthogonal_colored_dependence_threshold_analysis_cn.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    print(build_doc())
