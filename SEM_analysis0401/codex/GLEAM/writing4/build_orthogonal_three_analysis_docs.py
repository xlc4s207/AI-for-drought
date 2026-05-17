#!/usr/bin/env python3
"""Build three expanded orthogonal-decomposition SHAP analysis documents."""

from __future__ import annotations

from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path("/home/xulc/flash_drought/process/SEM_analysis0401/codex")
ORTHO = ROOT / "GLEAM/plots2/prepeak_shap_nomulticollinearity/orthogonal_decomposition"
FIG = ORTHO / "combined_figures"
OUT = ROOT / "GLEAM/writing4"

BIOMES = ["Forest", "Grassland", "Savanna", "Cropland", "Shrubland"]
METRICS = ["GPP", "RECO"]
FEATURE_ORDER = [
    "SSRD_z",
    "EVA_resid_after_SSRD_Pre_VPD",
    "TMP_resid_after_SSRD_STRD",
    "VPD_resid_after_SSRD_TMP_Wind",
    "SMrz_resid_after_Pre_EVA",
    "Pre_z",
    "STRD_resid_after_SSRD",
    "Wind_z",
    "Duration_z",
    "Intensity_z",
]
LABELS = {
    "SSRD_z": "SSRD",
    "Pre_z": "Pre",
    "Duration_z": "Duration",
    "Intensity_z": "Intensity",
    "Wind_z": "Wind",
    "STRD_resid_after_SSRD": "STRD_resid",
    "TMP_resid_after_SSRD_STRD": "TMP_resid",
    "VPD_resid_after_SSRD_TMP_Wind": "VPD_resid",
    "EVA_resid_after_SSRD_Pre_VPD": "EVA_resid",
    "SMrz_resid_after_Pre_EVA": "SMrz_resid",
}


def cn_font(run, name: str = "SimSun", size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_title(doc: Document, text: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    cn_font(r, "SimHei", 18, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    cn_font(r, "SimSun", 10)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        cn_font(r, "SimHei", 12 if level > 1 else 14, True)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.line_spacing = 1.18
    r = p.add_run(text)
    cn_font(r, "SimSun", 10)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        cn_font(r, "SimHei", 8, True)
        shade_cell(cell, "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            cn_font(r, "SimSun", 8)
    doc.add_paragraph()


def add_image(doc: Document, path: Path, caption: str, width: float = 6.4) -> None:
    if not path.exists():
        add_para(doc, f"图件缺失：{path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    cn_font(r, "SimSun", 9)


def new_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.85)
    sec.left_margin = Inches(0.82)
    sec.right_margin = Inches(0.82)
    doc.styles["Normal"].font.name = "SimSun"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    doc.styles["Normal"].font.size = Pt(10)
    return doc


def load_importance() -> pd.DataFrame:
    frames = []
    for metric in METRICS:
        for biome in BIOMES:
            p = ORTHO / metric / biome / "feature_importance.csv"
            df = pd.read_csv(p)
            df["metric"] = metric
            df["biome"] = biome
            frames.append(df)
    return pd.concat(frames, ignore_index=True)


def dependence_stats(metric: str, biome: str, feature: str) -> dict[str, object]:
    p = ORTHO / metric / biome / "dependence_plot_data.parquet"
    df = pd.read_parquet(p, columns=[f"feature__{feature}", f"shap__{feature}"])
    x = pd.to_numeric(df[f"feature__{feature}"], errors="coerce").to_numpy(dtype=float)
    y = pd.to_numeric(df[f"shap__{feature}"], errors="coerce").to_numpy(dtype=float)
    ok = np.isfinite(x) & np.isfinite(y)
    x = x[ok]
    y = y[ok]
    if len(x) < 30:
        return {"rho": np.nan, "low": np.nan, "high": np.nan, "direction": "样本不足", "threshold": ""}
    rho = float(pd.Series(x).corr(pd.Series(y), method="spearman"))
    low = float(np.nanmedian(y[x <= np.nanquantile(x, 0.25)]))
    high = float(np.nanmedian(y[x >= np.nanquantile(x, 0.75)]))
    delta = high - low
    if rho <= -0.35 and delta < 0:
        direction = "高值端倾向缩短恢复"
    elif rho >= 0.35 and delta > 0:
        direction = "高值端倾向延长恢复"
    else:
        direction = "非线性或分段响应"
    # Binned trend for approximate turning/zero-crossing in transformed units.
    edges = np.unique(np.nanquantile(x, np.linspace(0.02, 0.98, 25)))
    centers = []
    med = []
    for left, right in zip(edges[:-1], edges[1:]):
        mask = (x >= left) & (x <= right if right == edges[-1] else x < right)
        if mask.sum() >= 10:
            centers.append(float(np.nanmedian(x[mask])))
            med.append(float(np.nanmedian(y[mask])))
    threshold = ""
    if len(centers) >= 3:
        centers_arr = np.asarray(centers)
        med_arr = np.asarray(med)
        sign_change = np.where(np.signbit(med_arr[:-1]) != np.signbit(med_arr[1:]))[0]
        if len(sign_change):
            threshold = f"约 {centers_arr[sign_change[0]]:.2f} 至 {centers_arr[sign_change[0] + 1]:.2f}"
        else:
            slope = np.abs(np.diff(med_arr) / np.maximum(np.diff(centers_arr), 1e-9))
            idx = int(np.nanargmax(slope))
            threshold = f"主要转折约 {centers_arr[idx]:.2f}"
    return {"rho": rho, "low": low, "high": high, "direction": direction, "threshold": threshold}


def dependence_profile(metric: str, biome: str, feature: str) -> dict[str, object]:
    """Summarize binned SHAP dependence trend for document-level comparison."""
    p = ORTHO / metric / biome / "dependence_plot_data.parquet"
    df = pd.read_parquet(p, columns=[f"feature__{feature}", f"shap__{feature}"])
    x = pd.to_numeric(df[f"feature__{feature}"], errors="coerce").to_numpy(dtype=float)
    y = pd.to_numeric(df[f"shap__{feature}"], errors="coerce").to_numpy(dtype=float)
    ok = np.isfinite(x) & np.isfinite(y)
    x = x[ok]
    y = y[ok]
    if len(x) < 80:
        return {"direction": "样本不足", "threshold": np.nan, "threshold_note": "", "delta": np.nan, "range": np.nan}
    edges = np.unique(np.nanquantile(x, np.linspace(0.02, 0.98, 31)))
    centers: list[float] = []
    med: list[float] = []
    for left, right in zip(edges[:-1], edges[1:]):
        mask = (x >= left) & (x <= right if right == edges[-1] else x < right)
        if mask.sum() >= 30:
            centers.append(float(np.nanmedian(x[mask])))
            med.append(float(np.nanmedian(y[mask])))
    if len(centers) < 5:
        return {"direction": "样本不足", "threshold": np.nan, "threshold_note": "", "delta": np.nan, "range": np.nan}
    cx = np.asarray(centers)
    cy = np.asarray(med)
    y_range = float(np.nanmax(cy) - np.nanmin(cy))
    low = float(np.nanmedian(cy[cx <= np.nanquantile(cx, 0.25)]))
    high = float(np.nanmedian(cy[cx >= np.nanquantile(cx, 0.75)]))
    delta = high - low
    if abs(delta) < max(0.12 * y_range, 1e-9):
        direction = "非线性/弱方向"
    elif delta > 0:
        direction = "高值端延长恢复"
    else:
        direction = "高值端缩短恢复"

    threshold = np.nan
    note = "无明显零交叉"
    signs = np.sign(cy)
    signs[np.abs(cy) < max(0.02 * y_range, 1e-9)] = 0
    for i in range(len(cx) - 1):
        if signs[i] == 0:
            threshold = float(cx[i])
            note = "零交叉"
            break
        if signs[i] * signs[i + 1] < 0:
            denom = cy[i + 1] - cy[i]
            if abs(denom) > 1e-12:
                threshold = float(cx[i] - cy[i] * (cx[i + 1] - cx[i]) / denom)
            else:
                threshold = float((cx[i] + cx[i + 1]) / 2)
            note = "零交叉"
            break
    if not np.isfinite(threshold):
        slope = np.abs(np.diff(cy) / np.maximum(np.diff(cx), 1e-9))
        if np.isfinite(slope).any():
            idx = int(np.nanargmax(slope))
            threshold = float((cx[idx] + cx[idx + 1]) / 2)
            note = "最大斜率转折"
    return {"direction": direction, "threshold": threshold, "threshold_note": note, "delta": delta, "range": y_range}


def fmt_threshold(value: object) -> str:
    try:
        if pd.isna(value):
            return ""
        return f"{float(value):.2f}"
    except Exception:
        return ""


def direction_relation(g_dir: str, r_dir: str) -> str:
    if g_dir == r_dir:
        return "方向一致"
    if "非线性" in g_dir or "非线性" in r_dir:
        return "至少一方为非线性/弱方向"
    return "方向相反"


def dependence_compare_rows(imp: pd.DataFrame, biome: str, n: int = 5) -> list[list[str]]:
    g = imp[(imp.metric == "GPP") & (imp.biome == biome)].head(n)
    r = imp[(imp.metric == "RECO") & (imp.biome == biome)].head(n)
    feature_to_label = {}
    ordered_features = []
    for row in list(g.itertuples()) + list(r.itertuples()):
        if row.feature not in feature_to_label:
            feature_to_label[row.feature] = row.display_label
            ordered_features.append(row.feature)
    rows = []
    for feature in ordered_features:
        gs = dependence_profile("GPP", biome, feature)
        rs = dependence_profile("RECO", biome, feature)
        gt = gs["threshold"]
        rt = rs["threshold"]
        diff = ""
        if np.isfinite(gt) and np.isfinite(rt):
            diff = f"{abs(float(gt) - float(rt)):.2f}"
        rows.append(
            [
                feature_to_label[feature],
                gs["direction"],
                f"{fmt_threshold(gt)} {gs['threshold_note']}".strip(),
                rs["direction"],
                f"{fmt_threshold(rt)} {rs['threshold_note']}".strip(),
                direction_relation(str(gs["direction"]), str(rs["direction"])),
                diff,
            ]
        )
    return rows


def dependence_compare_sentence(imp: pd.DataFrame, biome: str) -> str:
    rows = dependence_compare_rows(imp, biome, 5)
    same = [r[0] for r in rows if r[5] == "方向一致"]
    opposite = [r[0] for r in rows if r[5] == "方向相反"]
    nonlinear = [r[0] for r in rows if "非线性" in r[5]]
    large_shift = []
    for r in rows:
        try:
            if r[6] and float(r[6]) >= 0.7:
                large_shift.append(r[0])
        except Exception:
            pass
    bits = []
    if same:
        bits.append(f"方向一致的主要特征包括{'、'.join(same[:4])}")
    if opposite:
        bits.append(f"方向相反的特征包括{'、'.join(opposite[:3])}")
    if nonlinear:
        bits.append(f"非线性或弱方向特征包括{'、'.join(nonlinear[:3])}")
    if large_shift:
        bits.append(f"转折位置偏移较大的特征包括{'、'.join(large_shift[:3])}")
    return "；".join(bits) + "。"


def top_rows(imp: pd.DataFrame, n: int = 5) -> list[list[str]]:
    rows = []
    for metric in METRICS:
        for biome in BIOMES:
            sub = imp[(imp.metric == metric) & (imp.biome == biome)].head(n)
            vals = [f"{int(r.rank)}.{r.display_label}({r.percent:.1f}%)" for r in sub.itertuples()]
            rows.append([metric, biome, "；".join(vals)])
    return rows


def biome_summary_sentence(imp: pd.DataFrame, metric: str, biome: str) -> str:
    sub = imp[(imp.metric == metric) & (imp.biome == biome)].head(5)
    names = [r.display_label for r in sub.itertuples()]
    first = sub.iloc[0]
    energy = sum(("SSRD" in n or "STRD" in n or "TMP" in n) for n in names)
    water = sum(("EVA" in n or "SMrz" in n or "Pre" in n or "VPD" in n) for n in names)
    event = sum(("Duration" in n or "Intensity" in n) for n in names)
    return (
        f"{biome} 的 {metric} 正交分解结果中，前五位特征依次为"
        f"{'、'.join(names)}。首位特征 {first.display_label} 的贡献占比为 {first.percent:.1f}%，"
        f"前五特征中能量/热量相关分量 {energy} 个，水分或大气干旱相关分量 {water} 个，事件属性分量 {event} 个。"
    )


def build_doc01(imp: pd.DataFrame) -> Path:
    doc = new_doc()
    add_title(doc, "01 正交分解 SHAP 的 Biome 特征贡献分析", "对应原始 01 文档：在降低共线性后重新审视 GPP 与 RECO 的主导因子")
    add_para(
        doc,
        "本文件对应原始 01_biome_feature_contribution_analysis 文档，但分析对象替换为正交分解后的 SHAP 结果。"
        "与原始结果不同，这里的 SSRD、Pre、Duration、Intensity 和 Wind 是标准化锚点，STRD、TMP、VPD、EVA 和 SMrz 是按物理顺序剥离上游共线信息后的残差分量。"
        "因此，本文件重点回答“在降低共线性后，哪些机制仍然稳定重要”，而不是直接读取原始物理单位阈值。"
    )
    add_image(doc, FIG / "orthogonal_beeswarm_comparison_5biomes_gpp_vs_reco.png", "图 1. 正交分解后五个 biome 中 GPP 与 RECO 的 SHAP beeswarm 对比。", 6.6)
    add_heading(doc, "1. 总体贡献格局", 1)
    add_para(
        doc,
        "最清晰的结果是 SSRD_z 的稳健性。GPP 和 RECO 的所有 biome 中，SSRD_z 均排在第一位，说明短波辐射在原始 SHAP 中的高贡献并不是由 STRD、TMP 或 VPD 的共线替代造成的。"
        "这为原始 01 文档中“SSRD 是多数 biome 的共同核心能量背景”的判断提供了稳健性支撑。"
    )
    add_para(
        doc,
        "第二个重要变化是 TMP、STRD、VPD、EVA 的解释方式发生改变。正交分解后，TMP_resid 表示剥离 SSRD 和 STRD 后的独立温度效应，STRD_resid 表示剥离 SSRD 后的独立长波热背景，VPD_resid 表示剥离辐射、温度和 Wind 后的大气干旱残差，EVA_resid 表示剥离能量、降水和 VPD 后的蒸散残差。"
        "这些残差分量仍然进入多个 biome 的前五，说明原始文档中关于能量、热量、水分和蒸散共同控制恢复时间的机制并未因共线性处理而消失。"
    )
    add_table(doc, ["Metric", "Biome", "前五重要特征"], top_rows(imp, 5))
    add_heading(doc, "2. GPP 恢复时间的 biome 差异", 1)
    for biome in BIOMES:
        add_heading(doc, f"GPP - {biome}", 2)
        add_para(doc, biome_summary_sentence(imp, "GPP", biome))
        if biome == "Forest":
            add_para(doc, "Forest 中 TMP_resid 位列第二，说明在短波和长波辐射背景之外，森林 GPP 恢复仍然保留独立温度调节。该结果与原始文档中森林受 SSRD、EVA、STRD、TMP、SMrz 共同影响的结论一致，但更强调温度残差的独立性。")
        elif biome == "Grassland":
            add_para(doc, "Grassland 中 VPD_resid 与 TMP_resid 靠前，表明草地 GPP 不只是响应辐射背景，还对独立大气干旱和温度分量敏感。这支持草地恢复受浅根层水分和大气蒸散需求共同调制的解释。")
        elif biome == "Savanna":
            add_para(doc, "Savanna 中 SSRD_z 占比最高，说明稀树草原 GPP 对短波辐射背景的依赖最集中。TMP_resid 和 STRD_resid 仍位于前列，说明热干边缘系统中能量和热量背景的分化效应非常明显。")
        elif biome == "Cropland":
            add_para(doc, "Cropland 中 EVA_resid 位列第二，说明即使剥离 SSRD、Pre 和 VPD 后，蒸散残差仍然是农田 GPP 恢复的重要独立信号。这与原始文档中 Cropland-GPP 的 EVA 贡献突出的结论高度一致。")
        elif biome == "Shrubland":
            add_para(doc, "Shrubland 中 Duration_z 进入前三，说明灌丛 GPP 恢复在降低共线性后仍显示事件记忆效应，持续时间可能通过累积水分亏缺和根区恢复滞后影响恢复速度。")
    add_heading(doc, "3. RECO 恢复时间的 biome 差异", 1)
    for biome in BIOMES:
        add_heading(doc, f"RECO - {biome}", 2)
        add_para(doc, biome_summary_sentence(imp, "RECO", biome))
        if biome == "Forest":
            add_para(doc, "Forest-RECO 中 TMP_resid 和 VPD_resid 靠前，说明森林呼吸恢复不仅依赖短波辐射所代表的底物和能量背景，还对独立温度与大气干旱残差敏感。")
        elif biome == "Grassland":
            add_para(doc, "Grassland-RECO 中 TMP_resid 的贡献高于 GPP，说明草地呼吸恢复比光合恢复更容易受独立温度效应控制，这与微生物和根系呼吸对热量条件敏感的过程一致。")
        elif biome == "Savanna":
            add_para(doc, "Savanna-RECO 中 SSRD_z、STRD_resid 和 TMP_resid 形成稳定的能量-热量组合，表明稀树草原呼吸恢复位于能量输入与热负荷共同控制之下。")
        elif biome == "Cropland":
            add_para(doc, "Cropland-RECO 中 EVA_resid 和 TMP_resid 同时靠前，提示农田呼吸恢复对蒸散重启和独立温度环境均有较强响应。")
        elif biome == "Shrubland":
            add_para(doc, "Shrubland-RECO 中 Pre_z 进入前三，说明降水补给对灌丛呼吸恢复具有更直接的独立解释力。")
    add_heading(doc, "4. 小结", 1)
    add_para(doc, "正交分解后的贡献格局证明：原始 SHAP 识别出的 SSRD 主导、Cropland 中 EVA 过程突出、Grassland/Savanna/Shrubland 中能量和热量过程重要等结论，在降低共线性后仍然成立。区别在于，残差变量的解释应强调“独立贡献”，而不是原始变量的总效应。")
    out = OUT / "01_orthogonal_biome_feature_contribution_analysis_cn.docx"
    doc.save(out)
    return out


def build_doc02(imp: pd.DataFrame) -> Path:
    doc = new_doc()
    add_title(doc, "02 正交分解 SHAP 的 Dependence Plot 与阈值机制分析", "对应原始 02 文档：从标准化锚点和残差分量解释非线性响应")
    add_para(
        doc,
        "本文件对应原始 02_dependence_threshold_analysis 文档，但这里的横轴为正交分解后的标准化变量或残差 z-score。"
        "因此，图中的转折点不能直接写作原始 SSRD、TMP 或 PRE 的物理阈值，而应解释为相对样本分布的标准化高低值或残差分量阈值。"
        "原始物理阈值仍应以原始 dependence plot 为主，正交分解 dependence plot 的作用是验证这些机制在去共线性后是否仍存在。"
    )
    add_heading(doc, "1. 图件说明", 1)
    add_para(doc, "每个 biome 的大图包含十个正交特征，左列为 GPP，右列为 RECO。散点表示 SHAP dependence 样本，红线为分箱中位趋势。无颜色映射版本主要用于判断主效应形状、方向和是否存在分段响应；新增的颜色映射版本则用于观察同一主效应是否受到另一个机制分量调制。")
    add_para(
        doc,
        "需要强调的是，正交分解 dependence plot 的横轴不是原始物理单位，而是标准化锚点或残差 z-score。"
        "因此，图中的 -1、0、1 更适合解释为“低于平均水平、接近平均水平、高于平均水平”的相对位置。"
        "当某个曲线在 0 附近发生零交叉，可以理解为该正交分量从负贡献转为正贡献；当曲线在某一区间快速上升或下降，则说明该区间是模型最敏感的响应带。"
    )
    add_para(
        doc,
        "颜色映射的目的不是重新计算交互 SHAP，而是辅助判断条件依赖关系。例如，SSRD_z 面板用 EVA_resid 着色，可以观察短波辐射效应是否随着蒸散残差不同而改变；"
        "TMP_resid 面板用 VPD_resid 着色，可以检查独立温度效应是否伴随大气干旱梯度变化；Pre_z 面板用 Duration_z 着色，可以判断降水残差与事件持续时间是否共同调节恢复时间。"
        "如果同一横轴范围内颜色分层明显，说明该主效应可能受到着色变量调制；如果颜色混合均匀，则说明该主效应更接近独立效应。"
    )
    add_table(
        doc,
        ["横轴特征", "颜色变量", "机制含义"],
        [
            ["SSRD_z", "EVA_resid", "短波辐射效应是否受到蒸散重启程度调制"],
            ["EVA_resid", "VPD_resid", "蒸散残差效应是否伴随大气干旱增强而改变"],
            ["TMP_resid", "VPD_resid", "独立温度效应是否与大气干旱共同塑造恢复"],
            ["VPD_resid", "SMrz_resid", "大气干旱效应是否受根区水分背景约束"],
            ["SMrz_resid", "Pre_z", "根区水分残差是否与降水补给共同作用"],
            ["Pre_z", "Duration_z", "降水残差效应是否受事件持续记忆调节"],
            ["STRD_resid", "TMP_resid", "长波热背景效应是否与独立温度分量耦合"],
            ["Wind_z", "VPD_resid", "风速效应是否体现大气干旱/蒸散需求梯度"],
            ["Duration_z", "Pre_z", "事件持续时间效应是否依赖降水补给"],
            ["Intensity_z", "Duration_z", "事件强度效应是否受持续时间共同调节"],
        ],
    )
    for biome in BIOMES:
        add_image(doc, FIG / "combined_by_biome" / f"{biome}_orthogonal_all_features_gpp_vs_reco.png", f"图 {BIOMES.index(biome)+1}. {biome} 正交分解后 GPP 与 RECO 的全部特征 dependence plot 对比。", 6.5)
    add_heading(doc, "2. 带颜色映射的 dependence plot", 1)
    add_para(
        doc,
        "带颜色映射的大图与无颜色版本使用同一批正交 SHAP 结果，但每个面板额外使用一个机制相关变量对散点着色。"
        "黑线为分箱中位趋势，表示该特征的主效应；颜色梯度用于判断该主效应是否在另一个正交分量的高低水平下发生分层。"
        "这类图比无颜色图更适合识别“同一个阈值在不同水热背景下是否提前或滞后出现”。"
    )
    colored_dir = FIG / "combined_by_biome_colored"
    for biome in BIOMES:
        add_image(
            doc,
            colored_dir / f"{biome}_orthogonal_all_features_gpp_vs_reco_colored.png",
            f"图 {BIOMES.index(biome)+1+len(BIOMES)}. {biome} 正交分解后带颜色映射的 GPP 与 RECO dependence plot。颜色变量按水热机制配对，用于辅助识别条件依赖和交互调制。",
            6.5,
        )
    add_heading(doc, "3. 各 biome 的关键 dependence 响应", 1)
    rows = []
    for metric in METRICS:
        for biome in BIOMES:
            sub = imp[(imp.metric == metric) & (imp.biome == biome)].head(5)
            desc = []
            for r in sub.itertuples():
                st = dependence_stats(metric, biome, r.feature)
                desc.append(f"{r.display_label}: {st['direction']}，{st['threshold']}")
                rows.append([metric, biome, r.display_label, f"{st['rho']:.2f}", st["direction"], st["threshold"]])
            add_heading(doc, f"{metric} - {biome}", 2)
            add_para(doc, "前五贡献特征的 dependence 形态为：" + "；".join(desc) + "。")
            if biome == "Forest":
                add_para(doc, "Forest 中 SSRD_z 和 TMP_resid 的趋势最值得关注。前者代表稳定的短波辐射背景，后者代表去除辐射背景后的独立温度效应；二者共同说明森林恢复并非单一光照限制，而是能量输入和热环境共同调节。")
                add_para(doc, "在彩色图中，若 SSRD_z 或 TMP_resid 面板中颜色随横轴或 SHAP 值呈分层，说明森林的能量或温度效应仍受到蒸散残差、大气干旱残差调制。森林冠层和深根系会缓冲短期水分波动，因此常见的表现不是单调强交互，而是某些区间内颜色分层更明显。")
            elif biome == "Grassland":
                add_para(doc, "Grassland 中 VPD_resid、TMP_resid 与 Pre_z 的分段形态说明草地恢复对大气干旱、温度和降水脉冲都具有敏感响应。由于这些变量已被正交化，其信号更接近各过程的独立边际效应。")
                add_para(doc, "草地彩色 dependence 图尤其需要关注 TMP_resid colored by VPD_resid 和 VPD_resid colored by SMrz_resid。如果高 VPD_resid 颜色集中在较高 SHAP 区域，说明温度促进恢复的同时伴随大气干旱导致恢复时间延长；如果 SMrz_resid 颜色分层明显，则说明根区水分背景决定了大气干旱是否真正转化为恢复滞后。")
            elif biome == "Savanna":
                add_para(doc, "Savanna 中 SSRD_z 的趋势最集中，反映稀树草原在热干边缘条件下对短波辐射背景高度敏感；STRD_resid 和 TMP_resid 的存在说明热量负荷仍会调节恢复路径。")
                add_para(doc, "稀树草原的彩色图用于验证能量-水分耦合。若 SSRD_z 的高 SHAP 区域同时对应较高 EVA_resid 或 VPD_resid，说明强能量背景下恢复时间的变化并不是单纯光照效应，而是同时受到蒸散需求和水分亏缺约束。")
            elif biome == "Cropland":
                add_para(doc, "Cropland 中 EVA_resid 的 dependence 响应说明，农田恢复不仅依赖辐射和降水，还取决于作物冠层蒸散重启中不能被这些上游变量解释的部分。")
                add_para(doc, "农田彩色图应重点阅读 EVA_resid colored by VPD_resid、SMrz_resid colored by Pre_z 和 Pre_z colored by Duration_z。若蒸散残差的正 SHAP 区间伴随较高 VPD_resid，说明农田恢复可能由水分供应恢复和大气蒸散需求共同塑造；若 Pre_z 的颜色受 Duration_z 分层，说明同样的降水残差在短持续和长持续闪旱事件中的含义不同。")
            elif biome == "Shrubland":
                add_para(doc, "Shrubland 中 Duration_z 或 Pre_z 的响应提示事件持续时间和降水补给仍具有独立意义，说明灌丛系统存在更强的事件记忆与水分补给约束。")
                add_para(doc, "灌丛彩色图的意义在于识别复杂分段响应背后的条件调制。SSRD_z colored by EVA_resid 可帮助判断高辐射区间的下降是否与蒸散耗水有关；VPD_resid colored by SMrz_resid 可检验大气干旱效应是否因根区水分不足而放大；Duration_z colored by Pre_z 则可判断事件持续时间的恢复滞后是否被降水补给缓解。")
    add_table(doc, ["Metric", "Biome", "特征", "Spearman rho", "方向", "近似转折/零交叉"], rows)
    add_heading(doc, "4. 与原始阈值分析的衔接", 1)
    add_para(doc, "原始 02 文档中的 SSRD、TMP、PRE、|EVA| 阈值用于解释实际生态水文过程，例如短波辐射由限制转为促进恢复、降水由无效补给进入有效补水区间等。正交分解图不能替代这些原始阈值，但可以判断这些阈值是否可能只是共线结构造成的表观现象。")
    add_para(doc, "本次结果显示，SSRD_z 在所有 biome 中仍具有稳定贡献和清晰 dependence 趋势，说明短波辐射机制具有稳健性；TMP_resid、STRD_resid、VPD_resid 和 EVA_resid 仍在若干 biome 中表现出分段响应，说明温度、长波热背景、大气干旱和蒸散过程并未被完全吸收到 SSRD 中。")
    add_para(
        doc,
        "因此，建议在论文中采用双层表述：原始 dependence plot 用于给出真实单位下的生态阈值，正交 dependence plot 用于说明这些阈值背后的独立机制是否稳健。"
        "当原始图和正交图在方向上相同，可以说该机制不依赖于共线变量共享信息；当方向相同但转折位置移动，应解释为原始变量的总效应被拆分为多个正交分量后，阈值在正交空间中发生重定位；"
        "当彩色图显示明显颜色分层，应进一步说明该阈值具有条件依赖性，即同一主效应在不同水热背景下会提前、滞后或强度改变。"
    )
    add_heading(doc, "5. 小结", 1)
    add_para(
        doc,
        "扩展后的正交 dependence 分析表明，去共线性并没有削弱原始 SHAP 对能量、水分和事件属性的核心判断。SSRD_z 的普遍高贡献和清晰响应说明短波辐射是共同背景；"
        "TMP_resid、STRD_resid、VPD_resid、EVA_resid、SMrz_resid 等残差分量的分段响应说明热量、大气干旱、蒸散和根区水分仍具有独立解释力；"
        "Pre_z、Duration_z 和 Intensity_z 的响应则提示事件属性和补水过程对恢复时间具有明显条件性。带颜色映射的图件进一步说明，这些主效应并不是孤立发生，而是在蒸散、大气干旱、土壤水分和事件持续时间共同调制下形成。"
    )
    out = OUT / "02_orthogonal_dependence_threshold_analysis_cn.docx"
    doc.save(out)
    return out


def diff_sentence(imp: pd.DataFrame, biome: str) -> str:
    g = imp[(imp.metric == "GPP") & (imp.biome == biome)].set_index("display_label")
    r = imp[(imp.metric == "RECO") & (imp.biome == biome)].set_index("display_label")
    labels = sorted(set(g.index) | set(r.index))
    diffs = []
    for label in labels:
        gp = float(g.loc[label, "percent"]) if label in g.index else 0.0
        rp = float(r.loc[label, "percent"]) if label in r.index else 0.0
        diffs.append((label, rp - gp, gp, rp))
    diffs = sorted(diffs, key=lambda x: abs(x[1]), reverse=True)[:4]
    text = []
    for label, delta, gp, rp in diffs:
        if delta > 0:
            text.append(f"{label} 在 RECO 中更高（RECO {rp:.1f}%，GPP {gp:.1f}%）")
        else:
            text.append(f"{label} 在 GPP 中更高（GPP {gp:.1f}%，RECO {rp:.1f}%）")
    return "；".join(text)


def build_doc04(imp: pd.DataFrame) -> Path:
    doc = new_doc()
    add_title(doc, "04 正交分解后的 GPP 与 RECO 差异对比分析", "对应原始 04 文档：在降低共线性后比较两类碳通量恢复机制")
    add_para(
        doc,
        "本文件对应原始 04_gpp_reco_difference_analysis 文档。原始 04 文档强调同一 biome 内 GPP 和 RECO 对共同水热背景的不同读取方式；本文件进一步检验，当输入变量被正交分解以后，这种 GPP-RECO 差异是否仍然存在。"
        "因此，本文的重点不是重新寻找原始物理阈值，而是比较 SSRD_z、TMP_resid、STRD_resid、VPD_resid、EVA_resid 等独立分量在 GPP 和 RECO 中的贡献差异。"
    )
    add_heading(doc, "0. ALE、PDP 和 ICE 的计算方法及其在本文中的作用", 1)
    add_para(
        doc,
        "PDP（Partial Dependence Plot）通过固定目标特征为一系列给定取值，同时保持其他特征为样本中的原始取值，然后对所有样本的模型预测取平均，得到目标特征的平均边际响应。"
        "其计算思想可以写作：对特征 x_j 的每个网格值 v，将数据集中每个样本的 x_j 替换为 v，得到模型预测 f(v, x_-j)，再对所有样本求均值。"
        "PDP 的优点是直观展示平均趋势，缺点是在特征相关性较强时可能生成现实中很少出现的变量组合，因此在本研究中主要作为辅助平均趋势，而不是阈值判定的唯一依据。"
    )
    add_para(
        doc,
        "ICE（Individual Conditional Expectation）与 PDP 使用相似的特征替换思想，但它不先对样本求平均，而是为每个样本保留一条响应曲线，显示个体样本在目标特征变化时预测值如何改变。"
        "PDP 可以看作 ICE 曲线的平均值。ICE 的优势是能够揭示 biome 内部异质性和交互作用，例如某些样本在高 VPD 下恢复延长，而另一些样本响应较弱。"
        "为了避免大量个体曲线拉伸纵坐标，本文件和验证图中主要使用 ICE mean，即 ICE 曲线的平均响应，用于判断平均方向是否支持 SHAP dependence。"
    )
    add_para(
        doc,
        "ALE（Accumulated Local Effects）不是把目标特征强行替换到所有样本上，而是在目标特征的相邻分箱内计算模型预测的局部差分，再沿特征取值方向累积这些局部效应。"
        "其核心步骤是：先按目标特征分位数划分区间，在每个区间内只使用实际落入该区间的样本，计算模型在区间右端点和左端点预测值的平均差，然后将这些局部差分累积并中心化。"
        "因此 ALE 更强调真实数据分布中的局部响应，通常比 PDP 更适合存在相关变量的生态气象数据。本文用 ALE 来验证 SHAP dependence 的方向和转折是否具有模型响应基础。"
    )
    add_para(
        doc,
        "在本文件中，SHAP 负责回答“哪些正交分量贡献最大、在不同取值下如何推高或压低恢复时间”；PDP、ICE mean 和 ALE 则用于验证这些趋势是否能在模型预测响应中复现。"
        "当 SHAP、ALE 和 ICE mean 在方向及转折附近一致时，可认为该特征的解释较稳健；当 PDP 与二者偏离时，更可能是边缘平均和相关变量外推造成的偏差，而不应直接否定 SHAP 结果。"
    )
    add_image(doc, FIG / "orthogonal_beeswarm_comparison_5biomes_gpp_vs_reco.png", "图 1. 正交分解后 GPP 与 RECO 的五 biome beeswarm 对比。", 6.6)
    add_heading(doc, "1. 总体差异", 1)
    add_para(doc, "GPP 和 RECO 的共同点非常明确：SSRD_z 在所有 biome 和两个 metric 中均位列第一。这说明两类碳通量恢复共享短波辐射所代表的背景能量约束。差异则主要体现在第二梯队特征：GPP 中某些 biome 更突出 EVA_resid、VPD_resid 或 Duration_z，而 RECO 中 TMP_resid、STRD_resid、Pre_z 往往更容易靠前。")
    add_para(doc, "这种差异与碳通量过程本身一致。GPP 更直接反映冠层光合恢复、气孔调节和蒸散过程；RECO 同时包含根系呼吸和微生物分解，因而更容易对独立温度效应、长波热背景和降水激活表现出敏感性。")
    add_image(
        doc,
        ORTHO.parent / "orthogonal_decomposition_importance_percent_bars_5biomes_gpp_vs_reco.png",
        "图 2. 正交分解后 GPP 与 RECO 在五个 biome 中的 SHAP 重要性百分比柱状图。柱状图按 mean(|SHAP|) 归一化后的百分比展示每个特征对恢复时间预测的相对贡献，可更直接比较 GPP 与 RECO 在同一 biome 中对能量、热量、水分和事件属性分量的依赖差异。",
        6.6,
    )
    add_para(
        doc,
        "与 beeswarm 图侧重显示单个样本的 SHAP 分布不同，百分比柱状图将每个特征的平均绝对 SHAP 值转化为贡献占比，因此更适合用于比较 GPP 与 RECO 的整体重要性结构。"
        "从该图可以看到，SSRD_z 在各 biome 中均保持最高或近最高贡献，说明短波辐射约束是两类碳通量恢复的共同背景；但第二梯队因子在 GPP 和 RECO 之间明显分化。"
        "GPP 更容易体现 VPD_resid、EVA_resid 或 Duration_z 等与光合调节、蒸散重启和事件记忆相关的贡献，而 RECO 往往更突出 TMP_resid、STRD_resid 或 Pre_z，反映呼吸过程对独立温度、热背景和降水补给的敏感性。"
    )
    add_para(
        doc,
        "从物理机制上看，GPP 的恢复首先受冠层光合能力、气孔导度、叶片水分状态和能量供给约束；因此当 VPD_resid 或 EVA_resid 在 GPP 中较高时，通常意味着光合恢复不仅取决于辐射量，"
        "还取决于大气蒸散需求是否过强以及植被-土壤连续体能否重新建立蒸散通量。RECO 的恢复则包含自养呼吸和异养呼吸，受温度、土壤湿度、底物供应和微生物活性共同控制；"
        "因此 TMP_resid、STRD_resid 或 Pre_z 在 RECO 中更高时，说明呼吸恢复更容易被热环境和降水补给激活。"
    )
    add_para(
        doc,
        "正交分解后的差异尤其有价值，因为它把共线变量中的共享部分剥离开来。例如 TMP_resid 不再代表总温度背景，而是剥离 SSRD 和 STRD 之后的独立温度分量；"
        "VPD_resid 不再只是温度和辐射共同升高的副产物，而是剥离辐射、温度和风速之后仍保留的大气干旱分量。"
        "因此，若这些残差分量在 GPP 或 RECO 中仍然靠前，就说明对应过程具有相对独立的解释力。"
    )
    add_table(doc, ["Metric", "Biome", "前五重要特征"], top_rows(imp, 5))
    add_heading(doc, "2. 分 biome 的 GPP-RECO 差异", 1)
    for biome in BIOMES:
        add_heading(doc, biome, 2)
        add_para(doc, f"{biome} 中最明显的贡献差异为：{diff_sentence(imp, biome)}。")
        if biome == "Forest":
            add_para(doc, "Forest 中 GPP 和 RECO 都由 SSRD_z 和 TMP_resid 主导，但 RECO 中 VPD_resid 更突出，说明森林呼吸恢复比光合恢复更敏感于独立大气干旱背景。")
            add_para(
                doc,
                "森林系统具有较强的冠层缓冲和深根水分调节能力，GPP 的恢复往往不会只受单一气象因子控制。SSRD_z 和 TMP_resid 同时靠前，说明森林光合恢复需要充足能量输入，"
                "但过强热背景也可能通过水分胁迫改变恢复速度。RECO 中 VPD_resid 更突出，可能反映高大气干旱背景下土壤-根系水分状态和呼吸底物可利用性的共同变化。"
            )
            add_para(
                doc,
                "Forest 的 dependence plot 显示，GPP 与 RECO 的关键差异不只体现在重要性排序上，也体现在响应转折位置上。EVA_resid 在 GPP 中呈现明显分段：当特征值低于约 -0.5 时，"
                "SHAP 值快速上升，说明低蒸散残差区间内蒸散重启对光合恢复的边际贡献很强；在 -0.5 至 0 附近，曲线基本进入平台期，表明中等蒸散残差下 GPP 恢复不再明显增加；"
                "超过 0 后才继续缓慢上升。相比之下，RECO 中 EVA_resid 基本在 -1 之后持续上升，说明森林呼吸恢复对蒸散/水热条件改善的响应更连续，没有 GPP 那样明显的平台段。"
            )
            add_para(
                doc,
                "Pre_z 的差异也很清楚。Forest-GPP 中，Pre_z 在 -1 至 0 之间对 SHAP 的影响基本平缓，超过 0 后 SHAP 值下降并进入负值区间，随后持续下降，说明较高降水残差可能对应光照受限、湿冷背景或非最优水分条件，"
                "从而不再促进 GPP 恢复。Forest-RECO 中，Pre_z 在约 -0.2 至 0.15 附近呈持续上升，之后仅轻微下降，SHAP 值整体接近 0，说明降水对森林呼吸恢复的边际作用较弱且更接近中性。"
                "这表明森林中降水对 GPP 和 RECO 的作用方向并不完全相同：GPP 更可能受到过湿或低辐射条件抑制，而 RECO 更接近轻微水分激活后趋于稳定。"
            )
        elif biome == "Grassland":
            add_para(doc, "Grassland 中 RECO 的 TMP_resid 高于 GPP，表明草地呼吸恢复对独立温度效应更敏感；GPP 中 VPD_resid 的重要性则强调草地光合恢复受气孔调节和大气干旱约束。")
            add_para(
                doc,
                "草地根系较浅、土壤水分周转快，因此 GPP 对 VPD_resid 和水分亏缺非常敏感；当大气干旱增强时，气孔关闭会直接限制光合作用恢复。"
                "RECO 对 TMP_resid 更敏感则符合呼吸过程的温度依赖性：在土壤水分允许的条件下，温度升高会加速根系和微生物代谢，使呼吸恢复更快进入活跃状态。"
            )
            add_para(
                doc,
                "Grassland 中 TMP_resid 的 GPP-RECO 差异非常典型。GPP 的 TMP_resid 在约 -0.5 之前持续上升，说明从偏低温度残差向中等温度残差过渡时，光合恢复受温度改善促进；"
                "但在 -0.25 至 0 附近基本不变，0 之后也只是缓慢上升，最高 SHAP 约为 10，表明草地 GPP 在达到适宜温度后逐渐受到水分或气孔限制，温度继续升高的边际收益有限。"
                "RECO 则在约 -1 之后持续快速上升，最高 SHAP 可接近 20，说明草地呼吸恢复对温度残差的响应更强、更近似连续升温激活。"
            )
            add_para(
                doc,
                "Duration_z 的转折进一步说明两类碳通量对事件记忆的读取不同。Grassland-GPP 中，Duration_z 在 -0.25 之前已经开始上升，-0.25 之后进入正 SHAP 区间并趋于平缓，"
                "说明较短到中等持续时间的变化就足以改变光合恢复状态，之后边际影响减弱。RECO 中 Duration_z 在 0.25 之前基本保持平缓，超过 0.25 后才上升到正值，"
                "说明草地呼吸恢复对持续时间的响应具有更晚的阈值，只有当闪旱持续时间足够长、土壤水分和底物过程受到更深影响时，RECO 恢复时间才明显增加。"
            )
        elif biome == "Savanna":
            add_para(doc, "Savanna 中两类通量都强烈受 SSRD_z 控制，但 RECO 的 STRD_resid 和 TMP_resid 更突出，说明呼吸恢复更容易体现热量负荷和暖背景效应。")
            add_para(
                doc,
                "稀树草原处在能量充足但水分受限的边缘环境中，因此 SSRD_z 的共同高贡献反映了强能量背景的重要性。"
                "但 RECO 对 STRD_resid 和 TMP_resid 更敏感，说明呼吸恢复更依赖夜间和近地面热背景、土壤温度以及降水后微生物活化。"
                "GPP 则更容易被水分亏缺和气孔限制截断，即使能量充足，也不一定线性促进光合恢复。"
            )
            add_para(
                doc,
                "Savanna 的 dependence plot 中，GPP 与 RECO 的主要曲线形态整体较为一致，说明稀树草原中两类碳通量更多共享同一套能量-水分耦合约束。"
                "在热干边缘系统中，SSRD、TMP 和 STRD 的变化往往同时改变冠层光合条件、土壤热环境和水分蒸发需求，因此 GPP 与 RECO 不容易像森林或灌丛那样出现明显分离的转折位置。"
                "这种一致性支持将 Savanna 解释为能量输入和水分限制共同控制的系统：GPP 和 RECO 的过程不同，但它们面对的是相同的热干背景阈值。"
            )
        elif biome == "Cropland":
            add_para(doc, "Cropland 中 GPP 和 RECO 都保留 EVA_resid 的高贡献，说明农田碳通量恢复高度依赖蒸散重启；但 RECO 中 TMP_resid 更高，提示土壤呼吸和根系活动对温度残差更敏感。")
            add_para(
                doc,
                "农田系统受作物生育期、灌溉、耕作和土壤管理影响较强，因此 EVA_resid 同时影响 GPP 和 RECO 是合理的：蒸散恢复往往意味着冠层功能、水分供应和能量交换重新建立。"
                "RECO 中 TMP_resid 的增强说明，在水分和底物条件恢复后，土壤呼吸及根系呼吸会更直接响应独立温度分量。"
            )
            add_para(
                doc,
                "Cropland 中 EVA_resid 的转折位置体现出 GPP 与 RECO 对蒸散重启的敏感阈值不同。GPP 的 EVA_resid 在特征值约 -1.5 左右就上升到正 SHAP 区间，说明农田光合恢复对蒸散残差的改善非常敏感，"
                "较低水平的蒸散恢复就可能对应冠层功能恢复和水分输送增强。RECO 中 EVA_resid 约在 -1 左右才进入正值，表明呼吸恢复对蒸散/水热恢复的阈值更高，"
                "需要更充分的土壤湿润、根系活化或微生物活动恢复后，RECO 才明显延长或增强响应。"
            )
            add_para(
                doc,
                "SMrz_resid 的差异进一步说明农田中水分过程对 GPP 更直接。GPP 中 SMrz_resid 在 -1 至 -0.5 之间快速上升，-0.5 之后在正 SHAP 区间缓慢上升，"
                "说明根区土壤水分残差从偏低到中等水平时对光合恢复具有明显促进或延长作用，之后边际变化减弱。RECO 中 SMrz_resid 基本平缓，说明在农田呼吸恢复中，"
                "土壤水分残差的独立作用可能被温度、蒸散重启、作物管理或底物供应吸收，未表现出与 GPP 同样清晰的分段响应。"
            )
        elif biome == "Shrubland":
            add_para(doc, "Shrubland 中 RECO 的 Pre_z 更突出，而 GPP 中 Duration_z 更靠前，说明灌丛光合恢复更受事件持续记忆影响，呼吸恢复则更直接受降水补给激活。")
            add_para(
                doc,
                "灌丛通常具有较强耐旱策略和较慢的冠层恢复过程，因此 GPP 对 Duration_z 更敏感，反映闪旱持续时间造成的水分亏缺累积和恢复滞后。"
                "RECO 对 Pre_z 更突出，则说明降水补给能够较快激活土壤湿润、根系活动和微生物分解过程，呼吸恢复对降水脉冲的响应可能比光合恢复更直接。"
            )
            add_para(
                doc,
                "Shrubland 的 GPP-RECO 差异最复杂。SSRD_z 在两类通量中都表现为 -1.5 至 -0.5 区间上升，说明低到中等短波辐射残差的改善对两类恢复都有促进作用；"
                "但在 -0.5 之后二者分化明显：GPP 在 -0.5 至 0.5 之间下降，随后再次上升并进入正值区间，呈现 U 型或分段型响应，可能反映灌丛光合恢复同时受光能促进和高辐射干旱胁迫约束；"
                "RECO 则在 -0.5 至 1 之间下降，之后趋于平缓，说明灌丛呼吸恢复对高 SSRD 背景并不持续增强，可能受土壤水分消耗或热干抑制限制。"
            )
            add_para(
                doc,
                "EVA_resid 和 VPD_resid 也显示出明显的 GPP-RECO 分离。GPP 的 EVA_resid 在约 0.7 附近开始下降，说明当蒸散残差过高时，光合恢复可能由水分可利用转向大气需求过强或植被耗水压力增加，"
                "从而出现边际效应下降；RECO 中 EVA_resid 更趋于平缓，说明呼吸过程对这一高蒸散残差区间的响应较弱。VPD_resid 在 GPP 中 0 之后仍继续上升，"
                "表明灌丛光合恢复时间可能随独立大气干旱增强而进一步延长；而 RECO 在 0 至 1 之间平缓，1 至 3 继续下降，说明极高 VPD 残差可能通过土壤干燥或微生物受限抑制呼吸恢复。"
            )
            add_para(
                doc,
                "Duration_z 在 Shrubland 中与 Grassland 类似，也表现出 GPP 阈值早、RECO 阈值晚的差异。GPP 在 -0.25 之前已经上升，-0.25 之后进入正值并趋于平缓，"
                "说明灌丛光合恢复对闪旱持续时间的记忆效应较早出现；RECO 在 0.25 之前基本不变，超过 0.25 后才上升到正值，说明呼吸恢复只有在事件持续足够长、土壤水分亏缺和底物过程被显著改变后才表现出明显延长。"
            )
        add_para(doc, f"{biome} 的 dependence plot 转折差异概括为：{dependence_compare_sentence(imp, biome)}")
        add_table(
            doc,
            ["特征", "GPP方向", "GPP转折", "RECO方向", "RECO转折", "方向关系", "转折差"],
            dependence_compare_rows(imp, biome, 5),
        )
        add_image(doc, FIG / "combined_by_biome" / f"{biome}_orthogonal_all_features_gpp_vs_reco.png", f"图 {BIOMES.index(biome)+3}. {biome} 中 GPP 与 RECO 的正交分解 dependence plot 对比。", 6.4)
    add_heading(doc, "3. 与原始 GPP-RECO 差异分析的关系", 1)
    add_para(doc, "原始 04 文档的核心判断是：GPP 与 RECO 共享水热背景，但二者通过不同生理和生物地球化学过程读取这些背景变量。正交分解结果支持这一判断。SSRD_z 的共同首位说明共同能量背景非常稳健；TMP_resid、STRD_resid、VPD_resid、EVA_resid 和 Pre_z 的差异排序，则说明两类碳通量在独立热效应、大气干旱、蒸散残差和降水补给方面仍存在清晰分化。")
    add_para(doc, "因此，正交分解后的 GPP-RECO 对比可以作为原始分析的共线性稳健性证据：原始图件负责解释真实物理变量和阈值，正交图件负责证明这些差异不是由变量相关性单独驱动。")
    add_heading(doc, "4. 综合总结", 1)
    add_para(
        doc,
        "综合 beeswarm、百分比柱状图和 dependence plot 可以得到三点结论。第一，SSRD_z 是 GPP 和 RECO 恢复时间最稳定的共同主控分量，说明短波辐射所代表的能量背景是两类碳通量恢复的共同底座。"
        "第二，GPP 与 RECO 的差异主要出现在第二梯队变量：GPP 更强调大气干旱、蒸散重启和事件持续记忆，RECO 更强调独立温度、长波热背景和降水补给。"
        "第三，dependence plot 的转折差异说明二者不仅重要性排序不同，响应形态也不同：同一特征在 GPP 中可能表现为阈值型光合限制，在 RECO 中则表现为温度或水分激活型响应。"
    )
    add_para(
        doc,
        "从生态过程角度看，这种分化符合 GPP 与 RECO 的机制差异。GPP 是植被冠层恢复的表征，受光能、气孔调节、叶片水分状态和事件持续胁迫控制；"
        "RECO 是生态系统呼吸恢复的表征，除植被自养呼吸外，还包含土壤微生物和有机质分解过程，因此更容易受到温度背景、土壤湿润和降水脉冲激活。"
        "因此，不能简单把 GPP 和 RECO 看作对同一气候因子的同向响应，而应将其理解为共享能量背景下的两个不同生理-生物地球化学通道。"
    )
    add_para(
        doc,
        "需要强调的是，正交分解结果并不替代原始物理单位的 dependence plot。原始图件用于讨论 SSRD、TMP、Pre、VPD 等变量在真实单位下的生态阈值；"
        "正交图件用于证明这些机制在降低共线性后仍然存在。若二者方向一致，说明机制稳健；若正交图中阈值位置发生移动，则说明原始阈值可能受到共线变量共享信息影响，需要在论文中表述为“相对阈值”或“正交空间中的转折”。"
    )
    out = OUT / "04_orthogonal_gpp_reco_difference_analysis_cn.docx"
    doc.save(out)
    return out


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    imp = load_importance()
    paths = [build_doc01(imp), build_doc02(imp), build_doc04(imp)]
    for p in paths:
        print(p)


if __name__ == "__main__":
    main()
