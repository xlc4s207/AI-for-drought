#!/usr/bin/env python3
"""Create a Chinese report for grouped PCA and orthogonal SHAP analyses."""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path("/home/xulc/flash_drought")
BASE = ROOT / "process/SEM_analysis0401/codex/GLEAM/plots2/prepeak_shap_nomulticollinearity"
OUT_DIR = ROOT / "process/SEM_analysis0401/codex/GLEAM/writing4"
DOCX_OUT = OUT_DIR / "11_groupPCA_orthogonal_SHAP_collinearity_robustness_analysis_cn.docx"
MD_OUT = OUT_DIR / "11_groupPCA_orthogonal_SHAP_collinearity_robustness_analysis_cn.md"


BIOMES = ["Forest", "Grassland", "Savanna", "Cropland", "Shrubland"]


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(text)


def add_table(doc: Document, df: pd.DataFrame, columns: list[str], headers: list[str], digits: int = 3) -> None:
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(columns):
            val = row[col]
            if isinstance(val, float):
                cells[i].text = f"{val:.{digits}f}"
            else:
                cells[i].text = str(val)


def add_picture(doc: Document, path: Path, caption: str, width: float = 6.3) -> None:
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def top_features_text(imp: pd.DataFrame, method: str, metric: str, biome: str, n: int = 3) -> str:
    sub = imp[(imp["method"] == method) & (imp["metric"] == metric) & (imp["biome"] == biome)].sort_values("rank").head(n)
    return "、".join(f"{r.display_label}（{r.percent:.1f}%）" for _, r in sub.iterrows())


def build_summary_tables() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    model = pd.read_csv(BASE / "nomulticollinearity_model_summary.csv")
    vif = pd.read_csv(BASE / "nomulticollinearity_vif_summary.csv")
    imp = pd.read_csv(BASE / "nomulticollinearity_feature_importance_all.csv")

    r2_summary = (
        model.groupby(["method", "metric"])["r2_holdout_split"]
        .agg(mean="mean", min="min", max="max")
        .reset_index()
        .replace({"group_pca": "分组PCA", "orthogonal_decomposition": "正交分解"})
    )
    vif_summary = (
        vif.groupby("method")["max_vif"]
        .agg(mean="mean", max="max")
        .reset_index()
        .replace({"group_pca": "分组PCA", "orthogonal_decomposition": "正交分解"})
    )
    mean_imp = (
        imp.groupby(["method", "metric", "display_label"])["percent"]
        .mean()
        .reset_index()
        .sort_values(["method", "metric", "percent"], ascending=[True, True, False])
        .groupby(["method", "metric"])
        .head(5)
        .replace({"group_pca": "分组PCA", "orthogonal_decomposition": "正交分解"})
    )
    return model, vif, imp, r2_summary, vif_summary, mean_imp


def build_docx() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    model, vif, imp, r2_summary, vif_summary, mean_imp = build_summary_tables()

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10.5)

    title = doc.add_heading("11 分组 PCA 与正交分解 SHAP 的共线性稳健性分析", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("一、分析目的与定位", level=1)
    add_para(doc, "前序 01、02 和 04 文档分别从 beeswarm 贡献格局、dependence plot 阈值响应以及 GPP-RECO 差异角度解释了 pre-event SHAP 结果；07 文档进一步将 SHAP、OPGD 和 SEM 连接成“预测贡献—空间解释力—路径机制”的证据链。上述分析已经说明 SSRD、EVA、TMP、STRD、SMrz、VPD、PRE、Wind 以及事件属性共同构成恢复时间的主要解释框架。不过，气象和生态水文变量天然存在共线性，特别是 SSRD、STRD、TMP 和 VPD 之间具有明确的热辐射耦合关系。如果不补充降共线性分析，审稿人可能会质疑单变量 SHAP 排名是否只是相关变量之间的贡献分摊。")
    add_para(doc, "因此，本文件基于同一套 prepeak 输入数据，限定使用 SSRD、EVA、TMP、STRD、SMrz、Wind、VPD、Duration、Pre、Intensity 这十个变量，重新构建两套低共线性 SHAP 分析。第一套为分组 PCA，将变量压缩为辐射热力、水分供给/储存、大气需水和事件属性四个机制主轴；第二套为正交分解，保留 SSRD、Pre、Duration、Intensity 和 Wind 等锚点变量，同时将 STRD、TMP、VPD、EVA 和 SMrz 转换为控制上游变量后的残差项。两套结果不是为了替代原始 SHAP，而是用于回答一个更稳健的问题：在显著降低共线性后，前序文档提出的能量—水分—大气干旱机制是否仍然成立。")

    doc.add_heading("二、方法概述：两种降共线性版本的含义", level=1)
    add_para(doc, "分组 PCA 版本把原始变量按生态物理意义分成四组：Energy_PC1 由 SSRD、STRD 和 TMP 构成，代表辐射输入与热量背景；Water_PC1 由 Pre、EVA 和 SMrz 构成，代表降水补给、蒸散耗散和根区储水的综合状态；AtmosDemand_PC1 由 VPD 和 Wind 构成，代表大气需水和湍流交换条件；Event_PC1 由 Duration 和 Intensity 构成，代表骤旱事件过程本身。每组仅保留第一主成分，因此该版本具有很强的降维和降共线性效果，但会牺牲一部分细粒度预测信息。")
    add_para(doc, "正交分解版本则更接近原始变量解释。它保留 SSRD 作为短波辐射锚点，保留 Pre 作为水分补给锚点，保留 Duration 和 Intensity 作为事件属性锚点，保留 Wind 作为交换条件锚点；随后依次构造 STRD_resid_after_SSRD、TMP_resid_after_SSRD_STRD、VPD_resid_after_SSRD_TMP_Wind、EVA_resid_after_SSRD_Pre_VPD 和 SMrz_resid_after_Pre_EVA。这些残差项表示控制上游变量之后仍然保留的独立变化，因此更适合回答“在去掉共线解释后，某个过程是否仍然重要”。")

    doc.add_heading("三、共线性控制是否有效", level=1)
    add_para(doc, "两套变换均显著降低了输入特征的 VIF。分组 PCA 版本中，各 metric-biome 组合的最大 VIF 均低于 2.43；正交分解版本中，最大 VIF 为 3.89，仍明显低于常用警戒线。相比原始诊断中 GPP_prepeak 的 TMP VIF 可达 38.08，这说明本轮变换已经从输入层面实质性缓解了共线性问题。")
    add_table(doc, vif_summary, ["method", "mean", "max"], ["方法", "最大VIF均值", "最大VIF上限"], digits=3)
    add_para(doc, "需要注意的是，两种方法的目标不同。分组 PCA 的最大 VIF 更低，适合证明机制组层面的结论稳健；正交分解的最大 VIF 略高但仍处于较安全范围，同时保留了更多变量信息，因此更适合与原始 SHAP、dependence plot 和 SEM 路径解释衔接。")

    doc.add_heading("四、模型解释力的变化", level=1)
    add_para(doc, "从 holdout R² 看，分组 PCA 由于只保留四个机制主轴，解释力低于正交分解版本。GPP 的分组 PCA 平均 holdout R² 为 0.264，RECO 为 0.313；正交分解版本中，GPP 平均 holdout R² 提高到 0.440，RECO 提高到 0.516。这一差异并不意味着分组 PCA 无效，而是说明恢复时间不仅受机制组主轴控制，还包含组内的细节差异，例如 SSRD 与 STRD 的分化、TMP 的独立热量异常、VPD 的独立大气干燥异常，以及 EVA/SMrz 在水分链条中的相对位置。")
    add_table(doc, r2_summary, ["method", "metric", "mean", "min", "max"], ["方法", "指标", "holdout R²均值", "最小值", "最大值"], digits=3)
    add_para(doc, "因此，在论文中更稳妥的写法是：分组 PCA 用于验证“机制组排序”的稳健性，正交分解用于验证“关键变量及其净效应”的稳健性。两者共同说明，原始 SHAP 的主结论不是共线性造成的偶然排序。")

    doc.add_heading("五、分组 PCA 结果：机制组层面的稳健结论", level=1)
    add_picture(doc, BASE / "group_pca_importance_percent_bars_5biomes_gpp_vs_reco.png", "图 1. 分组 PCA 版本中 GPP 与 RECO 在五类 biome 的机制主轴 SHAP 贡献占比。")
    add_para(doc, "分组 PCA 结果显示，Energy_PC1 是最稳定的主导机制轴。GPP 中，五个 biome 的首位特征均为 Energy_PC1，平均贡献约 46.4%；RECO 中，除 Cropland 由 Water_PC1 位列首位外，其余四个 biome 也均由 Energy_PC1 主导，平均贡献约 47.4%。这与 01 文档中 SSRD、STRD 和 TMP 共同靠前的单变量格局高度一致，也与 07 文档中“短波辐射—热量状态—大气干旱—蒸散”的机制链条一致。")
    add_para(doc, "从 loadings 看，Energy_PC1 在各 biome 中均由 SSRD、STRD 和 TMP 同向加载构成，解释率通常在 0.72-0.91 之间。这说明短波辐射、长波辐射和气温在统计上确实共享一个强热辐射背景，因此原始 SHAP 中 SSRD、STRD、TMP 同时重要时，不应简单理解为三个完全独立变量，而应理解为能量与热状态共同控制恢复时间。分组 PCA 的结果支持这种机制组解释。")
    add_para(doc, "GPP 的 biome 差异仍然清楚。Forest 中 Energy_PC1 贡献最高，达到 57.8%，Water_PC1 为 24.2%，说明森林恢复主要受辐射热力背景控制，但水分供给和蒸散状态仍是第二约束。Grassland 中 Energy_PC1 为 38.4%，AtmosDemand_PC1 为 29.8%，Event_PC1 为 18.4%，表明草地恢复在能量背景之外更容易受到 VPD/Wind 所代表的大气需水和事件过程共同调节。Savanna 中 Energy_PC1 达到 52.3%，说明稀树草原的恢复仍是典型热干边缘生态系统的能量—水分耦合响应。Cropland 中 Energy_PC1 和 Water_PC1 分别为 39.3% 和 35.8%，显示农田恢复更依赖水热匹配，而不是单纯辐射或单纯水分。Shrubland 中 Energy_PC1、AtmosDemand_PC1 和 Event_PC1 分别为 44.2%、22.6% 和 22.4%，说明灌丛恢复同时保留热干背景和事件记忆。")
    add_para(doc, "RECO 的分组 PCA 结果进一步强化了 GPP-RECO 差异。Forest、Grassland、Savanna 和 Shrubland 中 Energy_PC1 仍为首位，但 Shrubland 的 AtmosDemand_PC1 贡献高达 36.7%，明显高于 GPP 中的 22.6%，说明灌丛呼吸恢复对大气干燥和交换条件更敏感。Cropland 中 Water_PC1 位列首位，贡献为 43.4%，高于 Energy_PC1 的 33.0%，说明农田 RECO 恢复更强烈地受到蒸散—根区水分—管理补水背景的约束。这与 04 文档中 Cropland 的 VPD、SMrz 和 |EVA| 分化，以及 07 文档中 RECO 对空间水热分层更敏感的判断相一致。")

    doc.add_heading("六、正交分解结果：控制共线性后 SSRD 仍然稳定主导", level=1)
    add_picture(doc, BASE / "orthogonal_decomposition_importance_percent_bars_5biomes_gpp_vs_reco.png", "图 2. 正交分解版本中 GPP 与 RECO 在五类 biome 的 SHAP 贡献占比。")
    add_para(doc, "正交分解版本给出了更强的稳健性证据：在所有 GPP 和 RECO biome 中，SSRD 都是首位变量。GPP 中 SSRD 的平均贡献为 26.7%，RECO 中为 28.8%。这说明即使控制 STRD、TMP、VPD 等相关变量的共享信息，短波辐射本身仍然保留了独立预测贡献。换言之，原始 SHAP 中 SSRD 的重要性并不是完全由 TMP 或 STRD 共线性“带出来”的，而是具有独立的恢复时间解释价值。")
    add_para(doc, "控制 SSRD 和 STRD 后，TMP_resid 仍然在多数组合中位列前列。GPP 中 TMP_resid 的平均贡献为 13.8%，RECO 中为 15.7%，说明温度不仅是辐射背景的统计表现，也包含独立的热量异常信息。这个结果可以与 02 文档中的 TMP dependence 分段阈值联系起来：温度的作用不应被简单写成与辐射完全重复，而是代表在辐射背景之外的生理温度窗口、蒸散需求和代谢激活差异。")
    add_para(doc, "STRD_resid 在 GPP 和 RECO 中也保持较高贡献，平均贡献分别约为 12.7% 和 12.1%。这表明长波辐射中不被 SSRD 解释的部分仍具有独立意义，尤其在 Savanna 和 Shrubland 中较突出。结合前序文档的解释，SSRD 更接近光合可用能量，STRD_resid 更接近近地表热储存和大气保温效应。因此，正交分解结果支持“短波光能输入”和“长波热负荷”需要区分解释。")
    add_para(doc, "VPD_resid 和 EVA_resid 的保留贡献说明，大气需水和蒸散过程并没有在控制热辐射背景后消失。GPP 中 VPD_resid 和 EVA_resid 的平均贡献分别为 9.6% 和 8.4%；RECO 中分别为 9.0% 和 8.0%。这与 SEM 中 TMP -> VPD -> EVA 以及 SMrz/Pre -> EVA 的中介路径相互呼应：VPD 和 EVA 不只是 TMP 或 SSRD 的附属变量，而是将热量背景转化为水分亏缺、蒸散活动和恢复时间差异的重要中介过程。")

    doc.add_heading("七、分 biome 综合解释", level=1)
    for biome in BIOMES:
        add_para(
            doc,
            f"{biome} 中，分组 PCA 的 GPP 前三位为 {top_features_text(imp, 'group_pca', 'GPP', biome)}，RECO 前三位为 {top_features_text(imp, 'group_pca', 'RECO', biome)}；正交分解的 GPP 前三位为 {top_features_text(imp, 'orthogonal_decomposition', 'GPP', biome)}，RECO 前三位为 {top_features_text(imp, 'orthogonal_decomposition', 'RECO', biome)}。这表明该 biome 的原始 SHAP 结论在降共线性后仍有清晰延续：机制组层面以能量或水分主轴为核心，净效应层面则由 SSRD 及其下游热力、长波、VPD 或蒸散残差共同解释恢复时间。"
        )
    add_para(doc, "上述逐 biome 结果可进一步压缩为五点。Forest 表现为能量主轴稳定主导，但水分主轴仍是第二限制，符合森林冠层恢复需要辐射和根区供水共同满足的机制。Grassland 在 GPP 和 RECO 中均保留较强的能量主轴，同时 AtmosDemand 和事件轴贡献上升，说明草地浅根系和快速冠层周转使其更容易受到大气干旱和事件累积过程影响。Savanna 的 Energy_PC1 和 SSRD 贡献最高，支持热干边缘生态系统中能量—水分耦合主导恢复的判断。Cropland 的水分主轴和 EVA_resid 贡献突出，说明农田恢复需要放在蒸散、根区水分和管理补水背景中解释。Shrubland 中 AtmosDemand 和 Event 贡献较高，说明灌丛恢复不仅取决于后期水热状态，也保留较强事件记忆。")

    doc.add_heading("八、与原始 SHAP、Dependence Plot、OPGD 和 SEM 的关系", level=1)
    add_para(doc, "本轮降共线性分析并不推翻原始 SHAP 图，而是为其提供稳健性支撑。原始 beeswarm 图可以继续用于展示具体变量的贡献分布和高低值方向；分组 PCA 图用于说明这些变量背后共同指向哪些机制组；正交分解图则用于说明在控制共线性之后，哪些变量仍保留净贡献。三者结合后，可以避免把 SHAP 排名误写成完全独立因果效应，也能避免因为共线性而简单删除 SSRD、TMP、STRD 或 VPD 等具有明确物理意义的变量。")
    add_para(doc, "与 dependence plot 的关系上，分组 PCA 提醒我们，TMP、SSRD、STRD 和 VPD 的阈值不应孤立解释，而应看作热辐射和大气干燥背景的联合非线性响应；正交分解则进一步说明，某些变量在控制上游后仍有独立阈值解释价值。例如 SSRD 在正交版本中稳定首位，说明 SSRD dependence plot 中关于低辐射限制解除和高辐射水分调制的解释仍然必要；TMP_resid、VPD_resid 和 EVA_resid 的贡献则支持将温度窗口、大气需水和蒸散重启作为二级机制解释。")
    add_para(doc, "与 OPGD 和 SEM 的关系上，分组 PCA 对应 OPGD 中的空间分层机制组，正交分解则对应 SEM 中的路径分层思想。07 文档中提出的 STRD/SSRD -> TMP、TMP/Wind -> VPD、VPD/Pre/SMrz -> EVA、EVA/SSRD/TMP/事件属性 -> 恢复时间的机制链，在本轮结果中得到了输入层面的支持。尤其是 SSRD 在正交分解后仍为所有 biome 的首位变量，说明 SEM 中保留 SSRD 直接路径是合理的；而 TMP_resid、VPD_resid 和 EVA_resid 的持续贡献，则说明 SEM 不应只把这些变量作为共线冗余项删除，而应通过上游—中介—终点路径表达其机制作用。")

    doc.add_heading("九、可写入论文的结论段", level=1)
    add_para(doc, "为检验输入变量共线性是否影响 SHAP 解释，本研究进一步构建了分组 PCA 和正交分解两套低共线性 SHAP 模型。分组 PCA 将十个预测因子压缩为辐射热力、水分供给/储存、大气需水和事件属性四个机制主轴，所有组合的最大 VIF 均低于 2.43；正交分解在保留 SSRD 等关键变量的同时构造控制上游变量后的残差项，最大 VIF 低于 3.89。结果显示，辐射热力主轴在多数 GPP 和 RECO biome 中保持首位，正交分解后 SSRD 在所有 biome 和两个碳通量指标中均为最高贡献变量。这表明原始 SHAP 中 SSRD 及能量相关变量的重要性并非单纯由共线性造成，而是反映了干旱恢复过程中稳定的能量背景控制。同时，TMP、STRD、VPD、EVA 和 SMrz 的残差贡献说明热量、大气需水、蒸散和根区水分仍通过独立信息通道调节恢复时间。上述结果支持将原始 SHAP 的单变量解释提升为“能量输入—大气干旱—蒸散过程—根区水分—事件记忆”的综合机制框架，并为后续 OPGD 和 SEM 路径分析提供了共线性稳健性证据。")

    doc.save(DOCX_OUT)


def build_md() -> None:
    model, vif, imp, r2_summary, vif_summary, mean_imp = build_summary_tables()
    lines = [
        "# 11 分组 PCA 与正交分解 SHAP 的共线性稳健性分析",
        "",
        "本文件分析 `prepeak_shap_nomulticollinearity` 中的分组 PCA 和正交分解 SHAP 结果。核心结论是：降共线性后，辐射热力机制仍是多数 biome 的主导机制；正交分解后 SSRD 在所有 GPP/RECO biome 中仍位列首位，说明原始 SHAP 中 SSRD 的重要性不是单纯由共线性造成。",
        "",
        "## VIF 摘要",
        "",
        "```csv",
        vif_summary.to_csv(index=False).strip(),
        "```",
        "",
        "## R2 摘要",
        "",
        "```csv",
        r2_summary.to_csv(index=False).strip(),
        "```",
        "",
        "## 平均贡献 Top5",
        "",
        "```csv",
        mean_imp.to_csv(index=False).strip(),
        "```",
        "",
        "## 论文结论",
        "",
        "分组 PCA 和正交分解共同表明，能量输入、热量状态、大气需水、蒸散和根区水分构成恢复时间的稳定解释链条。分组 PCA 证明机制组层面的排序稳定，正交分解证明 SSRD 及其下游热力/水分残差信息在控制共线性后仍保留独立贡献。",
    ]
    MD_OUT.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    build_docx()
    build_md()
    print(f"Wrote {DOCX_OUT}")
    print(f"Wrote {MD_OUT}")


if __name__ == "__main__":
    main()
