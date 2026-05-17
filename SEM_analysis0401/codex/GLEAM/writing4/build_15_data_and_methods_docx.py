#!/usr/bin/env python3
"""Build the Data and Methods manuscript section."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT_DIR = Path("/home/xulc/flash_drought/process/SEM_analysis0401/codex/GLEAM/writing4")
OUT_DOCX = OUT_DIR / "15_data_and_methods_cn.docx"


def set_run_font(run, font: str = "SimSun", size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_cell_text(cell, text: str, font: str = "SimSun", size: float = 8.5, bold: bool = False) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, font, size, bold)


def set_cell_border(cell, **kwargs: dict[str, str]) -> None:
    """Set cell borders. kwargs keys: top, bottom, left, right, insideH, insideV."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs.get(edge) or {}
            tag = "w:{}".format(edge)
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key, value in edge_data.items():
                element.set(qn(f"w:{key}"), str(value))


def clear_cell_borders(cell) -> None:
    nil = {"val": "nil", "sz": "0", "space": "0", "color": "FFFFFF"}
    set_cell_border(cell, top=nil, bottom=nil, left=nil, right=nil, insideH=nil, insideV=nil)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_para(doc: Document, text: str, first_line: bool = True, size: float = 10.5) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.18
    p.paragraph_format.space_after = Pt(5)
    if first_line:
        p.paragraph_format.first_line_indent = Pt(21)
    r = p.add_run(text)
    set_run_font(r, "SimSun", size)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(level=level)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    set_run_font(r, "SimHei", 14 if level == 1 else 12, True)


def add_formula(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(21)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run_font(r, "Times New Roman", 10)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, "SimSun", 9)


def add_three_line_table(doc: Document) -> None:
    headers = ["数据集/来源", "变量", "用途", "时间跨度", "空间范围", "空间分辨率", "时间分辨率"]
    rows = [
        [
            "GLEAM",
            "SMrz、SMs、EVA",
            "识别骤旱事件；表征根区/表层水分状态和蒸散发过程",
            "1980-2024；主恢复分析匹配 1982-2022",
            "全球陆地",
            "统一至 0.1°",
            "日尺度",
        ],
        [
            "ERA5-Land",
            "2 m 气温、短波辐射、长波辐射、风速、2 m 露点温度",
            "构建热量、辐射和大气干燥度驱动因子；由气温和露点温度计算 VPD",
            "1982-2022；气候百分位和辅助诊断使用可用全时段",
            "全球陆地",
            "统一至 0.1°",
            "日尺度",
        ],
        [
            "MSWEP",
            "降水量",
            "表征事件前后水分输入和降水异常",
            "1982-2022",
            "全球陆地",
            "统一至 0.1°",
            "日尺度",
        ],
        [
            "MODIS MCD12",
            "土地利用/土地覆盖类型",
            "剔除非目标地表并划分 Cropland、Forest、Grassland、Savanna、Shrubland 等 biome",
            "基准土地覆盖年 2010",
            "全球陆地",
            "500 m；分析时重采样/聚合至统一网格",
            "年度/静态分类",
        ],
        [
            "BESS",
            "GPP、RECO",
            "主分析碳通量数据；计算 GPP 和 RECO 响应、峰值损伤和恢复时间",
            "1982-2022",
            "全球陆地",
            "0.1°",
            "日尺度",
        ],
        [
            "FluxSat",
            "GPP",
            "独立验证 BESS-GPP 响应和恢复时间结论",
            "2000-2019；有效验证期主要使用 2001-2018/2019 重叠时段",
            "全球陆地",
            "原始产品重采样至 0.25°，用于独立验证",
            "日尺度",
        ],
    ]

    add_caption(doc, "表 1  本研究使用的数据集及其用途")
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = True
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, header, "SimHei", 8.5, True)
        shade_cell(cell, "F2F2F2")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, "SimSun", 8, False)

    border_top = {"val": "single", "sz": "12", "space": "0", "color": "000000"}
    border_mid = {"val": "single", "sz": "8", "space": "0", "color": "000000"}
    border_bottom = {"val": "single", "sz": "12", "space": "0", "color": "000000"}
    nil = {"val": "nil", "sz": "0", "space": "0", "color": "FFFFFF"}
    for row in table.rows:
        for cell in row.cells:
            clear_cell_borders(cell)
            set_cell_border(cell, left=nil, right=nil)
    for cell in table.rows[0].cells:
        set_cell_border(cell, top=border_top, bottom=border_mid, left=nil, right=nil)
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom=border_bottom, left=nil, right=nil)
    doc.add_paragraph()


def build_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.start_type = WD_SECTION.NEW_PAGE
    doc.styles["Normal"].font.name = "SimSun"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    doc.styles["Normal"].font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Data and Methods")
    set_run_font(r, "Times New Roman", 18, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("骤旱恢复时间及其环境驱动机制的数据与方法")
    set_run_font(r, "SimHei", 16, True)

    add_heading(doc, "1 数据来源与预处理", 1)
    add_para(
        doc,
        "本研究面向全球陆地生态系统，构建了由土壤水分、气象强迫、碳通量和土地覆盖共同组成的日尺度事件数据库。数据组织遵循两个原则：第一，骤旱识别必须依赖能够反映植被可利用水分状态的土壤水分指标，因此以 GLEAM 根区土壤水分（SMrz）为主，同时保留表层土壤水分（SMs）作为定义敏感性和数据一致性检验；第二，碳循环响应和恢复时间必须在同一事件框架下比较 GPP 与 RECO，因此以 BESS 日尺度 GPP 和 RECO 作为主分析数据，并使用 FluxSat GPP 对 GPP 恢复结果进行独立验证。所有连续环境变量均被重投影、裁剪并重采样到统一分析网格，随后按事件发生位置和时间窗口提取事件前、骤旱发展期和恢复期的统计量。",
    )
    add_para(
        doc,
        "数据表中列出的 0.1° 表示本研究最终分析使用的统一空间网格，而不必然代表所有原始产品的原生分辨率。对于 MODIS 土地覆盖数据，原始 500 m 分类产品被用于确定像元主导土地利用类型，并在统一网格上生成 biome 掩膜；对于 FluxSat GPP，数据主要用于与 BESS-GPP 在重叠时段内进行恢复时间和趋势一致性验证，因此保持独立验证产品的重采样分辨率。土地覆盖分类剔除了冰雪、裸地、水体和城市等不适合解释植被碳恢复过程的像元，并将分析集中在 Cropland、Forest、Grassland、Savanna 和 Shrubland 五类主要生态系统。",
    )
    add_three_line_table(doc)

    add_heading(doc, "2 骤旱识别方法", 1)
    add_para(
        doc,
        "骤旱事件采用基于土壤水分百分位的两步识别框架。该框架与 Lu 等（2025）等近期研究中使用的百分位思想一致，即以土壤水分从相对正常状态快速跌落到干旱状态作为骤旱的核心判据。为削弱逐日土壤水分噪声并与经典骤旱研究中的 pentad 口径保持一致，本研究在事件识别阶段以 5 天为一个时间步长（pentad）刻画土壤水分变化。具体而言，先以 1981-2010 年为气候基准期，针对每个像元和每个 pentad 计算土壤水分百分位阈值，并对阈值序列进行平滑处理，以削弱短期样本波动带来的不连续性。土壤水分低于第 20 百分位（P20）时认为进入干旱状态；当土壤水分重新达到或高于 P20，并且该状态连续维持 2 个 pentad，即 10 天时，认为骤旱事件完成水分恢复。采用连续 2 个 pentad 的恢复判据，是为了避免土壤水分在 P20 附近短期反弹导致恢复时间被过早判定。",
    )
    add_para(
        doc,
        "在确定干旱段以后，事件起始的快速发展阶段通过向前回溯到土壤水分高于第 40 百分位（P40）的相对正常状态来确定。设 S(p) 为某像元第 p 个 pentad 的土壤水分百分位，p20 为首次低于 P20 的 pentad，ponset 为 p20 之前最近一次满足 S(p)>P40 的 pentad，则骤旱发展时间可表示为 D_onset=p20-ponset。只有当土壤水分在 4 个 pentad，即 20 天内从 P40 以上快速下降到 P20 以下时，该事件才被判定为骤旱。换言之，骤旱不仅要求达到低土壤水分状态，还要求下降过程具有足够快的突发性。骤旱强度和发展速率分别由土壤水分百分位下降幅度、P20 以下亏缺累积量及单位时间下降量刻画：",
    )
    add_formula(doc, "ΔS = S(ponset) - S(p20)")
    add_formula(doc, "R_onset = ΔS / D_onset")
    add_formula(doc, "Intensity = Σ_{p=p20}^{pend} max[0, P20(p) - S(p)]")
    add_para(
        doc,
        "其中 ΔS 表示从相对正常状态跌落至干旱状态的土壤水分百分位降幅，R_onset 表示以 pentad 为单位的发展速率，pend 表示事件恢复闭合的 pentad。事件持续时间 Duration 定义为从首次低于 P20 到重新达到 P20 并连续维持 2 个 pentad 的时间长度，Intensity 则用于描述事件期间 P20 以下的土壤水分亏缺累积程度。该定义兼顾了 Lu 等（2025）强调的快速下降、低土壤水分和最低持续时间三个要素，同时保留了后续日尺度碳通量数据对生态响应和恢复过程的精细刻画能力。",
    )

    add_heading(doc, "3 碳通量响应与恢复时间", 1)
    add_para(
        doc,
        "GPP 和 RECO 的响应与恢复时间均在同一骤旱事件库上计算。与只从骤旱开始日计时的做法不同，本研究将恢复时间定义为生态系统碳通量受损达到峰值之后恢复到灾前基线所需的时间。这样处理可以避免把“干旱发展到最大影响之前的滞后响应时间”与“系统从最大损伤状态恢复的时间”混在一起。对 GPP 而言，碳通量损伤峰值通常对应 5 日平滑后绝对 GPP 的最低点；对 RECO 而言，峰值响应可能表现为呼吸受抑或异常增强，因此其恢复时间采用与 GPP 一致的事件筛选、平滑和基线判定流程，但解释时强调其代表生态系统呼吸过程回到灾前水平的速度。",
    )
    add_para(
        doc,
        "参考 Lu 等（2025）对 GPP 标准化异常和恢复过程的处理，本研究首先构建日尺度碳通量异常。若 C(t) 表示 GPP 或 RECO 的日值，μC,d 和 σC,d 分别为对应像元同历日气候平均值与标准差，则标准化异常为：",
    )
    add_formula(doc, "Z_C(t) = [C(t) - μ_C,d] / σ_C,d")
    add_para(
        doc,
        "该异常用于判定事件是否存在清晰的碳通量响应。对 GPP 响应，要求骤旱开始后一定窗口内出现持续性负异常或持续下降信号，以保证恢复时间针对真实的生产力扰动而不是随机噪声计算。恢复基线采用灾前 30 天绝对碳通量平均值，而不是事件后局部反弹值。设 B_C 为灾前 30 天 5 日平滑碳通量的平均水平，t_peak 为碳通量损伤达到峰值的日期，则恢复终点 t_rec 定义为 t_peak 之后 5 日平滑碳通量连续 5 天达到或超过 0.95×B_C 的首次日期：",
    )
    add_formula(doc, "B_C = mean[C_s(t)],  t ∈ [tonset - 30, tonset - 1]")
    add_formula(doc, "t_rec = min{t > t_peak : C_s(t+k) ≥ 0.95 B_C, k = 0,1,2,3,4}")
    add_formula(doc, "T_recovery = N_GS(t_peak, t_rec)")
    add_para(
        doc,
        "其中 C_s(t) 为 5 日平滑后的碳通量，T_recovery 为最终用于 SHAP 和 SEM 的恢复时间，N_GS(t_peak,t_rec) 表示从峰值损伤到恢复终点之间累计的生长季有效天数。需要强调的是，恢复过程可以跨越非生长季，但恢复持续时间只累计生长季内的有效天数。生长季由逐像元逐年 ERA5-Land 2 m 气温确定：当某年首次出现连续 5 天日均气温高于 278.15 K（5°C）时，记为生长季开始；当随后首次出现连续 5 天日均气温低于 278.15 K 时，生长季结束。若骤旱事件持续期与生长季重叠比例不足，则该事件不纳入碳恢复时间分析。该处理保证了恢复时间主要反映植被能够实际生长和碳交换的时期，而不是被冬季或休眠期人为拉长。",
    )

    add_heading(doc, "4 驱动因子构建", 1)
    add_para(
        doc,
        "用于解释恢复时间的候选驱动因子包括能量、热量、水分、大气干燥度和事件属性五类。能量类变量包括短波向下辐射（SSRD）和长波向下辐射（STRD），热量变量为 2 m 气温（TMP），水分变量包括降水（Pre）、根区土壤水分（SMrz）和实际蒸散发（EVA），大气干燥度由水汽压亏缺（VPD）表示，事件属性包括 Duration 和 Intensity。VPD 由 2 m 气温和 2 m 露点温度计算：",
    )
    add_formula(doc, "e_s(T) = 0.6108 exp[17.27 T / (T + 237.3)]")
    add_formula(doc, "VPD = e_s(T_air) - e_s(T_dew)")
    add_para(
        doc,
        "其中 T_air 和 T_dew 以摄氏度表示，VPD 单位为 kPa。所有环境变量按照事件时间轴提取，重点使用骤旱发生前及峰值损伤前的 prepeak 窗口，以检验事件前环境背景和早期干旱状态对后续恢复时间的可预报性。对存在强共线性的变量，本研究另行构建了正交分解与分组 PCA 版本：正交分解通过线性残差化保留每个变量相对于前序相关变量的独立信息，分组 PCA 则将能量、水分和事件属性等相关变量压缩为低维综合轴。原始 SHAP 结果用于保持生态变量可解释性，正交和 PCA 结果用于检验结论是否受到共线性归因转移的影响。",
    )

    add_heading(doc, "5 SHAP 模型解释", 1)
    add_para(
        doc,
        "本研究以机器学习模型预测恢复时间，并使用 SHAP（SHapley Additive exPlanations）量化各驱动因子对单个事件恢复时间预测值的贡献。设 f(x_i) 为模型对第 i 个事件恢复时间的预测，x_i 包含该事件的全部环境和事件属性特征，则 SHAP 将预测值分解为全样本平均预测值 φ0 与各特征贡献 φij 的加和：",
    )
    add_formula(doc, "f(x_i) = φ_0 + Σ_{j=1}^{M} φ_{ij}")
    add_para(
        doc,
        "特征 j 的 SHAP 值来自 Shapley 博弈分解，其基本形式为：",
    )
    add_formula(
        doc,
        "φ_j = Σ_{S⊆F\\{j}} [|S|!(M-|S|-1)! / M!] × [f_{S∪{j}}(x_{S∪{j}}) - f_S(x_S)]",
    )
    add_para(
        doc,
        "其中 F 为全部特征集合，M 为特征数，S 表示不包含特征 j 的任意特征子集。φij>0 表示该特征在第 i 个事件上将预测恢复时间推高，即倾向于延长恢复；φij<0 表示倾向于缩短恢复。特征总体重要性使用平均绝对 SHAP 值衡量：",
    )
    add_formula(doc, "I_j = (1/n) Σ_{i=1}^{n} |φ_{ij}|")
    add_formula(doc, "P_j = I_j / Σ_{k=1}^{M} I_k × 100%")
    add_para(
        doc,
        "其中 I_j 表示特征 j 的平均贡献强度，P_j 表示其在所有特征总贡献中的百分比。蜂巢图用于同时展示重要性排序、贡献方向和特征取值分布；dependence plot 则以特征值为横坐标、对应 SHAP 值为纵坐标，用于识别阈值、非线性响应和不同 biome 中 GPP 与 RECO 恢复机制的差异。需要注意，SHAP 解释的是模型预测中的归因结构，并不直接等同于因果效应。因此，本研究进一步结合 ALE、ICE、PDP、地理探测器和 SEM 对方向、阈值、空间解释力和路径机制进行交叉验证。",
    )

    add_heading(doc, "6 结构方程模型（SEM）", 1)
    add_para(
        doc,
        "结构方程模型用于把 SHAP 揭示的非线性重要性进一步转化为可解释的路径机制。与 SHAP 主要回答“哪个变量对预测更重要”不同，SEM 侧重回答“变量之间如何通过直接和间接路径共同影响恢复时间”。本研究根据 SHAP 重要性排序、生态水热机制和共线性诊断设定路径结构，并分别在不同 biome 与 GPP/RECO 目标上拟合标准化路径系数。一般的 SEM 可写为：",
    )
    add_formula(doc, "η = Bη + Γξ + ζ")
    add_formula(doc, "y = Λ_y η + ε_y,    x = Λ_x ξ + ε_x")
    add_para(
        doc,
        "其中 η 表示内生变量，ξ 表示外生变量，B 为内生变量之间的路径系数矩阵，Γ 为外生变量对内生变量的路径系数矩阵，ζ 和 ε 为结构残差与测量残差。在本研究的实证模型中，各变量均标准化后进入模型，恢复时间方程可表达为：",
    )
    add_formula(
        doc,
        "T_recovery = β_0 + β_1 SSRD + β_2 STRD + β_3 TMP + β_4 VPD + β_5 SMrz + β_6 EVA + β_7 Pre + β_8 Duration + β_9 Intensity + ε",
    )
    add_para(
        doc,
        "根据具体 biome 的 SHAP 结果和变量共线性，部分变量也被设置为中介变量。例如，辐射和温度可以通过 EVA、VPD 或 SMrz 间接影响恢复时间，降水可以通过 SMrz 或 EVA 改变水分可利用性。若 a 表示 X→M 的路径系数，b 表示 M→T_recovery 的路径系数，c' 表示 X 对恢复时间的直接路径，则总效应可写为：",
    )
    add_formula(doc, "Effect_total(X) = c' + Σ_m a_m b_m")
    add_para(
        doc,
        "SEM 结果的显著性通过路径系数的标准误、z 值和 p 值判断，但由于全球事件样本量较大，极小 p 值并不必然代表机制解释力很强。因此本文同时报告模型拟合或预测解释力，并把 SEM 作为机制框架而不是单独的变量筛选工具。对于 SSRD、TMP、STRD 等高度相关的能量变量，SEM 路径设置采用“SHAP 重要性优先、生态机制约束、共线性诊断补充”的原则，避免仅因线性共线性而删除生态上关键的变量。",
    )

    add_heading(doc, "7 验证与辅助解释方法", 1)
    add_para(
        doc,
        "ALE、ICE、PDP 和地理探测器用于验证 SHAP dependence plot 中的方向、阈值和空间解释力。PDP（partial dependence plot）通过在保持其他样本特征分布不变的情况下，把某一特征固定为给定取值并对模型预测求平均，展示该特征的平均边际效应：",
    )
    add_formula(doc, "PDP_j(z) = (1/n) Σ_{i=1}^{n} f(z, x_{i,-j})")
    add_para(
        doc,
        "ICE（individual conditional expectation）保留每个样本的条件响应曲线，用于检查同一 biome 内部是否存在明显异质性：",
    )
    add_formula(doc, "ICE_{ij}(z) = f(z, x_{i,-j})")
    add_para(
        doc,
        "ALE（accumulated local effects）则在特征的局部区间内计算模型预测随该特征变化的平均差分并累积，较 PDP 更适合存在相关特征的环境变量场景：",
    )
    add_formula(doc, "ALE_j(z) = ∫_{z_0}^{z} E[∂f(X)/∂x_j | x_j=s] ds - constant")
    add_para(
        doc,
        "在解释上，如果 SHAP dependence、ALE 和 PDP 在同一特征区间内给出相似方向或相近转折点，说明该阈值具有较强稳健性；如果 ICE 曲线分散，则表明同一 biome 内部还存在背景气候或土地覆盖差异导致的异质性。地理探测器用于评估变量分层对恢复时间空间异质性的解释力，其核心统计量 q 为：",
    )
    add_formula(doc, "q = 1 - [Σ_{h=1}^{L} N_h σ_h^2] / [N σ^2]")
    add_para(
        doc,
        "其中 h 表示分层类别，N_h 和 σ_h^2 分别为第 h 层样本量和恢复时间方差，N 和 σ^2 为总体样本量和总体方差。q 的取值范围通常为 0-1，数值越大表示该因子对恢复时间空间分异的解释力越强。q 值本身没有正负方向，因此方向和阈值必须结合 SHAP、ALE/PDP/ICE 以及 risk detector 的分层均值共同解释。交互探测器进一步判断两个变量共同分层后的 q 值是否高于单因子 q 值，用于识别能量-水分耦合、热量-大气干燥度耦合等复合机制。",
    )

    add_heading(doc, "8 方法整合逻辑", 1)
    add_para(
        doc,
        "总体而言，本研究的方法链条从事件定义、碳恢复刻画、机器学习解释、路径机制和独立验证五个层次展开。骤旱定义确保所有事件具有可比的快速土壤水分亏缺过程；恢复时间定义确保目标变量对应生态系统从最大碳通量损伤中恢复的过程；SHAP 提供跨 biome、跨目标变量的非线性驱动因子排序和阈值识别；SEM 将关键驱动因子组织为直接与间接路径，以解释能量、水分和事件属性如何共同影响恢复；ALE、ICE、PDP 与地理探测器则分别从模型曲线稳健性、样本异质性、平均趋势和空间解释力角度验证 SHAP 结论。通过这一组合框架，本文既保留了机器学习对复杂非线性关系的识别能力，也避免把单一统计归因直接解释为因果机制。",
    )

    add_heading(doc, "参考文献（方法部分主要依据）", 1)
    refs = [
        "Lu 等（2025）Assessing recovery time of ecosystems in China: insights into flash drought impacts on gross primary productivity.",
        "Lundberg, S. M., & Lee, S.-I. (2017). A unified approach to interpreting model predictions.",
        "Apley, D. W., & Zhu, J. (2020). Visualizing the effects of predictor variables in black box supervised learning models.",
        "Goldstein, A., Kapelner, A., Bleich, J., & Pitkin, E. (2015). Peeking inside the black box: visualizing statistical learning with plots of individual conditional expectation.",
        "Wang, J.-F., Zhang, T.-L., & Fu, B.-J. (2016/2017). A measure of spatial stratified heterogeneity and geodetector-based attribution.",
        "Zhang 等（2021）Post-drought recovery time across global terrestrial ecosystems.",
        "Jiao 等（2022）Comprehensive quantification of the responses of ecosystem production and respiration to drought.",
    ]
    for ref in refs:
        add_para(doc, ref, first_line=False, size=9.5)

    return doc


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_doc()
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
