#!/usr/bin/env python3
"""Create a Chinese report for orthogonal-decomposition SHAP results."""

from __future__ import annotations

from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path("/home/xulc/flash_drought/process/SEM_analysis0401/codex")
BASE = ROOT / "GLEAM/plots2/prepeak_shap_nomulticollinearity"
ORTHO = BASE / "orthogonal_decomposition"
OUT_DIR = ROOT / "GLEAM/writing4"
OUT_DOCX = OUT_DIR / "01_orthogonal_decomposition_shap_collinearity_robustness_cn.docx"

BIOMES = ["Forest", "Grassland", "Savanna", "Cropland", "Shrubland"]
METRICS = ["GPP", "RECO"]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(9)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
        set_cell_shading(hdr[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(11)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "SimHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")


def load_importance() -> pd.DataFrame:
    rows = []
    for metric in METRICS:
        for biome in BIOMES:
            p = ORTHO / metric / biome / "feature_importance.csv"
            df = pd.read_csv(p)
            df["metric"] = metric
            df["biome"] = biome
            rows.append(df)
    return pd.concat(rows, ignore_index=True)


def load_vif() -> pd.DataFrame:
    rows = []
    for metric in METRICS:
        for biome in BIOMES:
            p = ORTHO / metric / biome / "vif_after_transform.csv"
            df = pd.read_csv(p)
            rows.append(
                {
                    "metric": metric,
                    "biome": biome,
                    "max_vif": float(df["vif"].max()),
                    "median_vif": float(df["vif"].median()),
                    "features_gt3": int((df["vif"] > 3).sum()),
                    "features_gt5": int((df["vif"] > 5).sum()),
                }
            )
    return pd.DataFrame(rows)


def load_removed_r2() -> pd.DataFrame:
    rows = []
    for metric in METRICS:
        for biome in BIOMES:
            p = ORTHO / metric / biome / "orthogonal_decomposition_models.csv"
            df = pd.read_csv(p)
            keep = (
                df.groupby(["orthogonal_feature", "source_feature"], as_index=False)["r2_removed_from_source"]
                .first()
                .assign(metric=metric, biome=biome)
            )
            rows.append(keep)
    return pd.concat(rows, ignore_index=True)


def load_model_summary() -> pd.DataFrame:
    summary = pd.read_csv(BASE / "nomulticollinearity_model_summary.csv")
    return summary[summary["method"] == "orthogonal_decomposition"].copy()


def dependence_direction(metric: str, biome: str, feature: str) -> str:
    p = ORTHO / metric / biome / "dependence_plot_data.parquet"
    df = pd.read_parquet(p, columns=[f"feature__{feature}", f"shap__{feature}"])
    x = pd.to_numeric(df[f"feature__{feature}"], errors="coerce")
    y = pd.to_numeric(df[f"shap__{feature}"], errors="coerce")
    valid = x.notna() & y.notna()
    x = x[valid]
    y = y[valid]
    if len(x) < 20:
        return "样本不足"
    rho = float(x.corr(y, method="spearman"))
    lo = y[x <= x.quantile(0.25)].median()
    hi = y[x >= x.quantile(0.75)].median()
    delta = float(hi - lo)
    if rho <= -0.35 and delta < 0:
        return f"高值端缩短恢复（rho={rho:.2f}）"
    if rho >= 0.35 and delta > 0:
        return f"高值端延长恢复（rho={rho:.2f}）"
    return f"非线性/分段响应（rho={rho:.2f}）"


def top_feature_rows(importance: pd.DataFrame, n: int = 5) -> list[list[str]]:
    rows = []
    for metric in METRICS:
        for biome in BIOMES:
            sub = importance[(importance["metric"] == metric) & (importance["biome"] == biome)].head(n)
            feats = []
            for r in sub.itertuples():
                feats.append(f"{int(r.rank)}.{r.display_label}({r.percent:.1f}%)")
            rows.append([metric, biome, "；".join(feats)])
    return rows


def model_rows(summary: pd.DataFrame) -> list[list[str]]:
    rows = []
    for metric in METRICS:
        for biome in BIOMES:
            r = summary[(summary["metric"] == metric) & (summary["biome"] == biome)].iloc[0]
            rows.append(
                [
                    metric,
                    biome,
                    f"{int(r.rows):,}",
                    f"{r.r2_train_split:.3f}",
                    f"{r.r2_holdout_split:.3f}",
                ]
            )
    return rows


def vif_rows(vif: pd.DataFrame) -> list[list[str]]:
    rows = []
    for metric in METRICS:
        for biome in BIOMES:
            r = vif[(vif["metric"] == metric) & (vif["biome"] == biome)].iloc[0]
            rows.append(
                [
                    metric,
                    biome,
                    f"{r.max_vif:.2f}",
                    f"{r.median_vif:.2f}",
                    str(int(r.features_gt3)),
                    str(int(r.features_gt5)),
                ]
            )
    return rows


def removed_r2_rows(removed: pd.DataFrame) -> list[list[str]]:
    mapping = [
        ("STRD_resid_after_SSRD", "STRD 去 SSRD"),
        ("TMP_resid_after_SSRD_STRD", "TMP 去 SSRD/STRD"),
        ("VPD_resid_after_SSRD_TMP_Wind", "VPD 去 SSRD/TMP/Wind"),
        ("EVA_resid_after_SSRD_Pre_VPD", "EVA 去 SSRD/Pre/VPD"),
        ("SMrz_resid_after_Pre_EVA", "SMrz 去 Pre/EVA"),
    ]
    rows = []
    for metric in METRICS:
        for feature, label in mapping:
            sub = removed[(removed["metric"] == metric) & (removed["orthogonal_feature"] == feature)]
            vals = ", ".join(f"{r.biome}:{r.r2_removed_from_source:.2f}" for r in sub.itertuples())
            rows.append([metric, label, vals])
    return rows


def add_image_if_exists(doc: Document, path: Path, caption: str, width: float = 6.2) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.name = "SimSun"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        run.font.size = Pt(9)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    importance = load_importance()
    vif = load_vif()
    removed = load_removed_r2()
    summary = load_model_summary()

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.85)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    styles = doc.styles
    styles["Normal"].font.name = "SimSun"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("正交分解后的 SHAP 共线性稳健性分析")
    run.bold = True
    run.font.name = "SimHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
    run.font.size = Pt(18)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("基于 prepeak 10 个核心特征的 GPP 与 RECO 恢复时间解释")
    r.font.name = "SimSun"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    r.font.size = Pt(11)

    add_heading(doc, "1. 写作目的与图件说明", 1)
    add_para(
        doc,
        "前期 01、02 和 04 文档已经分别从 beeswarm 特征贡献、dependence plot 阈值响应以及 GPP-RECO 差异三个角度说明了原始 SHAP 结果。"
        "这些结果保留了变量的原始物理含义，因此最适合用于解释具体阈值和生态水文过程；但由于 SSRD、STRD、TMP、VPD、EVA、SMrz 等水热变量之间存在明显耦合，审稿时可能会追问共线性是否改变了变量贡献排序。"
        "本文件专门使用正交分解后的 SHAP 结果作为稳健性分析，检验在显式削弱共线性后，主要机制是否仍然成立。"
    )
    add_para(
        doc,
        "正交分解版本仅使用十个核心变量：SSRD、EVA、TMP、STRD、SMrz、Wind、VPD、Duration、Pre 和 Intensity。"
        "其中 SSRD、Pre、Duration、Intensity 和 Wind 被保留为标准化锚点，分别记为 SSRD_z、Pre_z、Duration_z、Intensity_z 和 Wind_z；STRD、TMP、VPD、EVA 和 SMrz 则按物理顺序进行残差化，得到去除上游共线信息后的独立分量。"
        "因此，本文件中的 dependence plot 不再直接表示原始物理单位阈值，而表示标准化锚点或残差分量的独立贡献。"
    )

    add_heading(doc, "2. 正交分解方法与变量含义", 1)
    add_para(
        doc,
        "正交分解的核心思路是先保留具有明确机制优先级的基础变量，再将后续变量中可由这些基础变量线性解释的部分剥离。"
        "例如，STRD_resid_after_SSRD 表示 STRD 中不能被 SSRD_z 解释的剩余部分；TMP_resid_after_SSRD_STRD 表示 TMP 中不能被 SSRD_z 和已去 SSRD 的 STRD 分量解释的剩余部分。"
        "这种处理并不是为了替代原始变量，而是为了回答一个更严格的问题：在共线部分被去掉以后，该变量是否仍然对恢复时间具有独立解释力。"
    )
    add_table(
        doc,
        ["正交变量", "含义", "解释方式"],
        [
            ["SSRD_z", "SSRD 标准化锚点", "保留短波辐射的总体能量背景，作为后续热量变量的上游参照。"],
            ["STRD_resid_after_SSRD", "STRD 去 SSRD 后残差", "表示长波热量背景中独立于短波辐射的部分。"],
            ["TMP_resid_after_SSRD_STRD", "TMP 去 SSRD 和 STRD 后残差", "表示温度中不能由辐射背景解释的独立热效应。"],
            ["VPD_resid_after_SSRD_TMP_Wind", "VPD 去 SSRD/TMP/Wind 后残差", "表示大气干旱中独立于辐射、温度和通风条件的部分。"],
            ["EVA_resid_after_SSRD_Pre_VPD", "EVA 去 SSRD/Pre/VPD 后残差", "表示蒸散中不由能量、降水和大气干旱共同解释的部分。"],
            ["SMrz_resid_after_Pre_EVA", "SMrz 去 Pre/EVA 后残差", "表示根区水分中不由降水补给和蒸散消耗解释的剩余储水状态。"],
        ],
    )
    add_table(doc, ["Metric", "残差变量", "各 biome 被剥离方差比例 R2"], removed_r2_rows(removed))

    add_heading(doc, "3. 模型表现与共线性控制效果", 1)
    add_para(
        doc,
        "正交分解后，十个变量仍然保持了较好的预测解释力。GPP 各 biome 的 holdout R2 约为 0.355-0.496，RECO 的 holdout R2 约为 0.465-0.562。"
        "这说明去除显式共线结构后，模型并没有丢失主要恢复信息；相反，它将解释权重重新分配到更接近独立机制的变量分量上。"
    )
    add_table(doc, ["Metric", "Biome", "样本数", "Train R2", "Holdout R2"], model_rows(summary))
    add_para(
        doc,
        "VIF 结果表明正交分解有效降低了输入变量之间的线性共线性。所有组合的中位 VIF 均低于 2.5，大多数最大 VIF 低于 3，仅 RECO-Forest 的最大 VIF 接近 3.9，但仍显著低于通常需要警惕的 VIF=5 或 VIF=10 阈值。"
        "因此，正交分解结果可以作为原始 SHAP 的稳健性补充，用于回应“变量重要性是否只是共线变量替代效应”的问题。"
    )
    add_table(doc, ["Metric", "Biome", "最大 VIF", "中位 VIF", "VIF>3 个数", "VIF>5 个数"], vif_rows(vif))

    add_heading(doc, "4. 正交分解后的总体贡献格局", 1)
    add_para(
        doc,
        "最重要的结果是：在显式控制共线性后，SSRD 仍然是最稳定的主导变量。GPP 的五个 biome 中，SSRD_z 均位列第一；RECO 的五个 biome 中，SSRD_z 同样均位列第一。"
        "这与原始 beeswarm 分析中的核心结论一致，说明短波辐射对恢复时间的主导地位并不是由 STRD、TMP 或 VPD 等共线变量偶然替代造成的，而是具有稳健的独立解释力。"
    )
    add_para(
        doc,
        "与原始结果相比，正交分解后最明显的变化发生在 TMP、STRD、VPD 和 EVA。它们不再代表原始物理量的总效应，而是代表去掉上游共线信息后的残差效应。"
        "例如，TMP_resid 排名靠前并不意味着原始温度总效应强于辐射，而是说明在辐射背景已经被 SSRD 和 STRD 分量解释后，温度仍保留一部分独立热效应。"
    )
    add_image_if_exists(
        doc,
        BASE / "orthogonal_decomposition_importance_percent_bars_5biomes_gpp_vs_reco.png",
        "图 1. 正交分解后 GPP 与 RECO 在五个 biome 中的 SHAP 重要性百分比排序。",
        width=6.4,
    )
    add_table(doc, ["Metric", "Biome", "前五重要特征（括号为贡献百分比）"], top_feature_rows(importance, 5))

    add_heading(doc, "5. GPP 恢复时间的正交分解结果", 1)
    gpp_text = {
        "Forest": "Forest 中 SSRD_z 仍为首位，TMP_resid 位列第二，说明森林 GPP 恢复首先受短波辐射控制，但在去掉辐射共线背景后，温度仍保留明显独立调节作用。STRD_resid 和 EVA_resid 继续进入前五，表明森林冠层恢复不是单一能量过程，而是短波输入、独立热背景和蒸散交换共同作用的结果。",
        "Grassland": "Grassland 中 SSRD_z 排名第一，VPD_resid 和 TMP_resid 紧随其后。这与原始分析中草地受 SSRD、TMP、VPD 共同控制的判断一致，但正交分解进一步说明，大气干旱和温度不是简单地复制 SSRD 信息，而是在去除辐射背景后仍能解释草地 GPP 恢复差异。",
        "Savanna": "Savanna 的 SSRD_z 贡献占比超过 40%，是所有 GPP biome 中最集中的结果。这说明稀树草原的 GPP 恢复具有非常强的短波辐射依赖性；TMP_resid 和 STRD_resid 位列其后，说明热量背景仍会在能量-水分耦合边缘系统中调节恢复速度。",
        "Cropland": "Cropland 中 SSRD_z、EVA_resid 和 STRD_resid 位列前三。与原始结果中 Cropland-GPP 的 EVA 非常突出相互呼应，正交分解表明即使去掉 SSRD、Pre 和 VPD 对 EVA 的共同解释部分，蒸散残差仍保持较高贡献，说明农田恢复对作物冠层蒸散重启和水分利用状态非常敏感。",
        "Shrubland": "Shrubland 中 SSRD_z 排名第一，STRD_resid 排名第二，Duration_z 进入前三。该结果强调灌丛系统既受能量背景控制，也保留明显事件记忆效应；持续时间较长的骤旱可能通过累积水分亏缺和根区恢复滞后影响 GPP 恢复。",
    }
    for biome in BIOMES:
        add_heading(doc, f"5.{BIOMES.index(biome)+1} {biome}", 2)
        add_para(doc, gpp_text[biome])
        sub = importance[(importance["metric"] == "GPP") & (importance["biome"] == biome)].head(3)
        dirs = [f"{r.display_label}: {dependence_direction('GPP', biome, r.feature)}" for r in sub.itertuples()]
        add_para(doc, "前三特征的 dependence 方向可概括为：" + "；".join(dirs) + "。")

    add_heading(doc, "6. RECO 恢复时间的正交分解结果", 1)
    reco_text = {
        "Forest": "Forest 中 SSRD_z 和 TMP_resid 分别位列前两位，VPD_resid 也进入前三。与 GPP 相比，RECO 更直接体现温度和大气干旱对根系呼吸、微生物分解以及底物释放过程的调节。因此，正交分解后的 Forest-RECO 更清楚地显示出“短波能量背景 + 独立温度效应 + 大气干旱残差”的组合。",
        "Grassland": "Grassland 中 SSRD_z、TMP_resid 和 VPD_resid 共同主导 RECO，且 RECO 的 TMP_resid 贡献高于 GPP。这个结果支持原始 GPP-RECO 差异分析中的判断：草地呼吸恢复比光合恢复更容易受温度驱动，因为浅层根系和微生物活动对热量和土壤水分脉冲更敏感。",
        "Savanna": "Savanna 中 SSRD_z 的贡献仍最高，STRD_resid 和 TMP_resid 排名靠前。该组合符合热干边缘生态系统的能量-水分耦合特征：短波辐射提供恢复背景，长波和温度残差则进一步决定呼吸过程是否受到热负荷或暖夜背景调制。",
        "Cropland": "Cropland 中 SSRD_z、EVA_resid 和 TMP_resid 位列前三，说明农田 RECO 恢复同样依赖短波辐射背景，但呼吸过程对蒸散残差和独立温度效应更敏感。该结果与作物根系活动、土壤微生物分解以及灌溉或管理条件下水热重新匹配有关。",
        "Shrubland": "Shrubland 中 SSRD_z、STRD_resid 和 Pre_z 位列前三。Pre_z 的靠前说明灌丛 RECO 对降水补给和土壤水分激活具有更直接响应，而 STRD_resid 的重要性则提示长波热背景可能影响夜间或近地表呼吸环境。",
    }
    for biome in BIOMES:
        add_heading(doc, f"6.{BIOMES.index(biome)+1} {biome}", 2)
        add_para(doc, reco_text[biome])
        sub = importance[(importance["metric"] == "RECO") & (importance["biome"] == biome)].head(3)
        dirs = [f"{r.display_label}: {dependence_direction('RECO', biome, r.feature)}" for r in sub.itertuples()]
        add_para(doc, "前三特征的 dependence 方向可概括为：" + "；".join(dirs) + "。")

    add_heading(doc, "7. 与原始 SHAP 文档的关系", 1)
    add_para(
        doc,
        "本结果与 01 文档的贡献排序结论高度一致：SSRD 仍是最稳定的主导因子，Cropland 中 EVA 相关过程仍具有突出贡献，Grassland/Savanna/Shrubland 中能量和热量变量依然占据核心位置。"
        "不同之处在于，正交分解后的 TMP_resid、STRD_resid、VPD_resid、EVA_resid 和 SMrz_resid 不能再直接解释为原始物理量的总体效应，而应解释为去除上游共线部分后的独立边际信息。"
    )
    add_para(
        doc,
        "与 02 文档的 dependence threshold 分析相比，本文件不再强调原始单位阈值。原始 dependence plot 中的 SSRD、TMP、PRE、|EVA| 等阈值仍应作为生态水文解释的主要依据；正交分解 dependence plot 主要用于判断这些阈值背后的机制是否具有独立性。"
        "例如，SSRD_z 的稳定首位说明短波辐射阈值不是由温度或长波辐射共线造成的；而 TMP_resid 的靠前则说明在短波和长波背景之外，温度仍有剩余热效应。"
    )
    add_para(
        doc,
        "与 04 文档的 GPP-RECO 对比相比，正交分解强化了一个核心结论：GPP 和 RECO 共享短波辐射背景，但 RECO 更容易表现出独立温度、长波热背景和降水脉冲的贡献。"
        "这说明两类碳通量并非由完全不同的环境变量驱动，而是在同一水热背景下由不同生理和生物地球化学过程读取这些信号。"
    )

    add_heading(doc, "8. 可写入论文或回复审稿人的结论", 1)
    add_para(
        doc,
        "为检验输入变量共线性对 SHAP 解释的影响，我们进一步构建了基于十个核心 prepeak 特征的正交分解 SHAP 模型。"
        "结果显示，所有 metric-biome 组合的 VIF 均得到明显控制，且模型仍保持较好的 holdout 解释力。在去除显式共线结构后，SSRD_z 在 GPP 和 RECO 的所有 biome 中均保持最高贡献，说明短波辐射对恢复时间的主导作用具有稳健性，并非由 STRD、TMP 或 VPD 等相关变量的替代效应造成。"
        "同时，TMP、STRD、VPD、EVA 和 SMrz 的残差分量在不同 biome 中仍表现出可解释的独立贡献，表明原始 SHAP 识别的能量输入、热量背景、大气干旱、蒸散过程和根区储水机制在降低共线性后仍然成立。"
    )
    add_para(
        doc,
        "因此，原始 SHAP dependence plot 适合用于报告物理量阈值和生态机制解释，而正交分解 SHAP 则可作为共线性稳健性检验，证明主要结论不是由输入变量间的相关结构单独驱动。"
    )

    add_heading(doc, "附图：代表性 dependence plot", 1)
    examples = [
        ("GPP", "Savanna", "SSRD_z.png", "图 A1. Savanna-GPP 中 SSRD_z 的 dependence plot。"),
        ("GPP", "Cropland", "EVA_resid_after_SSRD_Pre_VPD.png", "图 A2. Cropland-GPP 中 EVA_resid 的 dependence plot。"),
        ("RECO", "Grassland", "TMP_resid_after_SSRD_STRD.png", "图 A3. Grassland-RECO 中 TMP_resid 的 dependence plot。"),
        ("RECO", "Shrubland", "Pre_z.png", "图 A4. Shrubland-RECO 中 Pre_z 的 dependence plot。"),
    ]
    for metric, biome, image, caption in examples:
        add_image_if_exists(doc, ORTHO / metric / biome / "dependence_plots" / image, caption, width=5.2)

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
