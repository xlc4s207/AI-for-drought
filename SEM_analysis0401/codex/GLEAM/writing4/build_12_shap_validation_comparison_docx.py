#!/usr/bin/env python3
"""Build a Chinese report comparing orthogonal SHAP with ALE/PDP/ICE and OPGD."""

from __future__ import annotations

from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path("/home/xulc/flash_drought/process/SEM_analysis0401/codex")
OUT = ROOT / "GLEAM/writing4"
ORTHO = ROOT / "GLEAM/plots2/prepeak_shap_nomulticollinearity/orthogonal_decomposition"
OPGD = ROOT / "GLEAM/validation/05_OPGD_Geodetector/orthogonal_comparison"
OVERLAY = OPGD / "validation_overlay_by_biome"

MATRIX_PNG = OPGD / "orthogonal_shap_opgd_reliability_matrix.png"
TOP3_CSV = OPGD / "orthogonal_shap_opgd_top3_consistency.csv"
FEATURE_CSV = OPGD / "orthogonal_shap_opgd_feature_comparison.csv"
INDEX_CSV = OVERLAY / "orthogonal_validation_opgd_overlay_index.csv"

BIOMES = ["Cropland", "Forest", "Grassland", "Savanna", "Shrubland"]
BIOME_CN = {
    "Cropland": "农田",
    "Forest": "森林",
    "Grassland": "草地",
    "Savanna": "稀树草原",
    "Shrubland": "灌丛",
}
METRIC_CN = {"GPP": "GPP 恢复时间", "RECO": "RECO 恢复时间"}
FEATURE_CN = {
    "SSRD_z": "SSRD_z",
    "EVA_resid_after_SSRD_Pre_VPD": "EVA_resid",
    "TMP_resid_after_SSRD_STRD": "TMP_resid",
    "VPD_resid_after_SSRD_TMP_Wind": "VPD_resid",
    "SMrz_resid_after_Pre_EVA": "SMrz_resid",
    "Pre_z": "Pre_z",
    "STRD_resid_after_SSRD": "STRD_resid",
    "Wind_z": "Wind_z",
    "Duration_z": "Duration_z",
    "Intensity_z": "Intensity_z",
}


def set_cn_font(run, font: str = "SimSun", size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def new_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    doc.styles["Normal"].font.name = "SimSun"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    doc.styles["Normal"].font.size = Pt(10)
    return doc


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("12 正交 SHAP 与 ALE、PDP、ICE 及地理探测器的对比验证分析")
    set_cn_font(r, "SimHei", 18, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("基于正交分解后的 SHAP dependence、模型响应曲线和 OPGD 空间解释力结果")
    set_cn_font(r, "SimSun", 10)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        set_cn_font(r, "SimHei", 14 if level == 1 else 12, True)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(20)
    p.paragraph_format.line_spacing = 1.18
    r = p.add_run(text)
    set_cn_font(r, "SimSun", 10)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-10)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run("• " + text)
    set_cn_font(r, "SimSun", 10)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill: str = "D9EAF7") -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        set_cn_font(r, "SimHei", 8, True)
        shade_cell(cell, header_fill)
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            set_cn_font(r, "SimSun", 8)
    doc.add_paragraph()


def add_image(doc: Document, path: Path, caption: str, width: float) -> None:
    if not path.exists():
        add_para(doc, f"图件缺失：{path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_cn_font(r, "SimSun", 9)


def fmt_num(x: float | int | str, digits: int = 2) -> str:
    try:
        if pd.isna(x):
            return ""
        return f"{float(x):.{digits}f}"
    except Exception:
        return str(x)


def load_data() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    top3 = pd.read_csv(TOP3_CSV)
    feat = pd.read_csv(FEATURE_CSV)
    index = pd.read_csv(INDEX_CSV)
    return top3, feat, index


def curve_summary_from_xy(x: np.ndarray, y: np.ndarray) -> dict[str, object]:
    ok = np.isfinite(x) & np.isfinite(y)
    x = x[ok]
    y = y[ok]
    if len(x) < 6:
        return {"direction": "样本不足", "threshold": np.nan, "threshold_type": "NA", "range": np.nan}
    order = np.argsort(x)
    x = x[order]
    y = y[order]
    y_range = float(np.nanmax(y) - np.nanmin(y))
    delta = float(np.nanmedian(y[x >= np.nanquantile(x, 0.75)]) - np.nanmedian(y[x <= np.nanquantile(x, 0.25)]))
    if y_range <= 1e-9 or abs(delta) < 0.12 * y_range:
        direction = "非线性/弱方向"
    elif delta > 0:
        direction = "高值端延长恢复"
    else:
        direction = "高值端缩短恢复"

    threshold = np.nan
    threshold_type = "无明显零交叉"
    signs = np.sign(y)
    signs[np.abs(y) < max(y_range * 0.02, 1e-9)] = 0
    for i in range(len(x) - 1):
        if signs[i] == 0:
            threshold = float(x[i])
            threshold_type = "零交叉"
            break
        if signs[i] * signs[i + 1] < 0:
            denom = y[i + 1] - y[i]
            if abs(denom) > 1e-12:
                threshold = float(x[i] - y[i] * (x[i + 1] - x[i]) / denom)
            else:
                threshold = float((x[i] + x[i + 1]) / 2)
            threshold_type = "零交叉"
            break
    if not np.isfinite(threshold) and len(x) >= 4:
        slope = np.abs(np.diff(y) / np.maximum(np.diff(x), 1e-9))
        if np.isfinite(slope).any():
            idx = int(np.nanargmax(slope))
            threshold = float((x[idx] + x[idx + 1]) / 2)
            threshold_type = "最大斜率转折"
    return {
        "direction": direction,
        "threshold": threshold,
        "threshold_type": threshold_type,
        "range": y_range,
        "delta": delta,
    }


def shap_trend_summary(metric: str, biome: str, feature: str) -> dict[str, object]:
    path = ORTHO / metric / biome / "dependence_plot_data.parquet"
    xcol = f"feature__{feature}"
    ycol = f"shap__{feature}"
    try:
        df = pd.read_parquet(path, columns=[xcol, ycol])
    except Exception:
        return {"direction": "缺失", "threshold": np.nan, "threshold_type": "NA", "range": np.nan}
    x = pd.to_numeric(df[xcol], errors="coerce").to_numpy(dtype=float)
    y = pd.to_numeric(df[ycol], errors="coerce").to_numpy(dtype=float)
    ok = np.isfinite(x) & np.isfinite(y)
    x = x[ok]
    y = y[ok]
    if len(x) < 100:
        return {"direction": "样本不足", "threshold": np.nan, "threshold_type": "NA", "range": np.nan}
    edges = np.unique(np.nanquantile(x, np.linspace(0.02, 0.98, 31)))
    centers, med = [], []
    for left, right in zip(edges[:-1], edges[1:]):
        mask = (x >= left) & (x <= right if right == edges[-1] else x < right)
        if mask.sum() >= 30:
            centers.append(float(np.nanmedian(x[mask])))
            med.append(float(np.nanmedian(y[mask])))
    return curve_summary_from_xy(np.asarray(centers), np.asarray(med))


def read_validation_curve(metric: str, biome: str, feature: str, method: str) -> dict[str, object]:
    path = OVERLAY / "curves" / metric / biome / f"{feature}_{method}.csv"
    if not path.exists():
        return {"direction": "缺失", "threshold": np.nan, "threshold_type": "NA", "range": np.nan}
    df = pd.read_csv(path)
    return curve_summary_from_xy(
        pd.to_numeric(df["feature_value"], errors="coerce").to_numpy(dtype=float),
        pd.to_numeric(df["effect"], errors="coerce").to_numpy(dtype=float),
    )


def same_direction(a: str, b: str) -> bool:
    if a in {"缺失", "样本不足"} or b in {"缺失", "样本不足"}:
        return False
    if a == b:
        return True
    if "非线性" in a and "非线性" in b:
        return True
    return False


def threshold_close(a: float, b: float, tol: float = 0.55) -> bool:
    if not np.isfinite(a) or not np.isfinite(b):
        return False
    return abs(a - b) <= tol


def build_curve_consistency(feat: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for r in feat.itertuples(index=False):
        metric = r.metric
        biome = r.biome
        feature = r.orthogonal_feature
        shap = shap_trend_summary(metric, biome, feature)
        curves = {m: read_validation_curve(metric, biome, feature, m) for m in ["ale", "ice_mean", "pdp"]}
        row = {
            "metric": metric,
            "biome": biome,
            "feature": feature,
            "label": r.orthogonal_label,
            "shap_rank": int(r.orthogonal_shap_rank),
            "shap_percent": float(r.shap_percent),
            "shap_direction": shap["direction"],
            "shap_threshold": shap["threshold"],
            "shap_threshold_type": shap["threshold_type"],
        }
        match_count = 0
        threshold_match_count = 0
        for method, summary in curves.items():
            dir_match = same_direction(str(shap["direction"]), str(summary["direction"]))
            thr_match = threshold_close(float(shap["threshold"]), float(summary["threshold"]))
            match_count += int(dir_match)
            threshold_match_count += int(thr_match)
            row[f"{method}_direction"] = summary["direction"]
            row[f"{method}_threshold"] = summary["threshold"]
            row[f"{method}_threshold_type"] = summary["threshold_type"]
            row[f"{method}_direction_match"] = dir_match
            row[f"{method}_threshold_match"] = thr_match
            row[f"{method}_threshold_diff"] = (
                abs(float(shap["threshold"]) - float(summary["threshold"]))
                if np.isfinite(float(shap["threshold"])) and np.isfinite(float(summary["threshold"]))
                else np.nan
            )
        if row["ale_direction_match"] and row["ice_mean_direction_match"] and (
            row["ale_threshold_match"] or row["ice_mean_threshold_match"]
        ):
            grade = "强一致"
        elif match_count >= 2:
            grade = "方向一致但阈值偏移"
        elif row["ale_direction_match"] or row["ice_mean_direction_match"]:
            grade = "部分一致"
        else:
            grade = "不一致/复杂"
        row["validation_grade"] = grade
        row["pdp_deviation"] = bool((row["ale_direction_match"] or row["ice_mean_direction_match"]) and not row["pdp_direction_match"])
        rows.append(row)
    out = pd.DataFrame(rows)
    out_path = OUT / "12_SHAP_ALE_PDP_ICE_threshold_consistency.csv"
    out.to_csv(out_path, index=False)
    return out


def add_methods_section(doc: Document) -> None:
    add_heading(doc, "1. 对比目的与方法定位", 1)
    add_para(
        doc,
        "本文件的核心目的不是重新替代前面已经完成的 SHAP 分析，而是把四类验证结果放在同一解释框架下："
        "正交 SHAP 用于判断降低共线性后各变量的相对贡献；SHAP dependence 用于显示单个正交变量取值变化时的局部贡献方向；"
        "ALE、PDP 和 ICE mean 用于验证这种方向与阈值是否也能被模型响应曲线支持；地理探测器或 OPGD 的 q 值用于检验对应原始因子的空间解释力。"
    )
    add_para(
        doc,
        "需要特别说明的是，正交分解后的变量并不等同于原始物理变量。SSRD_z、Pre_z、Wind_z、Duration_z 和 Intensity_z 是标准化后的锚点变量，"
        "而 STRD、TMP、VPD、EVA 和 SMrz 对应的是剥离上游共线信息后的残差分量。因此，图中的横坐标主要用于比较相对变化和方向，而不应直接解释为原始单位阈值。"
    )
    add_heading(doc, "1.1 ALE、PDP 和 ICE 的计算方法", 2)
    add_para(
        doc,
        "PDP（Partial Dependence Plot，部分依赖图）用于估计某一目标特征对模型预测值的平均边际影响。具体做法是：先选定目标特征的一组网格值，"
        "然后依次把所有样本中的该特征强制替换为某个网格值，其他特征保持样本原值不变，重新输入模型得到预测值，最后对所有样本的预测结果求平均。"
        "因此，PDP 曲线表示“如果所有样本都被设定为同一个目标特征水平，模型平均预测会如何变化”。其优点是直观、易解释；但在特征存在相关性时，"
        "这种强制替换可能产生现实数据中很少出现的组合，例如高温但低 VPD、强辐射但低蒸散等。因此，本研究中 PDP 主要作为平均趋势的辅助展示，不把它单独作为阈值判定依据。"
    )
    add_para(
        doc,
        "ICE（Individual Conditional Expectation，个体条件期望曲线）与 PDP 的计算基础相同，也是在一组网格值上替换目标特征并重新预测。不同之处在于，"
        "ICE 不立即对所有样本求平均，而是为每个样本保留一条响应曲线。也就是说，每一条 ICE 曲线表示某个具体样本在目标特征变化时模型预测如何变化。"
        "PDP 可以理解为所有 ICE 曲线的平均。ICE 的优势是能够显示 biome 内部异质性和潜在交互：如果不同样本曲线斜率或转折位置差异很大，说明同一变量在不同环境背景下的作用并不一致。"
        "在本文的叠加图中，为避免大量个体曲线压缩或拉伸纵坐标，只保留 ICE mean，用于检验平均响应方向是否与 SHAP dependence 和 ALE 一致。"
    )
    add_para(
        doc,
        "ALE（Accumulated Local Effects，累积局部效应）用于估计目标特征在真实样本分布附近的局部效应。它先按照目标特征的分位数把样本划分为若干区间，"
        "然后只使用实际落入某个区间的样本，分别把这些样本的目标特征设为该区间左端点和右端点，计算模型预测差值并取平均；最后沿着特征取值方向累积这些局部平均差值，并进行中心化。"
        "与 PDP 相比，ALE 不需要把所有样本强行替换到整个特征范围，而是只在样本真实存在的局部区间内计算差分，因此在气象、水文和生态变量高度相关时更稳健。"
        "在本文中，ALE 主要用于验证 SHAP dependence 所显示的方向、零交叉和主要转折是否具有模型局部响应基础。"
    )
    add_para(
        doc,
        "三者与 SHAP 的关系可以概括为：SHAP dependence 显示目标特征在不同取值下对单个预测的贡献如何变化；ALE 检查真实数据分布内的局部平均效应；"
        "ICE mean 检查个体响应平均后是否仍支持该方向；PDP 展示全样本平均边际趋势。若 SHAP、ALE 和 ICE mean 在方向及转折附近一致，则说明该机制较稳健；"
        "若 PDP 偏离但 ALE/ICE 与 SHAP 一致，则更可能是特征相关性和边缘外推影响了 PDP。"
    )
    add_table(
        doc,
        ["方法", "回答的问题", "与共线性的关系", "本文件中的使用方式"],
        [
            ["正交 SHAP", "降低共线性后，哪些变量仍然重要？", "通过标准化和残差化减少特征共享信息", "作为主结果，读取贡献排序和 dependence 形态"],
            ["ALE", "在真实数据分布内，局部平均效应方向是否稳定？", "比 PDP 更少依赖不现实的特征组合", "验证 SHAP dependence 的方向、转折和非线性"],
            ["PDP", "模型平均预测随某变量改变的总体趋势是什么？", "可能受特征相关性影响，因此只作辅助", "用于观察平均趋势是否与 SHAP/ALE 相反"],
            ["ICE mean", "biome 内个体曲线的平均响应是否支持总体趋势？", "个体差异被平均后更稳定，但会弱化异质性", "只保留 mean 曲线，避免个体 ICE 拉伸纵轴"],
            ["OPGD/地理探测器", "原始因子的空间分层解释力多强？", "q 值没有方向性，也不是模型贡献", "与正交 SHAP 做机制层面对照"],
        ],
    )


def add_reliability_section(doc: Document, top3: pd.DataFrame, feat: pd.DataFrame) -> None:
    add_heading(doc, "2. SHAP 与 OPGD 的总体一致性", 1)
    mean_overlap = top3["top3_overlap_count"].mean()
    mean_corr = top3["rank_correlation_on_shared10"].mean()
    strong = (top3["top3_overlap_count"] >= 2).sum()
    weak = (top3["top3_overlap_count"] == 0).sum()
    high_medium_top3 = feat[(feat["orthogonal_shap_rank"] <= 3) & (feat["opgd_reliability"].isin(["High", "Medium"]))]
    add_para(
        doc,
        f"从五个 biome 与 GPP/RECO 两个恢复指标的 10 组对比看，正交 SHAP 前三因子与 OPGD 前三原始因子的平均重叠数为 {mean_overlap:.1f}，"
        f"共享十个因子的平均排序相关约为 {mean_corr:.2f}。其中 {strong} 组达到至少两个前三因子重叠，{weak} 组没有前三因子重叠。"
        "这说明两类方法给出的主导因子并非完全一致，但在能量、温度和长波热背景等核心机制上存在较稳定的交叉支持。"
    )
    add_para(
        doc,
        f"在 SHAP 排名前三的正交变量中，共有 {len(high_medium_top3)} 个变量的 OPGD 可靠性达到 Medium 或 High。"
        "这类变量可以视为“模型贡献高、空间解释力也较强”的稳健机制；而 SHAP 高但 OPGD 低的变量，更适合解释为模型内部的预测贡献，不宜直接扩展为空间分异主因。"
    )
    add_para(
        doc,
        "整体上看，OPGD 与 SHAP 并不是相互否定的关系，而是从不同角度回答同一机制问题。SHAP 排名靠前，说明该变量在 LightGBM 预测恢复时间时贡献较大；"
        "OPGD q 值较高，说明该原始变量能够较好解释恢复时间的空间分层差异。二者一致时，代表该机制既被模型用于预测，也能解释空间异质性。"
        "二者不一致时，通常有三类原因：第一，正交 SHAP 使用残差变量，OPGD 使用原始变量，因此变量含义并非完全相同；第二，SHAP 反映模型内的局部非线性贡献，"
        "OPGD 反映离散分层后的方差解释力，对单调但空间分层不强的变量可能给出较低 q 值；第三，事件属性或局地管理因素可能提升模型预测贡献，但未必形成稳定空间分层。"
    )
    rows = []
    for r in top3.itertuples(index=False):
        rows.append(
            [
                r.metric,
                BIOME_CN.get(r.biome, r.biome),
                r.orthogonal_top3,
                r.opgd_top3_raw,
                str(int(r.top3_overlap_count)),
                "" if pd.isna(r.top3_overlap_raw_labels) else str(r.top3_overlap_raw_labels),
                fmt_num(r.rank_correlation_on_shared10, 2),
            ]
        )
    add_table(doc, ["指标", "Biome", "正交 SHAP 前三", "OPGD 前三", "重叠数", "重叠因子", "排序相关"], rows)
    add_para(
        doc,
        "具体差异需要分组理解。草地和稀树草原中，SSRD、TMP、STRD 等能量/热量变量在 SHAP 与 OPGD 中同时靠前，说明热量和辐射机制具有较强跨方法稳健性。"
        "森林中，SHAP 更突出 SSRD_z 与 TMP_resid 的预测贡献，而 OPGD 对 EVA、TMP、STRD 的空间解释力更强，说明森林恢复既有模型层面的辐射敏感性，"
        "也有空间层面的蒸散和热背景分异。农田中，SHAP 强调 SSRD_z 与 EVA_resid，OPGD 更强调 STRD、TMP、PRE 或 VPD，说明农田可能同时受管理背景、灌溉/水分条件和能量输入影响，"
        "导致预测贡献与空间分层主因不完全重合。灌丛中 GPP 的不一致最明显，提示灌丛 GPP 恢复更可能存在局地水分、群落结构或事件过程差异，不能只用单一空间 q 值解释。"
    )
    add_image(
        doc,
        MATRIX_PNG,
        "图 1. 正交 SHAP 与 OPGD 地理探测器结果的可靠性矩阵。颜色表示机制一致性/可靠性等级，q 值来自对应原始变量，SHAP 排名来自正交变量。",
        6.5,
    )


def add_curve_consistency_section(doc: Document, consistency: pd.DataFrame) -> None:
    add_heading(doc, "3. SHAP dependence 与 ALE、PDP、ICE 的方向和阈值一致性", 1)
    top5 = consistency[consistency["shap_rank"] <= 5].copy()
    grade_counts = top5["validation_grade"].value_counts().to_dict()
    strong = int(grade_counts.get("强一致", 0))
    shifted = int(grade_counts.get("方向一致但阈值偏移", 0))
    partial = int(grade_counts.get("部分一致", 0))
    complex_count = int(grade_counts.get("不一致/复杂", 0))
    ale_match = float(top5["ale_direction_match"].mean() * 100)
    ice_match = float(top5["ice_mean_direction_match"].mean() * 100)
    pdp_match = float(top5["pdp_direction_match"].mean() * 100)
    ale_thr = float(top5["ale_threshold_match"].mean() * 100)
    ice_thr = float(top5["ice_mean_threshold_match"].mean() * 100)
    pdp_dev = int(top5["pdp_deviation"].sum())
    add_para(
        doc,
        f"在每个 metric-biome 的前五个 SHAP 变量中，共比较了 {len(top5)} 个高贡献变量。"
        f"其中强一致 {strong} 个，方向一致但阈值存在偏移 {shifted} 个，部分一致 {partial} 个，不一致或响应复杂 {complex_count} 个。"
        f"方向一致率方面，ALE 与 SHAP 的一致率为 {ale_match:.1f}%，ICE mean 与 SHAP 的一致率为 {ice_match:.1f}%，PDP 与 SHAP 的一致率为 {pdp_match:.1f}%。"
        f"阈值一致率方面，ALE 的零交叉或最大斜率转折与 SHAP 接近的比例为 {ale_thr:.1f}%，ICE mean 为 {ice_thr:.1f}%。"
        f"另有 {pdp_dev} 个高贡献变量出现 ALE/ICE 与 SHAP 基本一致、但 PDP 偏离的情况，这类结果通常说明 PDP 受到相关变量边缘外推或平均化的影响。"
    )
    add_para(
        doc,
        "这里的“阈值一致”并不是指原始物理单位阈值完全相同，而是在正交变量坐标中，SHAP 平滑趋势与 ALE/PDP/ICE 曲线的零交叉点或最大斜率转折点距离较近。"
        "若 SHAP 与 ALE/ICE 在方向和转折位置上都接近，可以认为该特征的模型响应具有较高稳定性；若方向一致但转折点偏移，说明变量效应存在但具体阈值受曲线算法、样本分布或交互项影响；"
        "若 PDP 与 ALE/ICE 差异明显，应优先将 PDP 作为平均趋势参考，而不是作为阈值判断的主要依据。"
    )

    summary_rows = []
    for metric in ["GPP", "RECO"]:
        for biome in BIOMES:
            sub = top5[(top5["metric"] == metric) & (top5["biome"] == biome)]
            summary_rows.append(
                [
                    metric,
                    BIOME_CN[biome],
                    str(int((sub["validation_grade"] == "强一致").sum())),
                    str(int((sub["validation_grade"] == "方向一致但阈值偏移").sum())),
                    str(int((sub["validation_grade"] == "部分一致").sum())),
                    str(int((sub["validation_grade"] == "不一致/复杂").sum())),
                    f"{sub['ale_direction_match'].mean() * 100:.0f}%",
                    f"{sub['ice_mean_direction_match'].mean() * 100:.0f}%",
                    f"{sub['pdp_direction_match'].mean() * 100:.0f}%",
                ]
            )
    add_table(
        doc,
        ["指标", "Biome", "强一致", "方向同/阈值偏", "部分一致", "复杂", "ALE方向", "ICE方向", "PDP方向"],
        summary_rows,
        "EAF3DF",
    )

    top3 = consistency[consistency["shap_rank"] <= 3].copy()
    top3["best_threshold_diff"] = top3[["ale_threshold_diff", "ice_mean_threshold_diff", "pdp_threshold_diff"]].min(axis=1)
    top3 = top3.sort_values(["metric", "biome", "shap_rank"])
    rows = []
    for r in top3.itertuples(index=False):
        rows.append(
            [
                r.metric,
                BIOME_CN[r.biome],
                r.label,
                str(int(r.shap_rank)),
                r.shap_direction,
                fmt_num(r.shap_threshold, 2),
                r.ale_direction,
                fmt_num(r.ale_threshold, 2),
                r.ice_mean_direction,
                fmt_num(r.ice_mean_threshold, 2),
                r.pdp_direction,
                fmt_num(r.pdp_threshold, 2),
                r.validation_grade,
            ]
        )
    add_table(
        doc,
        ["指标", "Biome", "变量", "SHAP秩", "SHAP方向", "SHAP阈值", "ALE方向", "ALE阈值", "ICE方向", "ICE阈值", "PDP方向", "PDP阈值", "一致性"],
        rows,
        "F4E6D7",
    )

    add_para(
        doc,
        "从变量类型看，SSRD_z、TMP_resid 和 STRD_resid 是跨方法一致性最高的一组。它们在多个 biome 中不仅 SHAP 排名靠前，而且 ALE 与 ICE mean 往往能够复现 dependence plot 的主要方向，"
        "说明能量输入和热背景对恢复时间的影响不是 SHAP 单一解释器的产物。VPD_resid、EVA_resid 和 SMrz_resid 的一致性更依赖 biome：在草地、稀树草原等水热限制更强的系统中，"
        "它们更容易与 OPGD 或 ALE/ICE 形成一致证据；在森林或农田中，则可能由于冠层调节、管理干预或灌溉/土壤水分背景而表现为阈值偏移。"
        "Pre_z、Duration_z 和 Intensity_z 常表现为方向一致但阈值不稳定，说明事件属性确实参与恢复过程，但其响应更容易被样本分布、极端事件和局地条件改变。"
    )


def biome_feature_sentence(feat: pd.DataFrame, metric: str, biome: str) -> str:
    sub = feat[(feat["metric"] == metric) & (feat["biome"] == biome)].sort_values("orthogonal_shap_rank")
    top = sub.head(5)
    top_txt = "、".join([f"{r.orthogonal_label}({r.shap_percent:.1f}%)" for r in top.itertuples(index=False)])
    hm = top[top["opgd_reliability"].isin(["High", "Medium"])]
    hm_txt = "、".join([f"{r.orthogonal_label}-{r.opgd_reliability}" for r in hm.itertuples(index=False)])
    if not hm_txt:
        hm_txt = "前五因子中缺少 Medium/High 等级的 OPGD 支持"
    return f"{METRIC_CN[metric]} 的前五个正交 SHAP 因子为 {top_txt}；其中 OPGD 达到 Medium/High 的因子为 {hm_txt}。"


def add_biome_sections(doc: Document, feat: pd.DataFrame, index: pd.DataFrame, consistency: pd.DataFrame) -> None:
    add_heading(doc, "4. 分 biome 的 SHAP-dependence、验证曲线与 OPGD 综合解释", 1)
    add_para(
        doc,
        "下列五张大图在每个 biome 内同时绘制 GPP 与 RECO 的十个正交变量面板。灰色散点为 SHAP dependence，黑线为 SHAP 平滑趋势，"
        "橙线为 ALE，绿色为 ICE mean，紫色虚线为 PDP。若 SHAP、ALE 和 ICE mean 在主要取值区间方向一致，可以认为该变量的方向性较稳定；"
        "若 PDP 明显偏离而 ALE 与 SHAP 更一致，通常说明相关变量组合或边缘外推影响了 PDP。"
    )
    for biome in BIOMES:
        add_heading(doc, f"{BIOME_CN[biome]}（{biome}）", 2)
        add_para(doc, biome_feature_sentence(feat, "GPP", biome))
        add_para(doc, biome_feature_sentence(feat, "RECO", biome))
        cons = consistency[(consistency["biome"] == biome) & (consistency["shap_rank"] <= 5)]
        strong_features = cons[cons["validation_grade"] == "强一致"]
        shifted_features = cons[cons["validation_grade"] == "方向一致但阈值偏移"]
        complex_features = cons[cons["validation_grade"].isin(["部分一致", "不一致/复杂"])]
        strong_txt = "、".join(
            [f"{r.metric}-{r.label}" for r in strong_features.sort_values(["metric", "shap_rank"]).head(6).itertuples(index=False)]
        )
        shifted_txt = "、".join(
            [f"{r.metric}-{r.label}" for r in shifted_features.sort_values(["metric", "shap_rank"]).head(6).itertuples(index=False)]
        )
        complex_txt = "、".join(
            [f"{r.metric}-{r.label}" for r in complex_features.sort_values(["metric", "shap_rank"]).head(6).itertuples(index=False)]
        )
        if not strong_txt:
            strong_txt = "无特别突出的强一致前五因子"
        if not shifted_txt:
            shifted_txt = "较少"
        if not complex_txt:
            complex_txt = "较少"
        add_para(
            doc,
            f"从 ALE/ICE/PDP 与 SHAP dependence 的曲线一致性看，{BIOME_CN[biome]}中强一致的前五因子主要包括 {strong_txt}；"
            f"方向一致但阈值有偏移的因子主要包括 {shifted_txt}；需要谨慎解释的复杂或部分一致因子主要包括 {complex_txt}。"
            "这意味着该 biome 的解释不能只看 SHAP 排名，还需要结合响应曲线是否复现了 SHAP 的方向和转折位置。"
        )
        if biome == "Cropland":
            add_para(
                doc,
                "农田中，EVA_resid 和 TMP/STRD 相关分量在 SHAP 与 OPGD 间具有较强交叉信息。GPP 与 RECO 都显示 SSRD_z 为首要 SHAP 因子，"
                "但 OPGD 更强调 STRD、TMP、VPD 或降水的空间分层解释力，说明农田恢复过程同时受管理背景下的能量输入和水热空间差异控制。"
            )
            add_para(
                doc,
                "这种差异的关键在于尺度：SSRD_z 在模型中提供稳定预测贡献，但农田的空间格局可能被灌溉、作物类型、土壤背景和管理措施重新组织，"
                "使 STRD/TMP/Pre/VPD 这类原始变量在 OPGD 中获得更高 q 值。若 ALE 或 ICE mean 能复现 SSRD/EVA 的 SHAP 方向，则说明其模型响应可信；"
                "若 OPGD 不同步增强，则应写成“预测机制明显，但空间分层解释力受农田管理背景调节”。"
            )
        elif biome == "Forest":
            add_para(
                doc,
                "森林中，GPP 和 RECO 均表现为 SSRD_z 与 TMP_resid 的组合主导，且 RECO 中 TMP 与 EVA 的 OPGD 可靠性较高。"
                "这说明森林恢复的模型贡献主要来自辐射和独立温度效应，而空间分异更容易被蒸散和温度背景捕捉。"
            )
            add_para(
                doc,
                "森林的 SHAP-OPGD 差异不能简单解释为矛盾。森林冠层和根系缓冲会削弱某些气象变量的直接空间分层，但模型仍可能利用这些变量识别恢复时间的非线性变化。"
                "因此森林中更应关注 ALE/ICE 是否与 SHAP 在主要样本区间同向；若同向但阈值偏移，说明阈值受到冠层水分调节和局地气候背景影响。"
            )
        elif biome == "Grassland":
            add_para(
                doc,
                "草地中，GPP 和 RECO 都显示 SSRD_z、TMP_resid、STRD_resid 与 VPD_resid 的高贡献组合，其中 RECO 的 TMP/STRD/SSRD 在 OPGD 中也较强。"
                "这与草地水热限制和大气干旱敏感性一致，RECO 相比 GPP 更容易呈现热量控制。"
            )
            add_para(
                doc,
                "草地是多方法一致性较好的 biome。若 SHAP、ALE 与 ICE mean 同时显示 TMP 或 SSRD 的高值端效应增强，同时 OPGD q 值也较高，"
                "可以较有信心地将其解释为水热限制背景下的真实机制。PDP 若出现偏移，更多反映平均响应被相关变量组合拉动，而不是推翻 SHAP/ALE 的方向判断。"
            )
        elif biome == "Savanna":
            add_para(
                doc,
                "稀树草原中，SSRD_z 的 SHAP 占比最高，TMP_resid 与 STRD_resid 也保持前列；OPGD 对 STRD、TMP、EVA、VPD 的 q 值支持较强。"
                "这说明热干边缘生态系统中，能量输入、水分亏缺和热背景并不是单一因子作用，而是共同构成恢复时间差异。"
            )
            add_para(
                doc,
                "稀树草原的结果适合被写成能量-水分耦合机制：SSRD 提供强模型贡献，TMP/STRD/VPD/EVA 提供热干环境的空间分层解释。"
                "如果 ALE/ICE 的转折点接近 SHAP，则说明模型捕捉到的阈值具有较强稳健性；若阈值略有移动，也符合热干边缘系统中水分与热量交互导致响应分段的特征。"
            )
        else:
            add_para(
                doc,
                "灌丛中，GPP 与 RECO 的 SHAP 首位仍为 SSRD_z，但 GPP 的 SHAP 前三与 OPGD 前三重叠较弱，RECO 则在 SSRD、TMP、VPD 等机制上获得更明确的空间支持。"
                "因此灌丛中应谨慎把 GPP 的模型贡献直接解释为空间主导因子，而 RECO 的能量-温度-干旱机制更稳健。"
            )
            add_para(
                doc,
                "灌丛是差异解释最需要展开的一类。GPP 的正交 SHAP 中 Duration 或 STRD_resid 可能具有较高贡献，但 OPGD 更强调 TMP、VPD、SMrz 等原始水热背景，"
                "说明灌丛 GPP 恢复可能受事件过程、局地土壤水分和群落结构共同影响。此时应把 SHAP 写成模型识别到的预测贡献，把 OPGD 写成空间异质性的外部证据；"
                "二者不一致恰好提示灌丛恢复机制更复杂，而不是简单否认 SHAP 结果。"
            )
        png = Path(index.loc[index["biome"] == biome, "output_png"].iloc[0])
        add_image(
            doc,
            png,
            f"图 {BIOMES.index(biome) + 2}. {BIOME_CN[biome]}中正交 SHAP dependence 与 ALE、ICE mean、PDP 及 OPGD q 值的叠加对比。",
            4.85,
        )


def add_interpretation_section(doc: Document) -> None:
    add_heading(doc, "5. 综合解释：四类验证结果如何共同支撑结论", 1)
    add_para(
        doc,
        "正交 SHAP 与 ALE/PDP/ICE 的关系可以概括为“贡献排序”和“响应形态”的互证。SHAP beeswarm 和重要性排序回答哪个变量贡献更大，"
        "dependence plot 回答该变量在不同取值范围内如何推高或压低恢复时间，ALE/PDP/ICE 则检查这种方向是否是模型响应本身的稳定特征。"
        "因此，当某一变量同时满足 SHAP 排名靠前、dependence 趋势清晰、ALE 与 ICE mean 方向一致时，可以认为其模型解释较稳健。"
    )
    add_para(
        doc,
        "OPGD/地理探测器的作用不同。q 值衡量的是原始变量对恢复时间空间分层异质性的解释比例，不带正负方向，也不能单独证明因果关系。"
        "因此，OPGD 不应被用来替代 SHAP 的方向解释，而应作为空间层面的外部验证：若 SHAP 高贡献变量对应的原始因子也具有较高 q 值，"
        "则说明该机制不仅被机器学习模型使用，也能解释空间格局。"
    )
    add_para(
        doc,
        "从当前结果看，SSRD 的核心地位在正交 SHAP 中最稳健；TMP 和 STRD 在草地、稀树草原、森林及 RECO 相关结果中具有较强的跨方法支持；"
        "VPD 和 EVA 在不同 biome 中更多体现为水分胁迫和蒸散过程的补充解释；Pre、Duration 和 Intensity 的作用更依赖 biome 与响应指标，"
        "适合在事件属性或区域水分背景中解释，而不宜作为所有 biome 的统一主控因子。"
    )
    add_heading(doc, "6. 可写入论文或报告的结论段", 1)
    add_para(
        doc,
        "综合正交 SHAP、ALE、PDP、ICE mean 与 OPGD 地理探测器结果可以看出，闪旱后 GPP 与 RECO 恢复时间的差异主要由辐射能量输入、温度/长波热背景及水分胁迫共同塑造。"
        "在降低共线性后，SSRD 仍然是最稳定的模型贡献因子，说明其在原始 SHAP 中的重要性并非仅由与 STRD 或 TMP 的相关性造成。"
        "ALE 与 ICE mean 对多数高贡献因子的响应方向提供了模型层面的验证，而 OPGD q 值进一步表明 TMP、STRD、SSRD、VPD 和 EVA 等因子在不同 biome 中具有空间解释力。"
        "因此，本文可将 SHAP 解释为预测贡献和响应方向的证据，将 ALE/PDP/ICE 作为响应曲线稳健性检验，将 OPGD 作为空间异质性验证，共同支持能量-热量-水分耦合控制植被碳通量恢复时间的机制框架。"
    )


def build() -> Path:
    top3, feat, index = load_data()
    consistency = build_curve_consistency(feat)
    doc = new_doc()
    add_title(doc)
    add_methods_section(doc)
    add_reliability_section(doc, top3, feat)
    add_curve_consistency_section(doc, consistency)
    add_biome_sections(doc, feat, index, consistency)
    add_interpretation_section(doc)
    out = OUT / "12_SHAP_ALE_PDP_ICE_OPGD_validation_comparison_cn.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    path = build()
    print(path)
