#!/usr/bin/env python3
"""Write Chinese prepeak SHAP analysis documents for plots2 summary."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from scipy.stats import spearmanr


ROOT = Path("/home/xulc/flash_drought/process/SEM_analysis0401/codex/GLEAM")
RESULT_ROOT = ROOT / "results"
PLOTS_ROOT = ROOT / "plots2/prepeak_shap_summary_20260502"
DOC_ROOT = PLOTS_ROOT / "analysis_writing"

GPP_ROOT = RESULT_ROOT / "gpp_code1_flash_smrz_v20260401_onsetpeak_clean/prepeak_event_shap_sem_20260424/shap_by_biome"
RECO_ROOT = RESULT_ROOT / "reco_code1_flash_smrz_v20260401_mswepE_clean/prepeak_event_shap_sem_20260424/shap_by_biome"
BIOMES = ["Forest", "Grassland", "Savanna", "Cropland", "Shrubland"]
METRICS = {
    "GPP": GPP_ROOT,
    "RECO": RECO_ROOT,
}

SHORT_LABELS = {
    "prepeak_total_precipitation_mean": "PRE",
    "prepeak_total_evaporation_mean": "EVA",
    "prepeak_temperature_2m_mean": "TMP",
    "prepeak_VPD_mean": "VPD",
    "prepeak_SMrz_mean": "SMrz",
    "prepeak_lai_total_mean": "LAI",
    "prepeak_ssrd_mean": "SSRD",
    "prepeak_strd_mean": "STRD",
    "prepeak_wind_speed_mean": "WIND",
    "event_onset_days": "Onset",
    "event_duration": "Duration",
    "event_intensity": "Intensity",
}

BIOME_CN = {
    "Forest": "Forest",
    "Grassland": "Grassland",
    "Savanna": "Savanna",
    "Cropland": "Cropland",
    "Shrubland": "Shrubland",
}

METRIC_CN = {
    "GPP": "GPP 恢复时间",
    "RECO": "RECO 恢复时间",
}


@dataclass
class BiomeData:
    sample: pd.DataFrame
    shap: pd.DataFrame
    importance: pd.DataFrame


def short_label(name: str) -> str:
    return SHORT_LABELS.get(name, name.replace("prepeak_", ""))


def is_precip_feature(name: str) -> bool:
    return "precipitation" in name


def is_eva_feature(name: str) -> bool:
    return "evaporation" in name


def maybe_to_mm(values: np.ndarray) -> tuple[np.ndarray, str]:
    finite = values[np.isfinite(values)]
    if len(finite) == 0:
        return values, "mm"
    if np.nanquantile(np.abs(finite), 0.99) < 1.0:
        return values * 1000.0, "mm"
    return values, "mm"


def convert_feature_units(feature: str, values: np.ndarray) -> tuple[np.ndarray, str]:
    if is_precip_feature(feature) or is_eva_feature(feature):
        return maybe_to_mm(values)
    if feature == "prepeak_SMrz_mean":
        return values, "m3/m3"
    if feature == "prepeak_temperature_2m_mean":
        return values, "K"
    if feature == "prepeak_VPD_mean":
        return values, "kPa"
    if feature == "prepeak_wind_speed_mean":
        return values, "m/s"
    if feature in {"prepeak_ssrd_mean", "prepeak_strd_mean"}:
        return values, "W/m2"
    return values, ""


def load_biome(metric: str, biome: str) -> BiomeData:
    root = METRICS[metric] / biome
    return BiomeData(
        sample=pd.read_parquet(root / "dependence_sample_features.parquet"),
        shap=pd.read_parquet(root / "dependence_sample_shap_values.parquet"),
        importance=pd.read_csv(root / "feature_importance.csv"),
    )


def binned_curve(x: np.ndarray, y: np.ndarray, bins: int = 20) -> pd.DataFrame:
    valid = np.isfinite(x) & np.isfinite(y)
    x = x[valid]
    y = y[valid]
    if len(x) < 80:
        return pd.DataFrame(columns=["x", "y", "count"])
    qs = np.linspace(0, 1, bins + 1)
    edges = np.quantile(x, qs)
    edges = np.unique(edges)
    if len(edges) < 6:
        return pd.DataFrame(columns=["x", "y", "count"])
    cats = pd.cut(x, bins=edges, include_lowest=True, duplicates="drop")
    df = pd.DataFrame({"x": x, "y": y, "bin": cats})
    grp = df.groupby("bin", observed=True)
    out = grp.agg(x=("x", "median"), y=("y", "median"), count=("y", "size")).reset_index(drop=True)
    out["y_smooth"] = out["y"].rolling(window=3, center=True, min_periods=1).mean()
    return out


def classify_curve(curve: pd.DataFrame) -> tuple[str, float | None, str]:
    if curve.empty or len(curve) < 6:
        return "insufficient", None, "样本不足，无法稳定识别趋势。"
    xs = curve["x"].to_numpy(dtype=float)
    ys = curve["y_smooth"].to_numpy(dtype=float)
    amp = float(np.nanmax(ys) - np.nanmin(ys))
    if not np.isfinite(amp) or amp < 0.05:
        return "flat", None, "SHAP 响应振幅较小，整体呈弱变化。"
    diffs = np.diff(ys)
    eps = max(amp * 0.08, 1e-6)
    signs = np.sign(np.where(np.abs(diffs) < eps, 0.0, diffs))
    nz = signs[signs != 0]
    if len(nz) == 0:
        return "flat", None, "平滑曲线接近平坦。"
    change_idx = np.where(np.diff(nz) != 0)[0]
    overall = ys[-1] - ys[0]
    if len(change_idx) == 0:
        if overall > 0:
            return "monotonic_increase", None, "随特征值增大，SHAP 值整体上升。"
        return "monotonic_decrease", None, "随特征值增大，SHAP 值整体下降。"
    peak_idx = int(np.argmax(ys))
    trough_idx = int(np.argmin(ys))
    if peak_idx not in (0, len(ys) - 1) and ys[peak_idx] - max(ys[0], ys[-1]) > amp * 0.2:
        return "inverted_u", float(xs[peak_idx]), "中间区间 SHAP 最高，随后回落，表现为倒 U 型。"
    if trough_idx not in (0, len(ys) - 1) and min(ys[0], ys[-1]) - ys[trough_idx] > amp * 0.2:
        return "u_shape", float(xs[trough_idx]), "中间区间 SHAP 最低，随后回升，表现为 U 型。"
    turn_idx = int(np.argmax(np.abs(np.diff(ys))))
    turn_x = float(xs[min(turn_idx + 1, len(xs) - 1)])
    if overall > 0:
        return "threshold_rise", turn_x, "低值段变化较缓，越过阈值后正向效应增强。"
    return "threshold_drop", turn_x, "低值段变化较缓，越过阈值后负向效应增强。"


def format_num(val: float | None) -> str:
    if val is None or not np.isfinite(val):
        return "NA"
    aval = abs(val)
    if aval >= 100:
        return f"{val:.1f}"
    if aval >= 10:
        return f"{val:.2f}"
    if aval >= 1:
        return f"{val:.3f}"
    return f"{val:.4f}"


def feature_physics(label: str, trend_type: str, direction: str) -> str:
    if label == "PRE":
        return "降水通过补给土壤水分、抬升根区可利用水来缓冲干旱记忆，因此其效应通常要结合 SMrz 与 EVA 一起解释；当降水只停留在很低水平时，少量补水往往只能改变表层湿润状态，超过阈值后才更明显传导到恢复时间。"
    if label == "EVA":
        return "蒸散同时反映能量供给和植被耗水强度。若高 EVA 对应更短恢复时间，通常说明蒸散更多地代表活跃能量-碳交换背景；若高 EVA 对应更长恢复时间，则意味着耗水负担在土壤储水不足时放大了恢复拖尾。"
    if label == "SSRD":
        return "短波辐射控制冠层光合可用能量，但其作用并不总是单向。能量受限生态系统中更高 SSRD 往往有利于恢复，而水分受限生态系统中，高辐射会通过升温和蒸散需求把恢复推向更长尾。"
    if label == "STRD":
        return "长波辐射通常和暖湿背景、大气逆辐射及夜间保温有关，既可能缓解低温限制，也可能在高温高湿环境中增强呼吸负担，所以常与 TMP 联合解释。"
    if label == "TMP":
        return "温度一方面决定酶活性和生长节律，另一方面又通过提高饱和水汽压和蒸散需求强化干旱胁迫，因此温度曲线经常表现出阈值或双重效应。"
    if label == "VPD":
        return "VPD 直接表征大气干旱程度，升高后会抬升气孔关闭与蒸散失配风险；如果曲线在高值区突然转陡，通常意味着大气干旱已超过植被调节能力。"
    if label == "SMrz":
        return "根区土壤水是把前期降水、蒸散和辐射历史整合到恢复过程中的关键蓄水变量。较高 SMrz 往往缩短恢复，但在湿润生态系统中其边际效应可能较快饱和。"
    if label == "WIND":
        return "风速会通过增强湍流交换、提高蒸散需求和加快边界层耦合来放大能量-水分交换，因此只有在与 EVA、VPD 或 SMrz 一起看时才容易解释清楚。"
    if label == "Duration":
        return "事件持续时间刻画了胁迫记忆累积的长度，持续越久，非结构性碳储备和根区水库被消耗得越深，恢复通常越慢。"
    if label == "Intensity":
        return "事件强度反映胁迫幅度。强度升高通常意味着更深的碳同化损失或呼吸异常，恢复时间因此更容易拉长。"
    if label == "Onset":
        return "发生时间会改变事件与生长季阶段的重叠程度，从而决定植被还能否利用后续能量和降水窗口进行补偿。"
    if label == "LAI":
        return "LAI 代表冠层规模和叶面积投资，既影响蒸腾需求，也影响光能截获，因此其方向常取决于水分是否足以支撑较大的叶面积。"
    return "这一特征的物理含义需要结合其与水分、能量和根区储水的耦合关系共同解释。"


def compute_stats() -> pd.DataFrame:
    rows = []
    for metric in ["GPP", "RECO"]:
        for biome in BIOMES:
            data = load_biome(metric, biome)
            importance_lookup = dict(zip(data.importance["feature"], data.importance["importance"]))
            for rank, feature in enumerate(data.importance["feature"].tolist(), start=1):
                x = pd.to_numeric(data.sample[feature], errors="coerce").to_numpy(dtype=float)
                x, unit = convert_feature_units(feature, x)
                y = pd.to_numeric(data.shap[feature], errors="coerce").to_numpy(dtype=float)
                valid = np.isfinite(x) & np.isfinite(y)
                xv = x[valid]
                yv = y[valid]
                if len(xv) < 30:
                    continue
                rho = spearmanr(xv, yv, nan_policy="omit").statistic
                q30 = np.nanquantile(xv, 0.3)
                q70 = np.nanquantile(xv, 0.7)
                low_median = float(np.nanmedian(yv[xv <= q30]))
                high_median = float(np.nanmedian(yv[xv >= q70]))
                direction = "positive" if high_median > low_median else "negative"
                curve = binned_curve(xv, yv, bins=20)
                trend_type, threshold, trend_desc = classify_curve(curve)
                rows.append(
                    {
                        "metric": metric,
                        "biome": biome,
                        "feature": feature,
                        "label": short_label(feature),
                        "rank": rank,
                        "importance": float(importance_lookup[feature]),
                        "unit": unit,
                        "spearman_rho": float(rho) if np.isfinite(rho) else np.nan,
                        "low_shap_median": low_median,
                        "high_shap_median": high_median,
                        "direction": direction,
                        "trend_type": trend_type,
                        "threshold": threshold,
                        "trend_desc": trend_desc,
                        "physics_hint": feature_physics(short_label(feature), trend_type, direction),
                    }
                )
    out = pd.DataFrame(rows).sort_values(["metric", "biome", "rank"])
    return out


def ensure_doc_style(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Title"].font.size = Pt(16)
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 2"].font.size = Pt(12)


def add_title(doc: Document, title: str) -> None:
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.35
    p.add_run(text)


def add_image(doc: Document, path: Path, width_cm: float = 15.5, caption: str | None = None) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.line_spacing = 1.15
        c.add_run(caption)


def topn_rows(stats: pd.DataFrame, metric: str, biome: str, n: int = 5) -> pd.DataFrame:
    return stats[(stats["metric"] == metric) & (stats["biome"] == biome)].sort_values("rank").head(n)


def overall_intro(stats: pd.DataFrame, metric: str) -> str:
    sub = stats[stats["metric"] == metric]
    top1 = (
        sub[sub["rank"] == 1]
        .groupby("label")
        .size()
        .sort_values(ascending=False)
        .to_dict()
    )
    items = [f"{k} 在 {v} 个 biome 中位列第 1" for k, v in top1.items()]
    biome_details = []
    for biome in BIOMES:
        labels = topn_rows(stats, metric, biome, 3)["label"].tolist()
        biome_details.append(f"{biome} 的前三位分别为 {', '.join(labels)}")
    return (
        f"从 {METRIC_CN[metric]} 的五个 biome 整体排序看，"
        + "，".join(items)
        + "。这说明整个骤旱期间，"
        "恢复时间首先受能量供给与蒸散耗水背景控制，其次才由根区储水和事件属性调节。"
        + "；".join(biome_details)
        + "。前者更偏向能量受限或能量-耗水耦合控制，后者则说明不同 biome 对大气干旱、土壤蓄水和事件记忆的敏感性存在显著差异。"
    )


def biome_contribution_para(stats: pd.DataFrame, metric: str, biome: str) -> str:
    rows = topn_rows(stats, metric, biome, 5)
    pieces = []
    for _, row in rows.iterrows():
        pieces.append(
            f"{row['label']} 排名第 {int(row['rank'])}，平均绝对 SHAP 贡献为 {row['importance']:.2f}；"
            f"其高值段 SHAP 中位数为 {row['high_shap_median']:.2f}，低值段为 {row['low_shap_median']:.2f}，"
            f"表现为高值更易把恢复时间推向{'更长' if row['direction'] == 'positive' else '更短'}。"
        )
    summary = " ".join(pieces)
    dominant = rows.iloc[0]
    energy_features = {"SSRD", "STRD", "TMP", "VPD", "WIND"}
    water_features = {"PRE", "EVA", "SMrz"}
    energy_count = int(rows["label"].isin(energy_features).sum())
    water_count = int(rows["label"].isin(water_features).sum())
    same_feature_all = stats[(stats["metric"] == metric) & (stats["label"] == dominant["label"])]
    median_importance = float(same_feature_all["importance"].median())
    intensity_phrase = "高于" if dominant["importance"] > median_importance else "低于或接近"
    return (
        f"{BIOME_CN[biome]} 中，{summary} "
        f"其中最核心的控制项是 {dominant['label']}，这意味着该 biome 的恢复尾部主要由"
        f"{'能量输入增强后引发的耗水或呼吸负担' if dominant['label'] in {'SSRD', 'TMP', 'VPD', 'STRD'} else '水分补给与根区储水记忆'}"
        f"来塑造。该特征在 {metric} 五个 biome 内的贡献强度{intensity_phrase}其中位水平，"
        f"同时前五特征中能量相关变量有 {energy_count} 个，水分相关变量有 {water_count} 个，"
        f"说明该 biome 更偏向{'能量-大气干旱主导' if energy_count > water_count else '水分蓄存与耗水平衡主导'}。"
        f"{dominant['physics_hint']}"
    )


def dependence_para(stats: pd.DataFrame, metric: str, biome: str, feature_label: str) -> str:
    row = stats[(stats["metric"] == metric) & (stats["biome"] == biome) & (stats["label"] == feature_label)]
    if row.empty:
        return ""
    row = row.iloc[0]
    threshold_text = ""
    if pd.notna(row["threshold"]):
        threshold_text = f" 识别到的转折阈值约为 {format_num(row['threshold'])} {row['unit']}。"
    return (
        f"{BIOME_CN[biome]} 的 {feature_label} dependence plot 显示，曲线属于 {row['trend_type']} 类型，"
        f"Spearman 相关为 {row['spearman_rho']:.2f}；高值段相对低值段使 SHAP 中位数由 {row['low_shap_median']:.2f} 变化到 {row['high_shap_median']:.2f}。"
        f"{row['trend_desc']}{threshold_text} {row['physics_hint']}"
    )


def paired_compare_para(stats: pd.DataFrame, biome: str, feature_label: str) -> str:
    g = stats[(stats["metric"] == "GPP") & (stats["biome"] == biome) & (stats["label"] == feature_label)]
    r = stats[(stats["metric"] == "RECO") & (stats["biome"] == biome) & (stats["label"] == feature_label)]
    if g.empty or r.empty:
        return ""
    g = g.iloc[0]
    r = r.iloc[0]
    stronger = "RECO" if r["importance"] > g["importance"] else "GPP"
    rank_diff = int(g["rank"]) - int(r["rank"])
    rank_text = (
        "两者排序相同"
        if rank_diff == 0
        else f"{feature_label} 在 GPP 中排第 {int(g['rank'])}，在 RECO 中排第 {int(r['rank'])}"
    )
    threshold_bits = []
    if pd.notna(g["threshold"]):
        threshold_bits.append(f"GPP 的阈值约为 {format_num(g['threshold'])} {g['unit']}")
    if pd.notna(r["threshold"]):
        threshold_bits.append(f"RECO 的阈值约为 {format_num(r['threshold'])} {r['unit']}")
    threshold_text = "；".join(threshold_bits)
    return (
        f"在 {BIOME_CN[biome]} 中，{feature_label} 对 GPP 与 RECO 的作用并不完全一致。"
        f"GPP 侧的平均绝对 SHAP 为 {g['importance']:.2f}，RECO 侧为 {r['importance']:.2f}；"
        f"GPP 的高值段效应指向恢复时间{'延长' if g['direction'] == 'positive' else '缩短'}，"
        f"而 RECO 的高值段效应指向恢复时间{'延长' if r['direction'] == 'positive' else '缩短'}。"
        f"{rank_text}，且 {stronger} 一侧对该变量更敏感。"
        f"{(' ' + threshold_text) if threshold_text else ''}这意味着同一背景变量并不是简单同时控制碳同化和生态系统呼吸，而是通过不同的生理链条改变二者的恢复尾部。"
    )


def write_stats_files(stats: pd.DataFrame) -> None:
    DOC_ROOT.mkdir(parents=True, exist_ok=True)
    stats.to_csv(DOC_ROOT / "feature_direction_threshold_summary.csv", index=False)
    summary = (
        stats.groupby(["metric", "biome"], as_index=False)
        .agg(
            top5=("label", lambda s: ", ".join(s.head(5).tolist())),
            dominant_feature=("label", "first"),
            mean_top5_importance=("importance", lambda s: float(s.head(5).mean())),
        )
    )
    summary.to_csv(DOC_ROOT / "biome_top5_summary.csv", index=False)


def pick_dependence_features(stats: pd.DataFrame, metric: str, biome: str) -> list[str]:
    rows = topn_rows(stats, metric, biome, 5)
    preferred = []
    for feat in ["PRE", "SSRD", "EVA", "TMP", "VPD", "SMrz", "STRD"]:
        if feat in rows["label"].tolist():
            preferred.append(feat)
    for feat in rows["label"].tolist():
        if feat not in preferred:
            preferred.append(feat)
    return preferred[:3]


def write_doc1(stats: pd.DataFrame) -> Path:
    doc = Document()
    ensure_doc_style(doc)
    add_title(doc, "01 不同 Biome 的特征贡献分析")
    add_para(
        doc,
        "本文件基于整个骤旱期间的 SHAP beeswarm 结果，对五类 biome 中 GPP 与 RECO 恢复时间的主导特征进行归纳。分析重点不只是贡献排序，还包括高值与低值分别把恢复时间推向哪个方向，并结合能量、水分与根区储水过程给出物理解释。"
    )
    add_image(
        doc,
        PLOTS_ROOT / "beeswarm_comparison_5biomes_gpp_vs_reco.png",
        width_cm=16.2,
        caption="图 1. 五个 biome 中 GPP 与 RECO 的 SHAP beeswarm 对比图。",
    )
    for metric in ["GPP", "RECO"]:
        doc.add_heading(f"{METRIC_CN[metric]} 的 biome 差异", level=1)
        add_para(doc, overall_intro(stats, metric))
        for biome in BIOMES:
            add_image(
                doc,
                PLOTS_ROOT / metric / biome / f"{metric}_{biome}_beeswarm_redraw.png",
                width_cm=14.8,
                caption=f"{metric} - {biome} beeswarm 图",
            )
            add_para(doc, biome_contribution_para(stats, metric, biome))
    out = DOC_ROOT / "01_biome_feature_contribution_analysis_cn.docx"
    doc.save(out)
    return out


def write_doc2(stats: pd.DataFrame) -> Path:
    doc = Document()
    ensure_doc_style(doc)
    add_title(doc, "02 Dependence Plot 与阈值机制分析")
    add_para(
        doc,
        "本文件聚焦 dependence plot，识别不同 biome 中关键特征的单调响应、阈值响应以及 U 型或倒 U 型转折。阈值的判读来自分位数平滑曲线，而不是单看散点局部波动，因此更适合和物理机制一起解释。"
    )
    for metric in ["GPP", "RECO"]:
        doc.add_heading(f"{METRIC_CN[metric]} 的 dependence 模式", level=1)
        for biome in BIOMES:
            feats = pick_dependence_features(stats, metric, biome)
            doc.add_heading(f"{biome}", level=2)
            for feat in feats:
                add_image(
                    doc,
                    PLOTS_ROOT / metric / biome / "dependence_top5" / f"{feat}_colored_by_EVA.png",
                    width_cm=14.8,
                    caption=f"{metric} - {biome} - {feat}（EVA 着色）",
                )
                add_para(doc, dependence_para(stats, metric, biome, feat))
                add_image(
                    doc,
                    PLOTS_ROOT / metric / biome / "dependence_top5" / f"{feat}_colored_by_SMrz.png",
                    width_cm=14.8,
                    caption=f"{metric} - {biome} - {feat}（SMrz 着色）",
                )
                add_para(
                    doc,
                    f"将 {feat} 再用 SMrz 着色的目的，是检验相同特征值区间内，根区储水是否改变 SHAP 响应的离散带宽。"
                    f"如果同一横轴区间内高 SMrz 与低 SMrz 分别集中在 SHAP 的上下两侧，就说明该特征并不是独立作用，而是通过土壤储水状态重新分配其恢复效应。"
                )
    out = DOC_ROOT / "02_dependence_threshold_analysis_cn.docx"
    doc.save(out)
    return out


def write_doc3(stats: pd.DataFrame) -> Path:
    doc = Document()
    ensure_doc_style(doc)
    add_title(doc, "03 GPP 与 RECO 的差异比较分析")
    add_para(
        doc,
        "本文件比较 GPP 与 RECO 恢复时间在同一 biome 内对相同背景因子的响应差异。比较包括两部分：其一是 beeswarm 排位与贡献方向的差异，其二是 dependence plot 中单调性、阈值位置与交互离散度的差异。"
    )
    add_image(
        doc,
        PLOTS_ROOT / "beeswarm_comparison_5biomes_gpp_vs_reco.png",
        width_cm=16.2,
        caption="图 1. GPP 与 RECO 在五个 biome 中的 beeswarm 对比。",
    )
    for biome in BIOMES:
        doc.add_heading(f"{biome} 的 GPP-RECO 对比", level=1)
        g_top = topn_rows(stats, "GPP", biome, 5)["label"].tolist()
        r_top = topn_rows(stats, "RECO", biome, 5)["label"].tolist()
        common = [feat for feat in g_top if feat in r_top]
        if not common:
            common = list(dict.fromkeys(g_top[:3] + r_top[:3]))[:3]
        add_para(
            doc,
            f"{biome} 中，GPP 前五贡献特征为 {', '.join(g_top)}，RECO 前五贡献特征为 {', '.join(r_top)}。"
            f"共同进入前五的特征包括 {', '.join(common)}。这种重叠说明二者共享同一组水热背景约束，但排序先后和方向差异反映了碳同化与呼吸恢复并非同步。"
        )
        for feat in common[:3]:
            add_para(doc, paired_compare_para(stats, biome, feat))
            add_image(
                doc,
                PLOTS_ROOT / "GPP" / biome / "dependence_top5" / f"{feat}_colored_by_EVA.png",
                width_cm=14.8,
                caption=f"GPP - {biome} - {feat}（EVA 着色）",
            )
            add_image(
                doc,
                PLOTS_ROOT / "RECO" / biome / "dependence_top5" / f"{feat}_colored_by_EVA.png",
                width_cm=14.8,
                caption=f"RECO - {biome} - {feat}（EVA 着色）",
            )
    out = DOC_ROOT / "03_gpp_reco_difference_analysis_cn.docx"
    doc.save(out)
    return out


def main() -> None:
    stats = compute_stats()
    write_stats_files(stats)
    out1 = write_doc1(stats)
    out2 = write_doc2(stats)
    out3 = write_doc3(stats)
    print(out1)
    print(out2)
    print(out3)


if __name__ == "__main__":
    main()
