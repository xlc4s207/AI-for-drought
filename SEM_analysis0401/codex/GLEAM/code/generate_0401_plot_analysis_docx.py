#!/usr/bin/env python3
"""Generate three academic-style Word documents for 0401 SHAP/SEM plot analysis."""

from __future__ import annotations

from collections import Counter
from dataclasses import dataclass
from functools import lru_cache
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path("/home/xulc/flash_drought/process/SEM_analysis0401/codex/GLEAM/plots")
OUT_DIR = ROOT / "analysis_writing"
RESULT_ROOT = ROOT.parent / "results"
BIOMES = ["Forest", "Grassland", "Savanna", "Cropland", "Shrubland"]

GROUPS = [
    ("GPP_prepeak", "GPP", "prepeak"),
    ("GPP_recovery", "GPP", "recoverywin"),
    ("RECO_prepeak", "RECO", "prepeak"),
    ("RECO_recovery", "RECO", "recoverywin"),
]

GROUP_LABELS = {
    "GPP_prepeak": "GPP prepeak（峰值前背景）",
    "GPP_recovery": "GPP recovery（恢复期过程）",
    "RECO_prepeak": "RECO prepeak（峰值前背景）",
    "RECO_recovery": "RECO recovery（恢复期过程）",
}

PHASE_LABELS = {
    "prepeak": "峰值前背景",
    "recoverywin": "恢复期过程",
}

FEATURE_LABELS = {
    "PRE (mm)": "降水补给",
    "EVA (mm)": "蒸散耗水",
    "SSRD": "短波辐射背景",
    "STRD": "长波辐射背景",
    "TMP (K)": "气温",
    "SMrz": "根区土壤水",
    "VPD": "大气干旱需求",
    "WIND": "风速耦合",
    "DUR": "骤旱持续时间",
    "INT": "骤旱强度",
    "ONS": "事件发生时序",
    "LAI": "冠层状态",
}

BEESWARM_DIRS = {
    key: ROOT / "beeswarm" / key for key, *_ in GROUPS
}

DEPENDENCE_DIRS = {
    key: ROOT / "dependence_plot" / key for key, *_ in GROUPS
}

COMPARE_DIRS = {
    "GPP": ROOT / "gpp_code1_flash_smrz_compare_prepeak_vs_recoverywin_v20260401_20260424",
    "RECO": ROOT / "reco_code1_flash_smrz_compare_prepeak_vs_recoverywin_v20260401_20260424",
}

DEPENDENCE_DATA_DIRS = {
    "GPP_prepeak": RESULT_ROOT / "gpp_code1_flash_smrz_v20260401_onsetpeak_clean/prepeak_event_shap_sem_20260424/shap_by_biome",
    "GPP_recovery": RESULT_ROOT / "gpp_code1_flash_smrz_v20260401_recoverywin_clean/shap_process_recoverywin_precipEmean_sample50k_by_biome",
    "RECO_prepeak": RESULT_ROOT / "reco_code1_flash_smrz_v20260401_mswepE_clean/prepeak_event_shap_sem_20260424/shap_by_biome",
    "RECO_recovery": RESULT_ROOT / "reco_code1_flash_smrz_v20260401_mswepE_clean/shap_process_recoverywin_precipEmean_sample50k_by_biome",
}

TURNING_THRESHOLD_PATH = ROOT / "analysis_writing/dependence_turning_thresholds_0401.csv"

FEATURE_PRIORITY = [
    "SSRD",
    "PRE (mm)",
    "EVA (mm)",
    "TMP (K)",
    "VPD",
    "STRD",
    "WIND",
    "SMrz",
    "DUR",
    "INT",
    "ONS",
    "LAI",
]

TURNING_FEATURES = {"PRE", "EVA", "SSRD"}


@dataclass
class FigureItem:
    path: Path
    caption: str
    text: str


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def set_run_font(run, size: float, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "SimSun")


def set_doc_language(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    r_pr = normal.element.rPr
    r_fonts = r_pr.rFonts
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "SimSun")


def add_page_numbers(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_run_font(run, 10)


def make_document(title: str) -> Document:
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    add_page_numbers(section)
    set_doc_language(document)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_run_font(run, 16, bold=True)
    p.paragraph_format.space_after = Pt(12)
    return document


def add_heading(document: Document, text: str, level: int = 1) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6)
    if level == 1:
        p.style = document.styles["Heading 1"]
        size = 14
    elif level == 2:
        p.style = document.styles["Heading 2"]
        size = 13
    else:
        size = 12
    run = p.add_run(text)
    set_run_font(run, size, bold=True)


def add_paragraph(document: Document, text: str, first_line_cm: float = 0.74) -> None:
    p = document.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(first_line_cm)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, 12)


def add_figure(document: Document, fig: FigureItem, width_cm: float = 15.5) -> None:
    if not fig.path.exists():
        add_paragraph(document, f"缺图占位：{fig.path}")
        return
    if fig.text:
        add_paragraph(document, fig.text)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(fig.path), width=Cm(width_cm))
    p.paragraph_format.space_after = Pt(3)
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    run = cap.add_run(fig.caption)
    set_run_font(run, 10.5)


def save(document: Document, path: Path) -> None:
    ensure_dir(path.parent)
    document.save(path)


def decode_feature(label: str) -> str:
    return FEATURE_LABELS.get(label, label)


def read_beeswarm_summary(group_key: str) -> pd.DataFrame:
    return pd.read_csv(BEESWARM_DIRS[group_key] / "summary.csv")


@lru_cache(maxsize=None)
def read_turning_thresholds() -> pd.DataFrame:
    return pd.read_csv(TURNING_THRESHOLD_PATH)


def top_features_for_biome(group_key: str, biome: str, top_n: int = 3) -> list[str]:
    df = read_beeswarm_summary(group_key)
    row = df.loc[df["biome"] == biome].iloc[0]
    labels = [item.strip() for item in str(row["labels_used"]).split(",") if item.strip()]
    return labels[:top_n]


def top_feature_pairs_for_biome(group_key: str, biome: str, top_n: int = 3) -> list[tuple[str, str]]:
    df = read_beeswarm_summary(group_key)
    row = df.loc[df["biome"] == biome].iloc[0]
    features = [item.strip() for item in str(row["features_used"]).split(",") if item.strip()]
    labels = [item.strip() for item in str(row["labels_used"]).split(",") if item.strip()]
    return list(zip(features[:top_n], labels[:top_n]))


def group_recurrent_features(group_key: str, top_n: int = 4) -> list[str]:
    df = read_beeswarm_summary(group_key)
    counter: Counter[str] = Counter()
    for labels_used in df["labels_used"].tolist():
        labels = [item.strip() for item in str(labels_used).split(",") if item.strip()][:4]
        counter.update(labels)
    return [name for name, _ in counter.most_common(top_n)]


def beeswarm_image_path(group_key: str, biome: str, phase: str) -> Path:
    return BEESWARM_DIRS[group_key] / f"{biome}_{phase}_beeswarm_shortlabels.png"


def dependence_image_path(
    group_key: str,
    biome: str,
    phase: str,
    x_key: str,
    color_suffix: str = "SMrz",
) -> Path:
    return DEPENDENCE_DIRS[group_key] / f"{biome}_{phase}_{x_key}_colored_by_{color_suffix}.png"


def shared_intro_paragraphs() -> list[str]:
    return [
        "本组文档围绕全球不同植被类型在 SMrz 驱动骤旱后恢复时间的控制机制展开结果性分析。图件基础全部来自 0401 口径下的 SHAP 与 dependence plot 结果，恢复时间被视为事件后碳通量重新回到基线所需的持续时长，因此其驱动机制既包含峰值前背景形成的恢复记忆，也包含恢复期内部的近端环境约束。",
        "在解释框架上，beeswarm 图用于识别不同特征对恢复时间的整体贡献顺序和正负离散范围；dependence plot 用于识别关键特征在不同取值区间内对恢复时间的非线性影响，并通过颜色映射展示与根区土壤水状态的交互调制；比较分析则进一步区分峰值前背景控制与恢复期过程控制，并对 GPP 与 RECO 两类碳通量在恢复机制上的异同进行并列讨论。",
    ]


@lru_cache(maxsize=None)
def read_dependence_plot_data(group_key: str, biome: str) -> pd.DataFrame:
    path = DEPENDENCE_DATA_DIRS[group_key] / biome / "dependence_plot_data.parquet"
    return pd.read_parquet(path)


def feature_code(label: str) -> str:
    mapping = {
        "PRE (mm)": "PRE",
        "EVA (mm)": "EVA",
        "SSRD": "SSRD",
        "STRD": "STRD",
        "TMP (K)": "TMP",
        "VPD": "VPD",
        "WIND": "WIND",
        "SMrz": "SMrz",
        "DUR": "DUR",
        "INT": "INT",
        "ONS": "ONS",
        "LAI": "LAI",
    }
    return mapping.get(label, label)


def format_value(value: float, code: str) -> str:
    abs_value = abs(float(value))
    if code in {"SSRD", "STRD"}:
        return f"{value / 1e7:.2f}×10^7"
    if abs_value >= 100:
        return f"{value:.1f}"
    if abs_value >= 10:
        return f"{value:.2f}"
    if abs_value >= 1:
        return f"{value:.2f}"
    if abs_value >= 0.01:
        return f"{value:.3f}"
    return f"{value:.2e}"


def quantile_direction_stats(group_key: str, biome: str, feature_col: str) -> dict[str, float]:
    df = read_dependence_plot_data(group_key, biome)
    x_col = f"feature__{feature_col}"
    y_col = f"shap__{feature_col}"
    q10, q50, q90 = df[x_col].quantile([0.1, 0.5, 0.9]).tolist()
    q40, q60 = df[x_col].quantile([0.4, 0.6]).tolist()
    low = float(df.loc[df[x_col] <= q10, y_col].median())
    mid = float(df.loc[(df[x_col] >= q40) & (df[x_col] <= q60), y_col].median())
    high = float(df.loc[df[x_col] >= q90, y_col].median())
    return {
        "q10": float(q10),
        "q50": float(q50),
        "q90": float(q90),
        "low": low,
        "mid": mid,
        "high": high,
    }


def turning_row(group_key: str, biome: str, label: str) -> pd.Series | None:
    code = feature_code(label)
    if code not in TURNING_FEATURES:
        return None
    df = read_turning_thresholds()
    matched = df[(df["group"] == group_key) & (df["biome"] == biome) & (df["feature"] == code)]
    if matched.empty:
        return None
    return matched.iloc[0]


def direction_category(group_key: str, biome: str, feature_col: str, label: str) -> str:
    row = turning_row(group_key, biome, label)
    if row is not None and float(row["y_peak"]) > 0 and float(row["y_trough"]) < 0:
        if float(row["x_peak"]) < float(row["x_trough"]):
            return "high_negative"
        return "midhigh_positive"
    stats = quantile_direction_stats(group_key, biome, feature_col)
    if stats["low"] < 0 < stats["high"]:
        return "high_positive"
    if stats["low"] > 0 > stats["high"]:
        return "high_negative"
    if stats["high"] - stats["low"] >= 6:
        return "high_positive"
    if stats["low"] - stats["high"] >= 6:
        return "high_negative"
    return "mixed"


def feature_direction_sentence(group_key: str, biome: str, feature_col: str, label: str) -> str:
    code = feature_code(label)
    stats = quantile_direction_stats(group_key, biome, feature_col)
    row = turning_row(group_key, biome, label)
    if row is not None and float(row["y_peak"]) > 0 and float(row["y_trough"]) < 0:
        peak_x = format_value(float(row["x_peak"]), code)
        trough_x = format_value(float(row["x_trough"]), code)
        peak_y = float(row["y_peak"])
        trough_y = float(row["y_trough"])
        if float(row["x_peak"]) < float(row["x_trough"]):
            return (
                f"{label} 的方向上，高值端更容易落到负 SHAP 侧：曲线先在 {peak_x} 左右达到正峰"
                f"（约 {peak_y:.1f}），随后随着取值继续升高转向负侧，并在 {trough_x} 左右落到负谷"
                f"（约 {trough_y:.1f}）。按 10%/90% 分位比较，低值端 SHAP 中位数为 {stats['low']:.1f}，"
                f"高值端为 {stats['high']:.1f}。"
            )
        return (
            f"{label} 在低值端更偏负，但升到中高值后会明显转到正 SHAP 侧：负谷出现在 {trough_x} 左右"
            f"（约 {trough_y:.1f}），正峰出现在更高的 {peak_x} 左右（约 {peak_y:.1f}）。"
            f"对应地，10% 分位的 SHAP 中位数为 {stats['low']:.1f}，90% 分位提高到 {stats['high']:.1f}。"
        )
    if stats["low"] < 0 < stats["high"]:
        return (
            f"{label} 的高值端更偏正 SHAP 侧：10% 分位的 SHAP 中位数为 {stats['low']:.1f}，"
            f"而 90% 分位升到 {stats['high']:.1f}，说明该变量主要是在高值端开始表现为恢复拖尾因子。"
        )
    if stats["low"] > 0 > stats["high"]:
        return (
            f"{label} 的高值端更偏负 SHAP 侧：低值端 SHAP 中位数为 {stats['low']:.1f}，"
            f"高值端降到 {stats['high']:.1f}，说明变量增大后更容易把贡献重心压向缩短恢复时间的一侧。"
        )
    if stats["high"] - stats["low"] >= 6:
        return (
            f"{label} 虽然没有表现为严格过零翻转，但高值端明显更偏正侧：低值端与高值端的 SHAP 中位数"
            f"分别为 {stats['low']:.1f} 和 {stats['high']:.1f}，增幅已经足以说明方向性。"
        )
    if stats["low"] - stats["high"] >= 6:
        return (
            f"{label} 的方向更像是“随取值升高逐步转负”：低值端和高值端的 SHAP 中位数分别为"
            f" {stats['low']:.1f} 和 {stats['high']:.1f}，说明高值区更容易进入负贡献一侧。"
        )
    return (
        f"{label} 的正负两翼并不是完全单向展开，10%/90% 分位的 SHAP 中位数分别为"
        f" {stats['low']:.1f} 和 {stats['high']:.1f}，提示它更像是背景修饰项，而不是单一方向的硬阈值。"
    )


def feature_mechanism_sentence(group_key: str, biome: str, feature_col: str, label: str) -> str:
    metric, phase = group_key.split("_", 1)
    category = direction_category(group_key, biome, feature_col, label)

    if label == "SSRD":
        if phase == "prepeak" and category == "high_negative":
            return "这更符合“峰值前辐射记忆”机制：事件前较高短波输入通常意味着更强地表能量负荷和更积极的蒸散消耗，它本身未必直接决定恢复，但会通过放大根区失水和水热失衡，把系统带到一个更难恢复的起点。"
        if phase == "recoverywin" and category in {"high_positive", "midhigh_positive"}:
            return "这说明恢复期短波辐射并不是单向有利或不利，而更像一个窗口变量：中等偏高辐射可以维持冠层活动和蒸散重建，使系统仍处在“尚未恢复完毕但仍在运转”的阶段；只有再往更高端走时，额外能量负荷才会重新转写为耗水压力。"
        if phase == "recoverywin" and category == "high_negative":
            return "这提示在该 biome 中，恢复期高短波输入更容易越过“能量支撑”而进入“辐射放大耗水”的区间，也就是辐射不再主要帮助功能恢复，而是通过增强蒸散和叶面热负荷压缩有效恢复窗口。"
        return "总体上，SSRD 更适合被理解为能量约束变量，它的效应通常不是直接作用于恢复时间，而是通过改变蒸散需求、叶面能量平衡和根区失水速率来间接塑造恢复尾部。"

    if label == "STRD":
        if phase == "recoverywin" and category == "high_negative":
            return "从物理上看，这更像是长波热背景过强后的热负荷效应：夜间冷却受限、暖湿背景维持和净热积累增强，会使补水收益被热压力部分抵消，因此恢复并不会随着热背景增强而线性改善。"
        if phase == "recoverywin" and category in {"high_positive", "midhigh_positive"}:
            return "这说明中高长波背景在该情景下更像是维持恢复活动的热环境支撑，但它依然不是无限有利；一旦继续升高，往往意味着热背景开始压过补水和功能重建的收益。"
        return "STRD 在这里更接近整体热背景指标，它反映的不是瞬时地表光照，而是系统在较暖背景中维持恢复活动还是承受额外热负荷。"

    if label == "PRE (mm)":
        if phase == "recoverywin" and category in {"high_positive", "midhigh_positive"}:
            return "这最符合“恢复窗口补水”机制：极低降水对应的是恢复期内持续缺水，因此通常把 SHAP 压在负侧；当降水进入中等区间后，系统开始获得有效回补，但这种回补并不意味着立刻恢复完成，反而常对应一个仍需持续重配水分和功能重建的恢复阶段。"
        if phase == "prepeak" and category in {"high_positive", "midhigh_positive"}:
            return "放在峰值前背景里，这说明 PRE 更像是在刻画事件形成前的水分起点，而不是恢复期内即时补水本身。也就是说，降水的作用要先经过根区含水和后续蒸散过程重写，才会真正反映到恢复时间上。"
        if category == "high_negative":
            return "这类负向高值端通常不应机械理解为“降水越多越差”，更稳妥的解释是高降水样本可能同时携带低辐射、冷湿背景或事件组成差异，从而使额外降水的边际收益减弱。"
        return "因此，PRE 在这里代表的不是简单的“水越多越好”，而是有效补水是否真正打开了恢复窗口。"

    if label == "EVA (mm)":
        if phase == "prepeak" and category in {"high_positive", "midhigh_positive"}:
            return "这表明峰值前蒸散更像是一个耗水记忆量：当事件前已经存在持续或偏强的水分消耗时，根区储水缓冲更容易被提前削弱，后续一旦遭遇骤旱，恢复尾部就更容易被拉长。"
        if phase == "recoverywin" and category in {"high_positive", "midhigh_positive"}:
            return "放到恢复期内部，这说明蒸散并不只是“恢复成功”的标志。较强蒸散往往意味着系统一边恢复功能、一边继续耗水，因此如果补水与耗水不同步，恢复期反而会被延长。"
        if category == "high_negative":
            return "这里更可能反映的是蒸散减弱对应着能量不足、植被活动受限或系统已退出高耗水状态，因此高值端并没有继续把恢复尾部推长。"
        return "EVA 的机制意义通常在于它把能量条件和土壤失水过程真正连接起来，因此它更像是过程变量，而不是单纯背景量。"

    if label == "TMP (K)":
        if category in {"high_positive", "midhigh_positive"}:
            return "这说明较高温度更容易把系统推向恢复拖尾，其物理链条通常不是“升温本身”单独起作用，而是升温通过增强蒸散需求、改变酶促和呼吸活性、并加剧水热失配来共同放大恢复压力。"
        if category == "high_negative":
            return "若高温端反而偏负，则更可能表示该 biome 中温度升高并未单独构成主限制，或者高温样本更多对应较强能量输入但尚未越过耗水门槛的背景。"
        return "因此，TMP 更适合理解为热力背景变量，它决定的是恢复过程处在“活跃代谢”还是“热干失配”哪一侧。"

    if label == "VPD":
        if category in {"high_positive", "midhigh_positive"}:
            return "这类正向高值端很有物理一致性：较高 VPD 代表更强的大气需水拉力，即便有降水输入，若叶气和土壤供水无法同步跟上，冠层就会处在持续的水气失衡中，恢复时间因此被显著拉长。"
        if category == "high_negative":
            return "若高 VPD 没有继续推长恢复，往往意味着该 biome 的限制已转向别的环节，例如补水不足、能量不足或样本已进入极端干热下的受限区。"
        return "VPD 的关键不是它有多高，而是这种大气端拉力是否真的被转化成了实际耗水压力。"

    if label == "WIND":
        if biome == "Shrubland" and category in {"high_positive", "midhigh_positive"}:
            return "在灌丛里这尤其合理，因为稀疏冠层与近地层耦合更强，风速升高更容易增强湍流交换、边界层导度和近地表失水，从而把恢复过程直接暴露在空气动力控制下。"
        if biome == "Shrubland" and category == "high_negative":
            return "即便这里高风端转向负侧，也不能简单视为“风越大越有利”，更可能表示风场只在中等区间放大耗水和叶气耦合，而再更高端时样本进入了受限或低活动状态，边际效应减弱。"
        return "WIND 在机制上更接近空气动力调节变量，它决定大气端拉力能否高效传到冠层和地表。"

    if label == "SMrz":
        return "SMrz 代表根区可利用水的缓冲能力。它的重要性通常意味着系统恢复并不只是由即时天气决定，而是受到底层储水库是否足以承接前期耗水和后期补水波动的共同约束。"

    if label == "DUR":
        return "骤旱持续时间的贡献说明恢复并非只看峰值强度，事件拖得越久，植被和土壤系统经历的累计压力越大，恢复尾部也更容易被拉长。"

    if label == "INT":
        return "强度项的重要性意味着恢复时间不仅在记住背景条件，也在记住事件本身的冲击幅度；更强的瞬时冲击往往意味着更大的碳通量偏离和更长的回归路径。"

    if label == "ONS":
        return "事件发生时序进入高位，通常意味着季节相位正在调制恢复难度，同样强度的骤旱如果落在不同生长阶段，其后续恢复代价并不相同。"

    return "这一变量更适合被视为背景修饰项，它本身不一定提供单一方向的阈值，但会改变其它主控链条实际传递到恢复时间上的强弱。"


def compare_mechanism_sentence(metric: str, biome: str) -> str:
    if metric == "GPP" and biome == "Forest":
        return "从机制上看，这意味着森林同化恢复经历了由“峰值前能量-耗水记忆”向“恢复期补水-热负荷平衡”切换：前者决定系统带着什么样的亏缺起点进入事件后阶段，后者决定补水能否真正转化为冠层功能回升。"
    if metric == "GPP" and biome == "Grassland":
        return "这类切换说明草地同化恢复先受热干空气背景预设，再受恢复期内的补水和白天能量供给重塑，因此它不是单一路径，而是先记住大气背景、再接受近端水热条件重新筛分。"
    if metric == "GPP" and biome == "Savanna":
        return "因此稀树草原 GPP 的核心不是简单由某个单一气象量控制，而是辐射框架长期保持在前台，只是恢复窗口内补水开始决定这套辐射背景究竟表现为恢复支撑还是耗水压力。"
    if metric == "GPP" and biome == "Cropland":
        return "这里最关键的机制变化是：前期主要是耗水遗产在设定脆弱性，后期则变成大气需水与补水是否匹配在决定恢复能否收尾，这正是农田系统对管理和气象耦合高度敏感的表现。"
    if metric == "GPP" and biome == "Shrubland":
        return "这说明灌丛恢复并不是简单的“干旱结束后慢慢回升”，而是在恢复窗口内明显转入空气动力控制，风场会通过近地层交换和水分再分配直接改变恢复轨迹。"
    if metric == "RECO" and biome == "Forest":
        return "从呼吸机制看，这相当于从“事件前热耗水背景设定代谢起点”转向“恢复期热背景和大气需水决定回归节奏”，说明森林 RECO 的恢复比 GPP 更像是热湿调节下的代谢回落过程。"
    if metric == "RECO" and biome == "Grassland":
        return "也就是说，草地呼吸恢复先被热湿起点预设，真正进入恢复期后再由辐射和补水重新接管，这和 GPP 更直接暴露于近地表温度的路径并不完全一样。"
    if metric == "RECO" and biome == "Savanna":
        return "这说明稀树草原呼吸恢复比同化恢复更早、更彻底地转向补水主导，辐射背景仍重要，但更多是在规定补水效果能释放到什么程度。"
    if metric == "RECO" and biome == "Cropland":
        return "物理上这意味着农田呼吸恢复也存在从“前期耗水遗产”向“后期补水窗口”切换，但比 GPP 更快把补水推到主导位置，说明呼吸对有效回湿更敏感。"
    if metric == "RECO" and biome == "Shrubland":
        return "因此灌丛呼吸恢复虽然也开始受风场影响，但首要控制仍是补水是否到位，空气动力过程更像是在第二层放大或削弱这一补水效应。"
    return ""


def cross_metric_mechanism_sentence(phase: str, biome: str) -> str:
    if phase == "prepeak" and biome == "Forest":
        return "这说明在林地峰值前背景中，GPP 更像是在记住能量-耗水框架，而 RECO 更快把热状态本身写进恢复起点，因此两类通量虽然共享辐射背景，却并不共享同一条记忆链。"
    if phase == "prepeak" and biome == "Grassland":
        return "换句话说，同一草地事件进入恢复时，GPP 背着的是“热干空气记忆”，RECO 背着的是“热背景下的水分起点”，所以后续阈值和转折本来就不应完全一致。"
    if phase == "prepeak" and biome == "Savanna":
        return "这类差异说明在相同辐射框架下，同化恢复更容易被前期耗水背景塑形，而呼吸恢复更容易被前期补水背景塑形。"
    if phase == "prepeak" and biome == "Cropland":
        return "因此即便两类通量都把 EVA 放到高位，它们实际记住的也不是同一件事：GPP 更接近后续功能恢复难度，RECO 更接近底层水分和代谢起点。"
    if phase == "prepeak" and biome == "Shrubland":
        return "这意味着灌丛中的同化恢复更受热干空气记忆控制，而呼吸恢复更受热背景下的初始补水条件控制。"
    if phase == "recoverywin" and biome == "Forest":
        return "因此森林恢复期虽然两类通量排序接近，但 GPP 关注的是补水能否支撑冠层功能重建，RECO 关注的是同样热湿框架下代谢何时真正回到常态。"
    if phase == "recoverywin" and biome == "Grassland":
        return "这里的分歧说明，同化恢复仍紧贴近地表热环境，而呼吸恢复更像是在整体辐射背景中调整，二者对“热”的敏感层次并不相同。"
    if phase == "recoverywin" and biome == "Savanna":
        return "也就是说，在同一恢复窗口里，GPP 还要先跨过能量门槛，RECO 则已经更直接地表现为补水主导。"
    if phase == "recoverywin" and biome == "Cropland":
        return "这正是农田中同化与呼吸恢复分叉最清楚的地方：前者先看水气失衡是否解除，后者先看有效回湿是否发生。"
    if phase == "recoverywin" and biome == "Shrubland":
        return "因此风场虽然对两类通量都重要，但它在 GPP 中更像第一控制项，在 RECO 中更像第二控制项，反映出灌丛同化对空气动力耦合更敏感。"
    return ""


def group_direction_overview(group_key: str) -> str:
    recurrent = group_recurrent_features(group_key, top_n=4)
    chunks: list[str] = []
    for label in recurrent[:3]:
        pairs = {biome: dict((lbl, col) for col, lbl in top_feature_pairs_for_biome(group_key, biome, top_n=5)) for biome in BIOMES}
        categories: list[str] = []
        for biome in BIOMES:
            if label in pairs[biome]:
                categories.append(direction_category(group_key, biome, pairs[biome][label], label))
        high_negative = sum(cat == "high_negative" for cat in categories)
        high_positive = sum(cat in {"high_positive", "midhigh_positive"} for cat in categories)
        if high_negative >= 3:
            chunks.append(f"{label} 在多数 biome 中都表现为高值端更偏负侧")
        elif high_positive >= 3:
            chunks.append(f"{label} 在多数 biome 中都表现为高值端或中高值区更偏正侧")
    if not chunks:
        return ""
    return "方向上，" + "；".join(chunks) + "。这意味着同一变量即使都排在前列，其真正起作用的值域和符号也并不相同。"


def beeswarm_biome_text(group_key: str, biome: str) -> str:
    pairs = top_feature_pairs_for_biome(group_key, biome, top_n=3)
    labels = [label for _, label in pairs]
    rank_text = (
        f"{biome} 的排序不是泛泛的“几个变量都重要”，而是可以明确压缩成 {labels[0]}、{labels[1]}、{labels[2]} "
        f"三条主轴。前两位决定主控抓手，第三位更多负责把同一 biome 内部的事件分流。"
    )
    direction_text = " ".join(
        feature_direction_sentence(group_key, biome, feature_col, label)
        for feature_col, label in pairs[:2]
    )
    third_text = feature_direction_sentence(group_key, biome, pairs[2][0], pairs[2][1])
    mechanism_text = " ".join(
        feature_mechanism_sentence(group_key, biome, feature_col, label)
        for feature_col, label in pairs[:2]
    )
    return " ".join([rank_text, direction_text, third_text, mechanism_text])


def choose_shared_compare_feature(group_a: str, group_b: str, biome: str) -> tuple[str, str] | None:
    a_pairs = top_feature_pairs_for_biome(group_a, biome, top_n=6)
    b_pairs = top_feature_pairs_for_biome(group_b, biome, top_n=6)
    a_map = {label: col for col, label in a_pairs}
    b_map = {label: col for col, label in b_pairs}
    shared = [label for label in FEATURE_PRIORITY if label in a_map and label in b_map]
    if not shared:
        return None
    label = shared[0]
    return a_map[label], label


def phase_shift_sentence(metric: str, biome: str) -> str:
    pre_group = f"{metric}_prepeak"
    rec_group = f"{metric}_recovery"
    chosen = choose_shared_compare_feature(pre_group, rec_group, biome)
    pre_pairs = top_feature_pairs_for_biome(pre_group, biome, top_n=3)
    rec_pairs = top_feature_pairs_for_biome(rec_group, biome, top_n=3)
    head = (
        f"{biome} 的位次变化首先说明控制抓手已经换挡：prepeak 的前三位是 "
        f"{'、'.join(label for _, label in pre_pairs)}，而 recovery 的前三位改成了 "
        f"{'、'.join(label for _, label in rec_pairs)}。"
    )
    if chosen is None:
        return head
    pre_col, label = chosen
    rec_col = dict((lbl, col) for col, lbl in top_feature_pairs_for_biome(rec_group, biome, top_n=6))[label]
    pre_desc = feature_direction_sentence(pre_group, biome, pre_col, label)
    rec_desc = feature_direction_sentence(rec_group, biome, rec_col, label)
    mechanism = compare_mechanism_sentence(metric, biome)
    parts = [head, f"同一变量里最值得比较的是 {label}。", f"在 prepeak 中，{pre_desc}", f"到了 recovery 中，{rec_desc}"]
    if mechanism:
        parts.append(mechanism)
    return " ".join(parts)


def cross_metric_sentence(phase: str, biome: str) -> str:
    gpp_group = f"GPP_{'prepeak' if phase == 'prepeak' else 'recovery'}"
    reco_group = f"RECO_{'prepeak' if phase == 'prepeak' else 'recovery'}"
    chosen = choose_shared_compare_feature(gpp_group, reco_group, biome)
    gpp_pairs = top_feature_pairs_for_biome(gpp_group, biome, top_n=3)
    reco_pairs = top_feature_pairs_for_biome(reco_group, biome, top_n=3)
    head = (
        f"{biome} 在同一 {PHASE_LABELS[phase]} 下，GPP 与 RECO 的前三位分别是 "
        f"{'、'.join(label for _, label in gpp_pairs)} 和 {'、'.join(label for _, label in reco_pairs)}。"
    )
    if chosen is None:
        return head
    gpp_col, label = chosen
    reco_col = dict((lbl, col) for col, lbl in top_feature_pairs_for_biome(reco_group, biome, top_n=6))[label]
    gpp_desc = feature_direction_sentence(gpp_group, biome, gpp_col, label)
    reco_desc = feature_direction_sentence(reco_group, biome, reco_col, label)
    mechanism = cross_metric_mechanism_sentence(phase, biome)
    parts = [head, f"就共同变量 {label} 而言，GPP 这边，{gpp_desc}", f"RECO 这边，{reco_desc}"]
    if mechanism:
        parts.append(mechanism)
    return " ".join(parts)


DOC1_GROUP_OVERVIEW = {
    "GPP_prepeak": (
        "GPP 的峰值前背景结果在五个 biome 之间呈现出两类很清楚的结构。Forest、Savanna 和 Cropland 的前列变量都由辐射项与蒸散项占据，其中 Forest 和 Savanna 都以 SSRD、STRD 为核心，Cropland 则把 EVA 提升到首位，说明在这三类 biome 中，恢复时间首先对应的是事件前能量供给与耗水强度的组合背景，而不是单独的补水多少。相对地，Grassland 与 Shrubland 在 SSRD 之后更快地抬升了 TMP 和 VPD，表明开放植被对峰值前大气热力状态更敏感，恢复尾部更容易被高温和大气干旱需求预先设定。",
        "另一个值得强调的差异是，PRE 在五个 biome 中都没有进入前三位，SMrz 也只在 Cropland 进入前五。这说明在 GPP 的 prepeak 口径下，真正先被 SHAP 捕捉到的不是“静态水库大小”，而是驱动未来干化和能量失衡的背景条件本身。也正因为如此，Cropland 的 EVA 首位、Grassland/Shrubland 的 TMP-VPD 前移、以及 Forest/Savanna 的双辐射主导，实际上对应的是三种不同的峰值前恢复记忆类型。"
    ),
    "GPP_recovery": (
        "进入 recovery 口径后，GPP 的排序立刻变得更集中，最稳定的共同点是 PRE 在五个 biome 中全部进入前两位。这说明一旦把视角切到恢复窗口内部，补水过程不再只是背景变量，而直接成为恢复时间长短的近端控制项。不过，PRE 虽然普遍重要，它并没有在所有 biome 中处于同样的位置：Forest 仍由 STRD 居首，Cropland 由 VPD 居首，Shrubland 则由 WIND 居首，说明不同植被类型是在不同的水热框架下利用或消耗这部分补水。",
        "五个 biome 在 recovery 阶段可以进一步分成三组。Forest 体现为长波辐射加大气干旱需求的双重约束，说明封闭冠层中恢复时间受热背景和蒸散拉力共同限制；Grassland 与 Savanna 则都以 SSRD 和 PRE 为核心，只是 Grassland 把 TMP 提得更高，Savanna 更强调 STRD，反映开放区恢复更直接受白天地表能量输入控制；Cropland 与 Shrubland最有个性，前者把 VPD 放到第一位，说明恢复速度更受大气需水和补水是否匹配控制，后者则把 WIND 放到第一位，提示空气动力条件在稀疏冠层中对恢复期水分损耗更敏感。"
    ),
    "RECO_prepeak": (
        "RECO 的峰值前背景与 GPP 有相似之处，但也保留了更明显的呼吸特征。五个 biome 中，Forest 与 Cropland 都把 EVA 放进前两位，其中 Cropland 甚至同样由 EVA 居首，说明呼吸恢复在这两类 biome 中更容易记住事件前的耗水背景；Grassland 与 Shrubland 则把 TMP 和 PRE 提到更靠前的位置，表示开放植被的呼吸恢复对峰值前的热力状态和补水背景更敏感；Savanna 则介于两者之间，前三位是 SSRD、STRD、PRE，更像是一个以辐射背景控制为主、以补水背景补充的类型。",
        "和 GPP prepeak 相比，RECO prepeak 更早把 PRE 推到前三位，尤其在 Grassland、Savanna、Shrubland 中都如此，而 Forest 和 Cropland 仍保留 EVA 的高位。这说明呼吸恢复时间对峰值前背景的记忆更偏向“热湿收支”本身，而不只是能量与大气需求；也就是说，同样是 prepeak 口径，GPP 更容易先看到未来同化恢复会被什么样的能量-耗水背景拖慢，RECO 则更容易先看到呼吸系统在什么样的热湿起点上进入恢复。"
    ),
    "RECO_recovery": (
        "RECO 的 recovery 口径是四组结果里最整齐的一组：PRE 在五个 biome 中全部进入前两位，并且在 Savanna、Cropland、Shrubland 直接排到第一位。这说明一旦进入恢复阶段内部，呼吸系统对补水过程的依赖比 GPP 还更直接。不过，Forest 和 Grassland 并没有让 PRE 排到第一，Forest 首位是 STRD，Grassland 首位是 SSRD，这提示在这两类 biome 中，补水能否真正转化为呼吸恢复，还要先看热背景或辐射背景是否允许系统维持较高的代谢调整。",
        "如果把五个 biome 并列看，RECO recovery 又可分成两种主控模式。Forest、Grassland 更接近“辐射-补水联合控制”，前者偏 STRD，后者偏 SSRD；Savanna、Cropland、Shrubland 则更接近“补水主导，但被第二控制项分流”的模式，其中 Savanna 的第二、第三位是 SSRD 和 STRD，Cropland 是 SSRD 和 VPD，Shrubland 是 SSRD 和 WIND。也就是说，呼吸恢复时间虽然普遍离不开补水，但不同 biome 对这部分水分的利用路径并不相同，有的是先受能量约束，有的是先受大气需水或空气动力约束。"
    ),
}


DOC1_BIOME_TEXT = {
    "GPP_prepeak": {
        "Forest": "Forest 的前三位是 SSRD、EVA 和 STRD，这个排序在五个 biome 中最典型地体现了“辐射背景 + 耗水背景”双主导。与 Grassland 和 Shrubland 相比，Forest 并没有把 TMP、VPD 推到更靠前位置，说明林地 GPP 的恢复时间在峰值前阶段更先记住的是能量输入和耗水过程，而不是开放大气环境本身。这种排序也把 Forest 和 Savanna 放在同一类里，只不过 Forest 比 Savanna 更早抬升 EVA，意味着林地在事件发生前已经把蒸散耗水写进了恢复记忆。",
        "Grassland": "Grassland 的前三位是 SSRD、TMP 和 VPD，而不是 Forest/Cropland 那种以 EVA 或 STRD 为核心的结构。这个差异很关键，它说明草地 GPP 的峰值前恢复记忆更偏向大气热力状态本身，高温与高大气需水会更早进入主控序列。PRE 虽然也进了前五，但排位仍落后于 TMP、VPD，说明对 Grassland 来说，事件前是否已经处在热干空气背景下，比单纯多一点或少一点降水更能解释后续恢复时间的差异。",
        "Savanna": "Savanna 的前三位是 SSRD、STRD 和 EVA，和 Forest 一样属于辐射主导型，但它把 STRD 提到了 EVA 之前。这说明稀树草原的峰值前背景并不只是白天辐射或蒸散强弱的问题，整个热背景都更靠前地参与了恢复记忆形成。与 Cropland 相比，Savanna 没有让 EVA 占据第一位，说明其恢复时间更像是先受整体能量环境设定，再由耗水强度细化；与 Grassland 相比，它又没有把 TMP/VPD 推进前三，说明其主控并非单纯的开放大气热干背景。",
        "Cropland": "Cropland 是这一组里最鲜明的例子，因为它唯一把 EVA 放到了第一位，并且前三位是 EVA、SSRD、STRD。这个排序说明农田 GPP 的峰值前恢复记忆首先对应的是耗水过程本身，而不是单纯辐射强不强或温度高不高。PRE 只排到第四位，SMrz 排到第五位，说明对 Cropland 来说，进入恢复尾部之前最重要的不是已经存了多少水，而是事件前耗水是否已经把系统推到更脆弱的状态上。这也是它和 Forest/Savanna 的核心区别：后两者先看到辐射框架，Cropland 先看到耗水强度。",
        "Shrubland": "Shrubland 的前三位是 SSRD、TMP 和 VPD，和 Grassland 非常接近，说明稀疏植被的 GPP 恢复时间在峰值前阶段同样更容易受开放大气环境主导。不同于 Cropland 和 Forest，Shrubland 把 EVA 放到了第十位，说明蒸散耗水虽然存在，但并不是最先决定恢复记忆的因子；相反，气温和大气干旱需求的前移意味着这类 biome 更容易在事件发生前就暴露于热干空气背景，后续恢复慢更多是这种背景的延续，而不是单一耗水过程的结果。"
    },
    "GPP_recovery": {
        "Forest": "Forest 的前三位是 STRD、PRE 和 VPD，这与其他 biome 明显不同，因为它没有把 SSRD 放到最前面。这个排序说明森林 GPP 在恢复期内部首先受长波热背景控制，其次才是补水和大气需水。与 Grassland/Savanna 相比，Forest 的能量限制更偏向整体热负荷而非单纯短波输入；与 Cropland 相比，它虽然同样把 VPD 放进前三，但并没有让 VPD 居首，说明森林恢复更像是在热背景和补水之间寻找平衡，而不是直接被大气需水牵着走。",
        "Grassland": "Grassland 的前三位是 SSRD、PRE 和 TMP，是 recovery 组里最典型的“短波辐射 + 补水 + 温度”结构。这个排序说明草地 GPP 恢复时间对恢复窗口内部的白天能量输入非常敏感，补水虽然重要，但只有和适当的热力状态一起出现时才真正决定恢复快慢。与 Forest 不同，Grassland 把 STRD 压到第四位，把 TMP 抬到第三位，反映开放植被的恢复更多发生在近地表热环境直接控制之下，而不是深冠层内部的热背景调节。",
        "Savanna": "Savanna 的前三位是 SSRD、PRE 和 STRD，介于 Grassland 与 Forest 之间。它既像 Grassland 一样把 SSRD 放到首位，又像 Forest 一样把 STRD 保持在前三，说明稀树草原的恢复期控制并不是单一能量维度，而是白天短波输入和整体热背景同时重要。与 Grassland 相比，Savanna 没有让 TMP 进入前三，意味着其恢复差异更多由辐射结构本身塑造；与 Forest 相比，它又没有让 STRD 居首，说明补水和短波输入仍比长波热背景更靠前。",
        "Cropland": "Cropland 的前三位是 VPD、PRE 和 SSRD，是这一组里唯一由 VPD 居首的 biome。这个排序说明农田 GPP 在恢复阶段最敏感的不是“有没有水”，而是补水能否抵消大气端持续的需水拉力。PRE 排在第二位，说明补水仍然关键，但它必须和较低的蒸发需求同时出现才有利于缩短恢复尾部。与 Forest 的 STRD 首位和 Grassland/Savanna 的 SSRD 首位相比，Cropland 把大气干旱需求提前到了整个排序顶端，体现出其对水气耦合失衡的高敏感性。",
        "Shrubland": "Shrubland 的前三位是 WIND、PRE 和 TMP，是 recovery 组里最特别的一类，因为它把 WIND 直接推到了第一位。这个结果不能被忽略，它说明灌丛恢复期的 GPP 对空气动力条件极为敏感，风速变化可能通过增强湍流交换、加快近地表失水或放大叶气耦合，进而改变恢复速度。与 Cropland 的 VPD 首位相比，Shrubland 更像是由空气动力过程主导；与 Grassland 的 SSRD 首位相比，它又不是典型的辐射主控，而是一个“风场 + 补水 + 温度”共同控制的恢复结构。"
    },
    "RECO_prepeak": {
        "Forest": "Forest 的前三位是 SSRD、EVA 和 TMP，这个排序和 GPP prepeak 的 Forest 有相似性，但把 TMP 提到了 STRD 之前。它说明林地 RECO 的峰值前恢复记忆不只是能量和耗水问题，还更早吸收了热力状态信息。与 Cropland 一样，Forest 仍把 EVA 放在前两位，说明呼吸恢复在这两类 biome 中都会明显记住事件前的耗水背景；但 Forest 没有像 Cropland 那样让 EVA 居首，说明它仍然优先受辐射框架约束。",
        "Grassland": "Grassland 的前三位是 SSRD、TMP 和 PRE，这和同组的 Forest/Cropland 很不一样，因为它没有把 EVA 放进前三。这个结构说明草地 RECO 的峰值前差异主要来自热背景和补水背景，而不是耗水过程本身。PRE 能排到第三位，说明在呼吸恢复上，草地比 GPP 更早表现出对事件前水分输入的依赖；TMP 的第二位则提示热环境变化会更直接地改变呼吸系统进入恢复阶段时的起点。",
        "Savanna": "Savanna 的前三位是 SSRD、STRD 和 PRE，是 RECO prepeak 里最完整的“辐射双项 + 补水”结构。相比 Forest，它没有把 EVA 和 TMP 提得那么靠前；相比 Grassland 和 Shrubland，它又没有让 TMP 进入前三。这说明稀树草原的呼吸恢复时间在峰值前阶段主要受整体能量背景和补水条件共同设定，而大气热力需求更多是第二层因素。它与 GPP prepeak 的 Savanna 也不同，后者第三位是 EVA，这里则换成了 PRE，说明呼吸恢复比同化恢复更早显现出对补水背景的依赖。",
        "Cropland": "Cropland 的前三位是 EVA、SSRD 和 STRD，与 GPP prepeak 的 Cropland 几乎同型，但这里 PRE 和 SMrz 的位置互换为第六和第四位，说明 RECO 对事件前土壤水状态本身也更敏感一些。最重要的信息仍然是 EVA 居首，这意味着农田呼吸恢复时间首先记住的是耗水背景是否已经很强，然后才是辐射环境和补水条件。与 Grassland 和 Shrubland 相比，Cropland 明显不是热背景主导，而是耗水主导；与 Forest 相比，它又把耗水的权重推得更极端。",
        "Shrubland": "Shrubland 的前三位是 SSRD、TMP 和 PRE，和 Grassland 十分接近，说明灌丛 RECO 的峰值前恢复记忆同样主要来自开放环境中的热力状态和补水背景，而不是蒸散耗水本身。与 GPP prepeak 的 Shrubland 相比，这里 PRE 上升到了第三位，而 VPD 降到第六位，提示呼吸恢复比同化恢复更依赖事件前的水分输入条件。这种差异说明即便在同一 biome、同一 prepeak 口径下，GPP 和记住的是“热干空气”，RECO 更可能记住的是“热背景下的水分起点”。"
    },
    "RECO_recovery": {
        "Forest": "Forest 的前三位是 STRD、PRE 和 VPD，和 GPP recovery 的 Forest 高度一致，说明森林恢复期里无论看同化还是呼吸，长波热背景、补水和大气需水都是最核心的控制轴。但 RECO 的第四位是 SSRD，第五位才是 EVA，说明呼吸恢复比 GPP 更保留辐射背景本身，而不急于把它转写成耗水过程。也就是说，Forest 的 RECO 恢复比 GPP 更像是一个热背景主导的代谢调整问题。",
        "Grassland": "Grassland 的前三位是 SSRD、PRE 和 STRD，说明草地 RECO 的恢复期首先受辐射结构和补水过程控制，而 TMP 只排到第八位。这和 GPP recovery 的 Grassland 有明显差别，后者第三位是 TMP，这里则换成 STRD，表明呼吸恢复比同化恢复更少直接暴露在瞬时近地表温度下，而更受整体辐射背景控制。它也和 Forest 不同，因为 Forest 由 STRD 居首，Grassland 则仍由 SSRD 居首，说明开放植被的呼吸恢复更依赖白天能量输入。",
        "Savanna": "Savanna 的前三位是 PRE、SSRD 和 STRD，是这一组里最明确的“补水主导型”结构。与 Grassland 相比，它把 PRE 提到了第一位，说明稀树草原呼吸恢复对恢复窗口内部补水变化更直接；与 Cropland 和 Shrubland 相比，它虽然同样由 PRE 居首，但第二、第三位并不是 VPD 或 WIND，而仍然是 SSRD 和 STRD，表明这里的补水效应是嵌在强烈辐射背景里的，而不是主要由大气需水或空气动力过程重写。",
        "Cropland": "Cropland 的前三位是 PRE、SSRD 和 VPD，体现出一个很清楚的“补水 + 大气需水”双主导结构。PRE 居首说明农田呼吸恢复首先受补水控制，但 VPD 进入第三位又说明这部分补水是否有效，取决于大气端是否继续维持强蒸发需求。与 Savanna 的 PRE-SSRD-STRD 结构相比，Cropland 更少强调整体热背景，更多强调补水是否能抵消蒸发拉力；与 Forest 相比，它又不是长波热背景主导，而是更直接地暴露在水气失衡约束之下。",
        "Shrubland": "Shrubland 的前三位是 PRE、SSRD 和 WIND，这使它成为 RECO recovery 里另一种独特类型。与 Savanna/Cropland 一样，PRE 居首说明补水是第一控制项；但第二控制项之后它没有选择 VPD 或 STRD，而是把 WIND 提到第三位，提示灌丛呼吸恢复很可能更容易受到空气动力条件影响。这个排序与 GPP recovery 的 Shrubland 也形成呼应，只不过 GPP 那边 WIND 已经到第一位，这里则退到第三位，说明在灌丛中风场对同化恢复的影响比对呼吸恢复更直接、更靠前。"
    },
}


DOC3_METRIC_OVERVIEW = {
    "GPP": (
        "GPP 的两种口径差异非常清楚。prepeak 更强调事件发生前的能量-耗水背景，Forest、Savanna 和 Cropland 都把辐射项与蒸散项放在前列，而 Grassland 和 Shrubland 更早抬升 TMP、VPD，说明同化恢复在事件发生前就已经被不同类型的环境记忆分流。相对地，recovery 几乎在所有 biome 中都把 PRE 提到前两位，说明一旦进入恢复窗口，补水条件会迅速压过部分背景记忆，转而成为直接控制项。",
        "但 recovery 并不是简单地“PRE 最重要”就结束了。Forest 仍由 STRD 居首，Cropland 由 VPD 居首，Shrubland 则把 WIND 推到第一位，说明 GPP 恢复虽然普遍依赖补水，却是在不同的热背景、大气需水或空气动力框架下完成的。因此，GPP 的 prepeak/recovery 差异，本质上是“恢复记忆”向“恢复窗口直接控制”的位移，而不是同一机制在两个时段里的重复观测。"
    ),
    "RECO": (
        "RECO 的两种口径也呈现出清楚的分工，但与 GPP 不完全相同。prepeak 阶段，Forest 和 Cropland 都把 EVA 放进前两位，说明呼吸恢复更容易记住事件前的耗水背景；Grassland、Savanna、Shrubland 则更早把 PRE 抬进前三，表明呼吸系统在进入恢复期之前就已经受到热湿起点的预设。与 GPP prepeak 相比，这一组结果更早显出对水分输入条件的依赖。",
        "到了 recovery 阶段，RECO 比 GPP 更整齐地表现出补水主导特征：PRE 在五个 biome 中全部进入前两位，并在 Savanna、Cropland、Shrubland 居首。这意味着呼吸恢复比同化恢复更直接受恢复窗口内部的补水条件控制。不过，第二控制项仍在分流机制：Forest 偏 STRD，Grassland 偏 SSRD，Cropland 偏 VPD，Shrubland 偏 WIND。因此，RECO 的 prepeak/recovery 对比不是“换一张图再看一次”，而是从热湿背景记忆转向补水主导的真实机制切换。"
    ),
}


DOC3_METRIC_BIOME_COMPARE = {
    "GPP": {
        "Forest": "Forest 的对比图最能说明“背景记忆”和“恢复窗口控制”的分工。prepeak 中 Forest 更突出 SSRD、EVA 和 STRD，说明事件发生前的能量与耗水背景已经写进恢复尾部；recovery 中则改成 STRD、PRE 和 VPD，说明真正进入恢复窗口后，长波热背景、补水和大气需水共同决定恢复速度。也就是说，Forest 的恢复慢并不是单纯因为前面干过一次，而是因为后期仍处在高热负荷和高蒸发拉力框架里。",
        "Grassland": "Grassland 的变化比 Forest 更直接。prepeak 阶段它把 TMP 和 VPD 提得很靠前，说明草地同化恢复首先记住的是事件前热干空气背景；recovery 阶段则转成 SSRD、PRE 和 TMP，说明一旦进入恢复窗口，白天能量输入和补水条件迅速成为直接抓手。这个转变意味着 Grassland 的恢复时间并不是被单一干旱强度锁死，而是从“热干背景预设”切换到“补水与能量重新分配”。",
        "Savanna": "Savanna 在两种口径下都保留了强烈的辐射控制，但侧重点发生了位移。prepeak 中是 SSRD、STRD、EVA，强调整体能量环境和耗水背景；recovery 中则变成 SSRD、PRE、STRD，说明补水在恢复窗口内部上位，但辐射仍是框架变量。它不像 Grassland 那样把 TMP 推进前三，也不像 Cropland 那样把 VPD 或 EVA推到首位，因此 Savanna 更像是在稳定辐射框架内，由补水条件调节恢复速度。",
        "Cropland": "Cropland 是 GPP 里变化最戏剧化的 biome。prepeak 中 EVA 第一，说明农田恢复记忆首先来自事件前耗水是否已经把系统推向脆弱状态；recovery 中却变成 VPD、PRE、SSRD，说明真正控制恢复尾部的不是历史耗水本身，而是补水能否抵消当前大气需水拉力。这一前后转换很关键，它说明 Cropland 的恢复并不是简单由“缺水”决定，而是由‘历史耗水遗产’向‘当前水气失衡约束’的转换。 ",
        "Shrubland": "Shrubland 的对比最能体现开放稀疏冠层的特殊性。prepeak 中它和 Grassland 类似，更强调 SSRD、TMP 和 VPD，说明事件前热干空气背景是主要记忆来源；recovery 中却把 WIND 直接推到第一位，同时 PRE 升到第二，说明恢复窗口内部的空气动力条件开始主导近地表失水和叶气耦合。也就是说，Shrubland 并不是简单从热干背景恢复，而是在风场参与下完成恢复。"
    },
    "RECO": {
        "Forest": "Forest 的 RECO 对比与 GPP 有相似处，但更偏代谢背景。prepeak 中 Forest 强调 SSRD、EVA 和 TMP，说明呼吸恢复在进入恢复期前就记住了能量、耗水和热状态；recovery 中变成 STRD、PRE 和 VPD，说明恢复窗口内部由长波热背景、补水和大气需水共同决定呼吸何时回归基线。与 GPP 相比，Forest 的 RECO 更早吸收热背景信息，也更像是一个代谢调节过程而不只是同化重建过程。",
        "Grassland": "Grassland 的 RECO 在两种口径间表现出很典型的“热湿起点”向“辐射-补水控制”的切换。prepeak 阶段是 SSRD、TMP 和 PRE，说明事件前热环境和水分起点已经决定了部分恢复难度；recovery 阶段则转成 SSRD、PRE 和 STRD，TMP 显著后移，说明真正进入恢复期后，呼吸恢复更少受瞬时近地表温度直接控制，而更多受整体辐射环境与补水控制。",
        "Savanna": "Savanna 的 RECO 在 prepeak 和 recovery 中都保持了高位的 PRE，但意义不同。prepeak 中 PRE 只是和 SSRD、STRD 一起刻画事件前热湿背景；recovery 中 PRE 升到第一位，SSRD、STRD 跟在后面，说明恢复窗口内部补水已经成为第一控制项，辐射从主导背景转成约束补水效应的框架。这种前后位移说明 Savanna 的呼吸恢复比 GPP 更快转向补水主导。",
        "Cropland": "Cropland 的 RECO 对比也很有代表性。prepeak 中 EVA 第一，说明事件前耗水背景首先决定呼吸系统进入恢复期时的脆弱性；recovery 中则改成 PRE、SSRD、VPD，说明恢复阶段内部真正决定呼吸拖尾的是补水是否能抵消大气端蒸发需求。这与 GPP 的 Cropland 有呼应，但 RECO 更早让 PRE 居首，说明呼吸恢复对补水的依赖比同化恢复更直接。",
        "Shrubland": "Shrubland 的 RECO 对比显示出一个和 GPP 不同的细节。prepeak 阶段它更像 Grassland，主要受 SSRD、TMP 和 PRE 约束；recovery 阶段则变成 PRE、SSRD 和 WIND，说明补水上位的同时，空气动力条件也开始显著参与呼吸恢复。与 GPP 不同的是，GPP recovery 里 WIND 已经到第一位，这里 WIND 只到第三位，说明风场对灌丛同化恢复的影响比对呼吸恢复更直接。"
    },
}


DOC3_PHASE_OVERVIEW = {
    "prepeak": (
        "在 prepeak 口径下，GPP 与 RECO 的共同点是都明显受辐射背景控制，但分歧在于谁更早把水分输入和耗水过程抬到前列。GPP 更常把 SSRD、STRD、EVA 作为首批变量，说明同化恢复首先记住的是事件发生前能量-耗水框架；RECO 则在 Grassland、Savanna、Shrubland 中更早把 PRE 推进前三，表明呼吸恢复更快吸收了热湿起点信息。",
        "也正因为如此，prepeak 不能被简单理解为“干旱前的平均状态”。它更像是一套恢复记忆编码：GPP 记的是未来同化重建会在什么样的能量和耗水背景下启动，RECO 记的是呼吸系统会在什么样的热湿起点上进入恢复。因此，同样是 prepeak，GPP 和 RECO 的差异不是权重轻微变化，而是记忆内容本身不同。"
    ),
    "recoverywin": (
        "在 recovery 口径下，GPP 与 RECO 的共同点是 PRE 普遍前移，说明恢复窗口内部补水是无法回避的直接控制项。但两者的差异同样清楚：GPP 在 Forest、Cropland、Shrubland 中分别让 STRD、VPD、WIND 居首，说明同化恢复更容易被热背景、水气失衡和空气动力过程分流；RECO 则在 Savanna、Cropland、Shrubland 直接让 PRE 居首，说明呼吸恢复对补水的依赖更直接、更统一。",
        "因此，recovery 阶段并不能只说“水重要”。更准确的说法是：GPP 恢复是在不同的热力和空气动力框架下利用这部分补水，而 RECO 恢复则更像是在补水主导下，被第二控制项修饰。这个差异决定了同样的恢复降水条件，对 GPP 和 RECO 并不会产生同样的恢复轨迹。"
    ),
}


DOC3_CROSS_METRIC_BIOME = {
    "prepeak": {
        "Forest": "Forest 的 prepeak 对比说明，GPP 与 RECO 都重视辐射和耗水背景，但 RECO 比 GPP 更早抬升 TMP，GPP 则更早保留 STRD。这意味着林地同化恢复更偏向能量-耗水框架，呼吸恢复则更快吸收热状态本身。",
        "Grassland": "Grassland 的 prepeak 对比最能说明通量差异。GPP 的前三位是 SSRD、TMP、VPD，明显偏热干空气背景；RECO 的前三位是 SSRD、TMP、PRE，更早把补水背景写进恢复记忆。也就是说，同一草地事件，GPP 记住的是‘热干’，RECO 记住的是‘热湿起点’。",
        "Savanna": "Savanna 的 prepeak 对比显示，GPP 更强调 SSRD、STRD、EVA，RECO 则强调 SSRD、STRD、PRE。前者说明同化恢复更依赖事件前耗水背景，后者说明呼吸恢复更依赖事件前补水背景，这种差别已经决定了两类通量后续阈值不会完全一致。",
        "Cropland": "Cropland 的 prepeak 对比看起来最接近，因为 GPP 和 RECO 都由 EVA 领跑。但进一步看，RECO 把 SMrz 提得更靠前，GPP 则把 PRE 放得稍靠前，说明农田同化恢复与呼吸恢复虽然都记住了耗水背景，但一个更敏感于后续补水线索，一个更敏感于土壤水状态本身。",
        "Shrubland": "Shrubland 的 prepeak 对比也很清楚：GPP 强调 SSRD、TMP、VPD，RECO 强调 SSRD、TMP、PRE。这个差异说明灌丛中同化恢复更受热干空气背景控制，而呼吸恢复更受热背景下的水分起点控制。"
    },
    "recoverywin": {
        "Forest": "Forest 的 recovery 对比说明，GPP 和 RECO 都由 STRD、PRE、VPD 主导，但 GPP 更像在热背景和补水之间平衡同化重建，RECO 则更像在同一框架下完成代谢回归。表面排序相近，不代表过程相同。",
        "Grassland": "Grassland 的 recovery 对比里，GPP 是 SSRD、PRE、TMP，RECO 是 SSRD、PRE、STRD。这个差异很关键：同化恢复仍直接暴露在近地表热环境里，呼吸恢复却更多被整体辐射背景控制，因此不能把同一个热量阈值同时套给两类通量。",
        "Savanna": "Savanna 的 recovery 对比最能体现补水主导程度差别。GPP 仍由 SSRD 居首，PRE 第二；RECO 则直接由 PRE 居首。说明在稀树草原里，同化恢复仍要先看能量框架，呼吸恢复则已经转向补水主导。",
        "Cropland": "Cropland 的 recovery 对比非常鲜明。GPP 由 VPD、PRE、SSRD 主导，说明同化恢复最怕水气失衡；RECO 则由 PRE、SSRD、VPD 主导，说明呼吸恢复先看补水，再看蒸发拉力。排序只交换一个位置，机制却完全不同。",
        "Shrubland": "Shrubland 的 recovery 对比说明风场对两类通量的作用强度不同。GPP 里 WIND 直接排第一，RECO 里 WIND 退到第三，说明风场对同化恢复的直接性更高，而呼吸恢复仍以补水为第一控制项。"
    },
}


def build_doc1_beeswarm() -> Document:
    document = make_document("全球 GPP 与 RECO 恢复时间驱动机制分析（一）：SHAP Beeswarm 整体贡献分析")
    for para in shared_intro_paragraphs():
        add_paragraph(document, para)
    add_paragraph(
        document,
        "这次重写对 beeswarm 的使用原则比旧版更严格。排序当然重要，但排序只解决“谁经常重要”，并不自动回答“高值究竟把 SHAP 推向正侧还是负侧”。因此，下面每一组讨论都同时使用两层信息：第一层是 beeswarm 给出的贡献排序，第二层是对应 dependence 点云中 10% 与 90% 分位的 SHAP 中位数、以及在可识别情形下的正峰/负谷位置。也就是说，后文不会再停留在“某变量排第几”，而是进一步回答“这个变量是在哪个值域开始真正拉长恢复时间，还是在哪个高值端反而把贡献压向负侧”。",
    )

    figure_no = 1
    for group_key, metric, phase in GROUPS:
        add_heading(document, GROUP_LABELS[group_key], level=1)
        for para in DOC1_GROUP_OVERVIEW[group_key]:
            add_paragraph(document, para)
        overview = group_direction_overview(group_key)
        if overview:
            add_paragraph(document, overview)
        for biome in BIOMES:
            add_heading(document, biome, level=2)
            add_figure(
                document,
                FigureItem(
                    path=beeswarm_image_path(group_key, biome, phase),
                    text=" ".join([DOC1_BIOME_TEXT[group_key][biome], beeswarm_biome_text(group_key, biome)]),
                    caption=f"图 {figure_no}. {GROUP_LABELS[group_key]} 在 {biome} biome 的 SHAP beeswarm 图。",
                ),
            )
            figure_no += 1
        document.add_section(WD_SECTION_START.NEW_PAGE)

    add_heading(document, "综合讨论", level=1)
    add_paragraph(
        document,
        "综合全部 biome 的 rank 与 direction 可以看到，beeswarm 真正有用的地方在于它把“谁重要”和“怎么重要”同时压缩了出来。prepeak 阶段，高 SSRD 在 GPP 与 RECO 中大多都更容易落入负 SHAP 侧，说明峰值前高辐射更多是在写入一种未来会拖慢恢复的背景记忆；到了 recovery 阶段，PRE 在几乎所有 biome 中都表现为低值端强负、中高值区明显抬升，表明补水在恢复窗口内从背景信息变成了直接控制项。再往下看，GPP 更频繁地把 VPD、WIND 这类大气端和空气动力端变量推到正 SHAP 高值端，而 RECO 更常把 PRE 直接推到首位。这意味着同一个“高贡献变量”在不同通量和不同阶段下，其物理角色并不相同：有的变量是在高值端延长恢复尾部，有的则是在高值端把贡献压回负侧，真正的机制差异就藏在这一步里。",
    )
    return document


def dependence_text(metric: str, phase: str, biome: str, x_key: str) -> str:
    prefix = "峰值前背景" if phase == "prepeak" else "恢复期内部"
    metric_text = "GPP 恢复时间" if metric == "GPP" else "RECO 恢复时间"
    if x_key == "PRE":
        return (
            f"{biome} 中以 PRE 为横轴的 dependence plot 显示，{metric_text} 对降水补给的响应主要集中在低值到中值区间，"
            f"且散点颜色随 SMrz 梯度出现明显分层。这说明 {prefix} 的降水异常并不会孤立地影响恢复时间，"
            f"而是需要通过根区土壤水状态重新分配其作用强度；同等降水水平下，SMrz 更高或更低的背景都可能对应不同的恢复尾部长短。"
        )
    if x_key == "EVA":
        return (
            f"{biome} 中以 EVA 为横轴的 dependence plot 进一步表明，蒸散耗水是连接能量条件与土壤干化过程的重要桥梁。"
            f"图中 SHAP 值在不同 EVA 区间呈现出明显的非线性响应，并且在 SMrz 颜色映射下表现出垂向离散增强，"
            f"说明 {metric_text} 对蒸散的敏感性强烈依赖于土壤可利用水储量，蒸散增强往往意味着更强的恢复压力或更长的调整尾部。"
        )
    return (
        f"{biome} 中以 SSRD 为横轴的 dependence plot 反映了辐射背景对恢复时间的稳定调节作用。"
        f"无论在 {prefix} 还是恢复阶段内部，短波辐射都能够通过影响蒸散需求、叶面能量平衡和水热耦合强度间接改变 {metric_text}。"
        f"图中不同 SMrz 颜色带对应的点云包络并不重合，说明辐射控制始终受到水分背景约束，能量限制与水分限制共同决定了最终恢复轨迹。"
    )


def build_doc2_dependence() -> Document:
    document = make_document("全球 GPP 与 RECO 恢复时间驱动机制分析（二）：Dependence Plot 与 SMrz 交互分析")
    for para in shared_intro_paragraphs():
        add_paragraph(document, para)
    add_paragraph(
        document,
        "在初始版本中，本部分主要围绕 PRE、EVA 与 SSRD 三类特征展开，因为它们分别代表水分补给、蒸散耗水与短波辐射供能三条主轴。但在重新核对 beeswarm 排序之后可以发现，仅靠这三类图还不足以支撑可信的物理机制解释：TMP、STRD 与 VPD 在多数组合中同样持续位居高贡献位置，尤其 recovery 口径中的 STRD 以及 RECO 中的 VPD，如果不单独展开，就无法判断辐射控制究竟是短波主导、长波主导，还是经由热力状态与大气需水间接传递。基于这一点，我额外补绘了 TMP、STRD 与 VPD 的 dependence plot，并统一使用 SMrz 作为颜色映射。",
    )
    add_paragraph(
        document,
        "需要特别强调的是，下面对阈值的表述全部是基于已经过滤异常值后的实际曲线形态做出的近似诊断，目的在于描述“曲线在什么区间开始改变响应方向”，而不是把每一个阈值当作精确的生理常数。为保证可追溯性，我同时输出了一个转折点摘要表，保存于 analysis_writing 目录下的 dependence_turning_thresholds_0401.csv，可供后续继续核对。",
    )

    figure_no = 1
    add_heading(document, "一、为什么必须从 beeswarm 走向 dependence plot", level=1)
    add_paragraph(
        document,
        "beeswarm 图只能告诉我们某个变量整体上经常重要，却不能告诉我们它是在低值区间、过渡区间还是高值区间真正改变了恢复时间的方向。对于全球尺度骤旱恢复研究，这一点尤其关键，因为很多变量并不是单调控制项，而是在“先缓解、后拖延”或“先促进、后抑制”的结构中起作用。若不识别这些转折点，就很容易把本质上是阈值型的恢复过程误写成线性响应。",
    )
    add_paragraph(
        document,
        "因此，本节不再把所有 dependence plot 平铺直叙，而是围绕那些确实出现清晰转折的曲线来讨论物理机制，并把新增的 TMP、STRD 和 VPD 图作为必要补充，而不是装饰性的附图。判断逻辑也保持克制：只有当曲线转向在图上清楚、且与 SMrz 颜色分层一致时，才把它写成具有机制含义的现象；若只是局部噪声或稀疏尾部点云，则只保留为“可能的背景效应”，不强行解释。",
    )

    add_heading(document, "二、GPP：峰值前背景中的典型转折机制", level=1)
    add_figure(
        document,
        FigureItem(
            path=dependence_image_path("GPP_prepeak", "Cropland", "prepeak", "PRE"),
            text=(
                "GPP prepeak 中 Cropland 的 PRE 曲线是最典型的阈值型例子之一。曲线在极低降水区间先迅速上升，随后在大约 1.5–2.0 mm 左右由正值转向接近零或负值，之后在更高降水水平继续下降，8 mm 以后下行更明显。这个形态说明，极低峰值前降水对应的是更强水分短缺背景，因此少量补水会显著缓解恢复压力；但当峰值前背景已经较湿时，更多降水并不再简单代表“更容易恢复”，反而可能与更高云量、更低辐射、或更强的事件前异常幅度共同出现，从而延长后续恢复尾部。颜色上较高 SMrz 的点更多分布在中高 PRE 区间，提示降水效应并非直接作用于恢复时间，而是先重塑根区含水，再改变事件后的恢复轨迹。"
            ),
            caption=f"图 {figure_no}. GPP prepeak 中 Cropland 的 PRE-SMrz dependence plot。",
        ),
    )
    figure_no += 1
    add_figure(
        document,
        FigureItem(
            path=dependence_image_path("GPP_prepeak", "Cropland", "prepeak", "PRE", "EVA"),
            text=(
                "把同一条 PRE 曲线改用 EVA 着色之后，峰值前降水阈值背后的过程含义会更清楚。"
                "在极低 PRE 区间内，快速上升段对应的点并不主要集中在最强蒸散颜色带，"
                "说明这里首先反映的是水分短缺背景下对少量补水的高敏感，而不是单纯由蒸散增强驱动。"
                "但随着 PRE 进入中等及更高区间，颜色逐渐向更强蒸散带聚集，而 SHAP 同时由正转负，"
                "这提示中高降水之所以不再继续缩短恢复尾部，很可能是因为补水并没有脱离能量-耗水约束，"
                "反而进入了“补给增加但耗水同步增强”的耦合状态。换句话说，PRE 的转折不能只理解为水分输入阈值，"
                "它还包含一个由蒸散过程重写的有效补水阈值。"
            ),
            caption=f"图 {figure_no}. GPP prepeak 中 Cropland 的 PRE-EVA dependence plot。",
        ),
    )
    figure_no += 1
    add_figure(
        document,
        FigureItem(
            path=dependence_image_path("GPP_prepeak", "Cropland", "prepeak", "EVA"),
            text=(
                "同一 biome 中的 EVA 曲线进一步说明，单看水分输入并不足够，还必须看耗水强度。这里横轴 EVA 以负值呈现，因此“更靠左”意味着更强的蒸散耗水。曲线大致在 -3.0 至 -2.5 mm 附近升至最高，然后在约 -1.5 mm 左右开始明显回落，并在更弱蒸散区间转为负值。这个结果不宜被简单写成“蒸散越大恢复越慢”或“蒸散越大恢复越快”，更准确的说法是：中等偏强的峰值前蒸散最容易把系统推入恢复时间延长区，而过弱蒸散则更接近能量不足或植被活动较低的背景，恢复压力反而减轻。SMrz 颜色分层说明，强蒸散耗水只有在土壤水尚能支撑蒸散时才会表现为明显正 SHAP，否则它更像是干化后的结果而非主导因子本身。"
            ),
            caption=f"图 {figure_no}. GPP prepeak 中 Cropland 的 EVA-SMrz dependence plot。",
        ),
    )
    figure_no += 1
    add_figure(
        document,
        FigureItem(
            path=dependence_image_path("GPP_prepeak", "Cropland", "prepeak", "SSRD"),
            text=(
                "SSRD 曲线把能量控制的非线性表现得更清楚。短波辐射在大约 1.3×10^7–1.6×10^7 的范围内对应最高正 SHAP，随后在 1.8×10^7 左右附近开始跌破零值，并在 2.4×10^7 后出现第二次更陡的下滑。这个形态提示，短波辐射并不是越高越有利。适中的峰值前辐射可能代表活跃但尚未过热的生长背景，此时骤旱后恢复尾部更长；而过高辐射则更可能伴随强蒸散和土壤干化，使系统从“高活性背景”转向“强水分压力背景”，由此引发第二阶段的恢复拖延。"
            ),
            caption=f"图 {figure_no}. GPP prepeak 中 Cropland 的 SSRD-SMrz dependence plot。",
        ),
    )
    figure_no += 1
    add_figure(
        document,
        FigureItem(
            path=dependence_image_path("GPP_prepeak", "Cropland", "prepeak", "SSRD", "EVA"),
            text=(
                "当同一条 GPP prepeak 的 Cropland 短波辐射曲线改用 EVA 着色后，可以看到两个关键信息。"
                "第一，位于高正 SHAP 区的中等辐射段往往伴随中等偏强蒸散颜色带，说明此时辐射并不是单独发挥作用，而是在推动更活跃的地表耗水与冠层交换。"
                "第二，到了更高辐射区间之后，曲线快速转负，而这些点更多对应较强或持续增强的蒸散背景，这说明高 SSRD 的负效应并不是“光照太强”这么简单，"
                "而更可能代表辐射通过蒸散放大了根区失水和水热不匹配，最终把恢复时间推向更长。换句话说，SSRD 的后段转折可以被理解为从“辐射促进活动”切换到“辐射驱动耗水压力”的门槛。"
            ),
            caption=f"图 {figure_no}. GPP prepeak 中 Cropland 的 SSRD-EVA dependence plot。",
        ),
    )
    figure_no += 1

    add_heading(document, "三、GPP：恢复期内部的热辐射与长波门槛", level=1)
    add_figure(
        document,
        FigureItem(
            path=dependence_image_path("GPP_recovery", "Forest", "recoverywin", "PRE"),
            text=(
                "GPP recovery 中 Forest 的 PRE 曲线显示了另一种转折结构。恢复期降水在约 3–6 mm 区间内对应最高正 SHAP，而在约 9–10 mm 左右开始跨过零值，12 mm 以后转入持续负值。这里的物理含义更偏向恢复阶段内部的“补水窗口”：中等强度补水有助于维持较长的恢复调整尾部，说明森林系统仍在逐步回补根区水分与冠层功能；但当降水过高时，额外补水不再是主导限制，反而可能代表冷湿、低辐射、低光合利用效率的恢复环境，因此对恢复时间的边际作用下降甚至转负。"
            ),
            caption=f"图 {figure_no}. GPP recovery 中 Forest 的 PRE-SMrz dependence plot。",
        ),
    )
    figure_no += 1
    add_figure(
        document,
        FigureItem(
            path=dependence_image_path("GPP_recovery", "Forest", "recoverywin", "PRE", "EVA"),
            text=(
                "Forest 恢复期的 PRE 图改用 EVA 着色后，可以进一步判断中等补水为何对应最长的恢复调整尾部。"
                "从图上看，约 3–6 mm 的高正 SHAP 区并不是落在最低蒸散颜色上，"
                "而更多对应仍然保持明显耗水活动的颜色带，这说明这段降水并不是简单地“把系统浇活了”，"
                "而更像是在维持一个仍需持续耗水、持续重配水分的恢复阶段。"
                "随后当 PRE 继续升高并越过转折阈值时，曲线转负，但点云并未同步转向极弱蒸散，"
                "提示高降水端的边际减效并不只是补水过剩，而更可能包含冷湿、低辐射与蒸散效率下降共同作用。"
                "因此，这张图支持把恢复期森林的降水效应理解为“补水窗口”，而不是线性补得越多越有利。"
            ),
            caption=f"图 {figure_no}. GPP recovery 中 Forest 的 PRE-EVA dependence plot。",
        ),
    )
    figure_no += 1
    add_figure(
        document,
        FigureItem(
            path=dependence_image_path("GPP_recovery", "Forest", "recoverywin", "STRD"),
            text=(
                "新增的 STRD 图之所以必须加入，是因为 recovery 口径里 STRD 在 beeswarm 中反复位居前列，而单独看 SSRD 无法判断辐射控制究竟来自白天地表能量输入还是夜间/全天长波背景。Forest 的 STRD 曲线在约 2.6×10^7–3.1×10^7 范围内维持高正 SHAP，到了大约 3.15×10^7 后开始快速回落，并在 3.45×10^7 左右附近穿过零值。这个转折更像是从“温暖而仍可恢复”的长波背景切换到“过热、夜间冷却不足或大气逆辐射偏强”的环境状态，意味着恢复期的热负荷开始压过补水效应。由于高 STRD 区间下低 SMrz 颜色更常见，说明长波增强一旦叠加土壤偏干，恢复拖延会被进一步放大。"
            ),
            caption=f"图 {figure_no}. GPP recovery 中 Forest 的 STRD-SMrz dependence plot。",
        ),
    )
    figure_no += 1
    add_figure(
        document,
        FigureItem(
            path=dependence_image_path("GPP_recovery", "Forest", "recoverywin", "STRD", "EVA"),
            text=(
                "同一条 STRD 曲线若改用 EVA 着色，可以更直接看到长波热背景与实际耗水之间的传导关系。"
                "在维持高正 SHAP 的中高 STRD 区间内，点云多分布在较强蒸散颜色带上，"
                "说明这部分长波增强并不是抽象的热背景，而是在支持持续的能量交换和耗水活动。"
                "而当 STRD 再继续升高并越过回落门槛后，曲线快速下滑，但颜色仍主要停留在偏强蒸散区，"
                "这意味着恢复拖延的关键并不是“热背景变强”本身，而是热背景已经通过蒸散把系统推向了更高耗水负担。"
                "因此，Forest recovery 中 STRD 的阈值，可以更稳妥地理解为从“热量支撑恢复活动”切换到“热量放大耗水压力”的门槛。"
            ),
            caption=f"图 {figure_no}. GPP recovery 中 Forest 的 STRD-EVA dependence plot。",
        ),
    )
    figure_no += 1

    add_heading(document, "四、RECO：峰值前温度门槛与恢复期大气干旱需求", level=1)
    add_figure(
        document,
        FigureItem(
            path=dependence_image_path("RECO_prepeak", "Cropland", "prepeak", "TMP"),
            text=(
                "RECO prepeak 中 Cropland 的 TMP 曲线是新增图中最值得保留的一个例子。曲线在 288–294 K 区间变化较缓，到了约 295–297 K 附近转为明显负值，而在约 299–301 K 又陡然上冲至正峰，302 K 之后再度回落。这个“双转折”说明峰值前温度对呼吸恢复时间的影响不是单调增暖效应。较低到中等温度区间内，温度升高可能只是推动基质分解和呼吸活性，尚未形成明显恢复压力；一旦进入更高温区，温度与蒸散、大气干燥和土壤失水耦合，恢复尾部会急剧拉长；而更高温端再次回落，则更可能对应极端高温下样本稀疏或系统活动受抑后的边际效应减弱，因此这里只能做谨慎解释，不能写成确定性的单向机制。"
            ),
            caption=f"图 {figure_no}. RECO prepeak 中 Cropland 的 TMP-SMrz dependence plot。",
        ),
    )
    figure_no += 1
    add_figure(
        document,
        FigureItem(
            path=dependence_image_path("RECO_prepeak", "Cropland", "prepeak", "TMP", "EVA"),
            text=(
                "将同一条 TMP 曲线改用 EVA 着色后，可以进一步判断温度阈值是否只是蒸散阈值的表面投影。"
                "在 Cropland 这张图里，约 299–301 K 附近的高正 SHAP 区主要被更强蒸散耗水的颜色带占据，"
                "说明温度跃升之所以会把 RECO 恢复时间迅速推高，并不只是“高温本身”在起作用，"
                "而是高温触发了更强的蒸散拉动和水分消耗，使热力状态与耗水过程共同把系统推过恢复门槛。"
                "这也是为什么在只看 TMP-SMrz 图时，我们只能谨慎说存在双阈值；加入 EVA 映射后，才可以更有把握地指出其中至少一部分转折是经由蒸散中介实现的。"
            ),
            caption=f"图 {figure_no}. RECO prepeak 中 Cropland 的 TMP-EVA dependence plot。",
        ),
    )
    figure_no += 1
    add_figure(
        document,
        FigureItem(
            path=dependence_image_path("RECO_recovery", "Forest", "recoverywin", "PRE"),
            text=(
                "RECO recovery 中 Forest 的 PRE 曲线与 GPP recovery 一样存在明显补水阈值，但形态并不完全相同。曲线在约 3–6 mm 的恢复期降水区间达到高正值，9–10 mm 左右附近开始跌破零值，之后在高降水端维持弱负到中度负值。与 GPP 相比，RECO 对这一门槛的反应更可能体现为呼吸系统对冷湿恢复环境的持续敏感性：适度补水会延长呼吸回到基线所需的调整时间，但过高降水会把系统推向低温、低氧或低辐射背景，使得呼吸异常幅度不再继续放大。"
            ),
            caption=f"图 {figure_no}. RECO recovery 中 Forest 的 PRE-SMrz dependence plot。",
        ),
    )
    figure_no += 1
    add_figure(
        document,
        FigureItem(
            path=dependence_image_path("RECO_recovery", "Forest", "recoverywin", "VPD"),
            text=(
                "RECO recovery 中 Forest 的 VPD 图则说明为什么必须把 VPD 从补充变量提升为核心讨论对象。曲线在约 5–6 左右从负值转正，并在 12–15 左右达到高平台，17–18 之后才逐渐回落但仍保持正值。这说明恢复期呼吸时间并不是在低 VPD 下最长，而是在中高 VPD 条件下最容易被拉长。物理上，这意味着恢复中的森林并非单纯受补水控制，而是对大气干燥需求非常敏感：当 VPD 进入中高区间，蒸散拉动、气孔调节与水热不匹配共同增强，呼吸和同化的耦合恢复被拉长；但在极端更高的 VPD 端，响应边际开始减弱，可能意味着系统已进入受限状态，进一步增干并不会等比例增加恢复时间。SMrz 颜色由黄绿向蓝紫的变化也表明，同一 VPD 水平下，较低根区土壤水通常对应更高的恢复压力。"
            ),
            caption=f"图 {figure_no}. RECO recovery 中 Forest 的 VPD-SMrz dependence plot。",
        ),
    )
    figure_no += 1
    add_figure(
        document,
        FigureItem(
            path=dependence_image_path("RECO_recovery", "Forest", "recoverywin", "VPD", "EVA"),
            text=(
                "VPD 改用 EVA 着色后，恢复期森林中的物理链条变得更完整。"
                "中高 VPD 区间内的高正 SHAP 点更多集中在强蒸散颜色带上，说明这里的恢复拖延并不是抽象的大气干燥需求在单独起作用，"
                "而是大气需水经过实际蒸散这一通道被“落实”为持续耗水压力。"
                "换句话说，VPD 提供的是大气端拉力，而 EVA 则反映这种拉力是否真的转化为了生态系统耗水；"
                "只有当二者共同升高时，恢复尾部才最容易被显著拉长。"
            ),
            caption=f"图 {figure_no}. RECO recovery 中 Forest 的 VPD-EVA dependence plot。",
        ),
    )
    figure_no += 1
    add_figure(
        document,
        FigureItem(
            path=dependence_image_path("RECO_recovery", "Forest", "recoverywin", "SSRD", "EVA"),
            text=(
                "RECO recovery 中 Forest 的 SSRD-EVA 图把恢复期辐射控制的过程链进一步补完整。"
                "曲线在中等辐射区由负转正并达到平台时，点云主要落在较强蒸散颜色带上，说明此时短波辐射的增强首先表现为更活跃的能量驱动和更高的蒸散消耗；"
                "而当辐射继续升高后，曲线重新下滑甚至转负，高蒸散颜色仍然占优势，这表明恢复期短波辐射并不会无限制地促进系统恢复，"
                "相反，它会在一定阈值之后通过蒸散耗水加剧水分亏缺，使呼吸恢复尾部重新被拉长。"
                "因此，这张图支持一个更具体的机制判断：在恢复窗口内部，辐射效应之所以呈现非线性，关键不是辐射本身，而是它是否通过蒸散把系统推过了额外耗水的门槛。"
            ),
            caption=f"图 {figure_no}. RECO recovery 中 Forest 的 SSRD-EVA dependence plot。",
        ),
    )
    figure_no += 1

    add_heading(document, "五、蒸散着色为什么重要：SMrz 调制与 EVA 中介是两件不同的事", level=1)
    add_paragraph(
        document,
        "新增的 EVA 着色图帮助把此前已经观察到的非线性进一步拆成两个层面。SMrz 着色回答的是：同一环境变量水平下，不同土壤水背景是否会改变 SHAP 响应的强弱与方向；EVA 着色回答的则是：降水、温度、辐射或 VPD 这些环境变量，是否必须通过实际蒸散耗水这一过程，才会真正把恢复时间推向更长或更短。前者偏向“背景调制”，后者偏向“过程中介”，两者不能互相替代。",
    )
    add_paragraph(
        document,
        "因此，当讨论 PRE、TMP、STRD、SSRD 和 VPD 这类变量时，只保留 SMrz 映射是不够的。因为无论是补水增强、热量增强还是大气变干，都不必然意味着恢复时间一定改变；只有当这些异常确实转化为更强或更弱的蒸散、并进一步重塑根区含水状态时，曲线中的转折才真正具有物理机制意义。新增的 PRE-EVA、TMP-EVA、VPD-EVA、SSRD-EVA 与 STRD-EVA dependence plot 正是为了解决这个问题，它们让“环境异常是否只是表象，还是经由耗水过程变成真实压力”这一判断有了直接证据。",
    )

    add_heading(document, "六、为什么还要补画 TMP、STRD 和 VPD", level=1)
    add_paragraph(
        document,
        "补画这三类 dependence plot 不是为了增加篇幅，而是因为如果只保留 PRE、EVA 和 SSRD，就会留下三个解释缺口。第一，若不画 TMP，就无法判断温度效应是真正线性增强，还是像 RECO prepeak Cropland 那样存在冷暖双阈值。第二，若不画 STRD，就只能把辐射控制全部归于短波输入，而无法识别恢复期中长波热背景何时开始从“支持恢复”转为“增加热负荷”。第三，若不画 VPD，就会把大气干旱需求错误地压缩进 PRE 或 EVA 的综合效应里，从而失去对气孔调节和蒸散拉动机制的独立辨识能力。换言之，TMP、STRD 和 VPD 的新增 dependence plot 是让物理机制解释成立的必要条件，而不是可有可无的补图。",
    )
    add_heading(document, "综合讨论", level=1)
    add_paragraph(
        document,
        "综合这些曲线可以得到一个更真实、也更克制的认识。恢复时间的驱动机制不是“某变量越大越怎样”的简单逻辑，而是若干门槛依次接力的过程：低到中等补水阶段主要体现水分缓解效应，中高补水阶段则可能转向低辐射或高湿约束；中等偏强蒸散既可能代表活跃背景，也可能代表水分过度消耗；辐射变量在一定区间内支持活跃恢复，但一旦超过热负荷门槛就会转而拉长恢复尾部；VPD 和温度的作用也都不是线性的，而要通过 SMrz 背景来决定其真正方向。正因为存在这些转折，dependence plot 才是物理机制解释中不可替代的第二层证据，而新增的 TMP、STRD 和 VPD 图则是让这一解释不至于失真的关键补充。",
    )
    return document


def compare_beeswarm_path(metric: str) -> Path:
    return COMPARE_DIRS[metric] / "all_biomes_prepeak_vs_recoverywin_beeswarm.png"


def compare_dependence_path(metric: str, biome: str) -> Path:
    return COMPARE_DIRS[metric] / f"{biome}_prepeak_vs_recoverywin_dependence.png"


def build_doc3_comparison() -> Document:
    document = make_document("全球 GPP 与 RECO 恢复时间驱动机制分析（三）：prepeak/recovery 与 GPP/RECO 双层比较")
    for para in shared_intro_paragraphs():
        add_paragraph(document, para)
    add_paragraph(
        document,
        "本部分的比较包含两个层次。第一个层次是在同一碳通量内部比较 prepeak 与 recovery 两种解释口径，从而区分恢复记忆与恢复过程控制。第二个层次是在同一时间口径下比较 GPP 与 RECO，从而识别同化恢复与呼吸恢复在主控机制上的共性与差异。和旧版不同，下面的比较不再只停留在“变量位次换了”，而是专门追踪同一变量在两套口径下的 SHAP 方向是否发生翻转，例如是从“高值端压向负侧”变成“中高值区抬到正侧”，还是从“高值端正效应”改为“高值端负效应”。只有把这一步讲清楚，prepeak/recovery 与 GPP/RECO 的对比才真正有分析价值。",
    )

    fig_no = 1
    add_heading(document, "一、同一碳通量内部的 prepeak 与 recovery 比较", level=1)
    for metric in ["GPP", "RECO"]:
        add_heading(document, f"{metric}：前置记忆与恢复过程的口径差异", level=2)
        for para in DOC3_METRIC_OVERVIEW[metric]:
            add_paragraph(document, para)
        add_figure(
            document,
            FigureItem(
                path=compare_beeswarm_path(metric),
                text=(
                    f"{metric} 的全 biome beeswarm 对比图不是简单把两套图并排放在一起，而是用来检验同一个通量在“事件前背景记忆”和“恢复窗口直接控制”之间的排序位移。"
                    f"真正要看的不是某一个变量有没有出现，而是它在两种口径之间有没有明显前移、后移或被别的控制项替代。"
                ),
                caption=f"图 {fig_no}. {metric} 在 prepeak 与 recovery 两种口径下的全 biome beeswarm 对比图。",
            ),
        )
        fig_no += 1
        for biome in BIOMES:
            add_figure(
                document,
                FigureItem(
                    path=compare_dependence_path(metric, biome),
                    text=" ".join([DOC3_METRIC_BIOME_COMPARE[metric][biome], phase_shift_sentence(metric, biome)]),
                    caption=f"图 {fig_no}. {metric} 在 {biome} biome 中 prepeak 与 recovery 口径的 dependence 对比图。",
                ),
            )
            fig_no += 1

    add_heading(document, "二、同一口径下的 GPP 与 RECO 比较", level=1)
    for phase in ["prepeak", "recoverywin"]:
        add_heading(document, f"{PHASE_LABELS[phase]}下的 GPP 与 RECO 并列比较", level=2)
        for para in DOC3_PHASE_OVERVIEW[phase]:
            add_paragraph(document, para)
        for biome in BIOMES:
            add_paragraph(
                document,
                " ".join([DOC3_CROSS_METRIC_BIOME[phase][biome], cross_metric_sentence(phase, biome)])
            )
            gpp_key = f"GPP_{'prepeak' if phase == 'prepeak' else 'recovery'}"
            reco_key = f"RECO_{'prepeak' if phase == 'prepeak' else 'recovery'}"
            add_figure(
                document,
                FigureItem(
                    path=beeswarm_image_path(gpp_key, biome, phase),
                    text="",
                    caption=f"图 {fig_no}. GPP 在 {PHASE_LABELS[phase]}下 {biome} biome 的 SHAP beeswarm 图。",
                ),
            )
            fig_no += 1
            add_figure(
                document,
                FigureItem(
                    path=beeswarm_image_path(reco_key, biome, phase),
                    text="",
                    caption=f"图 {fig_no}. RECO 在 {PHASE_LABELS[phase]}下 {biome} biome 的 SHAP beeswarm 图。",
                ),
            )
            fig_no += 1

    add_heading(document, "三、关键证据对照：差异不是抽象排序，而是阈值和传导链不同", level=1)
    add_paragraph(
        document,
        "如果只看 beeswarm 排序，prepeak 与 recovery、GPP 与 RECO 的差异仍然容易被概括成“重要变量不一样”。但真正更关键的是，同样的变量在不同口径和不同通量下，其曲线达到峰值、开始回落、以及由正 SHAP 穿过零值的位置都不一样，而且这些变化常常伴随不同的蒸散颜色带。这意味着我们面对的不是简单的权重差异，而是恢复过程在不同情景下被不同的阈值和不同的传导链条所控制。",
    )
    add_figure(
        document,
        FigureItem(
            path=dependence_image_path("GPP_prepeak", "Cropland", "prepeak", "PRE", "EVA"),
            text=(
                "以 GPP prepeak 的 Cropland 为例，PRE-EVA 图显示峰值前降水在极低值区间先快速抬升，然后在进入中等区间后逐步转负。"
                "这里的重点不只是存在转折，而是转折后点云逐渐转向更强蒸散颜色带，说明峰值前补水效应会被随后的耗水过程重写。"
                "因此，prepeak 口径更像是在刻画事件形成前背景如何预设恢复尾部，而不是直接描述恢复窗口内部的即时补水收益。"
            ),
            caption=f"图 {fig_no}. GPP prepeak 中 Cropland 的 PRE-EVA dependence plot。",
        ),
    )
    fig_no += 1
    add_figure(
        document,
        FigureItem(
            path=dependence_image_path("GPP_recovery", "Forest", "recoverywin", "PRE", "EVA"),
            text=(
                "与之对照，GPP recovery 的 Forest PRE-EVA 图更像是在展示恢复期内部的“补水窗口”。"
                "中等 PRE 对应高正 SHAP，且仍伴随明显蒸散活动，说明这时系统并不是立刻恢复完毕，而是在继续进行耗水和功能重建；"
                "等到 PRE 更高、曲线转负之后，蒸散颜色并未同步消失，提示高降水端的边际减效还叠加了冷湿和低辐射约束。"
                "这与 prepeak 那种“背景预设”机制不同，体现的是恢复期近端环境对恢复尾部的直接塑形。"
            ),
            caption=f"图 {fig_no}. GPP recovery 中 Forest 的 PRE-EVA dependence plot。",
        ),
    )
    fig_no += 1
    add_figure(
        document,
        FigureItem(
            path=dependence_image_path("RECO_prepeak", "Cropland", "prepeak", "TMP", "EVA"),
            text=(
                "跨通量比较时，RECO prepeak 的 Cropland TMP-EVA 图提供了一个很好的反例。"
                "它并不是像 GPP 那样主要围绕补水阈值展开，而是在较高温区出现明显的高正 SHAP 突升，并且该突升与更强蒸散颜色带相伴。"
                "这说明呼吸恢复时间在峰值前背景下更容易受到热力状态和耗水耦合的共同放大，温度阈值在这里不是表层现象，而是通过蒸散与底物环境一起改变了恢复拖尾。"
            ),
            caption=f"图 {fig_no}. RECO prepeak 中 Cropland 的 TMP-EVA dependence plot。",
        ),
    )
    fig_no += 1
    add_figure(
        document,
        FigureItem(
            path=dependence_image_path("RECO_recovery", "Forest", "recoverywin", "VPD", "EVA"),
            text=(
                "再看 RECO recovery 的 Forest VPD-EVA 图，中高 VPD 区间的高正 SHAP 与强蒸散颜色紧密重叠，"
                "表明恢复期呼吸并不是简单受补水控制，而是强烈受大气需水拉力是否被转化为真实耗水压力所控制。"
                "这与 GPP recovery Forest 的 PRE-EVA 图形成鲜明对照：前者更强调大气端拉力与耗水中介，后者更强调补水窗口与热湿边际减效。"
                "因此，GPP 与 RECO 的差异并不只是一个偏同化、一个偏呼吸，而是两者在恢复期真正受控的过程链条不同。"
            ),
            caption=f"图 {fig_no}. RECO recovery 中 Forest 的 VPD-EVA dependence plot。",
        ),
    )
    fig_no += 1

    add_heading(document, "综合讨论", level=1)
    add_paragraph(
        document,
        "重写全部比较关系之后，可以把结论压缩成三点。第一，prepeak 与 recovery 的差异不是简单的“重要变量名单不同”，而是同一变量的符号方向都可能改写。例如 prepeak 中高 SSRD 更常把点云压向负侧，而 recovery 中不少 biome 的 SSRD 会先从低值端负侧抬升到中高值正峰，然后再在极高端回落，这意味着辐射从“背景记忆”转成“恢复过程门槛”时，物理角色已经变了。第二，GPP 与 RECO 的差异也不能只靠排序描述。GPP 更频繁地在高 VPD、高 WIND 端出现正 SHAP 拉升，说明同化恢复更容易被大气端和空气动力端拖尾；RECO 则更常在 PRE 的中高值区形成更强正峰，说明呼吸恢复更直接受补水窗口控制。第三，真正有用的比较必须同时交代 rank 与 direction：谁排在前列决定你该看哪条链，正负方向和翻转区间才决定这条链在不同场景下究竟是在延长恢复，还是在把恢复重新拉回去。"
    )
    return document


def main() -> None:
    ensure_dir(OUT_DIR)
    doc1 = build_doc1_beeswarm()
    save(doc1, OUT_DIR / "01_beeswarm_driver_analysis_cn.docx")
    doc2 = build_doc2_dependence()
    save(doc2, OUT_DIR / "02_dependence_interaction_analysis_cn.docx")
    doc3 = build_doc3_comparison()
    save(doc3, OUT_DIR / "03_comparison_analysis_cn.docx")
    print(OUT_DIR / "01_beeswarm_driver_analysis_cn.docx")
    print(OUT_DIR / "02_dependence_interaction_analysis_cn.docx")
    print(OUT_DIR / "03_comparison_analysis_cn.docx")


if __name__ == "__main__":
    main()
