#!/usr/bin/env python3
"""Generate an integrated academic summary docx from six existing documents."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path("/home/xulc/flash_drought")
OUT_DIR = ROOT / "process/SEM_analysis0401/codex/GLEAM/plots/writing2"
OUT_DOCX = OUT_DIR / "05_integrated_summary_cn.docx"
OUT_MD = OUT_DIR / "05_integrated_summary_cn.md"


def set_run_font(run, size: float, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "SimSun")


def set_doc_language(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    r_pr = normal.element.rPr
    r_fonts = r_pr.rFonts
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "SimSun")


def add_page_numbers(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_run_font(run, 10)


def make_document(title: str) -> Document:
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    add_page_numbers(section)
    set_doc_language(document)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_run_font(run, 16, bold=True)
    p.paragraph_format.space_after = Pt(12)
    return document


def add_heading(document: Document, text: str, level: int = 1) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6)
    if level == 1:
        p.style = document.styles["Heading 1"]
        size = 14
    else:
        p.style = document.styles["Heading 2"]
        size = 13
    run = p.add_run(text)
    set_run_font(run, size, bold=True)


def add_paragraph(document: Document, text: str, first_line_cm: float = 0.74) -> None:
    p = document.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(first_line_cm)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, 12)


def add_figure(document: Document, path: Path, caption: str, intro: str = "", width_cm: float = 15.4) -> None:
    if intro:
        add_paragraph(document, intro)
    if not path.exists():
        add_paragraph(document, f"缺图占位：{path}")
        return
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_run_font(run, 10.5)


def export_plain_text(document: Document, out_path: Path) -> None:
    lines: list[str] = []
    for p in document.paragraphs:
        text = p.text.strip()
        if text:
            lines.append(text)
            lines.append("")
    out_path.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = make_document("全球骤旱后 GPP 与 RECO 恢复时间驱动机制分析：综合总结版")

    add_paragraph(
        doc,
        "本总结版文档并不是对既有六份长文的简单压缩，而是在统一研究主线下重新组织证据。全文只保留对后续写作最关键的内容：骤旱事件如何被定义，为什么主事件库采用 GLEAM、主碳通量产品采用 BESS，GPP 恢复时间在 0401 版本下如何计算，GPP 与 RECO 在 prepeak 和记忆阶段及 recovery 过程阶段分别呈现出怎样的驱动模式，以及这些模式最终如何在 dependence plot 和 SEM 路径中收束为可解释的物理过程链。"
    )
    add_paragraph(
        doc,
        "从材料结构上看，这份总结版整合了两类证据。第一类是方法和数据源合理性验证，包括 GLEAM 与 ERA5 的骤旱识别对照，以及 BESS 与 FluxSat 的 GPP 恢复结果对照。第二类是机制解释证据，包括 SHAP beeswarm 给出的贡献排序、dependence plot 给出的阈值与交互，以及 SEM 给出的路径传导关系。这样组织的目的，是让“事件定义”“响应恢复规则”“统计驱动模式”和“机制传导链”四层内容形成闭环，而不是各自分散成孤立图件。"
    )

    add_heading(doc, "1. 方法基础：骤旱定义与主数据源选择", level=1)
    add_paragraph(
        doc,
        "本研究采用的骤旱识别不是经验式拼接，而是一个可回溯的两步法过程。第一步先在逐日百分位框架下识别全部干旱事件：土壤湿度滑动平均值低于对应日的第 20 百分位被视为进入干旱，只有当连续 10 天回升到第 40 百分位及以上时事件才结束，同时还要求事件内低于 P20 的天数不少于 15 天。第二步再向前回溯最多 90 天，寻找最近一次高于 P40 的湿润起点，并以 onset days 是否落在 5–20 天之间将事件归类为 flash drought。也就是说，本研究中的骤旱必须同时满足“已经足够干”和“由湿转干足够快”两层约束。"
    )
    add_paragraph(
        doc,
        "在这一统一规则下，GLEAM 与 ERA5 可直接平行比较。结果显示，两者在空间热点上总体一致，但一致性主要体现在频次场而非强度场。根区事件频次上，GLEAM 和 ERA5 的空间相关为 0.709，表层为 0.851；而平均强度场的相关仅约 0.121–0.127。这意味着不同土壤湿度产品对“哪里容易发生骤旱”给出的答案较接近，但对“同一位置的骤旱绝对强度有多大”差异更明显。对后续生态响应研究而言，更稳定、也更值得依赖的是事件出现的空间结构和事件库样本规模，而不是不同产品间难以无偏互换的强度绝对值。"
    )
    add_figure(
        doc,
        ROOT / "process/result_analysis/result_weighted/conclusion/gleam_era5_flash_frequency_intensity_spatial/gleam_vs_era5_flash_frequency_global.png",
        "图 1. GLEAM 与 ERA5 识别的全球骤旱频次空间分布对照。上排为根区，下排为表层。",
        intro="图 1 展示了采用统一判定规则后 GLEAM 与 ERA5 在全球骤旱频次场上的大尺度一致性。尽管根区情景下 GLEAM 识别出的事件更多，但热点主要仍集中在相近区域，这也是后续保留 GLEAM 作为主事件库的重要前提。"
    )
    add_paragraph(
        doc,
        "之所以最终以 GLEAM 作为主事件源，是因为它同时满足三个条件：一是与 ERA5 在全球频次格局上具有较高一致性，因此并没有偏离大尺度骤旱热点；二是在根区层面提供了更充分的事件样本，更适合支持 biome 分组、SHAP 建模和 SEM 机制分析；三是在下游碳通量结果中更容易保留恢复阶段的长期拖延信号，因此更有利于解释骤旱后的生态后果。"
    )

    add_heading(doc, "2. 0401 版本下 GPP 恢复时间计算及 BESS–FluxSat 验证", level=1)
    add_paragraph(
        doc,
        "0401 版本的关键并不在于改写骤旱事件本身，而在于重新定义 GPP 恢复时间的计量口径。首先，只有当某个骤旱事件在干旱持续期内超过一半的天数落在像元生长季内时，该事件才进入 GPP 响应与恢复分析。其次，GPP 被定义为负向 drought response，程序会同时记录从 onset start 和 drought start 出发到首次稳定进入负异常区的时间，并继续向后搜索冲击峰值。最后，恢复日期允许跨越休眠季，但恢复历时不再按日历日简单累计，而只统计峰值之后到恢复日期之间落在生长季掩膜内的有效日。这个变化非常关键，因为它剔除了冬季静止期对 GPP 恢复时间的人为放大。"
    )
    add_paragraph(
        doc,
        "在这一统一框架下，BESS 的优势在于它同时提供 GPP 与 RECO 的长时间连续序列，并且二者共享同一空间网格、时间轴和异常计算方法。这使得后续关于“骤旱后碳吸收与碳释放是否具有不同恢复节律”的讨论，能够主要反映生态过程本身，而不是不同数据源预处理链路不一致造成的伪差异。为了验证 BESS 的 GPP 结果是否可信，本研究又引入独立 GPP 产品 FluxSat 作为外部对照。"
    )
    add_paragraph(
        doc,
        "BESS 与 FluxSat 的比较表明，两者在恢复时间绝对值和可进入统计的样本集合上确有差异，但在最重要的方向性和空间格局上是一致的。在 2001–2018 共同覆盖期、统一 100 天恢复上限的口径下，SMrz 情景中 BESS 与 FluxSat 的 GPP 响应变慢趋势分别为 9.96 d/10a 和 10.26 d/10a；SMs 情景中分别为 10.75 d/10a 和 10.79 d/10a。共同期平均恢复时间场的空间相关也达到 0.556（SMrz）和 0.601（SMs）。因此，BESS 给出的 GPP 恢复结果并非孤立产品特征，而具有独立数据源支持。"
    )
    add_figure(
        doc,
        ROOT / "process/result_analysis/result_weighted/conclusion/BESS_Fluxsat_valid/smrz_recovery_mean_bess0401_vs_fluxsat_fixlon.png",
        "图 2. 根区闪旱情景下 BESS 与 FluxSat 的平均恢复时间空间对照（共同期统一口径）。",
        intro="图 2 用共同期统一口径直接比较 BESS 与 FluxSat 的恢复时间场。两者在主要高值区和低值区的位置上保持中等偏强一致性，说明 BESS 作为后续 GPP 与 RECO 并行分析主数据源在空间结构上是可靠的。"
    )

    add_heading(doc, "3. Beeswarm：不同 biome 在 prepeak 与 recovery 阶段的主导模式", level=1)
    add_paragraph(
        doc,
        "SHAP beeswarm 的作用不是只给出变量排名，而是帮助识别不同 biome 在不同阶段究竟优先受什么类型的控制。对 GPP 来说，prepeak 阶段最鲜明的特征是 biome 分化非常明显。Forest 和 Savanna 主要由辐射背景和蒸散背景主导，Cropland 则最早把 EVA 推到第一位，说明农田更容易记住峰值前耗水遗产；相对地，Grassland 与 Shrubland 更快抬升 TMP 和 VPD，表明开放植被在骤旱爆发前更容易被热干空气背景预设。进入 recovery 阶段后，这种分化明显收束，PRE 在五个 biome 中全部进入前两位，说明恢复窗口内部的补水条件迅速成为直接控制项，但不同 biome 又分别被 STRD、VPD 或 WIND 进一步分流。"
    )
    add_figure(
        doc,
        ROOT / "process/SEM_analysis0401/codex/GLEAM/plots/gpp_code1_flash_smrz_compare_prepeak_vs_recoverywin_v20260401_20260424/all_biomes_prepeak_vs_recoverywin_beeswarm.png",
        "图 3. GPP 在 prepeak 与 recovery 两种口径下的全 biome beeswarm 对比图。",
        intro="图 3 最核心的信息不是某个变量有没有出现，而是控制抓手在阶段间如何换挡。GPP 在 prepeak 更像记录能量-耗水背景记忆，而在 recovery 则迅速转向“补水是否有效”的直接控制。"
    )
    add_paragraph(
        doc,
        "RECO 的 beeswarm 结构与 GPP 不完全相同。prepeak 阶段，RECO 的结果更加整齐，多个 biome 都由 SSRD 与 TMP 领跑，说明呼吸恢复的峰值前记忆更接近一个统一的热力背景框架；进入 recovery 阶段后，PRE 在 Savanna、Cropland 和 Shrubland 中直接排到第一位，表明呼吸恢复对恢复窗口内部的有效补水依赖更直接。与 GPP 相比，RECO 较少把高 VPD 或高 WIND 推到首位，因此它更像是以补水主导、再由第二控制项修饰的恢复结构。"
    )
    add_figure(
        doc,
        ROOT / "process/SEM_analysis0401/codex/GLEAM/plots/reco_code1_flash_smrz_compare_prepeak_vs_recoverywin_v20260401_20260424/all_biomes_prepeak_vs_recoverywin_beeswarm.png",
        "图 4. RECO 在 prepeak 与 recovery 两种口径下的全 biome beeswarm 对比图。",
        intro="图 4 显示 RECO 的阶段转换更集中：prepeak 主要记录热力背景，recovery 则更清晰地转向补水主导。这也是后续 GPP 与 RECO 机制比较的出发点。"
    )

    add_heading(doc, "4. Dependence plot：关键变量的阈值、转折与阶段差异", level=1)
    add_paragraph(
        doc,
        "如果说 beeswarm 解决的是“谁重要”，dependence plot 解决的就是“在什么值域开始真正改变恢复时间”。总结既有结果后，可以把最具代表性的阈值证据压缩成四类。第一类是峰值前降水并不是线性缓解因子。以 GPP prepeak 的 Cropland 为例，PRE 在极低值区先快速抬升，随后在约 1.5–2.0 mm 左右转向，在 8 mm 之后进一步下行，说明少量补水可以缓解水分短缺背景，但当峰值前背景已经较湿时，额外降水反而可能与更低辐射、更高云量或更强前置异常共同出现。第二类是恢复期降水呈现“补水窗口”。以 GPP recovery 的 Forest 为例，PRE 在约 3–6 mm 区间对应最高正 SHAP，而在约 9–10 mm 左右开始跨过零值，说明恢复阶段存在一个既能维持功能重建、又尚未被冷湿低辐射效应抵消的有效补水窗口。"
    )
    add_figure(
        doc,
        ROOT / "process/SEM_analysis0401/codex/GLEAM/plots/dependence_plot/GPP_prepeak/Cropland_prepeak_PRE_colored_by_EVA.png",
        "图 5. GPP prepeak 中 Cropland 的 PRE-EVA dependence plot。",
        intro="图 5 是峰值前背景里最典型的阈值图之一。它说明 PRE 的作用要经过蒸散过程重写，而不是单纯表现为“降水越多越容易恢复”。"
    )
    add_figure(
        doc,
        ROOT / "process/SEM_analysis0401/codex/GLEAM/plots/dependence_plot/GPP_recovery/Forest_recoverywin_PRE_colored_by_EVA.png",
        "图 6. GPP recovery 中 Forest 的 PRE-EVA dependence plot。",
        intro="图 6 则代表恢复期内部的补水窗口机制。中等 PRE 仍对应明显蒸散活动和较高正 SHAP，说明这时系统并不是立刻恢复完毕，而是在继续进行耗水和功能重建。"
    )
    add_paragraph(
        doc,
        "第三类是温度和大气干燥需求并不是简单单调因子。RECO prepeak 的 Cropland TMP 曲线在约 295–297 K 附近转入负值、在 299–301 K 左右又快速冲到正峰，说明峰值前温度更像一个由热力状态和蒸散中介共同驱动的双阈值变量；而 RECO recovery 的 Forest VPD 曲线在约 5–6 左右由负转正，并在 12–15 左右形成高平台，说明恢复期呼吸时间最容易在中高 VPD 条件下被显著拉长。第四类是短波与长波辐射都具有门槛性质。GPP prepeak 的 Cropland SSRD 大约在 1.3×10^7–1.6×10^7 区间出现最高正 SHAP，之后在约 1.8×10^7 后跌破零值，说明辐射会从“支持活跃背景”切换到“放大耗水压力”。而 GPP recovery 的 Forest STRD 则在约 3.15×10^7 后快速回落，反映长波热背景可能从“维持恢复活动”转向“增加热负荷”。"
    )
    add_figure(
        doc,
        ROOT / "process/SEM_analysis0401/codex/GLEAM/plots/dependence_plot/RECO_prepeak/Cropland_prepeak_TMP_colored_by_EVA.png",
        "图 7. RECO prepeak 中 Cropland 的 TMP-EVA dependence plot。",
        intro="图 7 代表峰值前温度阈值。约 299–301 K 的高正 SHAP 区与更强蒸散颜色带重叠，说明高温并不是单独起作用，而是经由耗水过程把系统推过恢复门槛。"
    )
    add_figure(
        doc,
        ROOT / "process/SEM_analysis0401/codex/GLEAM/plots/dependence_plot/RECO_recovery/Forest_recoverywin_VPD_colored_by_EVA.png",
        "图 8. RECO recovery 中 Forest 的 VPD-EVA dependence plot。",
        intro="图 8 则说明恢复期大气干燥需求如何转化为真实耗水压力。中高 VPD 区间的高正 SHAP 与强蒸散颜色紧密重叠，表明恢复拖延来自大气端拉力与生态系统耗水的同步增强。"
    )

    add_heading(doc, "5. SEM：从统计贡献走向过程链", level=1)
    add_paragraph(
        doc,
        "SEM 的作用不是重复 SHAP 排名，而是把“哪个变量重要”进一步组织为“哪条过程链更主导”。GPP 与 RECO 在这一层的差异非常清楚。GPP prepeak 的路径在不同 biome 之间高度分化：Forest、Grassland 和 Shrubland分别保留了不同形式的辐射或热背景控制，Cropland 更突出 VPD、TMP 和 EVA 的耗水遗产，Savanna 则明显转向由 P-ET、PRE 和 EVA 领跑的水分收支链。进入 recovery 阶段后，GPP 的五个 biome 几乎全部围绕 P-ET 收束，说明真正决定恢复窗口内部 GPP 尾部长度的，不再是单个气象量，而是补水与耗水的净结果；Forest、Grassland 和 Shrubland 再由热湿背景和长波辐射继续分流。"
    )
    add_figure(
        doc,
        ROOT / "process/SEM_analysis0401/codex/GLEAM/plots/SEM/SEM路径机制图/GPP_cropland_prepeak.png",
        "图 9. GPP 在 Cropland biome 的 prepeak SEM 路径机制图。",
        intro="图 9 选取 GPP prepeak 中最有代表性的农田路径图。它清楚展示了 Cropland 在峰值前阶段不是先被单独补水控制，而是更早暴露出 VPD、TMP 与蒸散耗水共同设定的脆弱背景，因此非常适合作为“耗水遗产型” prepeak 机制的代表。"
    )
    add_figure(
        doc,
        ROOT / "process/SEM_analysis0401/codex/GLEAM/plots/SEM/SEM路径机制图/GPP_forest_recovery.png",
        "图 10. GPP 在 Forest biome 的 recovery SEM 路径机制图。",
        intro="图 10 则代表 GPP recovery 阶段的正式路径展示。相较于 prepeak，Forest 在恢复窗口内部已经明显转向补水与耗水净结果主导，并继续受到热湿背景的修饰，体现了 GPP 从“背景记忆”到“恢复过程控制”的阶段切换。"
    )
    add_paragraph(
        doc,
        "RECO 的路径图则更整齐。prepeak 阶段，多个 biome 都由 SSRD 的负效应和 TMP 的正效应共同定义恢复起点，说明呼吸恢复的背景记忆更接近一个统一的热力框架；进入 recovery 阶段后，除 Cropland 之外，Forest、Grassland、Savanna 与 Shrubland 的最强直接路径都切换为 TMP 的负效应，而 SSRD 退居第二层修饰。这说明 RECO 并不是在每个阶段都由同一条链主导，而是从峰值前的辐射背景记忆切换为恢复阶段的高温拖尾机制。"
    )
    add_figure(
        doc,
        ROOT / "process/SEM_analysis0401/codex/GLEAM/plots/SEM/SEM路径机制图/RECO_grass_prepeak.png",
        "图 11. RECO 在 Grassland biome 的 prepeak SEM 路径机制图。",
        intro="图 11 代表 RECO prepeak 的统一热力背景结构。Grassland 这张图最适合展示 SSRD 与 TMP 如何共同定义呼吸恢复起点，因为它保留了最典型的“辐射背景为主、温度为辅”的 prepeak 模式。"
    )
    add_figure(
        doc,
        ROOT / "process/SEM_analysis0401/codex/GLEAM/plots/SEM/SEM路径机制图/RECO_forest_recovery.png",
        "图 12. RECO 在 Forest biome 的 recovery SEM 路径机制图。",
        intro="图 12 则作为 RECO recovery 的正式代表图。与 prepeak 阶段由 SSRD 领跑不同，这里最重要的变化是温度路径前移并主导恢复拖尾，说明恢复窗口内部的高温持续比单纯辐射输入更直接地控制呼吸回到基线的节奏。"
    )

    add_heading(doc, "6. 综合结论", level=1)
    add_paragraph(
        doc,
        "综合六份材料后，可以把本研究的主线压缩为四个层次。第一，骤旱事件的定义是可追溯的两步法：先识别持续干旱，再约束由湿转干的快速爆发过程。第二，在统一事件定义下，GLEAM 与 ERA5 在空间热点上总体一致，但 GLEAM 在根区提供了更充分的事件样本，也在下游恢复阶段保留了更清楚的生态拖延信号，因此更适合作为主事件库。第三，0401 版本下的 GPP 恢复时间采用了生长季有效日计量规则，使恢复历时更贴近植被真正的功能重建；BESS 在这一框架下既能与 FluxSat 的 GPP 结果保持时空一致性，又能同时为 GPP 与 RECO 提供结构一致的输入，因此适合作为主碳通量数据源。第四，从 SHAP 到 dependence 再到 SEM 的证据链共同表明，prepeak 阶段更像是在记录不同 biome 以什么方式进入脆弱状态，而 recovery 阶段则更像是在回答补水、热背景和耗水过程如何共同决定恢复尾部。"
    )
    add_paragraph(
        doc,
        "若只保留一句总括性的认识，那么可以表述为：骤旱后的碳通量恢复并不是由单一水分亏缺线性控制，而是由阶段不同、biome 不同、通量类型不同的多条过程链共同塑造。GPP 更容易在 recovery 阶段表现为有效水分收支主导，并被热湿背景与空气动力条件进一步分流；RECO 则更容易在 prepeak 阶段表现出统一的热力背景记忆、在 recovery 阶段转向高温拖尾。也正因为这两类通量的控制层级并不相同，使用统一事件库和统一响应恢复框架开展并行分析，才具有真正的机制解释价值。"
    )

    doc.save(OUT_DOCX)
    export_plain_text(doc, OUT_MD)
    print(OUT_DOCX)
    print(OUT_MD)


if __name__ == "__main__":
    main()
