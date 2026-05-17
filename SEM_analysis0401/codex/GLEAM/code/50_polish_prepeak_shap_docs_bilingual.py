#!/usr/bin/env python3
"""Generate polished bilingual SHAP writing docs inspired by local SHAP/SEM literature style."""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path("/home/xulc/flash_drought/process/SEM_analysis0401/codex/GLEAM")
PLOTS_ROOT = ROOT / "plots2/prepeak_shap_summary_20260502"
WRITE_ROOT = ROOT / "plots2/writing"
SUMMARY_CSV = WRITE_ROOT / "feature_direction_threshold_summary.csv"

BIOMES = ["Forest", "Grassland", "Savanna", "Cropland", "Shrubland"]
METRICS = ["GPP", "RECO"]

METRIC_CN = {"GPP": "GPP 恢复时间", "RECO": "RECO 恢复时间"}
METRIC_EN = {"GPP": "GPP recovery time", "RECO": "RECO recovery time"}


def load_stats() -> pd.DataFrame:
    return pd.read_csv(SUMMARY_CSV)


def topn(stats: pd.DataFrame, metric: str, biome: str, n: int = 5) -> pd.DataFrame:
    sub = stats[(stats["metric"] == metric) & (stats["biome"] == biome)].copy()
    sub = sub[sub["label"] != "LAI"]
    return sub.sort_values("rank").head(n)


def feature_reason_cn(label: str, direction: str, biome: str) -> str:
    if label == "SSRD":
        if direction == "negative":
            return (
                "这通常说明该 biome 的恢复阶段并不是单纯缺光，而是更容易落入“高辐射伴随高叶温、高蒸散需求和更强大气干旱”的组合情景。"
                "在这种背景下，较强短波辐射虽然提高了可用能量，但同时放大了水分亏缺和光氧化压力，因此反而有助于更快结束恢复尾部，"
                "意味着高 SSRD 对应的是较少受持续湿冷或辐射受限拖累的事件。"
            )
        return (
            "这表明该 biome 中恢复更受能量受限约束，高辐射首先提供了冠层恢复所需的光能和热量，"
            "其促进作用超过了附带的耗水代价，因此高 SSRD 更容易对应较短的恢复时间。"
        )
    if label == "EVA":
        if direction == "positive":
            return (
                "高 EVA 在这里更像是持续耗水和地表-冠层耦合增强的信号，而不是单纯代表植被活跃。"
                "当根区储水尚未恢复时，较强蒸散会继续抽干土壤水并维持较高的大气需求，"
                "从而把系统锁定在恢复未完成的状态，导致恢复时间延长。"
            )
        return (
            "高 EVA 在这里更可能代表能量供应充足且植被-大气交换已经重新活跃，"
            "即系统已重新进入较强的碳水通量状态。此时蒸散并不是额外负担，而是恢复已经展开的结果，"
            "因此对应更短的恢复时间。"
        )
    if label == "TMP":
        if direction == "positive":
            return (
                "这意味着升温的负面效应占了上风。较高温度一方面提高蒸散需求和 VPD，另一方面增强呼吸消耗，"
                "使得水分赤字和碳损耗在恢复期继续累积，因此恢复更容易被拉长。"
            )
        return (
            "这说明该 biome 仍受到一定的低温或冷湿背景限制。温度升高后，酶活性、光合恢复和地表能量交换得到改善，"
            "其正面作用超过了附带的耗水效应，因此恢复时间缩短。"
        )
    if label == "VPD":
        if direction == "positive":
            return (
                "高 VPD 直接意味着更强的大气蒸发需求和更高的气孔关闭风险。"
                "一旦大气干旱超过植被调节能力，碳吸收恢复会被压制，土壤水恢复也更难积累，因此恢复时间延长。"
            )
        return (
            "这里的负向关系往往说明低 VPD 并不一定代表有利条件，反而可能对应持续阴湿、辐射不足或边界层交换偏弱的背景。"
            "适度升高的 VPD 可能与更清晰的晴空能量条件同步出现，因此更有利于恢复推进。"
        )
    if label == "PRE":
        if direction == "positive":
            return (
                "这并不意味着降水本身有害，而更可能说明高降水事件与恢复拖尾共现。"
                "在这类 biome 中，较大的降水往往出现在低辐射、低温或土壤通气受限的背景下，"
                "因此虽然补了水，却没有同步提高有效恢复速率。"
            )
        return (
            "这符合水分补给逻辑。较高降水可以补充根区土壤水、降低后续 VPD 敏感性，并为冠层重新建立蒸散和光合活动提供缓冲，"
            "因此整体上会缩短恢复时间。"
        )
    if label == "SMrz":
        if direction == "positive":
            return (
                "若高 SMrz 仍对应更长恢复时间，通常说明该 biome 的慢恢复并不是由即时水分短缺主导，"
                "而更可能与低能量、过湿抑制或较强的生态系统结构惯性有关。"
            )
        return (
            "这表明根区土壤水是该 biome 恢复最直接的缓冲器。"
            "较高的 SMrz 不仅缓解水力限制，还能降低高温和高 VPD 对植被的放大作用，因此恢复更快完成。"
        )
    if label == "STRD":
        if direction == "positive":
            return (
                "较高 STRD 往往对应暖湿且逆辐射较强的大气背景。"
                "如果这类背景同时伴随夜间保温和呼吸维持成本上升，那么 RECO 或整体恢复尾部都可能被拉长。"
            )
        return (
            "负向关系说明较高 STRD 更接近一个温暖且能量条件较稳的背景指标，"
            "它减少了低温抑制并有利于恢复过程连续推进，因此恢复时间缩短。"
        )
    if label == "WIND":
        if direction == "positive":
            return (
                "较强风速会增强湍流交换并提高蒸散需求，尤其在土壤水未恢复时更容易放大耗水压力，"
                "因此恢复时间被拉长。"
            )
        return (
            "适度风速可能代表边界层交换改善，使热量和水汽更容易扩散，避免近地表滞湿或过热积累，因此恢复时间缩短。"
        )
    if label == "Duration":
        return "持续时间越长，说明胁迫记忆累积越深，非结构性碳储备、根区水库和叶片功能恢复都需要更长时间，因此对恢复时间通常表现为正向拉长。"
    if label == "Intensity":
        return "强度越高，表示骤旱对生理过程的冲击越深，光合体系和呼吸体系都更容易出现持续偏离，因此恢复更慢。"
    return "这一方向反映的是该变量在当前 biome 中与能量、水分和生态系统结构背景共同作用后的净效应，而不是单一物理过程的孤立结果。"


def feature_reason_en(label: str, direction: str, biome: str) -> str:
    # concise English counterpart for bilingual consistency
    if label == "SSRD":
        return "This suggests that radiation acts through its coupling with thermal and evaporative stress rather than as a pure light benefit." if direction == "negative" else "This indicates that energy limitation remains important, so higher radiation primarily accelerates canopy recovery."
    if label == "EVA":
        return "Here, stronger evaporation is better interpreted as sustained water loss under incomplete hydrological recovery." if direction == "positive" else "Here, stronger evaporation more likely reflects the reactivation of land-atmosphere exchange during recovery."
    if label == "TMP":
        return "Higher temperature mainly amplifies evaporative demand and respiratory costs in this biome." if direction == "positive" else "A warmer background appears to relieve thermal or cold-wet constraints in this biome."
    if label == "VPD":
        return "Higher VPD directly strengthens atmospheric water stress and suppresses recovery." if direction == "positive" else "A moderate rise in VPD may co-occur with clearer energy conditions rather than purely harmful atmospheric dryness."
    if label == "PRE":
        return "Large precipitation likely co-occurs with low-radiation or poorly ventilated conditions rather than acting as a purely beneficial water subsidy." if direction == "positive" else "Higher precipitation replenishes root-zone water and buffers subsequent atmospheric stress."
    if label == "SMrz":
        return "This implies that slow recovery is not controlled by immediate water shortage alone." if direction == "positive" else "Root-zone soil moisture acts as a direct hydrological buffer that shortens recovery."
    if label == "STRD":
        return "Higher longwave radiation likely marks a warm, humid background that raises respiration-related costs." if direction == "positive" else "Higher longwave radiation likely reflects a more thermally stable background that supports continuous recovery."
    return "The directional effect should be interpreted as a net outcome of energy, water, and ecosystem structural constraints."


def dominant_groups(rows: pd.DataFrame) -> tuple[int, int, int]:
    energy = {"SSRD", "STRD", "TMP", "VPD", "WIND"}
    water = {"PRE", "EVA", "SMrz"}
    event = {"Duration", "Intensity", "Onset"}
    labels = set(rows["label"])
    return len(labels & energy), len(labels & water), len(labels & event)


def top_share_text(rows: pd.DataFrame) -> str:
    top1 = float(rows.iloc[0]["importance"])
    top5sum = float(rows["importance"].sum())
    top2sum = float(rows.iloc[0]["importance"] + rows.iloc[1]["importance"])
    return f"首位特征单独贡献了前五总解释量的 {top1 / top5sum * 100:.1f}%，前两位合计贡献 {top2sum / top5sum * 100:.1f}%"


def ensure_style(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    for name in ["Title", "Heading 1", "Heading 2"]:
        st = styles[name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Title"].font.size = Pt(16)
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 2"].font.size = Pt(12)


def add_title(doc: Document, title: str) -> None:
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.style = "Normal"
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.35
    p.add_run(text)


def add_image(doc: Document, path: Path, caption: str, width_cm: float = 15.0) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.add_run(caption)


def literature_style_note() -> str:
    return (
        "本轮润色参考了本地 SHAP_SEM 文献目录中论文在结果与讨论部分的共同写法，即先明确主导因子与方向，"
        "再识别关键阈值或转折区间，随后从水分供给、蒸散需求、能量约束和生态系统功能差异四条链条解释 biome 间差异，"
        "避免停留在图形描述层面。当前图件同时对三处显示逻辑做了统一：一是移除了 LAI，二是 beeswarm 中 EVA 颜色映射使用其绝对值，三是 dependence plot 中凡是 EVA 作为横轴或颜色条时均写作 |EVA|，以避免负号造成视觉歧义。"
    )


def literature_style_note_en() -> str:
    return (
        "The revised writing follows the common Results-plus-Discussion style observed in the local SHAP/SEM literature set: "
        "the narrative first identifies the dominant drivers and their effect directions, then highlights critical thresholds or turning points, "
        "and finally interprets biome contrasts through water supply, evaporative demand, energy limitation, and ecosystem functional differences. "
        "In the current figures, LAI is removed, EVA is color-mapped by its absolute magnitude in beeswarm plots, and dependence plots display EVA as |EVA| whenever it appears on the x-axis or in the color bar."
    )


def contribution_overview_cn(stats: pd.DataFrame, metric: str) -> str:
    s = stats[stats["metric"] == metric]
    counts = s[s["rank"] == 1].groupby("label").size().sort_values(ascending=False)
    order = "，".join([f"{k} 在 {v} 个 biome 中位列首位" for k, v in counts.items()])
    biome_bits = []
    for biome in BIOMES:
        rows = topn(stats, metric, biome, 5)
        e, w, ev = dominant_groups(rows)
        biome_bits.append(
            f"{biome} 的前五特征中能量相关变量 {e} 个、水分相关变量 {w} 个、事件属性 {ev} 个，{top_share_text(rows)}"
        )
    max_biome = (
        s.groupby("biome")["importance"].apply(lambda x: x.sort_values(ascending=False).head(5).mean()).sort_values(ascending=False)
    )
    return (
        f"从整体贡献格局看，{METRIC_CN[metric]} 的主导因子并非随机分布，而是围绕能量输入、蒸散耗水和根区储水三类过程展开。"
        f"{order}。这种排序格局说明，整个骤旱期间的恢复时间主要由背景水热状态决定，而不是单一事件强度所能完全解释。"
        f"进一步看，{'; '.join(biome_bits)}。"
        f"其中前五特征平均贡献最高的 biome 为 {max_biome.index[0]}，说明该 biome 的恢复时间对关键背景变量最为敏感；"
        f"而平均贡献最低的 biome 为 {max_biome.index[-1]}，说明其恢复更可能受到多因子共同分摊控制，而不是被单一变量强烈主导。"
    )


def contribution_overview_en(stats: pd.DataFrame, metric: str) -> str:
    s = stats[stats["metric"] == metric]
    counts = s[s["rank"] == 1].groupby("label").size().sort_values(ascending=False)
    order = ", ".join([f"{k} ranked first in {v} biomes" for k, v in counts.items()])
    biome_bits = []
    for biome in BIOMES:
        rows = topn(stats, metric, biome, 5)
        e, w, ev = dominant_groups(rows)
        biome_bits.append(
            f"{biome} includes {e} energy-related, {w} water-related, and {ev} event-related variables in its top five; {top_share_text(rows)}"
        )
    return (
        f"At the overall level, the contribution structure of {METRIC_EN[metric]} is organized around three coupled processes: "
        f"energy supply, evaporative water loss, and root-zone water storage. {order}. "
        f"This pattern indicates that recovery time during the drought period is primarily governed by background hydroclimatic states rather than by event severity alone. "
        f"More specifically, {'; '.join(biome_bits)}."
    )


def biome_para_cn(stats: pd.DataFrame, metric: str, biome: str) -> str:
    rows = topn(stats, metric, biome, 5)
    top_labels = "、".join(rows["label"].tolist())
    lead = rows.iloc[0]
    second = rows.iloc[1]
    fifth = rows.iloc[4]
    e, w, ev = dominant_groups(rows)
    structure = (
        "能量约束占主导"
        if e > w and e > ev
        else "水分补给与耗水平衡占主导"
        if w > e and w > ev
        else "事件属性与背景状态共同控制"
    )
    return (
        f"{biome} 中，前五位特征依次为 {top_labels}。其中，{lead['label']} 的平均绝对 SHAP 值为 {lead['importance']:.2f}，"
        f"明显高于第 {int(fifth['rank'])} 位特征的 {fifth['importance']:.2f}，表明恢复时间的解释权重高度集中在少数关键背景变量上。"
        f"从组成上看，前五特征中能量相关变量有 {e} 个，水分相关变量有 {w} 个，事件属性变量有 {ev} 个，说明该 biome 的恢复过程整体上呈现出“{structure}”的格局。"
        f"首位特征 {lead['label']} 的方向并不能机械理解为“变量越大恢复越快”或“越慢”，更准确地说，它记录的是该变量所代表的背景状态在当前 biome 中对应哪一种恢复环境。"
        f"就本结果而言，{lead['label']} 升高更倾向于使恢复时间{'延长' if lead['direction']=='positive' else '缩短'}，{feature_reason_cn(lead['label'], lead['direction'], biome)}"
        f"排名第 2 的 {second['label']} 则进一步限定了这种解释：它与 {lead['label']} 联合出现时，决定了恢复拖尾究竟是由持续的能量-耗水耦合放大，还是由土壤储水不足与事件记忆累积所维持。"
    )


def biome_para_en(stats: pd.DataFrame, metric: str, biome: str) -> str:
    rows = topn(stats, metric, biome, 5)
    top_labels = ", ".join(rows["label"].tolist())
    lead = rows.iloc[0]
    second = rows.iloc[1]
    fifth = rows.iloc[4]
    e, w, ev = dominant_groups(rows)
    return (
        f"In {biome}, the top five features are {top_labels}. The leading driver, {lead['label']}, has a mean absolute SHAP contribution of {lead['importance']:.2f}, "
        f"substantially exceeding the {fifth['importance']:.2f} of the fifth-ranked feature, which indicates that explanatory power is concentrated in a limited set of background controls. "
        f"The top-five structure contains {e} energy-related, {w} water-related, and {ev} event-related variables, implying that recovery is shaped by a structured background regime rather than by isolated predictors. "
        f"In this biome, higher {lead['label']} tends to {'lengthen' if lead['direction']=='positive' else 'shorten'} recovery time. {feature_reason_en(lead['label'], lead['direction'], biome)} "
        f"The second-ranked factor, {second['label']}, further constrains this interpretation by indicating whether recovery tails are maintained by persistent energy-water coupling or by insufficient root-zone storage and event memory."
    )


def dep_para_cn(row: pd.Series) -> str:
    threshold = ""
    if pd.notna(row["threshold"]):
        threshold = f"识别到的转折阈值约为 {row['threshold']:.3f} {row['unit']}。"
    return (
        f"{row['label']} 的 dependence 曲线表现为 {row['trend_type']}，Spearman 相关系数为 {row['spearman_rho']:.2f}。"
        f"从低值段到高值段，SHAP 中位数由 {row['low_shap_median']:.2f} 变化为 {row['high_shap_median']:.2f}，"
        f"说明该变量并不是简单线性响应，而是在特定区间内明显改变了恢复时间的边际效应。"
        f"{threshold}{row['physics_hint']}"
    )


def dep_para_en(row: pd.Series) -> str:
    threshold = ""
    if pd.notna(row["threshold"]):
        threshold = f"The estimated turning threshold is approximately {row['threshold']:.3f} {row['unit']}. "
    return (
        f"The dependence curve of {row['label']} is classified as {row['trend_type']}, with a Spearman correlation of {row['spearman_rho']:.2f}. "
        f"From the low to the high value range, the median SHAP effect shifts from {row['low_shap_median']:.2f} to {row['high_shap_median']:.2f}, "
        f"indicating that the response is not purely linear but involves a substantial reorganization of the marginal effect on recovery time within specific value ranges. "
        f"{threshold}{row['physics_hint']}"
    )


def compare_para_cn(g: pd.Series, r: pd.Series, biome: str, feature: str) -> str:
    return (
        f"在 {biome} 中，{feature} 同时进入 GPP 与 RECO 的前列，但其解释重心并不完全相同。"
        f"GPP 的平均绝对 SHAP 值为 {g['importance']:.2f}，RECO 为 {r['importance']:.2f}；"
        f"对应的高值段效应分别指向恢复时间{'延长' if g['direction']=='positive' else '缩短'}和{'延长' if r['direction']=='positive' else '缩短'}。"
        f"这表明同一环境背景变量可以通过不同生理通道影响碳同化恢复与生态系统呼吸恢复：前者更直接体现冠层光合受限，后者则更容易受底物供给、温度敏感性及呼吸维持成本共同调制。"
    )


def compare_para_en(g: pd.Series, r: pd.Series, biome: str, feature: str) -> str:
    return (
        f"In {biome}, {feature} ranks among the leading predictors for both GPP and RECO, but the explanatory emphasis is not identical. "
        f"The mean absolute SHAP contribution is {g['importance']:.2f} for GPP and {r['importance']:.2f} for RECO, while the high-value regime tends to {'lengthen' if g['direction']=='positive' else 'shorten'} GPP recovery but {'lengthen' if r['direction']=='positive' else 'shorten'} RECO recovery. "
        f"This contrast indicates that the same environmental background variable propagates through different physiological pathways: canopy photosynthetic limitation dominates the GPP response, whereas substrate availability, temperature sensitivity, and maintenance respiration jointly modulate the RECO response."
    )


def write_cn_docs(stats: pd.DataFrame) -> None:
    doc1 = Document(); ensure_style(doc1)
    add_title(doc1, "01 不同 Biome 的特征贡献分析（润色版）")
    add_para(doc1, literature_style_note())
    add_para(doc1, "本文件将蜂巢图结果改写为更接近论文结果与讨论的表达方式，重点强调主导因子的排序、方向及 biome 间差异，而非停留在散点分布的表面描述。")
    add_image(doc1, PLOTS_ROOT / "beeswarm_comparison_5biomes_gpp_vs_reco.png", "图 1. 五个 biome 中 GPP 与 RECO 的 SHAP beeswarm 对比图。", 16.2)
    for metric in METRICS:
        doc1.add_heading(f"{METRIC_CN[metric]} 的整体贡献格局", level=1)
        add_para(doc1, contribution_overview_cn(stats, metric))
        for biome in BIOMES:
            add_image(doc1, PLOTS_ROOT / metric / biome / f"{metric}_{biome}_beeswarm_redraw.png", f"{metric} - {biome} beeswarm 图", 14.8)
            add_para(doc1, biome_para_cn(stats, metric, biome))
    doc1.save(WRITE_ROOT / "01_biome_feature_contribution_analysis_cn_polished.docx")

    doc2 = Document(); ensure_style(doc2)
    add_title(doc2, "02 Dependence Plot 与阈值机制分析（润色版）")
    add_para(doc2, literature_style_note())
    add_para(doc2, "本文件将 dependence plot 的解读重点放在非线性响应、阈值位置及其生态水文含义上。对于每个 biome，仅展示最具代表性的特征，以突出不同生态系统在水分补给、蒸散需求与辐射负荷之间的差异。")
    focus = ["PRE", "SSRD", "EVA", "TMP", "VPD", "SMrz", "STRD"]
    for metric in METRICS:
        doc2.add_heading(f"{METRIC_CN[metric]} 的关键非线性响应", level=1)
        for biome in BIOMES:
            doc2.add_heading(biome, level=2)
            sub = topn(stats, metric, biome, 5)
            labels = [x for x in focus if x in sub["label"].tolist()][:3]
            if not labels:
                labels = sub["label"].tolist()[:3]
            for label in labels:
                row = stats[(stats["metric"] == metric) & (stats["biome"] == biome) & (stats["label"] == label)].iloc[0]
                dep_label = f"|{label}|" if label == "EVA" else label
                add_image(doc2, PLOTS_ROOT / metric / biome / "dependence_top5" / f"{label}_colored_by_EVA.png", f"{metric} - {biome} - {dep_label}（|EVA| 着色）", 14.6)
                add_para(doc2, dep_para_cn(row))
                add_image(doc2, PLOTS_ROOT / metric / biome / "dependence_top5" / f"{label}_colored_by_SMrz.png", f"{metric} - {biome} - {label}（SMrz 着色）", 14.6)
                add_para(doc2, f"进一步使用 SMrz 着色，是为了检验该阈值响应是否受根区蓄水状态调制。如果相同 {dep_label} 值区间内的点云随着 SMrz 升高而系统性偏移，则说明该特征的边际效应需要放在土壤水储量背景下解释。对于 EVA，这里的横轴与颜色条都按 |EVA| 展示，因此数值越大表示蒸散发强度越强。")
    doc2.save(WRITE_ROOT / "02_dependence_threshold_analysis_cn_polished.docx")

    doc3 = Document(); ensure_style(doc3)
    add_title(doc3, "03 GPP 与 RECO 的差异比较分析（润色版）")
    add_para(doc3, literature_style_note())
    add_para(doc3, "本文件强调同一 biome 内 GPP 与 RECO 对共同背景变量的不同响应方式，重点比较贡献强度、方向以及 dependence 曲线的转折特征，以支撑对恢复机制差异的解释。")
    add_image(doc3, PLOTS_ROOT / "beeswarm_comparison_5biomes_gpp_vs_reco.png", "图 1. 五个 biome 中 GPP 与 RECO 的 beeswarm 对比图。", 16.2)
    for biome in BIOMES:
        doc3.add_heading(f"{biome} 中 GPP 与 RECO 的差异", level=1)
        gtop = topn(stats, "GPP", biome, 5)
        rtop = topn(stats, "RECO", biome, 5)
        common = [x for x in gtop["label"].tolist() if x in rtop["label"].tolist()][:3]
        add_para(doc3, f"{biome} 中，GPP 前五位为 {', '.join(gtop['label'])}，RECO 前五位为 {', '.join(rtop['label'])}。共同高频出现的变量包括 {', '.join(common)}，说明二者共享相同的背景水热约束，但并不共享完全相同的恢复生理过程。")
        for feat in common:
            g = stats[(stats["metric"] == "GPP") & (stats["biome"] == biome) & (stats["label"] == feat)].iloc[0]
            r = stats[(stats["metric"] == "RECO") & (stats["biome"] == biome) & (stats["label"] == feat)].iloc[0]
            add_para(doc3, compare_para_cn(g, r, biome, feat))
            add_image(doc3, PLOTS_ROOT / "GPP" / biome / "dependence_top5" / f"{feat}_colored_by_EVA.png", f"GPP - {biome} - {feat}", 14.6)
            add_image(doc3, PLOTS_ROOT / "RECO" / biome / "dependence_top5" / f"{feat}_colored_by_EVA.png", f"RECO - {biome} - {feat}", 14.6)
    doc3.save(WRITE_ROOT / "03_gpp_reco_difference_analysis_cn_polished.docx")


def write_en_docs(stats: pd.DataFrame) -> None:
    doc1 = Document(); ensure_style(doc1)
    add_title(doc1, "01 Feature Contribution Across Biomes (Polished Version)")
    add_para(doc1, literature_style_note_en())
    add_para(doc1, "This document rewrites the beeswarm interpretation in a more article-style form, emphasizing driver hierarchy, effect direction, and inter-biome contrasts rather than only describing point distributions.")
    add_image(doc1, PLOTS_ROOT / "beeswarm_comparison_5biomes_gpp_vs_reco.png", "Figure 1. SHAP beeswarm comparison of GPP and RECO across five biomes.", 16.2)
    for metric in METRICS:
        doc1.add_heading(f"Overall contribution structure of {METRIC_EN[metric]}", level=1)
        add_para(doc1, contribution_overview_en(stats, metric))
        for biome in BIOMES:
            add_image(doc1, PLOTS_ROOT / metric / biome / f"{metric}_{biome}_beeswarm_redraw.png", f"{metric} - {biome} beeswarm", 14.8)
            add_para(doc1, biome_para_en(stats, metric, biome))
    doc1.save(WRITE_ROOT / "01_biome_feature_contribution_analysis_en_polished.docx")

    doc2 = Document(); ensure_style(doc2)
    add_title(doc2, "02 Dependence Plots and Threshold Mechanisms (Polished Version)")
    add_para(doc2, literature_style_note_en())
    add_para(doc2, "This document focuses on nonlinear responses, threshold locations, and their ecohydrological meaning. Only the most representative features are shown for each biome in order to highlight differences in water supply, evaporative demand, and radiation load.")
    focus = ["PRE", "SSRD", "EVA", "TMP", "VPD", "SMrz", "STRD"]
    for metric in METRICS:
        doc2.add_heading(f"Key nonlinear responses of {METRIC_EN[metric]}", level=1)
        for biome in BIOMES:
            doc2.add_heading(biome, level=2)
            sub = topn(stats, metric, biome, 5)
            labels = [x for x in focus if x in sub["label"].tolist()][:3]
            if not labels:
                labels = sub["label"].tolist()[:3]
            for label in labels:
                row = stats[(stats["metric"] == metric) & (stats["biome"] == biome) & (stats["label"] == label)].iloc[0]
                dep_label = f"|{label}|" if label == "EVA" else label
                add_image(doc2, PLOTS_ROOT / metric / biome / "dependence_top5" / f"{label}_colored_by_EVA.png", f"{metric} - {biome} - {dep_label} (colored by |EVA|)", 14.6)
                add_para(doc2, dep_para_en(row))
                add_image(doc2, PLOTS_ROOT / metric / biome / "dependence_top5" / f"{label}_colored_by_SMrz.png", f"{metric} - {biome} - {label} (colored by SMrz)", 14.6)
                add_para(doc2, f"The additional SMrz coloring is used to test whether the threshold-like response of {dep_label} is conditioned by root-zone water storage. A systematic vertical separation of points within the same {dep_label} range would imply that the marginal effect of this driver must be interpreted under a soil-water background. For EVA, both the x-axis and the EVA color bar are displayed as |EVA|, so larger values consistently indicate stronger evaporation.")
    doc2.save(WRITE_ROOT / "02_dependence_threshold_analysis_en_polished.docx")

    doc3 = Document(); ensure_style(doc3)
    add_title(doc3, "03 Comparative Analysis of GPP and RECO (Polished Version)")
    add_para(doc3, literature_style_note_en())
    add_para(doc3, "This document emphasizes how GPP and RECO respond differently to shared environmental backgrounds within the same biome, with particular attention to contribution strength, effect direction, and turning behavior in dependence curves.")
    add_image(doc3, PLOTS_ROOT / "beeswarm_comparison_5biomes_gpp_vs_reco.png", "Figure 1. Beeswarm comparison of GPP and RECO across five biomes.", 16.2)
    for biome in BIOMES:
        doc3.add_heading(f"GPP versus RECO in {biome}", level=1)
        gtop = topn(stats, "GPP", biome, 5)
        rtop = topn(stats, "RECO", biome, 5)
        common = [x for x in gtop["label"].tolist() if x in rtop["label"].tolist()][:3]
        add_para(doc3, f"In {biome}, the top five features are {', '.join(gtop['label'])} for GPP and {', '.join(rtop['label'])} for RECO. The shared leading variables, {', '.join(common)}, indicate that both processes are constrained by similar hydroclimatic backgrounds, while their recovery pathways remain physiologically distinct.")
        for feat in common:
            g = stats[(stats["metric"] == "GPP") & (stats["biome"] == biome) & (stats["label"] == feat)].iloc[0]
            r = stats[(stats["metric"] == "RECO") & (stats["biome"] == biome) & (stats["label"] == feat)].iloc[0]
            add_para(doc3, compare_para_en(g, r, biome, feat))
            add_image(doc3, PLOTS_ROOT / "GPP" / biome / "dependence_top5" / f"{feat}_colored_by_EVA.png", f"GPP - {biome} - {feat}", 14.6)
            add_image(doc3, PLOTS_ROOT / "RECO" / biome / "dependence_top5" / f"{feat}_colored_by_EVA.png", f"RECO - {biome} - {feat}", 14.6)
    doc3.save(WRITE_ROOT / "03_gpp_reco_difference_analysis_en_polished.docx")


def main() -> None:
    stats = load_stats()
    WRITE_ROOT.mkdir(parents=True, exist_ok=True)
    (WRITE_ROOT / "literature_style_basis_cn.md").write_text(
        "\n".join([
            "# SHAP/SEM 文献写法提炼",
            "",
            "- 结果段先给出主导变量及其相对强弱，而不是先解释图形外观。",
            "- 对非线性变量，优先说明阈值、转折点或高低值区间效应差异。",
            "- 机制解释通常围绕水分供给、蒸散需求、辐射能量和生态系统结构差异展开。",
            "- 比较分析强调不同生态系统或不同过程对同一变量的敏感性差异。",
            "- 本次润色据此将原有文本改写为更接近论文 Results/Discussion 的风格。"
        ]),
        encoding="utf-8",
    )
    write_cn_docs(stats)
    write_en_docs(stats)


if __name__ == "__main__":
    main()
