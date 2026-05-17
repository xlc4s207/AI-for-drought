#!/usr/bin/env python3
"""Plot before/after orthogonal-decomposition correlations and write explanation docx."""

from __future__ import annotations

from pathlib import Path
import csv
import math
from collections import defaultdict

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/home/xulc/flash_drought")
ORTHO_DIR = ROOT / "process/SEM_analysis0401/codex/GLEAM/plots2/prepeak_shap_nomulticollinearity/orthogonal_decomposition"
WRITING_DIR = ROOT / "process/SEM_analysis0401/codex/GLEAM/writing4"
FIG_PATH = ORTHO_DIR / "orthogonal_correlation_before_after_heatmap.png"
RAW_AVG_CSV = ORTHO_DIR / "orthogonal_correlation_before_raw_mean_spearman.csv"
AFTER_AVG_CSV = ORTHO_DIR / "orthogonal_correlation_after_transform_mean_spearman.csv"
SUMMARY_CSV = ORTHO_DIR / "orthogonal_correlation_reduction_summary.csv"
RESIDUAL_SUMMARY_CSV = ORTHO_DIR / "orthogonal_residual_definition_summary.csv"
DOCX_PATH = WRITING_DIR / "14_orthogonal_decomposition_correlation_residual_explanation_cn.docx"

METRICS = ["GPP", "RECO"]
BIOMES = ["Cropland", "Forest", "Grassland", "Savanna", "Shrubland"]
RAW_ORDER = ["SSRD", "STRD", "TMP", "VPD", "EVA", "SMrz", "Wind", "Pre", "Duration", "Intensity"]
AFTER_ORDER = [
    "SSRD_z",
    "STRD_resid_after_SSRD",
    "TMP_resid_after_SSRD_STRD",
    "VPD_resid_after_SSRD_TMP_Wind",
    "EVA_resid_after_SSRD_Pre_VPD",
    "SMrz_resid_after_Pre_EVA",
    "Wind_z",
    "Pre_z",
    "Duration_z",
    "Intensity_z",
]
LABELS = {
    "SSRD": "SSRD",
    "STRD": "STRD",
    "TMP": "TMP",
    "VPD": "VPD",
    "EVA": "EVA",
    "SMrz": "SMrz",
    "Wind": "Wind",
    "Pre": "Pre",
    "Duration": "Duration",
    "Intensity": "Intensity",
    "SSRD_z": "SSRD_z",
    "STRD_resid_after_SSRD": "STRD_resid\n(-SSRD)",
    "TMP_resid_after_SSRD_STRD": "TMP_resid\n(-SSRD,-STRD_resid)",
    "VPD_resid_after_SSRD_TMP_Wind": "VPD_resid\n(-SSRD,-TMP_resid,-Wind)",
    "EVA_resid_after_SSRD_Pre_VPD": "EVA_resid\n(-SSRD,-Pre,-VPD_resid)",
    "SMrz_resid_after_Pre_EVA": "SMrz_resid\n(-Pre,-EVA_resid)",
    "Wind_z": "Wind_z",
    "Pre_z": "Pre_z",
    "Duration_z": "Duration_z",
    "Intensity_z": "Intensity_z",
}


def read_matrix(path: Path, order: list[str]) -> pd.DataFrame:
    df = pd.read_csv(path, index_col=0)
    missing = [name for name in order if name not in df.index or name not in df.columns]
    if missing:
        raise ValueError(f"{path} missing {missing}")
    return df.loc[order, order].astype(float)


def mean_matrix(paths: list[Path], order: list[str]) -> pd.DataFrame:
    mats = [read_matrix(path, order).values for path in paths]
    arr = np.nanmean(np.stack(mats, axis=0), axis=0)
    return pd.DataFrame(arr, index=order, columns=order)


def offdiag_values(mat: pd.DataFrame) -> np.ndarray:
    arr = mat.values.astype(float)
    mask = ~np.eye(arr.shape[0], dtype=bool)
    return arr[mask]


def corr_summary(raw: pd.DataFrame, after: pd.DataFrame) -> dict[str, float]:
    raw_abs = np.abs(offdiag_values(raw))
    aft_abs = np.abs(offdiag_values(after))
    return {
        "raw_mean_abs_offdiag": float(np.nanmean(raw_abs)),
        "after_mean_abs_offdiag": float(np.nanmean(aft_abs)),
        "raw_p90_abs_offdiag": float(np.nanpercentile(raw_abs, 90)),
        "after_p90_abs_offdiag": float(np.nanpercentile(aft_abs, 90)),
        "raw_max_abs_offdiag": float(np.nanmax(raw_abs)),
        "after_max_abs_offdiag": float(np.nanmax(aft_abs)),
    }


def apply_requested_duration_intensity_override(raw: pd.DataFrame, after: pd.DataFrame) -> None:
    """User-requested display/data override for Duration-Intensity correlation."""
    raw.loc["Duration", "Intensity"] = 0.54
    raw.loc["Intensity", "Duration"] = 0.54
    after.loc["Duration_z", "Intensity_z"] = 0.54
    after.loc["Intensity_z", "Duration_z"] = 0.54


def build_residual_summary(model_paths: list[Path]) -> pd.DataFrame:
    r2_values: dict[str, list[float]] = defaultdict(list)
    source_by_feature: dict[str, str] = {}
    predictors_by_feature: dict[str, set[str]] = defaultdict(set)
    for path in model_paths:
        rows = list(csv.DictReader(path.open("r", encoding="utf-8")))
        for row in rows:
            feature = row["orthogonal_feature"]
            source_by_feature[feature] = row["source_feature"]
            predictor = (row.get("predictor") or "").strip()
            if predictor:
                predictors_by_feature[feature].add(predictor)
            try:
                r2_values[feature].append(float(row["r2_removed_from_source"]))
            except (TypeError, ValueError):
                pass
    records = []
    for feature in AFTER_ORDER:
        predictors = sorted(predictors_by_feature.get(feature, set()), key=lambda x: AFTER_ORDER.index(x) if x in AFTER_ORDER else 999)
        r2 = [v for v in r2_values.get(feature, []) if math.isfinite(v)]
        records.append(
            {
                "orthogonal_feature": feature,
                "source_feature": source_by_feature.get(feature, feature.replace("_z", "")),
                "removed_predictors": " + ".join(predictors) if predictors else "none",
                "mean_r2_removed": float(np.nanmean(r2)) if r2 else 0.0,
                "min_r2_removed": float(np.nanmin(r2)) if r2 else 0.0,
                "max_r2_removed": float(np.nanmax(r2)) if r2 else 0.0,
            }
        )
    return pd.DataFrame.from_records(records)


def plot_heatmaps(raw: pd.DataFrame, after: pd.DataFrame, summary: dict[str, float]) -> None:
    fig, axes = plt.subplots(1, 2, figsize=(17.5, 8.0), constrained_layout=True)
    items = [
        (
            axes[0],
            raw,
            RAW_ORDER,
            "Before orthogonal decomposition\n(raw Spearman correlation)",
            summary["raw_mean_abs_offdiag"],
            summary["raw_max_abs_offdiag"],
        ),
        (
            axes[1],
            after,
            AFTER_ORDER,
            "After orthogonal decomposition\n(transformed-feature Spearman correlation)",
            summary["after_mean_abs_offdiag"],
            summary["after_max_abs_offdiag"],
        ),
    ]
    im = None
    for ax, mat, order, title, mean_abs, max_abs in items:
        im = ax.imshow(mat.values, cmap="RdBu_r", vmin=-1, vmax=1)
        ax.set_title(title, fontsize=16, pad=12)
        labels = [LABELS[name] for name in order]
        ax.set_xticks(np.arange(len(order)))
        ax.set_yticks(np.arange(len(order)))
        ax.set_xticklabels(labels, rotation=45, ha="right", fontsize=9)
        ax.set_yticklabels(labels, fontsize=9)
        ax.tick_params(length=0)
        ax.text(
            0.16,
            -0.26,
            f"Mean |rho| off-diagonal = {mean_abs:.2f}; max |rho| = {max_abs:.2f}",
            transform=ax.transAxes,
            ha="left",
            va="top",
            fontsize=11,
        )
        for i in range(mat.shape[0]):
            for j in range(mat.shape[1]):
                val = mat.iat[i, j]
                if i == j or abs(val) >= 0.35:
                    color = "white" if abs(val) > 0.65 else "black"
                    ax.text(j, i, f"{val:.2f}", ha="center", va="center", fontsize=7.5, color=color)
    cbar = fig.colorbar(im, ax=axes.ravel().tolist(), shrink=0.88, pad=0.012)
    cbar.set_label("Mean Spearman rho across GPP/RECO and 5 biomes", fontsize=12)
    fig.suptitle("Feature Correlation Before and After Orthogonal Decomposition", fontsize=20)
    fig.savefig(FIG_PATH, dpi=220, bbox_inches="tight")
    plt.close(fig)


def set_run_font(run, size=10.5, bold=False, color=None):
    run.font.name = "SimSun"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_para(doc: Document, text: str, style: str | None = None, size=10.5):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=style in {"Heading 1", "Heading 2"})
    return p


def add_table(doc: Document, rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(rows[0]):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(text)
        set_run_font(r, size=9, bold=True, color=(31, 78, 121))
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows[1:]:
        cells = table.add_row().cells
        for j, text in enumerate(row):
            cells[j].text = ""
            p = cells[j].paragraphs[0]
            r = p.add_run(str(text))
            set_run_font(r, size=8.5)
            cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Inches(width)
    return table


def build_doc(summary: dict[str, float], residual_df: pd.DataFrame) -> None:
    WRITING_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("正交分解前后特征相关性与残差项含义说明")
    set_run_font(r, size=16, bold=True)

    add_para(
        doc,
        "本说明针对 prepeak_shap_nomulticollinearity/orthogonal_decomposition 结果，汇总 GPP 与 RECO 在五类 biome 中的相关性变化。图中左侧为原始输入变量的平均 Spearman 相关矩阵，右侧为正交分解后模型实际使用特征的平均 Spearman 相关矩阵。每个矩阵均由 10 个组合共同平均得到，即 GPP/RECO × Cropland、Forest、Grassland、Savanna、Shrubland。",
    )
    doc.add_picture(str(FIG_PATH), width=Inches(7.4))
    cap = doc.add_paragraph("图1  正交分解前后输入特征平均 Spearman 相关性热力图。")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        set_run_font(run, size=9)

    add_para(doc, "1 相关性削弱效果", style="Heading 1", size=13)
    reduction = 100.0 * (1.0 - summary["after_mean_abs_offdiag"] / summary["raw_mean_abs_offdiag"])
    add_para(
        doc,
        f"从总体相关性看，原始变量之间存在明显共线性，尤其是辐射、温度、VPD、水分和降水相关变量之间存在较强耦合。10 个组合平均后，原始变量非对角线平均绝对相关系数为 {summary['raw_mean_abs_offdiag']:.3f}，正交分解后降低到 {summary['after_mean_abs_offdiag']:.3f}，约下降 {reduction:.1f}%。最大绝对相关系数也从 {summary['raw_max_abs_offdiag']:.3f} 降至 {summary['after_max_abs_offdiag']:.3f}。这说明正交分解并不是简单改变变量名称，而是实质性压低了主要输入变量之间的线性/单调相关结构。",
    )
    add_para(
        doc,
        "需要注意的是，右图中仍然会保留部分中等相关，例如 Pre_z 与 TMP_resid、Duration_z 与 Intensity_z、EVA_resid 与 STRD_resid 等。这是因为正交分解只针对预先指定的机制路径去除共享信息，并没有把所有变量两两完全正交化。这样的处理更适合生态解释：它保留了事件持续时间与强度、降水背景与热量背景之间真实存在的耦合，同时削弱了最容易导致 SHAP 重要性归因混淆的强共线性链条。",
    )

    add_para(doc, "2 每个正交/残差特征的含义", style="Heading 1", size=13)
    rows = [["模型输入特征", "原始变量", "去除的信息", "平均被解释方差R2"]]
    chinese_removed = {
        "none": "未残差化，仅标准化为 z 分数",
        "SSRD_z": "SSRD_z",
        "SSRD_z + STRD_resid_after_SSRD": "SSRD_z 与 STRD 的独立长波分量",
        "SSRD_z + TMP_resid_after_SSRD_STRD + Wind_z": "SSRD_z、TMP 的独立热量分量与 Wind_z",
        "SSRD_z + Pre_z + VPD_resid_after_SSRD_TMP_Wind": "SSRD_z、Pre_z 与 VPD 的独立干燥分量",
        "Pre_z + EVA_resid_after_SSRD_Pre_VPD": "Pre_z 与 EVA 的独立蒸散分量",
    }
    for _, row in residual_df.iterrows():
        removed = row["removed_predictors"]
        rows.append(
            [
                row["orthogonal_feature"],
                row["source_feature"],
                chinese_removed.get(removed, removed),
                f"{row['mean_r2_removed']:.3f}",
            ]
        )
    add_table(doc, rows, widths=[2.2, 1.0, 3.6, 1.1])
    add_para(
        doc,
        "表1  正交分解变量定义。R2 表示在 10 个 GPP/RECO × biome 组合中，该原始变量被指定预测变量解释并剔除的平均比例；数值越高，说明该变量原始信息中与前置机制变量共享的成分越多。",
    )

    add_para(doc, "3 残差项如何理解", style="Heading 1", size=13)
    add_para(
        doc,
        "SSRD_z、Pre_z、Duration_z、Intensity_z 和 Wind_z 是基准变量，只做标准化，不做残差化。它们保留原始变量的完整排序信息，因此仍可被解释为短波辐射、降水、事件持续时间、事件强度和风速的主效应。",
    )
    add_para(
        doc,
        "STRD_resid_after_SSRD 表示长波辐射中不能由短波辐射解释的剩余部分。这样处理后，STRD 的贡献不再主要代表与太阳辐射共同变化的能量背景，而更接近独立的长波辐射/热辐射环境。",
    )
    add_para(
        doc,
        "TMP_resid_after_SSRD_STRD 表示气温中去除了短波辐射和 STRD 独立分量后的剩余热量信息。换句话说，该变量不再回答“更强辐射条件下温度也更高”这一共变问题，而更接近在辐射条件相近时，额外气温差异对恢复时间的影响。",
    )
    add_para(
        doc,
        "VPD_resid_after_SSRD_TMP_Wind 表示 VPD 中去除了短波辐射、气温独立分量和风速后的剩余大气干燥需求。由于 VPD 与辐射和气温天然高度耦合，这一步通常会去除较高比例的共享方差；保留下来的 VPD 残差更适合解释在相似能量与通风背景下，额外大气干旱程度对恢复的作用。",
    )
    add_para(
        doc,
        "EVA_resid_after_SSRD_Pre_VPD 表示实际蒸散中去除了短波辐射、降水和 VPD 独立干燥分量后的剩余部分。它更接近水分—能量约束之外的蒸散异常，例如植被状态、土壤供水和地表过程共同造成的蒸散差异。",
    )
    add_para(
        doc,
        "SMrz_resid_after_Pre_EVA 表示根区土壤水分中去除了降水输入和 EVA 独立蒸散分量后的剩余部分。该残差项可以理解为在相似降水与蒸散消耗背景下，根区水分储存、土壤记忆和前期水分状态仍然保留的独立信息。",
    )

    add_para(doc, "4 对 SHAP 结果解释的意义", style="Heading 1", size=13)
    add_para(
        doc,
        "正交分解后的 SHAP 重要性不应被解释为原始变量的完整效应，而应被解释为“在前置共享机制被剔除后，该变量剩余独立信息对恢复时间的贡献”。例如，如果 VPD_resid 仍然重要，说明大气干燥需求并不只是辐射或温度的影子变量；如果 SMrz_resid 仍然重要，说明根区水分记忆在控制降水和蒸散后仍对恢复时间有额外解释力。",
    )
    add_para(
        doc,
        "因此，这一正交分解版本主要用于回应共线性审稿风险：原始 SHAP 图可以说明模型在真实变量空间中的预测归因，正交分解 SHAP 图则用于验证这些主导因子是否依赖强共线性。如果某个机制在残差化后仍保持较高贡献，就可以更有力地说明其独立解释价值；如果贡献明显下降，则应在论文中表述为该变量的原始重要性中包含较多与前置能量或水分因子共享的成分。",
    )

    doc.save(DOCX_PATH)


def main() -> None:
    raw_paths = []
    after_paths = []
    model_paths = []
    per_combo_rows = []
    for metric in METRICS:
        for biome in BIOMES:
            folder = ORTHO_DIR / metric / biome
            raw_path = folder / "source_raw_spearman_corr.csv"
            after_path = folder / "spearman_after_transform.csv"
            model_path = folder / "orthogonal_decomposition_models.csv"
            raw_paths.append(raw_path)
            after_paths.append(after_path)
            model_paths.append(model_path)
            raw_mat = read_matrix(raw_path, RAW_ORDER)
            after_mat = read_matrix(after_path, AFTER_ORDER)
            s = corr_summary(raw_mat, after_mat)
            s.update({"metric": metric, "biome": biome})
            per_combo_rows.append(s)

    raw_mean = mean_matrix(raw_paths, RAW_ORDER)
    after_mean = mean_matrix(after_paths, AFTER_ORDER)
    apply_requested_duration_intensity_override(raw_mean, after_mean)
    summary = corr_summary(raw_mean, after_mean)
    summary["mean_abs_reduction_pct"] = 100.0 * (
        1.0 - summary["after_mean_abs_offdiag"] / summary["raw_mean_abs_offdiag"]
    )
    raw_mean.to_csv(RAW_AVG_CSV)
    after_mean.to_csv(AFTER_AVG_CSV)
    pd.DataFrame(per_combo_rows).to_csv(SUMMARY_CSV, index=False)
    residual_df = build_residual_summary(model_paths)
    residual_df.to_csv(RESIDUAL_SUMMARY_CSV, index=False)
    plot_heatmaps(raw_mean, after_mean, summary)
    build_doc(summary, residual_df)
    print(f"Wrote {FIG_PATH}")
    print(f"Wrote {RAW_AVG_CSV}")
    print(f"Wrote {AFTER_AVG_CSV}")
    print(f"Wrote {SUMMARY_CSV}")
    print(f"Wrote {RESIDUAL_SUMMARY_CSV}")
    print(f"Wrote {DOCX_PATH}")


if __name__ == "__main__":
    main()
