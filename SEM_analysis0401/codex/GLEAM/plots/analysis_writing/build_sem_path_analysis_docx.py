#!/home/xulc/.local/share/mamba/envs/Flash_dra/bin/python
from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE = Path("/home/xulc/flash_drought/process/SEM_analysis0401/codex/GLEAM")
OUT_DIR = BASE / "plots/analysis_writing"
OUT_DOCX = OUT_DIR / "04_SEM_path_analysis_cn.docx"

SEM_FIG_DIR = BASE / "plots/SEM/SEM路径机制图"
GPP_DIR = BASE / "plots/SEM/gpp_code1_flash_smrz_unified11_20260426"
RECO_DIR = BASE / "plots/SEM/reco_code1_flash_smrz_sem_simplified5_20260426"

BIOMES = ["Cropland", "Forest", "Grassland", "Savanna", "Shrubland"]

SCENARIOS = {
    "GPP_prepeak": {
        "label": "GPP prepeak",
        "metric": "GPP",
        "phase": "prepeak",
        "r2_csv": GPP_DIR / "unified11_r2_summary.csv",
        "path_csv": GPP_DIR / "prepeak_path_effect_strengths.csv",
    },
    "GPP_recovery": {
        "label": "GPP recoverywin",
        "metric": "GPP",
        "phase": "recoverywin",
        "r2_csv": GPP_DIR / "unified11_r2_summary.csv",
        "path_csv": GPP_DIR / "recoverywin_path_effect_strengths.csv",
    },
    "RECO_prepeak": {
        "label": "RECO prepeak",
        "metric": "RECO",
        "phase": "prepeak",
        "r2_csv": RECO_DIR / "reco_prepeak_simplified5_r2_summary.csv",
        "path_csv": RECO_DIR / "reco_prepeak_simplified5_path_effect_strengths.csv",
    },
    "RECO_recovery": {
        "label": "RECO recoverywin",
        "metric": "RECO",
        "phase": "recoverywin",
        "r2_csv": RECO_DIR / "reco_recoverywin_simplified5_r2_summary.csv",
        "path_csv": RECO_DIR / "reco_recoverywin_simplified5_path_effect_strengths.csv",
    },
}


SCENARIO_OVERVIEW = {
    "GPP_prepeak": [
        "GPP 的 prepeak 路径图最明显的特征，是 biome 之间的分化程度非常高，并不存在一个可以机械套用到全部 biome 的统一主控链。Forest、Grassland 与 Cropland 都把大气干燥需求、温度或辐射放在直接控制的前列，而 Savanna 则明显转向以水分收支链主导，Shrubland 又重新回到辐射背景占优的结构。这说明 GPP 在峰值前阶段更像是在记录“不同 biome 分别以什么方式进入脆弱状态”，而不是在重复描述同一个恢复机制。",
        "从解释度看，这一组并不低，Forest 和 Grassland 的 holdout R² 分别达到 0.276 和 0.299，Savanna 和 Cropland 虽然较低，但路径强度非常集中，尤其 Savanna 的 P-ET、PRE 和 EVA 构成了极强的直接链。这意味着 GPP prepeak 虽然是“背景记忆模型”，但并不是松散的背景噪声，而是在不同 biome 中记录了性质不同的恢复起点。"
    ],
    "GPP_recovery": [
        "GPP 的 recovery 路径图和 prepeak 完全不是一个味道。五个 biome 的最强直接路径全部围绕恢复窗口内部的水分收支、温度和热湿状态展开，其中 P-ET 在五个 biome 中都排到首位，说明一旦进入恢复期，真正决定 GPP 恢复尾部的不是单个气象变量，而是补水与耗水的净结果。与前面 SHAP 文档里 PRE 普遍前移的现象放在一起看，可以把这里理解为：PRE 很重要，但在 SEM 里它进一步被压缩成一个更物理的“有效水分收支”抓手。",
        "这一组的解释度也是四套模型里最高的一档，Forest 达到 0.744，Savanna 达到 0.730，Grassland 和 Cropland 也都在 0.65 以上。也就是说，GPP recovery 不是一个低解释度下的粗糙故事，而是真正能把恢复时间的大部分差异组织成稳定路径网络的阶段。真正需要解释的重点，不是 PRE 在不在，而是为什么不同 biome 在同一个水分收支主轴之外，又分别被温度、露点、STRD 或蒸散重新分流。"
    ],
    "RECO_prepeak": [
        "RECO 的 prepeak 路径图比 GPP 明显整齐。五个 biome 中，SSRD 都是指向恢复时间的最强直接负路径，TMP 都是第二强正路径，SMrz 和 PRE 只在个别 biome 中轻微前移。这说明呼吸恢复在峰值前阶段的记忆内容并不复杂，核心就是“能量背景如何改变热状态，再把系统带入不同的恢复起点”。和 GPP prepeak 那种 biome 分化很大的结构相比，RECO prepeak 更像一个热力背景主导的统一框架。",
        "但整齐不等于解释度高。RECO prepeak 的 holdout R² 仍普遍偏低，Cropland 和 Savanna 只有 0.04–0.05，Grassland 和 Shrubland 稍高。这意味着它更适合用来识别“谁在主导”，而不是用来宣称已经解释了绝大部分恢复时间差异。因此，这一组的正确打开方式不是把系数写成确定律，而是把它当成热背景主导的证据链。"
    ],
    "RECO_recovery": [
        "RECO 的 recovery 路径图同样比 GPP 简洁，但主导链已经从 prepeak 的“SSRD 正负配温度”切换成“温度与 SSRD 的双主轴”。这里最容易写错的地方是机械宣称 SSRD 在所有 biome 中都最强。实际并不是这样：只有 Cropland 的最强直接路径是 SSRD，Forest、Grassland、Savanna 和 Shrubland 的最强直接路径全部变成了温度的负效应。这说明恢复窗口内部的 RECO 更像是在回答‘高温会不会持续拖延呼吸回基线’，而不是只回答‘辐射够不够’。",
        "这组模型的解释度仍然不高，但与 prepeak 相比没有显著改善，说明对 RECO 而言，恢复窗口内虽然主控路径更清楚，却仍有大量未被这套简化结构捕捉的方差。真正稳定的结论不是“模型很好”，而是阶段差异非常清楚：prepeak 以 SSRD 为核心负路径，recovery 则在多数 biome 中转为温度主导的负路径。"
    ],
}


BIOME_TEXT = {
    "GPP_prepeak": {
        "Cropland": "Cropland 的 GPP prepeak 路径最像“热干背景 + 耗水背景”共同锁定恢复尾部。最强直接路径是 VPD 的负效应，其次是 TMP 和 EVA 的正效应，这说明农田在峰值前并不是简单受补水多少控制，而是更先受大气干燥需求、热状态和耗水强度共同设定。值得注意的是，PRE 的直接路径几乎为零，说明它在这里更多通过土壤水和水分收支链间接起作用，而不是直接决定恢复时间。",
        "Forest": "Forest 的 GPP prepeak 更像一个能量背景主导型结构。TMP、SSRD 和 VPD 三条路径强度接近，说明林地恢复时间在峰值前首先记住的是热背景与辐射背景如何共同塑造干燥需求，而不是像 Cropland 那样先暴露出强耗水主导。EVA 仍然靠前，但没有压过温度和短波辐射，说明林地在事件前的耗水信号重要，却仍嵌在能量框架之内。",
        "Grassland": "Grassland 的 GPP prepeak 是这一组里最典型的多路径并列结构。SSRD 是最强负路径，TMP 是最强正路径，SMrz 和 STRD 又同时进入前三到前四位，说明草地并不是单纯由热干空气或单纯由土壤水记忆控制，而是白天辐射、热背景、根区水分和长波背景一起定义了恢复起点。相比 Forest，它更把 SMrz 推到前面；相比 Cropland，它又没有让 EVA 进入前三，体现了开放草地与农田完全不同的背景记忆类型。",
        "Savanna": "Savanna 的 GPP prepeak 是最不该被套话处理的一张图，因为它的主导直接路径不是 TMP、VPD 或 SSRD，而是 P-ET、PRE 和 EVA 这条水分收支链。P-ET 的绝对效应值最大，PRE 和 EVA 紧随其后，说明稀树草原在峰值前阶段首先记住的是净水分盈亏，而不是单个气象变量。这一点和 Forest、Grassland、Cropland 都不同，意味着 Savanna 的恢复记忆更像是‘事件前系统已经亏了多少水’，而不是‘事件前有多热、多亮、多干’。",
        "Shrubland": "Shrubland 的 GPP prepeak 重新回到辐射主导，但又不是 Forest 的复制品。最强直接路径是 STRD 和 SSRD，VPD 排第三，说明灌丛在峰值前更受整体辐射背景控制，其中长波热背景比温度本身更靠前。与 Grassland 相比，它没有把 SMrz 推到很前面；与 Cropland 相比，也没有让 EVA 成为主轴。这意味着灌丛恢复记忆更像是一种‘热辐射框架’遗产。"
    },
    "GPP_recovery": {
        "Cropland": "Cropland 的 GPP recovery 明确是水分收支主导，但第二层修饰项非常有特征。P-ET 居首，TMP、PRE 和 STRD 紧随其后，说明恢复期农田首先受有效水分盈亏控制，但同样的水分收支结果会被温度和长波背景进一步放大或抵消。与 prepeak 中的 VPD 首位相比，这里已经从‘大气干燥需求’转成‘净水分结果’，说明控制层级确实换了挡位。",
        "Forest": "Forest 的 GPP recovery 解释度最高，也是最适合拿来讲恢复窗口机制的一类。最强直接路径是露点温度、TMP 和 P-ET，随后是 STRD 与 PRE，说明林地恢复速度不仅受水分收支控制，还高度依赖热湿背景是否允许冠层重新进入有效交换状态。它不是简单的‘水够不够’，而是‘在什么样的热湿环境下，补水能不能真正转化为恢复’。",
        "Grassland": "Grassland 的 GPP recovery 里，P-ET、TMP 和 STRD 构成了前三强路径，PRE 排在第四。这个排序告诉我们，草地恢复期里 PRE 虽然重要，但它已经被折叠进更上层的水分收支抓手，真正直接拖长恢复尾部的是高温和长波背景。与 prepeak 阶段的 SSRD/TMP/SMrz 并列结构相比，Grassland 已经显著转入‘收支主导、热背景放大’的恢复过程模型。",
        "Savanna": "Savanna 的 GPP recovery 继续保留了非常强的水分收支主导：P-ET 居首，随后是露点温度和 PRE。与 prepeak 时那种直接由 PRE、EVA、P-ET 共管的结构相比，recovery 阶段已经把这些水分信息收敛成更明确的净效应抓手。这说明稀树草原不是在每个阶段都由同一组变量直接控制，而是在恢复窗口内把复杂的前置水分记忆收束成一个更直接的恢复过程链。",
        "Shrubland": "Shrubland 的 GPP recovery 仍由 P-ET 主导，但第二层路径比其他 biome 更偏向露点温度和 STRD，而不是 PRE 或 TMP。这意味着灌丛恢复期里同样的净水分结果，是否会继续拖长恢复尾部，很大程度取决于热湿背景和长波环境，而不是单独的补水多少。和 prepeak 时的辐射框架记忆相比，这里已经转成‘热湿修饰下的水分收支控制’。"
    },
    "RECO_prepeak": {
        "Cropland": "Cropland 的 RECO prepeak 结构非常清楚：SSRD 是最强负路径，TMP 是最强正路径，SMrz 和 PRE 只剩下很弱的直接作用。也就是说，农田呼吸恢复在峰值前阶段首先受能量背景和热状态控制，而不是直接受补水或根区水分控制。与 GPP prepeak 中 Cropland 由 VPD/TMP/EVA 三强并列相比，RECO 明显更简洁。",
        "Forest": "Forest 的 RECO prepeak 和 Cropland 很像，但温度次强路径更弱一些，SMrz 的负效应稍微更明显。这说明林地呼吸恢复同样是热力框架主导，只是相对农田，它还保留了少量土壤水状态信息。最重要的事实仍然没有变：SSRD 负路径排第一，说明峰值前林地呼吸恢复首先记住的是辐射背景，而不是复杂的植被结构变量。",
        "Grassland": "Grassland 的 RECO prepeak 是这一组里最标准的模板：SSRD 最强负路径，TMP 次强正路径，其余路径全部明显偏弱。这种整齐结构本身就是一个结果，说明草地呼吸恢复在峰值前并不依赖多条并列机制，而更像由热力背景统一驱动。与 GPP prepeak 的 Grassland 复杂结构相比，RECO 明显更集中。",
        "Savanna": "Savanna 的 RECO prepeak 也延续了这一热力主导框架，只是 TMP 的正效应比 Grassland 更强一些，SMrz 的负效应稍大。也就是说，稀树草原呼吸恢复在峰值前更明显地保留了热背景放大效应，但它依旧不是一个由 PRE 或 EVA 直接领跑的结构。这点和 GPP prepeak 的 Savanna 形成鲜明对照：后者以水分收支链为主，这里则是辐射-温度链更稳。",
        "Shrubland": "Shrubland 的 RECO prepeak 是这一组里唯一把 PRE 推到第二梯队前列的 biome，但最强路径仍然是 SSRD，随后才是 PRE 和 TMP 的弱正效应。这说明灌丛呼吸恢复虽然比其他 biome 更容易保留峰值前补水痕迹，但总体仍被辐射背景主导。不能因为 PRE 往前动了一点，就把它写成补水主导。"
    },
    "RECO_recovery": {
        "Cropland": "Cropland 的 RECO recovery 是五个 biome 里唯一由 SSRD 作为最强直接路径的情形，TMP 负效应排第二。这个结构意味着农田呼吸恢复在恢复窗口内部首先受即时短波供能控制，而高温随后再把恢复尾部往回拖。也就是说，它不是纯温度主导，也不是纯补水主导，而是一个辐射优先、温度跟进的恢复过程结构。",
        "Forest": "Forest 的 RECO recovery 与旧文档写法正好相反：最强路径不是 SSRD，而是 TMP 的负效应，SSRD 只排第二。这个差异必须写清楚，因为它说明林地呼吸恢复在恢复阶段内更直接受高温拖延，而短波辐射只是第二层修饰。若把它写成‘Forest recovery 由 SSRD 主导’，那就是把图和数据都写反了。",
        "Grassland": "Grassland 的 RECO recovery 同样由 TMP 负效应居首，SSRD 第二，SMrz 第三。说明草地呼吸恢复在恢复窗口内部主要受高温压制，而辐射供能只是在第二层定义方向。与 prepeak 阶段的 SSRD 绝对主导相比，这里已经是明显的阶段转换：从辐射背景记忆，切到恢复期高温拖尾。",
        "Savanna": "Savanna 的 RECO recovery 也不是 SSRD 第一，而是 TMP 负效应更强，SSRD 第二。它和 Forest、Grassland 一样，都表明恢复阶段的呼吸回基线更怕高温持续，而不是单纯依赖更强辐射。这一点和 prepeak 阶段‘SSRD 第一’的统一结构形成强烈反差，恰恰说明阶段效应很真实。",
        "Shrubland": "Shrubland 的 RECO recovery 是五个 biome 中温度负效应最强的一类，TMP 绝对值最大，SSRD 次之，SMrz 再次。说明灌丛呼吸恢复在恢复窗口内部最直接受高温压制，而补水和土壤水状态都只是后排因素。与 GPP recovery 里灌丛更受 P-ET 和露点温度分流不同，这里体现的是更简洁但更强烈的温度拖尾结构。"
    },
}


def read_csv(path: Path) -> list[dict[str, str]]:
    with open(path, "r", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def set_run_font(run, size_pt=11, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "SimSun")


def add_paragraph(doc: Document, text: str, *, bold=False, center=False, size=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    if not center:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, size_pt=size, bold=bold)
    return p


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    run = p.add_run(text)
    set_run_font(run, size_pt={1: 16, 2: 14, 3: 12}.get(level, 11), bold=True)
    return p


def add_figure(doc: Document, image: Path, caption: str, width_cm: float = 15.5):
    doc.add_picture(str(image), width=Cm(width_cm))
    add_paragraph(doc, caption, center=True, size=10)


def get_r2_map(key: str) -> dict[str, float]:
    rows = read_csv(SCENARIOS[key]["r2_csv"])
    if key.startswith("GPP"):
        phase = SCENARIOS[key]["phase"]
        rows = [r for r in rows if r["scope"] == phase]
        return {r["biome"]: float(r["holdout_r2"]) for r in rows if r["biome"] in BIOMES}
    return {r["biome"]: float(r["holdout_r2"]) for r in rows if r["biome"] in BIOMES}


def strongest_paths(key: str) -> dict[str, list[dict[str, str]]]:
    rows = read_csv(SCENARIOS[key]["path_csv"])
    rows = [r for r in rows if r["to"] == "t_recover_to_baseline_abs_peak" and r["biome"] in BIOMES]
    out: dict[str, list[dict[str, str]]] = {}
    for biome in BIOMES:
        sub = [r for r in rows if r["biome"] == biome]
        sub.sort(key=lambda r: abs(float(r["estimate"])), reverse=True)
        out[biome] = sub
    return out


def fmt_name(name: str) -> str:
    return (
        name.replace("prepeak_", "")
        .replace("recoverywin_", "")
        .replace("_mean", "")
        .replace("_2m", "")
        .replace("total_precipitation", "PRE")
        .replace("total_evaporation", "EVA")
        .replace("temperature", "TMP")
        .replace("wind_speed", "Wind")
        .replace("strd", "STRD")
        .replace("ssrd", "SSRD")
        .replace("lai_total", "LAI")
        .replace("p_minus_et", "P-ET")
        .replace("dewpoint_temperature", "Dewpoint")
        .replace("SMrz", "SMrz")
        .replace("VPD", "VPD")
    )


def fig_path(metric: str, biome: str, phase: str) -> Path:
    biome_token = "grass" if metric == "RECO" and biome == "Grassland" else biome.lower()
    phase_token = "recovery" if phase == "recoverywin" else "prepeak"
    return SEM_FIG_DIR / f"{metric}_{biome_token}_{phase_token}.png"


def build_document() -> Document:
    r2_maps = {k: get_r2_map(k) for k in SCENARIOS}
    top_paths = {k: strongest_paths(k) for k in SCENARIOS}

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    add_paragraph(doc, "全球 GPP 与 RECO 恢复时间的 SEM 路径机制分析", bold=True, center=True, size=18)
    add_paragraph(doc, "基于 0401 口径下既有 SEM 路径机制图的结构化解释", center=True, size=12)

    add_heading(doc, "1. 资料基础与本次重写原则", 1)
    add_paragraph(
        doc,
        "本次文档重写只使用已经整理完成的 SEM 路径机制图，不再额外重画任何新图。具体图源全部来自 plots/SEM/SEM路径机制图 目录下的 20 张 biome 级路径图；定量解释则只使用与这套图对应的两组底表：GPP 使用 unified11 版本的 path effect strengths 与 R² 汇总，RECO 使用 simplified5 版本的 path effect strengths 与 R² 汇总。这样处理的目的，是把“展示图”“路径系数”“文字解释”三者口径完全拉齐，避免出现图和文说的不是同一版结果。",
    )
    add_paragraph(
        doc,
        "因此，本文件不再沿用旧版中那种先抽象概括、再用少量例子补一句的写法，而是按四个场景分别比较五个 biome 的主导路径差异。重点不是把所有路径都念一遍，而是识别每个 biome 指向恢复时间的最强直接控制项、第二层修饰项以及它与其他 biome 的不同之处。只有这样，SEM 路径图才真正承担“机制分流”的角色，而不是沦为对 SHAP 结果的图像化复述。"
    )

    add_heading(doc, "2. 模型解释度概览", 1)
    add_paragraph(
        doc,
        "四套 SEM 的解释度差异本身已经说明了阶段性机制的不同。GPP 的 unified11 结果在 recovery 阶段解释度显著提高，Forest、Savanna 和 Grassland 的 holdout R² 均达到 0.69 以上，说明恢复窗口内部的水分收支与热湿背景链可以解释大量恢复时间差异。相反，RECO 的 simplified5 结果无论在 prepeak 还是 recovery 阶段都处于较低水平，表示它更适合识别主导方向，而不适合被写成对全部方差的充分解释。",
    )
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Biome", "GPP prepeak", "GPP recovery", "RECO prepeak", "RECO recovery"]):
        cell.text = text
    for biome in BIOMES:
        row = table.add_row().cells
        row[0].text = biome
        row[1].text = f"{r2_maps['GPP_prepeak'][biome]:.3f}"
        row[2].text = f"{r2_maps['GPP_recovery'][biome]:.3f}"
        row[3].text = f"{r2_maps['RECO_prepeak'][biome]:.3f}"
        row[4].text = f"{r2_maps['RECO_recovery'][biome]:.3f}"
    add_paragraph(doc, "表 1. 四套 SEM 在五个共同 biome 中的 holdout R²。", center=True, size=10)

    section_fig = 1
    for key in ["GPP_prepeak", "GPP_recovery", "RECO_prepeak", "RECO_recovery"]:
        if key != "GPP_prepeak":
            doc.add_section(WD_SECTION_START.NEW_PAGE)
        label = SCENARIOS[key]["label"]
        metric = SCENARIOS[key]["metric"]
        phase = SCENARIOS[key]["phase"]
        phase_cn = "峰值前背景" if phase == "prepeak" else "恢复期过程"
        add_heading(doc, f"3.{['GPP_prepeak','GPP_recovery','RECO_prepeak','RECO_recovery'].index(key)+1} {label} 的路径机制差异", 1)
        for para in SCENARIO_OVERVIEW[key]:
            add_paragraph(doc, para)

        for biome in BIOMES:
            add_heading(doc, biome, 2)
            r2 = r2_maps[key][biome]
            ranked = top_paths[key][biome]
            top2 = "；".join(
                f"{fmt_name(r['from'])}={float(r['estimate']):.3f}" for r in ranked[:2]
            )
            add_paragraph(
                doc,
                f"{biome} 的 {metric} {phase_cn} 路径图中，指向恢复时间的前两条最强直接路径分别为 {top2}，holdout R²={r2:.3f}。"
                + BIOME_TEXT[key][biome]
            )
            image = fig_path(metric, biome, phase)
            add_figure(
                doc,
                image,
                f"图 {section_fig}. {metric} 在 {biome} biome 的 {phase_cn} SEM 路径机制图。",
            )
            section_fig += 1

    add_heading(doc, "4. 综合问题诊断与结论", 1)
    add_paragraph(
        doc,
        "把四套路径图放在一起，最稳的结论有两层。第一，GPP 与 RECO 的阶段切换方式并不一样。GPP 在 prepeak 阶段高度分化，不同 biome 会分别被辐射、热背景、VPD 或水分收支记忆主导；进入 recovery 阶段后，则普遍收束到以 P-ET 为主轴的恢复过程控制。RECO 则相反，它在 prepeak 阶段结构更整齐，几乎统一由 SSRD 和 TMP 控制；到 recovery 阶段，多数 biome 又切换到以温度负效应主导、SSRD 次级修饰的结构。",
    )
    add_paragraph(
        doc,
        "第二，旧版写法里最容易出错的地方，恰恰是把‘看起来差不多’误写成‘机制一致’。以 RECO recovery 为例，只有 Cropland 是 SSRD 最强直接路径，其他四个 biome 都是温度负效应更强；如果不逐个 biome 对照，文字就会把整套图写歪。同样，GPP prepeak 里 Savanna 由 P-ET、PRE 和 EVA 领跑，而不是简单的辐射链；GPP recovery 里五个 biome 全部由 P-ET 居首，也不能再写成单个 SSRD 或 TMP 主导。因此，这份重写后的文档最重要的价值，不是多讲了几句，而是把图、系数和结论重新拉到了同一条线上。"
    )
    return doc


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(str(OUT_DOCX))
    print(f"Wrote {OUT_DOCX}")


if __name__ == "__main__":
    main()
