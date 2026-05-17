#!/usr/bin/env python3
"""Create a Word document comparing OPGD Geodetector and SHAP results."""

from __future__ import annotations

from pathlib import Path
import sys

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
from common_validation import SHAP_ROOTS, SHORT_LABELS  # noqa: E402


OUT_NAME = "opgd_shap_comparison_cn.docx"


def set_east_asian_font(style, font_name: str = "SimSun") -> None:
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def short_label(feature: str) -> str:
    return SHORT_LABELS.get(feature, feature)


def format_opgd_top_features(rows: pd.DataFrame) -> str:
    return "; ".join(
        f"{row.label} (q={row.q:.3f}, {row.method}/{int(row.bins)})" for row in rows.itertuples()
    )


def format_shap_top_features(rows: pd.DataFrame) -> str:
    return "; ".join(
        f"{short_label(row.feature)} ({float(row.importance):.2f})" for row in rows.itertuples()
    )


def overlap_note(overlap: list[str]) -> str:
    if len(overlap) >= 2:
        return "高度一致：共同变量同时具有模型预测贡献和空间分层解释力。"
    if len(overlap) == 1:
        return "部分一致：共同变量较稳健，差异主要反映 SHAP 与 OPGD 的解释尺度不同。"
    return "差异较大：需要结合 ALE/ICE/PDP 和 risk detector 进一步判断方向与机制。"


def build_comparison_rows(work_dir: Path) -> list[list[str]]:
    opgd = pd.read_csv(work_dir / "opgd_factor_q.csv")
    candidate_features = set(opgd["feature"])
    rows: list[list[str]] = []
    for (metric, biome), group in opgd.groupby(["metric", "biome"], sort=False):
        shap = pd.read_csv(SHAP_ROOTS[metric] / biome / "feature_importance.csv")
        shap = shap[shap["feature"].isin(candidate_features)].copy()
        opgd_top = group.head(3).copy()
        shap_top = shap.head(3).copy()
        opgd_labels = list(opgd_top["label"])
        shap_labels = [short_label(feature) for feature in shap_top["feature"]]
        overlap = [label for label in opgd_labels if label in shap_labels]
        rows.append(
            [
                metric,
                biome,
                format_opgd_top_features(opgd_top),
                format_shap_top_features(shap_top),
                ", ".join(overlap) if overlap else "无 Top3 重叠",
                overlap_note(overlap),
            ]
        )
    return rows


def method_summary(work_dir: Path) -> str:
    opgd = pd.read_csv(work_dir / "opgd_factor_q.csv")
    counts = opgd["method"].value_counts()
    return "；".join(f"{method} {count} 次" for method, count in counts.items())


def row_scope_summary(work_dir: Path) -> str:
    opgd = pd.read_csv(work_dir / "opgd_factor_q.csv")
    rows = sorted(int(value) for value in opgd["rows"].dropna().unique())
    if len(rows) == 1:
        return f"当前文档基于每个 metric-biome 子集抽样 {rows[0]} 行的 OPGD 结果。"
    return "当前文档基于不同 metric-biome 子集的可用样本行数生成，行数范围为 " + f"{min(rows)}-{max(rows)}。"


def reliability_counts_text(reliability: pd.DataFrame) -> str:
    counts = reliability["reliability_grade"].value_counts()
    ordered = ["High", "Medium", "Low"]
    return "；".join(f"{grade} {int(counts.get(grade, 0))} 个" for grade in ordered)


def figure_paths(work_dir: Path) -> list[Path]:
    fig_dir = work_dir / "figures"
    return [
        fig_dir / "shap_opgd_reliability_matrix.png",
        fig_dir / "opgd_interaction_heatmaps.png",
    ]


def add_figure_if_exists(doc: Document, path: Path, caption: str) -> None:
    if not path.exists():
        add_paragraph(doc, f"未找到图件：{path}")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.8))
    caption_paragraph = doc.add_paragraph(caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption_paragraph.runs:
        run.font.size = Pt(9)


def add_reliability_table(doc: Document, work_dir: Path) -> None:
    path = work_dir / "reliability" / "reliability_score.csv"
    if not path.exists():
        add_paragraph(doc, "尚未发现 reliability/reliability_score.csv，因此本节未加入变量级可靠性表。")
        return
    reliability = pd.read_csv(path)
    high = reliability[reliability["reliability_grade"] == "High"].sort_values("q_opgd", ascending=False)
    headers = ["指标", "Biome", "变量", "q", "CV", "Top3频率", "SHAP Top3", "最小组占比"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = True
    for idx, header in enumerate(headers):
        run = table.rows[0].cells[idx].paragraphs[0].add_run(header)
        run.bold = True
    for row in high.itertuples(index=False):
        cells = table.add_row().cells
        values = [
            row.metric,
            row.biome,
            row.label,
            f"{row.q_opgd:.3f}",
            f"{row.bootstrap_q_cv:.3f}",
            f"{row.top3_frequency:.2f}",
            "是" if row.in_shap_top3 else "否",
            f"{row.opgd_min_group_share:.3f}",
        ]
        for idx, value in enumerate(values):
            cells[idx].text = value
    format_table(table)


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.first_line_indent = Pt(22)
    paragraph.paragraph_format.line_spacing = 1.25


def format_table(table) -> None:
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "SimSun"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
                    run.font.size = Pt(8)


def add_comparison_table(doc: Document, work_dir: Path) -> None:
    headers = ["指标", "Biome", "OPGD Top 3", "SHAP Top 3", "重叠变量", "解释"]
    rows = build_comparison_rows(work_dir)
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = True
    for idx, header in enumerate(headers):
        run = table.rows[0].cells[idx].paragraphs[0].add_run(header)
        run.bold = True
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cells[idx].text = value
    format_table(table)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style=None)
        paragraph.paragraph_format.left_indent = Pt(18)
        paragraph.paragraph_format.first_line_indent = Pt(-9)
        paragraph.add_run(f"- {item}")


def main() -> None:
    work_dir = Path(__file__).resolve().parent
    docx_path = work_dir / OUT_NAME

    doc = Document()
    styles = doc.styles
    set_east_asian_font(styles["Normal"])
    styles["Normal"].font.size = Pt(11)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        set_east_asian_font(styles[style_name])

    title = doc.add_heading("OPGD 地理探测器与 SHAP 结果对比分析", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. 分析目的", level=1)
    add_paragraph(
        doc,
        "本文件基于 05_OPGD_Geodetector 的试验结果，比较改进参数地理探测器（OPGD）与 SHAP 特征重要性在 GPP 和 RECO 恢复时间解释中的一致性与差异。OPGD 关注环境因子能否解释恢复时间的空间分层异质性，SHAP 关注机器学习模型预测中变量的平均贡献强度。二者回答的问题不同，因此更适合互为验证，而不是简单互相替代。",
    )

    doc.add_heading("2. 方法口径", level=1)
    add_paragraph(
        doc,
        "OPGD 使用与 04_Geodetector 相同的目标变量、metric、biome 和候选特征，但不再固定六分位数分层，而是在 quantile、equal_interval、geometric_interval 和 standard_deviation 四类方法以及 3-10 个分层数之间搜索最优 q 值。" + row_scope_summary(work_dir),
    )
    add_paragraph(doc, f"本次 OPGD 最优分层方法分布为：{method_summary(work_dir)}。")

    doc.add_heading("3. OPGD 与 SHAP Top3 对比", level=1)
    add_comparison_table(doc, work_dir)

    doc.add_heading("4. 图形化呈现", level=1)
    main_fig, interaction_fig = figure_paths(work_dir)
    add_paragraph(
        doc,
        "图形展示采用同一行列结构对齐 SHAP、OPGD 和可靠性等级。左图表示 SHAP 平均绝对贡献，中图表示 OPGD q 值，右图表示可靠性等级并叠加 SHAP Top3 与 OPGD Top3 的重叠标记。交互热图用于展示 SHAP 重要变量之间的 OPGD 交互解释力。",
    )
    add_figure_if_exists(
        doc,
        main_fig,
        "Figure 1. SHAP prediction contribution, OPGD q statistic, and reliability/Top3 overlap matrix.",
    )
    add_figure_if_exists(
        doc,
        interaction_fig,
        "Figure 2. OPGD interaction detector heatmaps for SHAP top-feature pairs across metric-biome subsets.",
    )

    doc.add_heading("5. 可靠性与准确度评估", level=1)
    reliability_path = work_dir / "reliability" / "reliability_score.csv"
    if reliability_path.exists():
        reliability = pd.read_csv(reliability_path)
        add_paragraph(
            doc,
            "地理探测器不是预测模型，因此不适合用分类准确率或回归误差来评价。本文将准确度表述为结果可靠性，主要从四个方面判断：q 值解释力、bootstrap q 稳定性、最优分层样本量是否充足，以及是否与 SHAP Top3 变量一致。",
        )
        add_paragraph(doc, f"综合可靠性等级分布为：{reliability_counts_text(reliability)}。")
        add_paragraph(
            doc,
            "其中 High 等级表示该变量 q 值较高或中等、bootstrap 变异系数较低、Top3 出现频率稳定，并且没有依赖过小样本分层；Medium 表示该变量有一定解释力但存在 SHAP 不一致、q 值偏低或分层敏感性问题；Low 表示其空间分层解释力或稳健性不足，论文中不宜作为主导机制变量。",
        )
        add_reliability_table(doc, work_dir)
    else:
        add_paragraph(doc, "尚未运行 reliability 检查，因此本文档只包含 OPGD 与 SHAP 的基本对比。")

    doc.add_heading("6. 综合比较结论", level=1)
    add_bullets(
        doc,
        [
            "总体上，OPGD 与 SHAP 的共同主线集中在热量、辐射和蒸散相关变量，即 TMP、STRD、SSRD 和 |EVA|。这说明恢复时间既受模型预测层面的能量背景影响，也具有可检测的空间分层异质性。",
            "SHAP 中 SSRD 经常排在前列，尤其在 Grassland、Savanna、Shrubland 和 RECO 中较突出；但 OPGD 在部分 biome 中更倾向于把 STRD、TMP 或 |EVA| 识别为空间分层主导因子。这说明 SSRD 对模型预测有强贡献，但空间格局未必总是由单独短波辐射分层主导。",
            "PRE 在 SHAP 中可进入 RECO 的 Grassland、Savanna 和 Shrubland Top3，但在 OPGD Top3 中只稳定出现在 GPP Cropland 或部分水分敏感系统之外。这提示降水可能更像局部预测贡献或事件状态变量，而不是所有 biome 中最强的空间分层变量。",
            "RECO 的 OPGD q 值整体更高，特别是 Savanna 和 Shrubland，说明 RECO 恢复时间的空间分异比 GPP 更容易被热量、辐射和大气干旱条件分层解释。",
            "OPGD 和 SHAP 高度一致的组合，如 GPP Forest 的 |EVA|/STRD、GPP Grassland 的 TMP/SSRD、RECO Forest 的 |EVA|/TMP、RECO Shrubland 的 SSRD/TMP，适合作为后续论文机制解释中的稳健变量。",
        ],
    )

    doc.add_heading("7. 分 biome 解释", level=1)
    add_paragraph(
        doc,
        "Cropland 中，GPP 和 RECO 都只有 STRD 与 SHAP Top3 重叠，说明农田恢复时间的空间分层更偏向长波辐射和温度背景，而 SHAP 还会把 |EVA|、SSRD 等过程变量视为重要预测因子。该类结果适合解释为管理背景和能量-水分过程共同作用，而不是单一变量控制。",
    )
    add_paragraph(
        doc,
        "Forest 中，GPP 的 |EVA| 与 STRD 同时进入 OPGD 和 SHAP Top3，RECO 的 |EVA| 与 TMP 同时重叠，说明森林恢复时间的预测贡献和空间分层都与蒸散需求、热量条件和辐射环境密切相关。相比之下，SSRD 在 SHAP 中更强，但在 OPGD 中不一定是最高 q 因子，表明其预测贡献可能通过与其他能量变量共同表达。",
    )
    add_paragraph(
        doc,
        "Grassland 中，GPP 和 RECO 均显示 TMP 与 SSRD 的一致性，OPGD 还将 STRD 识别为重要空间分层因子。这说明草地系统的恢复时间具有清楚的热量-辐射控制特征，且这种特征同时反映在模型贡献和空间异质性上。",
    )
    add_paragraph(
        doc,
        "Savanna 中，GPP 的 STRD 与 |EVA| 同时被 OPGD 和 SHAP 识别，RECO 则主要重叠 STRD。SHAP 中 SSRD 贡献很高，但 OPGD 更强调 TMP、STRD 和 VPD 的空间分层解释力，说明稀树草原恢复过程更像热干背景下的能量-蒸散耦合，而不是单一短波辐射效应。",
    )
    add_paragraph(
        doc,
        "Shrubland 中，GPP 的 TMP 与 VPD 重叠，RECO 的 SSRD 与 TMP 重叠，是 OPGD 与 SHAP 一致性较高的干旱系统。该结果支持灌丛区恢复时间受热量、大气干旱和辐射条件共同控制；其中 RECO 的空间分层解释力尤其强。",
    )

    doc.add_heading("8. 写作建议", level=1)
    add_paragraph(
        doc,
        "论文中建议将 SHAP 表述为识别模型预测中的关键变量，将 OPGD 表述为空间层面对这些关键变量的独立验证。若某变量同时位于 SHAP Top3 和 OPGD Top3，可称其既具有预测贡献，也具有空间分层解释力；若 SHAP 高但 OPGD 低，应解释为局部预测贡献强但未必主导空间格局；若 OPGD 高但 SHAP 低，则说明该变量更偏空间分区背景，作用方向仍需结合 ALE、ICE、PDP 和 risk detector 判断。",
    )

    doc.core_properties.title = "OPGD Geodetector 与 SHAP 对比分析"
    doc.core_properties.subject = "GPP/RECO recovery time validation"
    doc.save(docx_path)
    print(docx_path)


if __name__ == "__main__":
    main()
