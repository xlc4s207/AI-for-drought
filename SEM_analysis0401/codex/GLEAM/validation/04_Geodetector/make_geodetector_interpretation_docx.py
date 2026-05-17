#!/usr/bin/env python3
"""Create a Word document from the Geodetector interpretation markdown."""

from pathlib import Path
import sys

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
from common_validation import SHAP_ROOTS, SHORT_LABELS  # noqa: E402


TABLE_MARKER = "<!-- SHAP_GEODETECTOR_COMPARISON_TABLE -->"


def set_east_asian_font(style, font_name: str = "SimSun") -> None:
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def short_label(feature: str) -> str:
    return SHORT_LABELS.get(feature, feature)


def build_comparison_rows(work_dir: Path) -> list[list[str]]:
    geo = pd.read_csv(work_dir / "geodetector_factor_q.csv")
    rows: list[list[str]] = []
    for (metric, biome), group in geo.groupby(["metric", "biome"], sort=False):
        shap = pd.read_csv(SHAP_ROOTS[metric] / biome / "feature_importance.csv")
        geo_top = group.head(3).copy()
        shap_top = shap.head(3).copy()

        geo_text = "; ".join(f"{row.label} (q={row.q:.3f})" for row in geo_top.itertuples())
        shap_text = "; ".join(
            f"{short_label(row.feature)} ({row.importance:.2f})" for row in shap_top.itertuples()
        )

        geo_labels = set(geo_top["label"])
        shap_labels = {short_label(feature) for feature in shap_top["feature"]}
        overlap = sorted(geo_labels & shap_labels)
        overlap_text = ", ".join(overlap) if overlap else "无 Top3 重叠"
        if len(overlap) >= 2:
            note = "高度一致：关键变量同时具有模型贡献和空间解释力。"
        elif len(overlap) == 1:
            note = "部分一致：共同变量较稳健，差异反映 SHAP 与空间分层关注尺度不同。"
        else:
            note = "差异较大：需要结合 ALE/ICE/risk detector 判断局部方向和空间机制。"
        rows.append([metric, biome, geo_text, shap_text, overlap_text, note])
    return rows


def format_table(table) -> None:
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "SimSun"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
                    run.font.size = Pt(8.5)


def add_comparison_table(doc: Document, work_dir: Path) -> None:
    headers = [
        "指标",
        "Biome",
        "Geodetector Top 3",
        "SHAP Top 3",
        "Top3 重叠",
        "解释",
    ]
    rows = build_comparison_rows(work_dir)
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = True
    for idx, header in enumerate(headers):
        run = table.rows[0].cells[idx].paragraphs[0].add_run(header)
        run.bold = True
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cells[idx].text = value
    format_table(table)


def main() -> None:
    md_path = Path(__file__).with_name("geodetector_results_interpretation_cn.md")
    docx_path = md_path.with_suffix(".docx")
    work_dir = Path(__file__).resolve().parent

    doc = Document()
    styles = doc.styles
    set_east_asian_font(styles["Normal"])
    styles["Normal"].font.size = Pt(11)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        set_east_asian_font(styles[style_name])

    in_code = False
    code_lines: list[str] = []
    for raw in md_path.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                paragraph = doc.add_paragraph("\n".join(code_lines))
                paragraph.style = "No Spacing"
                for run in paragraph.runs:
                    run.font.name = "Consolas"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
                    run.font.size = Pt(9.5)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue
        if line == TABLE_MARKER:
            add_comparison_table(doc, work_dir)
            continue
        if not line:
            continue

        if line.startswith("# "):
            paragraph = doc.add_heading(line[2:], level=0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        else:
            paragraph = doc.add_paragraph(line.replace("`", ""))
            paragraph.paragraph_format.first_line_indent = Pt(22)
            paragraph.paragraph_format.line_spacing = 1.25

    doc.core_properties.title = "Geodetector 结果解释"
    doc.core_properties.subject = "GPP/RECO Geodetector validation interpretation"
    doc.save(docx_path)
    print(docx_path)


if __name__ == "__main__":
    main()
