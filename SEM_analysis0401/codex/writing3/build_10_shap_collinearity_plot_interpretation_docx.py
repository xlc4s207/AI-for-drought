#!/usr/bin/env python3
"""Build a note on interpreting SHAP plots under collinearity and mitigation options."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path("/home/xulc/flash_drought")
OUT = ROOT / "process/SEM_analysis0401/codex/writing3/10_SHAP_plot_interpretation_and_collinearity_mitigation_cn.docx"
MD_OUT = ROOT / "process/SEM_analysis0401/codex/GLEAM/plots2/prepeak_shap_summary_20260502/collinearity_grouped_shap_20260506/shap_plot_interpretation_and_collinearity_mitigation_20260506.md"


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(text)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def build_docx() -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10.5)

    title = doc.add_heading("共线性背景下 SHAP 蜂巢图与 Dependence Plot 的解释及降共线性方案", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("一、核心判断", level=1)
    add_para(doc, "当前 SHAP 图不能简单按“单变量独立效应”解释。蜂巢图和 dependence plot 仍然有效，但它们展示的是模型在原始相关变量空间中的贡献分配，而不是在控制其他高度相关变量后的净因果效应。因此，图像解释应从“变量 A 单独导致恢复时间变化”改为“变量 A 所在的机制组携带了模型预测所需的信息”。")
    add_para(doc, "如果审稿人追问“有没有去除共线性”，仅用文字说明是不够的。真正去除或降低共线性，需要在建模前处理输入变量，并重新训练模型、重新计算 SHAP。画图阶段不能把已经存在于模型输入中的共线性消掉。")

    doc.add_heading("二、蜂巢图应该如何解释", level=1)
    add_para(doc, "蜂巢图中的横轴是每个样本上该变量对模型预测的 SHAP 贡献，点的颜色表示该变量原始取值。共线性存在时，高排名变量说明该变量携带了对预测有用的信息，但这部分信息可能与同组变量共享。例如 SSRD、STRD、TMP 同时高，并不意味着三者分别独立控制恢复时间，而是说明能量输入与热状态这一机制组是主导信息来源。")
    add_para(doc, "建议在正文中保留单变量蜂巢图，但图注或结果段应增加一句限制性说明：由于部分气象变量存在物理耦合，蜂巢图中的单变量 SHAP 贡献表示模型归因分配，不应被解释为完全独立的因果效应；本文进一步采用机制组贡献和 SEM 分层路径进行稳健解释。")

    doc.add_heading("三、Dependence Plot 应该如何解释", level=1)
    add_para(doc, "Dependence plot 展示某个特征取值与其 SHAP 贡献之间的关系，可用于识别方向、阈值和非线性。但在共线性背景下，横轴变量变化往往伴随相关变量共同变化。例如 TMP 上升通常伴随 STRD 或 VPD 变化，因此 TMP dependence plot 中的阈值更应解释为热力-干燥背景增强后的联合响应，而不是 TMP 的纯净独立阈值。")
    add_para(doc, "如果 dependence plot 用另一个变量着色，例如 SSRD colored by EVA 或 TMP colored by VPD，颜色梯度可以帮助判断交互或耦合方向，但不能证明被着色变量是独立调节因子。更严谨的写法是：颜色映射显示该 dependence pattern 在不同水分或大气干燥背景下存在分层差异，提示潜在交互，但其独立性需要 ALE、ICE、PDP、地理探测器和 SEM 共同验证。")

    doc.add_heading("四、几种降共线性方法的比较", level=1)
    rows = [
        ["机制组 SHAP", "不改变原模型；按机制组汇总 mean(|SHAP|)", "最适合当前论文主解释；保留变量名称；易与生态机制对应", "没有真正改变模型输入，共线性仍存在", "正文主结果与补充说明"],
        ["变量筛选", "每个强相关组只保留 1-2 个代表变量并重新训练", "最简单，图最容易解释", "会删掉 SSRD/STRD/TMP 等有意义变量，可能被质疑遗漏机制", "敏感性分析，不建议替代主结果"],
        ["PCA", "把相关变量转成正交主成分后建模", "数学上最直接消除线性共线性；VIF 接近消失", "主成分不是原变量，蜂巢图解释变成 PC1/PC2；生态解释依赖载荷表", "适合补充稳健性，而不是主图完全替换"],
        ["分组 PCA", "在辐射热力、水分、大气需水组内部做 PCA", "比全变量 PCA 更可解释；可得到能量PC、水分PC等", "仍需解释载荷；dependence plot 不再对应原始变量", "推荐作为强稳健性分析"],
        ["残差化/正交分解", "保留关键原变量，其他变量转为相对上游变量的残差", "可保留 SSRD 等核心变量名称；降低 VIF；适合 SEM", "残差变量含义更抽象，如 TMP_resid 是控制辐射后的温度异常", "推荐用于 SEM 或补充 SHAP"],
        ["条件 SHAP/Owen values", "按相关结构或特征聚类计算条件归因", "理论上更贴近相关变量归因；可减少不合理分摊", "实现复杂，计算成本高；与现有图口径不同", "可作为方法升级，但不一定必要"],
        ["ALE/ICE/PDP 验证", "用扰动或局部累积效应看方向和阈值", "能验证 dependence 方向和非线性是否稳定", "不能完全消除共线性，PDP 在强相关下可能生成不现实组合", "作为解释验证，而非去共线主方法"],
    ]
    add_table(doc, ["方法", "做法", "优点", "缺点", "建议用途"], rows)

    doc.add_heading("五、PCA 是否适合当前研究", level=1)
    add_para(doc, "PCA 可以使用，但不建议直接把全部变量做一次 PCA 后完全替代原始 SHAP 主结果。原因是本文的目标不仅是预测恢复时间，还要解释 SSRD、TMP、VPD、EVA、SMrz 等具体生态水文变量的作用。全变量 PCA 会让蜂巢图变成 PC1、PC2 的贡献排序，虽然统计上正交，但读者难以直接理解某个具体变量的作用。")
    add_para(doc, "更合适的是分组 PCA：对 SSRD、STRD、TMP 构建辐射热力 PC；对 VPD、WIND 构建大气需水 PC；对 PRE、EVA、SMrz 构建水分供给/储存 PC；事件属性单独保留。这样既能显著降低组内共线性，又能保持生态机制可解释性。论文中可以把原始 SHAP 作为主结果，把分组 PCA-SHAP 作为稳健性验证：若主导机制组排序仍一致，就说明结论不依赖单个相关变量的归因分摊。")

    doc.add_heading("六、正交分解/残差化是否适合当前研究", level=1)
    add_para(doc, "正交分解比 PCA 更适合保留关键变量，尤其是用户当前强调 SSRD 必须保留。可采用“保留核心变量 + 其他变量残差化”的方式，例如保留 SSRD 原值，把 STRD 分解为 SSRD 无法解释的剩余长波辐射信息，把 TMP 分解为在 SSRD 和 STRD 之外的温度异常，把 VPD 分解为在 TMP 和 WIND 之外的大气干燥异常。这样模型仍包含 SSRD，同时其他变量不再与 SSRD/TMP 过度争夺解释权。")
    add_para(doc, "不过残差化后的变量不能再解释为原始 TMP 或 VPD，而应写成“控制上游辐射或热力背景后的 TMP 异常/VPD 异常”。这类方法非常适合 SEM，因为 SEM 本身强调路径层级；也可用于补充 SHAP 模型，检验原始 SHAP 中的方向和主导机制是否稳定。")

    doc.add_heading("七、推荐的实际执行方案", level=1)
    add_para(doc, "建议采用三级方案。第一级是当前必须做的论文解释修订：保留原始蜂巢图和 dependence plot，但图注和正文明确说明它们反映相关变量空间中的模型归因；同时展示机制组 SHAP 贡献。第二级是推荐补充的稳健性分析：构建分组 PCA-SHAP，查看辐射热力组、水分组、大气需水组的贡献排序是否与原始 SHAP 一致。第三级是用于 SEM 或更严格补充分析的正交分解：保留 SSRD 等关键变量，将 STRD、TMP、VPD、EVA 等按物理路径进行残差化，然后重新估计路径系数。")
    add_para(doc, "因此，最稳妥的论文策略不是宣称已经完全消除了共线性，而是写成：我们识别并量化了共线性；主 SHAP 图用于展示模型归因；机制组 SHAP 用于避免过度单变量解释；PCA/残差化敏感性分析用于检验主结论是否依赖共线变量的贡献分摊；SEM 使用分层路径减少并列直连导致的系数不稳定。")

    doc.add_heading("八、可直接写入图注的表述", level=1)
    add_para(doc, "蜂巢图图注建议：点表示单个样本中该变量对恢复时间预测的 SHAP 贡献，颜色表示变量取值。由于部分气象和生态水文变量存在物理耦合，单变量 SHAP 排名反映模型在相关输入空间中的贡献分配，不被解释为完全独立的因果效应；本文进一步采用机制组 SHAP、ALE/ICE/PDP 和 SEM 分层路径进行稳健解释。")
    add_para(doc, "Dependence plot 图注建议：曲线和散点展示变量取值与 SHAP 贡献之间的非线性关系，用于识别潜在方向和阈值。对于 TMP、STRD、SSRD、VPD、EVA 和 SMrz 等相关变量，dependence pattern 应理解为相应机制组的联合响应；颜色映射用于提示背景调节或交互，而非证明被着色变量的独立因果效应。")

    doc.save(OUT)


def build_md() -> None:
    text = """# 共线性背景下 SHAP 图解释与降共线性方案

## 核心判断

蜂巢图和 dependence plot 仍然可以使用，但不能解释为完全独立的单变量因果效应。它们表示模型在原始相关变量空间中的 SHAP 归因分配。真正要降低共线性，必须重新构造输入变量并重新训练模型、重新计算 SHAP。

## 蜂巢图解释

SSRD、STRD、TMP 同时重要时，应解释为辐射热力机制组重要，而不是三个变量分别独立控制恢复时间。VPD 和 WIND 代表大气需水组，PRE、EVA、SMrz 代表水分供给/储存组。

## Dependence plot 解释

Dependence plot 可用于看方向、阈值和非线性，但 TMP 的阈值可能同时反映 STRD/VPD 的共同变化。因此应写为热力-干燥背景的联合响应。颜色映射说明分层或潜在交互，不等于因果调节。

## 方法选择

| 方法 | 是否真正降低共线性 | 解释性 | 建议 |
|---|---|---|---|
| 机制组 SHAP | 否 | 强 | 主文解释必须加入 |
| 变量筛选 | 是 | 强 | 作为敏感性分析 |
| 全变量 PCA | 是 | 弱 | 不建议替代主结果 |
| 分组 PCA | 是 | 中等偏强 | 推荐做稳健性分析 |
| 残差化/正交分解 | 是 | 中等 | 推荐用于 SEM 或补充 SHAP |
| 条件 SHAP/Owen values | 部分 | 中等 | 方法升级，可选 |

## 推荐路线

1. 主结果保留原始 SHAP 蜂巢图和 dependence plot。
2. 增加机制组 SHAP 贡献，避免单变量过度解释。
3. 补充分组 PCA-SHAP，检验机制组排序是否稳定。
4. 对 SEM 使用正交分解或分层路径，降低路径系数不稳定。
5. 图注中明确：SHAP 单变量贡献不是独立因果效应。
"""
    MD_OUT.write_text(text, encoding="utf-8")


def main() -> None:
    MD_OUT.parent.mkdir(parents=True, exist_ok=True)
    build_docx()
    build_md()
    print(f"Wrote {OUT}")
    print(f"Wrote {MD_OUT}")


if __name__ == "__main__":
    main()
