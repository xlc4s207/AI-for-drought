#!/usr/bin/env python3
"""Build document 07: integrated SHAP, OPGD, and SEM mechanism analysis."""

from __future__ import annotations

from pathlib import Path
import sys

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path("/home/xulc/flash_drought/process/SEM_analysis0401/codex")
WRITING_DIR = ROOT / "writing3"
VALIDATION_DIR = ROOT / "GLEAM/validation/05_OPGD_Geodetector"
SHAP_PLOT_DIR = ROOT / "GLEAM/plots2/prepeak_shap_summary_20260502"
OUT_DOCX = WRITING_DIR / "07_SHAP_OPGD_SEM_integrated_mechanism_analysis_cn.docx"


def set_east_asian_font(style, font_name: str = "SimSun") -> None:
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.first_line_indent = Pt(22)
    paragraph.paragraph_format.line_spacing = 1.25


def add_image(doc: Document, image_path: Path, caption: str, width: float = 6.7) -> None:
    if not image_path.exists():
        add_paragraph(doc, f"未找到图件：{image_path}")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    caption_paragraph = doc.add_paragraph(caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption_paragraph.runs:
        run.font.size = Pt(9)


def format_table(table) -> None:
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "SimSun"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
                    run.font.size = Pt(8)


def add_dataframe_table(doc: Document, df: pd.DataFrame, headers: list[str], widths: list[int] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = True
    for idx, header in enumerate(headers):
        run = table.rows[0].cells[idx].paragraphs[0].add_run(header)
        run.bold = True
    for row in df.itertuples(index=False):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    format_table(table)


def shap_top5_table() -> pd.DataFrame:
    path = SHAP_PLOT_DIR / "analysis_writing/biome_top5_summary.csv"
    df = pd.read_csv(path)
    return df.rename(
        columns={
            "metric": "指标",
            "biome": "Biome",
            "top5": "SHAP Top5",
            "dominant_feature": "主导变量",
            "mean_top5_importance": "Top5平均贡献",
        }
    ).round({"Top5平均贡献": 2})


def reliability_high_table() -> pd.DataFrame:
    path = VALIDATION_DIR / "reliability/reliability_score.csv"
    df = pd.read_csv(path)
    high = df[df["reliability_grade"] == "High"].copy()
    high = high.sort_values(["metric", "biome", "q_opgd"], ascending=[True, True, False])
    high["q_opgd"] = high["q_opgd"].round(3)
    high["bootstrap_q_cv"] = high["bootstrap_q_cv"].round(3)
    high["top3_frequency"] = high["top3_frequency"].round(2)
    high["SHAP Top3"] = high["in_shap_top3"].map({True: "是", False: "否"})
    return high[["metric", "biome", "label", "q_opgd", "bootstrap_q_cv", "top3_frequency", "SHAP Top3"]].rename(
        columns={
            "metric": "指标",
            "biome": "Biome",
            "label": "变量",
            "q_opgd": "OPGD q",
            "bootstrap_q_cv": "Bootstrap CV",
            "top3_frequency": "Top3频率",
        }
    )


def sem_direct_path_table() -> pd.DataFrame:
    path = VALIDATION_DIR / "sem_path_mechanism_from_image.csv"
    df = pd.read_csv(path)
    direct = df[df["path_type"] == "direct_to_recovery"].copy()
    pivot = direct.pivot_table(index=["biome_assumed", "metric"], columns="source", values="coefficient").reset_index()
    pivot = pivot.rename(columns={"biome_assumed": "Biome", "metric": "指标"})
    for col in ["TMP", "SSRD", "EVA", "Duration", "Intensity"]:
        if col in pivot.columns:
            pivot[col] = pivot[col].round(3)
    return pivot[["Biome", "指标", "TMP", "SSRD", "EVA", "Duration", "Intensity"]]


def consistency_text() -> str:
    path = VALIDATION_DIR / "reliability/shap_opgd_consistency.csv"
    df = pd.read_csv(path)
    overlap2 = int((df["top3_overlap_count"] >= 2).sum())
    overlap1 = int((df["top3_overlap_count"] == 1).sum())
    corr_mean = df["spearman_rank_correlation"].mean()
    return (
        f"在 10 个 metric-biome 组合中，有 {overlap2} 个组合的 SHAP Top3 与 OPGD Top3 至少重叠 2 个变量，"
        f"{overlap1} 个组合重叠 1 个变量；二者排序相关均值约为 {corr_mean:.2f}。"
    )


def add_shap_opgd_consistency_analysis(doc: Document) -> None:
    consistency = pd.read_csv(VALIDATION_DIR / "reliability/shap_opgd_consistency.csv")
    reliability = pd.read_csv(VALIDATION_DIR / "reliability/reliability_score.csv")
    add_paragraph(
        doc,
        "从一致性角度看，SHAP 与 OPGD 的关系不是完全重合，而是形成了“预测贡献”和“空间解释力”的互补。SHAP 反映机器学习模型在样本层面依赖哪些变量进行恢复时间预测；OPGD 则检验这些变量能否把恢复时间划分出具有内部一致性和组间差异的空间层。因此，二者一致时说明该变量既是预测变量，也是空间分异变量；二者不一致时，则说明该变量可能只在局部预测或空间背景分层中发挥作用。",
    )
    for row in consistency.itertuples(index=False):
        high = reliability[
            (reliability["metric"] == row.metric)
            & (reliability["biome"] == row.biome)
            & (reliability["reliability_grade"] == "High")
        ].sort_values("q_opgd", ascending=False)
        high_text = "、".join(
            f"{item.label}(q={item.q_opgd:.3f})" for item in high.itertuples(index=False)
        ) or "无 High 等级变量"
        add_paragraph(
            doc,
            f"{row.metric} {row.biome} 中，SHAP Top3 与 OPGD Top3 的重叠变量为 {row.top3_overlap_labels or '无'}，"
            f"Top3 重叠数量为 {row.top3_overlap_count}，排序相关为 {row.spearman_rank_correlation:.2f}。"
            f"OPGD 高可靠变量为 {high_text}。这一结果说明该组合中可优先解释的机制变量应从重叠变量和 High 等级变量中选择；"
            f"若某变量仅在 SHAP 中突出，则更适合解释为模型预测贡献，若仅在 OPGD 中突出，则更适合解释为空间背景分层因子。",
        )


def add_sem_physical_mechanism_analysis(doc: Document) -> None:
    sem = pd.read_csv(VALIDATION_DIR / "sem_path_mechanism_from_image.csv")
    direct = sem[sem["path_type"] == "direct_to_recovery"].copy()
    pivot = direct.pivot_table(index=["biome_assumed", "metric"], columns="source", values="coefficient")
    add_paragraph(
        doc,
        "SEM 图中的物理机制可以分成能量输入、大气干旱、土壤水分补给、蒸散中介和事件属性五个部分。STRD 和 SSRD 共同指向 TMP，说明长波辐射和短波辐射通过热量状态影响后续过程；TMP 进一步正向驱动 VPD，表示升温增强大气水汽亏缺；VPD 和 SMrz 又共同影响 EVA，说明蒸散状态同时受大气干旱需求和土壤供水条件调节。这个结构比单变量解释更合理，因为干旱后恢复过程本质上是能量供给、水分亏缺和植被-土壤水分状态共同变化的结果。",
    )
    add_paragraph(
        doc,
        "SSRD 对恢复时间的直接路径在所有 biome 和两个通量指标中均为负值。这意味着在本研究的恢复时间定义下，较高短波辐射通常对应较短恢复时间。物理上可以理解为：较强辐射提高可用能量，促进植被光合活动恢复，也可能增强地表能量交换，使系统更快回到基准状态。但这种负向路径并不表示辐射没有胁迫作用，而是在控制其他路径后，SSRD 的直接效应主要体现为恢复过程中的能量供给效应。",
    )
    add_paragraph(
        doc,
        "EVA 对恢复时间也表现出稳定负向直接效应。由于 EVA 位于 VPD 和 SMrz 的下游，它更像综合水热状态的中介指标：当土壤水分和大气需求能够支撑较高蒸散时，生态系统通常已经具备更强的水分交换和能量耗散能力，因此恢复时间缩短。相反，若 EVA 偏低，则可能代表水分供应不足或植被活动受限，恢复过程更慢。",
    )
    add_paragraph(
        doc,
        "TMP 的直接路径具有明确的 biome 差异。Forest、Grassland、Savanna 和 Shrubland 中，TMP 对 GPP-RT 或 RECO-RT 均为正值，说明更高温度背景通常延长恢复时间；但 Cropland 中 TMP -> GPP-RT 和 TMP -> RECO-RT 均为负值，分别为 -0.081 和 -0.092。这个结果不能再写成“Cropland 温度效应较弱但同向”，而应明确写成“Cropland 中温度直接路径为负向”。物理上，TMP 仍通过 TMP -> VPD -> EVA 参与大气干旱和蒸散中介链条，但在控制 SSRD、EVA、Duration、Intensity 以及中介路径后，TMP 对农田恢复时间的剩余直接效应表现为缩短恢复时间，这可能与灌溉、作物物候、播收制度和人为管理调节有关。",
    )
    add_paragraph(
        doc,
        "Duration 和 Intensity 的直接效应也具有 biome 依赖性，而不是所有系统都为正向。Forest、Grassland 和 Cropland 中，Duration/Intensity -> RT 为负向；Savanna 和 Shrubland 中则为正向。这说明事件持续时间和强度不能被简单解释为统一的恢复延迟因子。在干旱或半干旱系统中，持续更久、强度更高的事件更容易加深水分亏缺并延长恢复；但在 Forest、Grassland 和 Cropland 中，控制水热背景和蒸散状态后，事件属性的剩余直接效应为负，可能反映事件后水热条件补偿、恢复判定边界差异或农田管理干预。",
    )
    for biome in ["Forest", "Grassland", "Savanna", "Cropland", "Shrubland"]:
        gpp = pivot.loc[(biome, "GPP")]
        reco = pivot.loc[(biome, "RECO")]
        if biome == "Cropland":
            interpretation = (
                "Cropland 的关键特征是 TMP、SSRD、EVA、Duration 和 Intensity 对 RT 的直接路径均为负向，"
                "表明农田恢复时间受到管理、灌溉和作物物候调节后，温度与事件属性的剩余直接效应不同于自然生态系统。"
            )
        elif biome in ["Forest", "Grassland"]:
            interpretation = (
                "该 biome 中 TMP 为正向而 SSRD、EVA、Duration 和 Intensity 多为负向，"
                "说明热量背景倾向于延长恢复，而辐射、蒸散状态和事件属性的剩余直接效应倾向于缩短恢复。"
            )
        else:
            interpretation = (
                "该 biome 中 TMP、Duration 和 Intensity 为正向，而 SSRD 和 EVA 为负向，"
                "说明热量背景和事件属性会延长恢复，辐射能量供给和蒸散水热交换状态会缩短恢复。"
            )
        add_paragraph(
            doc,
            f"{biome} 中，GPP 的直接路径为 TMP={gpp['TMP']:.3f}、SSRD={gpp['SSRD']:.3f}、EVA={gpp['EVA']:.3f}、"
            f"Duration={gpp['Duration']:.3f}、Intensity={gpp['Intensity']:.3f}；"
            f"RECO 的直接路径为 TMP={reco['TMP']:.3f}、SSRD={reco['SSRD']:.3f}、EVA={reco['EVA']:.3f}、"
            f"Duration={reco['Duration']:.3f}、Intensity={reco['Intensity']:.3f}。{interpretation}"
        )


def add_shap_opgd_sem_synthesis(doc: Document) -> None:
    add_paragraph(
        doc,
        "将 SHAP、OPGD 和 SEM 放在一起，可以把变量分成三类。第一类是三者共同支持的稳健机制变量，例如 SSRD、TMP 和部分 biome 中的 EVA/STRD。这些变量既在 SHAP 中贡献高，又在 OPGD 中具有空间分层解释力，并且在 SEM 中具有明确路径。第二类是 SHAP 贡献强但 OPGD 空间解释力较弱的变量，这类变量可能有助于个体样本预测，但未必主导全球空间格局。第三类是 OPGD 高但 SHAP 排名不高的变量，这类变量更像空间背景或环境分区因子，需要结合 SEM 路径方向解释。",
    )
    add_paragraph(
        doc,
        "例如 SSRD 在 SHAP 结果中频繁位居前列，在 OPGD 中也在 RECO Shrubland、RECO Grassland、GPP Grassland 等组合中表现为 High reliability，同时 SEM 中 SSRD -> RT 稳定为负向。这说明 SSRD 是最适合写入主机制的变量之一：它既解释模型预测，又解释空间异质性，还具有一致的直接路径方向。",
    )
    add_paragraph(
        doc,
        "TMP 的证据链显示出明显的 biome 依赖性。SHAP 和 OPGD 均说明 TMP 在草地、稀树草原和灌丛中重要，SEM 则进一步表明 TMP -> RT 在 Forest、Grassland、Savanna 和 Shrubland 中为正向，但在 Cropland 中为负向。因此 TMP 应被写成“受生态系统类型和管理背景调节的热量路径变量”，而不是统一方向的简单驱动因子。",
    )
    add_paragraph(
        doc,
        "EVA 和 VPD 的证据链强调过程调节。VPD 在 SEM 中通过 EVA 间接影响恢复时间，EVA 又对 RT 具有稳定负向直接路径。这说明大气干旱并非只通过直接胁迫影响恢复，还通过改变蒸散状态和水热交换过程间接调节恢复速度。因此，解释 EVA 时应避免写成单一蒸散量效应，而应强调其代表综合水热耦合状态。",
    )


def build_document() -> Path:
    doc = Document()
    styles = doc.styles
    set_east_asian_font(styles["Normal"])
    styles["Normal"].font.size = Pt(11)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        set_east_asian_font(styles[style_name])

    title = doc.add_heading("07 SHAP-OPGD-SEM 综合机制路径分析", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. 分析目标与证据链", level=1)
    add_paragraph(
        doc,
        "本文综合 SHAP、改进参数地理探测器（OPGD）和 SEM 路径模型，解释 GPP 与 RECO 恢复时间的主要控制机制。SHAP 用于识别模型预测中的关键变量，OPGD 用于检验这些变量是否能够解释恢复时间的空间分层异质性，SEM 则进一步给出变量之间的直接路径和中介路径方向。",
    )
    add_paragraph(
        doc,
        "因此，本文采用的解释逻辑是：先由 SHAP 判断变量预测贡献，再由 OPGD 判断空间解释力和可靠性，最后由 SEM 判断变量如何通过热量、辐射、大气干旱、蒸散和土壤水分路径影响恢复时间。",
    )

    doc.add_heading("2. SHAP 结果：预测贡献中的主控变量", level=1)
    add_image(
        doc,
        SHAP_PLOT_DIR / "beeswarm_comparison_5biomes_gpp_vs_reco.png",
        "图 1. GPP 与 RECO 在五类 biome 中的 pre-event SHAP beeswarm 对比图。",
    )
    add_image(
        doc,
        SHAP_PLOT_DIR / "shap_importance_percent_bars_5biomes_gpp_vs_reco.png",
        "图 2. GPP 与 RECO 在五类 biome 中的 SHAP 变量贡献占比图。",
    )
    add_paragraph(
        doc,
        "SHAP 结果显示，SSRD 是最稳定的高贡献变量，几乎在 Forest、Grassland、Savanna、Shrubland 以及 RECO 多个 biome 中占据前列。EVA 在 Cropland 和 Forest 中贡献较高，TMP 在 Grassland、Savanna 和 Shrubland 中贡献突出，说明恢复时间预测主要受到辐射、热量和蒸散状态控制。",
    )
    add_paragraph(
        doc,
        "从物理意义上看，SSRD 代表恢复期前事件阶段的短波辐射和可用能量背景，其高 SHAP 贡献说明模型在判断恢复时间长短时高度依赖辐射条件。TMP 代表热量背景，同时会影响 VPD 和蒸散需求；EVA 则反映水热交换状态。因此，SHAP 的高贡献变量本身已经指向一个能量-水分耦合机制，而不是单纯的降水控制机制。",
    )
    add_dataframe_table(doc, shap_top5_table(), ["指标", "Biome", "SHAP Top5", "主导变量", "Top5平均贡献"])

    doc.add_heading("3. OPGD 结果：空间分层解释力与可靠性", level=1)
    add_image(
        doc,
        VALIDATION_DIR / "figures/shap_opgd_reliability_matrix.png",
        "图 3. SHAP 预测贡献、OPGD q 值和可靠性等级的综合对比图。",
    )
    add_paragraph(
        doc,
        "OPGD 结果表明，SHAP 中重要的能量和水分变量并不都具有同等强度的空间分层解释力。总体上，RECO 的 OPGD q 值和高可靠变量数量高于 GPP，尤其在 Savanna 和 Shrubland 中，SSRD、TMP、VPD 等变量具有较高空间解释力。",
    )
    add_paragraph(doc, consistency_text())
    add_paragraph(
        doc,
        "可靠性分析进一步区分了稳定机制变量和低稳健变量。High 等级变量表示其 q 值、bootstrap 稳定性、Top3 出现频率和分层样本量均较可靠；Low 等级变量不宜作为主导机制解释。"
    )
    add_dataframe_table(doc, reliability_high_table(), ["指标", "Biome", "变量", "OPGD q", "Bootstrap CV", "Top3频率", "SHAP Top3"])
    doc.add_heading("3.1 SHAP 与 OPGD 一致性逐 biome 解释", level=2)
    add_shap_opgd_consistency_analysis(doc)

    doc.add_heading("4. SEM 结果：路径方向与中介机制", level=1)
    add_image(
        doc,
        WRITING_DIR / "SEM总图2.png",
        "图 4. GPP 与 RECO 恢复时间的 SEM 路径机制图。",
    )
    add_paragraph(
        doc,
        "SEM 路径图显示，恢复时间不是由单一变量控制，而是由热量-辐射-大气干旱-蒸散-土壤水分构成的路径网络共同调节。核心中介路径包括 STRD -> TMP、SSRD -> TMP、TMP -> VPD、VPD -> EVA、PRE -> SMrz 和 SMrz -> EVA。其中 TMP 是由长波辐射和短波辐射共同驱动的热量中介节点，而 EVA 是连接 VPD、SMrz 和恢复时间的重要过程变量。",
    )
    add_paragraph(
        doc,
        "直接进入恢复时间的路径主要包括 TMP -> RT、SSRD -> RT、EVA -> RT、Duration -> RT 和 Intensity -> RT。SSRD -> RT 和 EVA -> RT 在所有 biome 和两个通量指标中均为负向，说明短波辐射增强和较高蒸散状态通常对应恢复时间缩短。TMP -> RT、Duration -> RT 和 Intensity -> RT 则具有 biome 差异：Cropland 中 TMP -> RT 为负向；Forest、Grassland 和 Cropland 中 Duration/Intensity -> RT 为负向，而 Savanna 和 Shrubland 中为正向。",
    )
    add_dataframe_table(doc, sem_direct_path_table(), ["Biome", "指标", "TMP", "SSRD", "EVA", "Duration", "Intensity"])
    doc.add_heading("4.1 SEM 路径的物理机制解释", level=2)
    add_sem_physical_mechanism_analysis(doc)

    doc.add_heading("5. 三类证据的综合机制解释", level=1)
    add_paragraph(
        doc,
        "SSRD 是三类证据中最一致的变量之一。SHAP 显示 SSRD 在多个 biome 中具有高预测贡献，OPGD 显示 SSRD 在 RECO Shrubland、RECO Grassland、GPP Grassland 等系统中具有较高空间分层解释力，SEM 则显示 SSRD 对恢复时间具有稳定负向直接效应。因此 SSRD 可被解释为同时具有预测贡献、空间解释力和明确路径方向的稳健机制变量。",
    )
    add_paragraph(
        doc,
        "TMP 的解释需要强调方向差异。SHAP 和 OPGD 都表明 TMP 在 Grassland、Savanna 和 Shrubland 中重要，SEM 则显示 TMP 对自然和半自然 biome 的恢复时间多为正向直接效应，说明热量背景升高通常会加重恢复负担并延长 RT；但 Cropland 中 TMP -> RT 为负向，提示农田系统中的灌溉、作物物候、播收制度和管理措施可能改变温度对恢复时间的直接作用方向。",
    )
    add_paragraph(
        doc,
        "EVA 和 VPD 更适合作为过程链条变量解释。SHAP 中 EVA 在 Cropland 和 Forest 中贡献较高，SEM 中 VPD -> EVA 和 SMrz -> EVA 路径稳定，说明蒸散状态综合了大气干旱和土壤水分条件。EVA -> RT 的稳定负向路径表明，当蒸散状态较高时，系统恢复时间通常缩短。",
    )
    add_paragraph(
        doc,
        "PRE 和 SMrz 的作用更偏水分补给链条。SEM 显示 PRE -> SMrz -> EVA 是稳定路径，但在 SHAP 和 OPGD 中 PRE 并非所有 biome 的最高贡献变量。这说明降水和土壤水分对恢复时间的影响更多通过中介过程体现，而不是始终作为最强直接预测因子出现。",
    )
    doc.add_heading("5.1 SHAP-OPGD-SEM 证据链整合", level=2)
    add_shap_opgd_sem_synthesis(doc)

    doc.add_heading("6. GPP 与 RECO 差异", level=1)
    add_paragraph(
        doc,
        "RECO 的空间分层解释力和 SEM 路径强度整体高于 GPP，尤其在 Savanna 和 Shrubland 中更明显。RECO 的 SSRD -> RT 负向效应强于 GPP，说明生态系统呼吸恢复对辐射和能量背景更敏感。GPP 也受到相同路径控制，但其路径强度和 OPGD q 值通常低于 RECO，表明生产力恢复可能还受到冠层状态、物候和水分利用策略等更复杂因素调节。",
    )

    doc.add_heading("7. 可用于论文的结论表述", level=1)
    add_paragraph(
        doc,
        "综合 SHAP、OPGD 和 SEM 结果可见，GPP 和 RECO 恢复时间主要受热量、辐射、大气干旱、蒸散和土壤水分共同调节。SHAP 结果首先识别出 SSRD、TMP、EVA、VPD 等关键预测变量；OPGD 进一步证明其中一部分变量具有显著空间分层解释力；SEM 则揭示这些变量通过 STRD/SSRD -> TMP -> VPD -> EVA -> RT 以及 PRE -> SMrz -> EVA -> RT 等路径影响恢复时间。SSRD 和 EVA 对恢复时间表现为稳定负向直接效应，TMP、Duration 和 Intensity 的方向具有 biome 依赖性，其中 Cropland 的 TMP -> RT 为负向，是区别于其他 biome 的关键路径。RECO 的路径强度和空间解释力整体高于 GPP，说明呼吸恢复对能量和水分背景更敏感。"
    )

    doc.core_properties.title = "SHAP-OPGD-SEM 综合机制路径分析"
    doc.core_properties.subject = "GPP/RECO recovery time mechanism"
    doc.save(OUT_DOCX)
    return OUT_DOCX


def main() -> None:
    print(build_document())


if __name__ == "__main__":
    main()
