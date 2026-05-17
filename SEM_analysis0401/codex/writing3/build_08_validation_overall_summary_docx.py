#!/usr/bin/env python3
"""Build document 08: validation and overall all-biome SHAP/SEM synthesis."""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path("/home/xulc/flash_drought/process/SEM_analysis0401/codex")
WRITING_DIR = ROOT / "writing3"
GLEAM = ROOT / "GLEAM"
VALIDATION = GLEAM / "validation"
OVERALL_SHAP = GLEAM / "plots2/prepeak_shap_summary_20260502/overall_all_biomes"
OVERALL_SEM = GLEAM / "plots2/SEM/sem_prepeak_overall_all_biomes_20260506"
OUT_DOCX = WRITING_DIR / "08_validation_and_overall_SHAP_SEM_synthesis_cn.docx"


def set_east_asian_font(style, font_name: str = "SimSun") -> None:
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.first_line_indent = Pt(22)
    paragraph.paragraph_format.line_spacing = 1.25


def add_image(doc: Document, image_path: Path, caption: str, width: float = 6.6) -> None:
    if not image_path.exists():
        add_paragraph(doc, f"未找到图件：{image_path}")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size = Pt(9)


def format_table(table) -> None:
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "SimSun"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
                    run.font.size = Pt(8)


def add_dataframe_table(doc: Document, df: pd.DataFrame) -> None:
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.autofit = True
    for idx, col in enumerate(df.columns):
        run = table.rows[0].cells[idx].paragraphs[0].add_run(str(col))
        run.bold = True
    for row in df.itertuples(index=False):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    format_table(table)


def validation_model_table() -> pd.DataFrame:
    df = pd.read_csv(VALIDATION / "01_ALE/model_fit_summary.csv")
    out = df.copy()
    out["R2_train"] = out["r2_train"].round(3)
    out["R2_test"] = out["r2_test"].round(3)
    return out[["metric", "biome", "train_rows", "test_rows", "R2_train", "R2_test"]].rename(
        columns={"metric": "指标", "biome": "Biome", "train_rows": "训练样本", "test_rows": "测试样本"}
    )


def validation_output_table() -> pd.DataFrame:
    rows = []
    specs = [
        ("ALE", VALIDATION / "01_ALE/ale_validation_index.csv", "验证 SHAP dependence 的局部方向与阈值"),
        ("ICE", VALIDATION / "02_ICE/ice_validation_index.csv", "检查 biome 内部响应异质性"),
        ("PDP", VALIDATION / "03_PDP/pdp_validation_index.csv", "展示平均模型响应趋势"),
        ("Geodetector", VALIDATION / "04_Geodetector/geodetector_factor_q.csv", "验证空间分异解释力"),
        ("OPGD", VALIDATION / "05_OPGD_Geodetector/opgd_factor_q.csv", "优化分箱后的空间解释力与可靠性检验"),
    ]
    for method, path, purpose in specs:
        df = pd.read_csv(path)
        rows.append({"方法": method, "结果行数": len(df), "主要作用": purpose})
    return pd.DataFrame(rows)


def geodetector_top_table(path: Path, q_col: str = "q") -> pd.DataFrame:
    df = pd.read_csv(path)
    rows = []
    for (metric, biome), group in df.groupby(["metric", "biome"], sort=False):
        top = group.sort_values(q_col, ascending=False).head(3)
        rows.append(
            {
                "指标": metric,
                "Biome": biome,
                "Top1": f"{top.iloc[0]['label']} ({top.iloc[0][q_col]:.3f})",
                "Top2": f"{top.iloc[1]['label']} ({top.iloc[1][q_col]:.3f})",
                "Top3": f"{top.iloc[2]['label']} ({top.iloc[2][q_col]:.3f})",
            }
        )
    return pd.DataFrame(rows)


def reliability_summary() -> pd.DataFrame:
    path = VALIDATION / "05_OPGD_Geodetector/reliability/reliability_score.csv"
    df = pd.read_csv(path)
    counts = df["reliability_grade"].value_counts().reindex(["High", "Medium", "Low"]).fillna(0).astype(int)
    return pd.DataFrame(
        {
            "可靠性等级": counts.index,
            "变量数": counts.values,
            "解释": ["可优先写入主机制", "可作为支持性证据", "谨慎解释或仅作背景参考"],
        }
    )


def shap_overall_importance_table() -> pd.DataFrame:
    df = pd.read_csv(OVERALL_SHAP / "overall_all_biomes_feature_importance.csv")
    rows = []
    for metric, group in df.groupby("metric", sort=False):
        top = group.sort_values("rank").head(8)
        for item in top.itertuples(index=False):
            rows.append(
                {
                    "指标": metric,
                    "Rank": int(item.rank),
                    "变量": item.display_label,
                    "mean(|SHAP|)": round(float(item.importance), 3),
                    "贡献百分比": f"{float(item.percent):.1f}%",
                }
            )
    return pd.DataFrame(rows)


def sem_r2_table() -> pd.DataFrame:
    df = pd.read_csv(OVERALL_SEM / "tables/overall_all_biomes_sem_r2.csv")
    out = df.copy()
    out["Holdout R2"] = out["holdout_r2"].round(3)
    out["Train R2"] = out["train_r2"].round(3)
    return out[["metric", "rows", "Holdout R2", "Train R2", "predictor_count", "biome_control_count"]].rename(
        columns={
            "metric": "指标",
            "rows": "样本数",
            "predictor_count": "生态路径变量数",
            "biome_control_count": "Biome固定效应数",
        }
    )


def sem_direct_table() -> pd.DataFrame:
    df = pd.read_csv(OVERALL_SEM / "tables/overall_all_biomes_target_direct_paths.csv")
    pivot = df.pivot_table(index="metric", columns="from_label", values="estimate", aggfunc="first").reset_index()
    cols = ["metric", "SSRD", "STRD", "TMP", "VPD", "|EVA|", "SMrz", "Duration", "Intensity"]
    pivot = pivot[[c for c in cols if c in pivot.columns]]
    for col in pivot.columns:
        if col != "metric":
            pivot[col] = pivot[col].round(3)
    return pivot.rename(columns={"metric": "指标"})


def curve_strength_table() -> pd.DataFrame:
    ale = pd.read_csv(VALIDATION / "01_ALE/ale_validation_index.csv")
    ice = pd.read_csv(VALIDATION / "02_ICE/ice_validation_index.csv")
    pdp = pd.read_csv(VALIDATION / "03_PDP/pdp_validation_index.csv")
    ale["ALE幅度"] = (ale["ale_max"] - ale["ale_min"]).round(2)
    pdp["PDP幅度"] = (pdp["pdp_max"] - pdp["pdp_min"]).round(2)
    ice["ICE异质性IQR"] = ice["effect_iqr"].round(2)
    merged = ale[["metric", "biome", "label", "ALE幅度"]].merge(
        pdp[["metric", "biome", "label", "PDP幅度"]], on=["metric", "biome", "label"], how="left"
    ).merge(
        ice[["metric", "biome", "label", "ICE异质性IQR"]], on=["metric", "biome", "label"], how="left"
    )
    top = merged.sort_values("ALE幅度", ascending=False).head(10)
    return top.rename(columns={"metric": "指标", "biome": "Biome", "label": "变量"}).reset_index(drop=True)


def add_intro(doc: Document) -> None:
    doc.add_heading("1. 文档目标", level=1)
    add_paragraph(
        doc,
        "前序分析已经分别完成了 GPP 和 RECO 在不同 biome 中的 SHAP、dependence plot、SEM、Geodetector 和 OPGD 解释。本文件进一步总结 validation 路径下的 ALE、ICE、PDP、Geodetector、OPGD 结果，并整合最新补充的 all-biome pooled SHAP 与整体 SEM 结果，形成一套可写入论文或报告的验证性证据链。",
    )
    add_paragraph(
        doc,
        "本文的核心逻辑是：SHAP 负责回答“哪些变量对模型预测最重要”；ALE/PDP/ICE 负责回答“这些变量的响应方向、阈值和异质性是否稳定”；Geodetector/OPGD 负责回答“这些变量是否解释恢复时间的空间分异”；整体 SHAP 和整体 SEM 则用于补充跨 biome 的总体机制图景。",
    )


def add_validation_methods(doc: Document) -> None:
    doc.add_heading("2. Validation 方法与输出概况", level=1)
    add_paragraph(
        doc,
        "validation 目录下的 ALE、ICE 和 PDP 都不是直接从已有 SHAP 图中抽取曲线，而是基于同一批 prepeak feature table，针对每个 metric-biome 重新训练 LightGBM 回归模型，再计算响应曲线。因此它们可以作为相对独立的模型响应验证。Geodetector 和 OPGD 则不依赖预测模型，而是从空间分层异质性的角度检验解释力。",
    )
    add_dataframe_table(doc, validation_output_table())
    add_paragraph(
        doc,
        "ALE 使用分位数区间内的局部预测差值累积，较适合存在强共线性的气候生态变量；ICE 固定单个样本的其他变量，只改变目标变量，因此用于识别 biome 内部是否存在响应方向差异；PDP 将所有样本的某个变量统一替换为给定网格值，展示平均响应趋势，但在 SSRD、STRD、TMP、VPD 等高度相关变量中需要谨慎解释。",
    )
    doc.add_heading("2.1 LightGBM 验证模型拟合情况", level=2)
    add_dataframe_table(doc, validation_model_table())
    add_paragraph(
        doc,
        "验证模型的测试集 R2 显示，RECO 的可预测性整体高于 GPP，尤其 RECO Shrubland、Grassland 和 Savanna 的 R2 较高。这意味着 ALE/ICE/PDP 对 RECO 的响应解释相对更稳定；GPP Forest 和 Shrubland 的 R2 较低，说明其恢复时间可能受更多未显式纳入的生理、结构或管理因素影响。",
    )


def add_response_curve_summary(doc: Document) -> None:
    doc.add_heading("3. ALE、ICE 与 PDP 对 SHAP 结果的验证", level=1)
    add_paragraph(
        doc,
        "ALE、ICE 和 PDP 的共同结论是：SSRD、TMP、STRD、VPD、EVA 和 PRE 是响应幅度最明显的变量，且多个变量呈现非线性或阈值型响应。与 SHAP 相比，这些曲线不用于重新排序变量重要性，而是用于验证 SHAP dependence 所显示的方向和阈值是否在重新训练模型中仍然存在。",
    )
    add_dataframe_table(doc, curve_strength_table())
    add_paragraph(
        doc,
        "从 ALE 幅度看，RECO Grassland 的 SSRD/TMP、RECO Savanna 的 TMP/SSRD、GPP Savanna 的 TMP/SSRD 等组合响应最强，这与 SHAP 中能量和热量变量主导恢复时间的判断一致。ICE 的 IQR 较高说明同一 biome 内部存在明显异质性，即平均趋势不能代表所有样本。PDP 则提供更直观的平均响应，但由于可能构造现实中较少出现的变量组合，应作为 ALE 的辅助展示而不是主证据。",
    )
    add_paragraph(
        doc,
        "因此，本文建议在正文中优先引用 ALE 验证方向和阈值，用 ICE 说明 biome 内部响应异质性，用 PDP 作为读者更容易理解的平均趋势图。三者与 SHAP 的关系不是替代，而是从不同角度支撑 SHAP 发现的非线性响应结构。",
    )


def add_geodetector_summary(doc: Document) -> None:
    doc.add_heading("4. Geodetector 与 OPGD：空间解释力验证", level=1)
    add_paragraph(
        doc,
        "Geodetector 的 q 值衡量的是某个变量分层后能够解释目标变量空间分异的比例。q 值越大，说明该变量对恢复时间空间分层的解释力越强。需要强调的是，q 值没有方向性，也不能直接解释为因果效应；方向仍需要结合 SHAP dependence、ALE/PDP 曲线和 SEM 路径系数判断。",
    )
    doc.add_heading("4.1 原始 Geodetector Top3 因子", level=2)
    add_dataframe_table(doc, geodetector_top_table(VALIDATION / "04_Geodetector/geodetector_factor_q.csv"))
    doc.add_heading("4.2 OPGD 优化分箱 Top3 因子", level=2)
    add_dataframe_table(doc, geodetector_top_table(VALIDATION / "05_OPGD_Geodetector/opgd_factor_q.csv"))
    add_paragraph(
        doc,
        "OPGD 在原始地理探测器基础上搜索不同离散化方法和分箱数，通常能提高 q 值并减少人为分箱选择带来的不确定性。结果显示，TMP、STRD、SSRD、VPD 和 EVA 是最稳定的空间解释变量。其中 Grassland、Savanna 和 Shrubland 中 TMP/SSRD/STRD/VPD 的空间解释力尤其强，说明热干边缘或半干旱系统中恢复时间的空间差异更受能量-水分耦合控制。",
    )
    doc.add_heading("4.3 OPGD 可靠性分级", level=2)
    add_dataframe_table(doc, reliability_summary())
    add_image(
        doc,
        VALIDATION / "05_OPGD_Geodetector/figures/shap_opgd_reliability_matrix.png",
        "图 1. SHAP-OPGD 可靠性矩阵。High 变量可作为主机制证据，Medium 变量作为支持证据，Low 变量谨慎解释。",
        width=6.6,
    )
    add_image(
        doc,
        VALIDATION / "05_OPGD_Geodetector/figures/opgd_interaction_heatmaps.png",
        "图 2. OPGD 交互探测热图。交互项用于判断变量组合是否增强恢复时间的空间解释力。",
        width=6.6,
    )
    add_paragraph(
        doc,
        "interaction 结果的重要性在于，它说明恢复时间并非由单个变量孤立控制，而是由能量、温度、大气干旱、蒸散和土壤水分共同决定。若两个变量的交互 q 值明显高于单变量 q 值，说明它们的组合能够更好解释空间异质性，这与 SHAP dependence 图中的颜色映射和 SEM 中的中介路径是一致的。",
    )


def add_overall_shap_summary(doc: Document) -> None:
    doc.add_heading("5. All-biome pooled SHAP 整体分析", level=1)
    add_paragraph(
        doc,
        "此前 SHAP 图主要按 biome 分别展示，有利于识别生态系统差异，但缺少一个跨 biome 的整体视角。为此，本文补充了 all-biome pooled SHAP 分析，将 Cropland、Forest、Grassland、Savanna 和 Shrubland 合并后绘制整体蜂巢图和整体 dependence plot。该分析用于回答：在五类生态系统整体上，哪些变量仍然是恢复时间预测的主导变量。",
    )
    add_image(
        doc,
        OVERALL_SHAP / "overall_all_biomes_beeswarm_gpp_vs_reco.png",
        "图 3. GPP 与 RECO 的 all-biome pooled SHAP beeswarm 图。",
        width=6.8,
    )
    add_dataframe_table(doc, shap_overall_importance_table())
    add_paragraph(
        doc,
        "整体 SHAP 结果显示，SSRD 在 GPP 和 RECO 中均为第一主导变量，贡献分别约为 25.1% 和 27.9%。TMP 分别位列第二，EVA 和 STRD 也稳定保持在前列。这说明在跨 biome 总体尺度上，短波辐射、温度、蒸散和长波辐射构成了恢复时间预测的核心变量组合。与分 biome 结果相比，整体结果强化了 SSRD 的主导地位，但也会弱化某些 biome 特异变量，例如 Cropland 中的 EVA/SMrz 或 Shrubland 中的 VPD/PRE。",
    )
    add_image(
        doc,
        OVERALL_SHAP / "gpp_overall_all_biomes_dependence_top10.png",
        "图 4. GPP all-biome pooled SHAP dependence top10 图。",
        width=6.8,
    )
    add_image(
        doc,
        OVERALL_SHAP / "reco_overall_all_biomes_dependence_top10.png",
        "图 5. RECO all-biome pooled SHAP dependence top10 图。",
        width=6.8,
    )
    add_paragraph(
        doc,
        "整体 dependence plot 的作用是展示跨 biome 平均响应，而不是替代分 biome 图。由于不同 biome 的变量分布和响应方向可能不同，整体 dependence 图中出现的宽散点云和非单调趋势本身就是生态异质性的证据。因此，整体图适合写总体机制，分 biome 图适合解释生态系统差异。",
    )


def add_overall_sem_summary(doc: Document) -> None:
    doc.add_heading("6. All-biome pooled SEM 机制路径", level=1)
    add_paragraph(
        doc,
        "整体 SEM 使用 SSRD event-aware 机制结构，并将五个 biome 合并后拟合。为了避免 biome 间平均差异直接污染生态路径，模型在每个结构方程中加入 biome fixed effects 控制项；路径图中只展示生态变量之间的机制路径，不绘制 biome dummy。该设置使整体 SEM 可以被解释为“控制 biome 平均差异后的总体机制路径”。",
    )
    add_dataframe_table(doc, sem_r2_table())
    add_paragraph(
        doc,
        "整体 SEM 的 holdout R2 为 GPP 0.132、RECO 0.177，低于重新训练 LightGBM 验证模型，但高于一些过度简化的线性 SEM。这个差异说明恢复时间具有明显非线性和空间异质性，SEM 更适合作为机制路径检验，而不是替代机器学习模型进行高精度预测。",
    )
    add_image(
        doc,
        OVERALL_SEM / "figures/overall_all_biomes_sem_path_mechanism_gpp_vs_reco.png",
        "图 6. All-biome pooled SEM 路径机制图。红色为正路径，蓝色为负路径，biome fixed effects 已控制但未绘制。",
        width=6.8,
    )
    add_dataframe_table(doc, sem_direct_table())
    add_paragraph(
        doc,
        "整体 SEM 直接路径显示，TMP 对恢复时间具有最强正向路径，GPP 和 RECO 中分别为 1.490 和 1.692；SSRD 和 STRD 对恢复时间均为负向路径，说明在控制 TMP、VPD、EVA、SMrz 和事件属性后，辐射变量的直接效应主要表现为缩短恢复时间。RECO 中 EVA 的直接路径接近 0 且不显著，而 GPP 中 EVA 为正向，说明蒸散对光合恢复和呼吸恢复的作用并不完全一致。",
    )
    add_image(
        doc,
        OVERALL_SEM / "figures/overall_all_biomes_target_direct_coefficients_heatmap.png",
        "图 7. All-biome SEM 中指向 recovery time 的直接路径系数热图。",
        width=6.2,
    )
    add_image(
        doc,
        OVERALL_SEM / "figures/overall_all_biomes_total_effects_heatmap.png",
        "图 8. All-biome SEM 总效应热图，包含直接效应和通过中介变量形成的间接效应。",
        width=6.2,
    )


def add_synthesis(doc: Document) -> None:
    doc.add_heading("7. 综合解释：如何共同使用 validation、整体 SHAP 和整体 SEM", level=1)
    add_paragraph(
        doc,
        "综合各类证据后，可以形成三层解释框架。第一层是 SHAP 和整体 SHAP 给出的预测贡献层：SSRD、TMP、EVA、STRD、PRE/VPD 是恢复时间预测中最稳定的变量。第二层是 ALE/PDP/ICE 给出的响应验证层：这些变量具有明显非线性、阈值和 biome 内部异质性。第三层是 Geodetector/OPGD 与 SEM 给出的空间机制层：TMP、STRD、SSRD、VPD、EVA 等变量不仅在预测中重要，也具有空间分异解释力，并通过辐射-温度-VPD-蒸散-土壤水分路径影响恢复时间。",
    )
    add_paragraph(
        doc,
        "整体结果与分 biome 结果之间存在互补关系。整体 SHAP 和整体 SEM 适合写总体机制，即恢复时间主要由能量输入、热量背景、大气干旱、蒸散状态和事件属性共同控制；分 biome 结果则用于说明这种机制在不同生态系统中如何改变。例如 Grassland/Savanna/Shrubland 更表现出 TMP、SSRD、VPD 的强响应，而 Cropland 更容易受到 EVA、PRE、SMrz 以及人为管理背景的调节。",
    )
    add_paragraph(
        doc,
        "需要避免的解释误区有三点。第一，Geodetector 和 OPGD 的 q 值没有方向性，不能用 q 值判断变量增加会延长还是缩短恢复时间。第二，PDP 曲线在变量共线性强时可能生成生态上不常见的变量组合，因此应优先参考 ALE。第三，整体 pooled 图不能替代 biome-specific 图，因为整体图会把不同生态系统的响应混合在一起，可能掩盖局部阈值和方向差异。",
    )
    doc.add_heading("8. 可写入论文或报告的结论段", level=1)
    add_paragraph(
        doc,
        "综合 SHAP、ALE、ICE、PDP、Geodetector、OPGD 和 SEM 结果表明，GPP 与 RECO 恢复时间主要受辐射能量、热量背景、大气干旱、蒸散状态和事件属性的共同控制。整体 SHAP 分析显示 SSRD 是跨 biome 最稳定的主导预测因子，TMP、EVA、STRD 和 PRE/VPD 构成次级关键变量；ALE 和 PDP 进一步验证这些变量具有明显非线性和阈值响应，ICE 则揭示同一 biome 内部仍存在显著响应异质性。Geodetector 和 OPGD 结果表明，TMP、STRD、SSRD、VPD 和 EVA 也是恢复时间空间分异的重要解释因子，而 SEM 路径进一步说明这些变量通过“辐射-温度-VPD-蒸散-土壤水分-恢复时间”的耦合链条发挥作用。因此，本研究的机制解释不应依赖单一变量或单一方法，而应将 SHAP 的预测贡献、响应曲线的方向阈值、地理探测器的空间解释力和 SEM 的路径结构结合起来，形成对闪旱后 GPP/RECO 恢复过程的综合解释。",
    )


def build_document() -> Path:
    doc = Document()
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name in doc.styles:
            set_east_asian_font(doc.styles[style_name])
    doc.styles["Normal"].font.size = Pt(11)

    title = doc.add_heading("08 Validation 与整体 SHAP-SEM 综合分析", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_intro(doc)
    add_validation_methods(doc)
    add_response_curve_summary(doc)
    add_geodetector_summary(doc)
    add_overall_shap_summary(doc)
    add_overall_sem_summary(doc)
    add_synthesis(doc)

    doc.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    print(build_document())
