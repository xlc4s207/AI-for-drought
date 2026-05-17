#!/usr/bin/env python3
"""Build a SHAP collinearity handling note and grouped SHAP summaries."""

from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from matplotlib import font_manager


ROOT = Path("/home/xulc/flash_drought")
SHAP_DIR = ROOT / "process/SEM_analysis0401/codex/GLEAM/plots2/prepeak_shap_summary_20260502"
COL_DIR = ROOT / "process/SEM_analysis0401/codex/GLEAM/results/overall_shap_results_20260502/collinearity_analysis_20260502"
OUT_DIR = SHAP_DIR / "collinearity_grouped_shap_20260506"
DOCX_OUT = ROOT / "process/SEM_analysis0401/codex/writing3/09_SHAP_collinearity_handling_and_reviewer_response_cn.docx"
MD_OUT = OUT_DIR / "collinearity_handling_for_shap_interpretation_20260506.md"


GROUP_BY_LABEL = {
    "SSRD": "辐射热力组",
    "STRD": "辐射热力组",
    "TMP": "辐射热力组",
    "VPD": "大气需水组",
    "WIND": "大气需水组",
    "Wind": "大气需水组",
    "PRE": "水分供给/储存组",
    "Pre": "水分供给/储存组",
    "EVA": "水分供给/储存组",
    "|EVA|": "水分供给/储存组",
    "SMrz": "水分供给/储存组",
    "Duration": "事件属性组",
    "DUR": "事件属性组",
    "Intensity": "事件属性组",
    "INT": "事件属性组",
    "Onset": "事件属性组",
    "ONS": "事件属性组",
    "LAI": "植被状态组",
}

GROUP_ORDER = ["辐射热力组", "水分供给/储存组", "大气需水组", "事件属性组", "植被状态组"]
GROUP_COLORS = {
    "辐射热力组": "#d55e00",
    "水分供给/储存组": "#0072b2",
    "大气需水组": "#009e73",
    "事件属性组": "#cc79a7",
    "植被状态组": "#7f7f7f",
}


def normalize_label(row: pd.Series) -> str:
    for col in ("display_label", "label"):
        val = row.get(col)
        if pd.notna(val):
            text = str(val)
            if text == "wind":
                return "Wind"
            if text == "Pre":
                return "Pre"
            return text
    return str(row["feature"])


def add_group(df: pd.DataFrame, scope: str) -> pd.DataFrame:
    df = df.copy()
    df["feature_label"] = df.apply(normalize_label, axis=1)
    df["mechanism_group"] = df["feature_label"].map(GROUP_BY_LABEL).fillna("其他")
    df["scope"] = scope
    if "importance" not in df.columns:
        df["importance"] = df.get("mean_abs_shap")
    if "mean_abs_shap" not in df.columns:
        df["mean_abs_shap"] = df.get("importance")
    return df


def build_grouped_shap() -> tuple[pd.DataFrame, pd.DataFrame]:
    overall = pd.read_csv(SHAP_DIR / "overall_all_biomes/overall_all_biomes_feature_importance.csv")
    overall = add_group(overall, "Overall")

    by_biome = pd.read_csv(SHAP_DIR / "shap_importance_percent_bars_5biomes_gpp_vs_reco.csv")
    by_biome = add_group(by_biome, "By biome")
    by_biome["scope"] = by_biome["biome"]

    all_importance = pd.concat([overall, by_biome], ignore_index=True, sort=False)
    grouped = (
        all_importance.groupby(["scope", "metric", "mechanism_group"], as_index=False)
        .agg(mean_abs_shap=("importance", "sum"), percent=("percent", "sum"))
        .sort_values(["scope", "metric", "percent"], ascending=[True, True, False])
    )
    grouped["group_rank"] = grouped.groupby(["scope", "metric"])["percent"].rank(ascending=False, method="first").astype(int)

    all_importance.to_csv(OUT_DIR / "feature_importance_with_mechanism_group.csv", index=False)
    grouped.to_csv(OUT_DIR / "grouped_shap_importance_by_metric_scope.csv", index=False)
    return all_importance, grouped


def plot_grouped_shap(grouped: pd.DataFrame) -> None:
    cjk_font_path = "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"
    if Path(cjk_font_path).exists():
        font_manager.fontManager.addfont(cjk_font_path)
        cjk_font = font_manager.FontProperties(fname=cjk_font_path).get_name()
    else:
        cjk_font = "DejaVu Sans"
    plt.rcParams.update({
        "font.family": cjk_font,
        "axes.spines.top": False,
        "axes.spines.right": False,
        "figure.dpi": 160,
    })

    overall = grouped[grouped["scope"] == "Overall"].copy()
    fig, ax = plt.subplots(figsize=(8.5, 4.6))
    x_positions = []
    labels = []
    values = []
    colors = []
    xpos = 0
    for metric in ["GPP", "RECO"]:
        sub = overall[overall["metric"] == metric].set_index("mechanism_group").reindex(GROUP_ORDER).dropna(subset=["percent"])
        for group, row in sub.iterrows():
            x_positions.append(xpos)
            labels.append(f"{metric}\n{group}")
            values.append(row["percent"])
            colors.append(GROUP_COLORS.get(group, "#333333"))
            xpos += 1
        xpos += 0.8
    ax.bar(x_positions, values, color=colors, width=0.72)
    for x, y in zip(x_positions, values):
        ax.text(x, y + 0.8, f"{y:.1f}%", ha="center", va="bottom", fontsize=8)
    ax.set_ylabel("Grouped mean(|SHAP|) contribution (%)")
    ax.set_title("Overall all-biome SHAP contributions by mechanism group")
    ax.set_xticks(x_positions)
    ax.set_xticklabels(labels, rotation=35, ha="right", fontsize=8)
    ax.set_ylim(0, max(values) * 1.22)
    fig.tight_layout()
    fig.savefig(OUT_DIR / "overall_grouped_shap_importance.png", dpi=300)
    plt.close(fig)

    scopes = [s for s in grouped["scope"].unique() if s != "Overall"]
    for metric in ["GPP", "RECO"]:
        pivot = (
            grouped[(grouped["metric"] == metric) & (grouped["scope"].isin(scopes))]
            .pivot(index="scope", columns="mechanism_group", values="percent")
            .reindex(index=["Forest", "Grassland", "Savanna", "Cropland", "Shrubland"])
            .reindex(columns=GROUP_ORDER)
            .fillna(0.0)
        )
        fig, ax = plt.subplots(figsize=(8.8, 4.8))
        bottom = None
        for group in GROUP_ORDER:
            vals = pivot[group].values
            ax.bar(pivot.index, vals, bottom=bottom, label=group, color=GROUP_COLORS[group])
            bottom = vals if bottom is None else bottom + vals
        ax.set_ylim(0, 100)
        ax.set_ylabel("Grouped mean(|SHAP|) contribution (%)")
        ax.set_title(f"{metric} SHAP contributions by biome and mechanism group")
        ax.legend(loc="upper center", bbox_to_anchor=(0.5, -0.15), ncol=3, frameon=False)
        fig.tight_layout()
        fig.savefig(OUT_DIR / f"{metric.lower()}_by_biome_grouped_shap_importance.png", dpi=300)
        plt.close(fig)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(text)


def add_df_table(doc: Document, df: pd.DataFrame, columns: list[str], headers: list[str], max_rows: int | None = None) -> None:
    data = df[columns].head(max_rows) if max_rows else df[columns]
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for _, row in data.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(columns):
            val = row[col]
            if isinstance(val, float):
                cells[i].text = f"{val:.3f}"
            else:
                cells[i].text = str(val)


def build_docx(grouped: pd.DataFrame) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_heading("SHAP 输入特征共线性的处理方案与审稿应对说明", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("基于 prepeak SHAP、ALE/ICE/PDP、地理探测器与 SEM 机制分析的补充说明")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "一、为什么需要主动处理共线性", 1)
    add_para(doc, "当前 SHAP 模型使用的输入变量具有明确的生态物理耦合关系，例如短波辐射、长波辐射、气温、VPD、蒸散和土壤水分并非彼此独立。若在论文中只展示 SHAP 排名而不说明共线性，审稿人可能会质疑单变量贡献是否被相关变量分摊或替代。这个问题不意味着 SHAP 结果无效，但意味着 SHAP 重要性不能被写成完全独立的因果效应。更稳妥的表述是：SHAP 反映机器学习模型在现有相关结构下对预测的边际贡献分配，而机制解释应结合变量组、依赖图、ALE/ICE/PDP、地理探测器和 SEM。")
    add_para(doc, "因此，本研究不建议删除 SSRD、STRD、TMP、VPD、EVA 或 SMrz 等具有明确机理意义的变量，而是采用“诊断共线性、机制组解释、稳健性验证、SEM 分层路径”的策略，把单变量 SHAP 排名转化为对热辐射、大气需水、水分供给和事件属性等机制组的解释。")

    add_heading(doc, "二、已有共线性诊断结果", 1)
    summary = pd.read_csv(COL_DIR / "all_scenarios_overall_summary.csv")
    add_df_table(
        doc,
        summary,
        ["scenario", "n_samples", "top_vif_label", "top_vif", "top_pearson_pair", "top_pearson_abs_r", "high_vif_features", "moderate_vif_features", "condition_number"],
        ["场景", "样本量", "最高VIF变量", "最高VIF", "最强Pearson组合", "|r|", "高VIF变量", "中等VIF变量", "条件数"],
    )
    add_para(doc, "总体诊断显示，四类场景中最稳定的强共线关系均集中在 TMP 与 STRD 之间，Pearson 相关绝对值约为 0.845-0.923，Spearman 相关绝对值约为 0.847-0.929。GPP_prepeak 的 TMP VIF 最高，达到 38.08，同时 STRD 和 VPD 也进入高 VIF 范围。其他场景中 TMP 和 STRD 仍是主要高 VIF 变量，SSRD、EVA、VPD 在不同场景中多处于中等 VIF 范围。")
    add_para(doc, "这说明模型输入确实包含一条稳定的热辐射耦合链条。对审稿人而言，关键不是证明没有共线性，而是证明我们已经识别该问题，并且没有把相关变量的 SHAP 排名误读为彼此独立的线性效应。")

    add_heading(doc, "三、机制组 SHAP 贡献作为主解释口径", 1)
    overall_group = grouped[grouped["scope"] == "Overall"].sort_values(["metric", "group_rank"])
    add_df_table(
        doc,
        overall_group,
        ["metric", "mechanism_group", "mean_abs_shap", "percent", "group_rank"],
        ["指标", "机制组", "mean(|SHAP|)合计", "贡献百分比", "组内排序"],
    )
    add_para(doc, "为了避免单个相关变量之间争夺解释权，本研究将变量归并为机制组：SSRD、STRD 和 TMP 归入辐射热力组；VPD 和 WIND 归入大气需水组；PRE、EVA 和 SMrz 归入水分供给/储存组；Duration、Intensity 和 Onset 归入事件属性组；LAI 单独作为植被状态组。这样处理后，SHAP 结果的主结论从“某个变量独立控制恢复时间”转为“某类生态水热机制对模型预测贡献最大”。")
    add_para(doc, "在整体 all-biome 结果中，GPP 和 RECO 的主导贡献均来自辐射热力组，其中 SSRD 仍然是单变量排序中的核心变量，但其解释应与 TMP 和 STRD 共同放在能量输入与热状态框架下。水分供给/储存组提供第二层解释，尤其 EVA、PRE 和 SMrz 共同反映水分补给、耗散和根区储水状态。大气需水组主要通过 VPD 和风速表征空气干燥需求。")

    for image_name, caption in [
        ("overall_grouped_shap_importance.png", "图 1. 整体 all-biome 的机制组 SHAP 贡献。"),
        ("gpp_by_biome_grouped_shap_importance.png", "图 2. GPP 分 biome 的机制组 SHAP 贡献。"),
        ("reco_by_biome_grouped_shap_importance.png", "图 3. RECO 分 biome 的机制组 SHAP 贡献。"),
    ]:
        img = OUT_DIR / image_name
        if img.exists():
            doc.add_picture(str(img), width=Inches(6.3))
            p = doc.add_paragraph(caption)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "四、结果解释应如何修改", 1)
    add_para(doc, "第一，在 beeswarm 和重要性柱状图中，仍可以展示单变量 SHAP 排名，因为它是模型解释的直观结果；但正文中不应将 TMP、STRD、SSRD、VPD 的高贡献分别写成完全独立机制。更合适的写法是：SSRD、STRD 和 TMP 共同表明能量输入与热状态是恢复时间预测的主导信息来源；VPD 和 WIND 代表热干背景下的大气需水压力；PRE、EVA 和 SMrz 则描述水分补给、耗散和可利用储水。")
    add_para(doc, "第二，在 dependence plot 中，如果 TMP、STRD 或 VPD 出现相似阈值，应解释为热辐射-大气干燥链条的联合非线性响应，而不是三个互不相关的阈值。对于 PRE、EVA、SMrz 的关系，也应强调它们对应水分供给、蒸散耗散和根区储水的连续过程。")
    add_para(doc, "第三，在 SEM 中不建议把所有 SHAP 前五变量全部平行直连恢复时间。更稳妥的做法是使用分层路径，例如 STRD 指向 TMP，TMP 和 WIND 指向 VPD，PRE 和 VPD 指向 EVA，PRE 和 EVA 指向 SMrz，再由 TMP、VPD、EVA/SMrz 和必要事件属性解释恢复时间。这样可以保留 SSRD 等重要变量，同时降低高相关变量在终点方程中的并列竞争。")

    add_heading(doc, "五、建议增加的稳健性分析", 1)
    add_para(doc, "建议在现有结果基础上增加三类补充分析。第一，保留当前 VIF、Pearson/Spearman 相关矩阵和条件数作为补充材料，并在方法中说明共线性诊断阈值。第二，加入机制组 SHAP 贡献表，作为正文或补充结果，证明主要结论并不依赖某个单变量在相关变量组内的排序。第三，可选做 reduced-feature 或 residualized-feature 敏感性模型：在热辐射组中分别保留 SSRD、TMP 或残差化后的 SSRD/TMP，比较主导机制组和 dependence 方向是否稳定。")
    add_para(doc, "如果时间有限，最小可接受方案是：保留原始 SHAP 图；新增共线性诊断表；新增机制组 SHAP 贡献；在正文中把 SHAP 解释从单变量因果改写为机制组贡献；在 SEM 中采用分层路径而非所有变量平行直连。这个组合已经能够回应大多数关于共线性的审稿意见。")

    add_heading(doc, "六、可直接写入论文的方法段", 1)
    add_para(doc, "为评估输入变量之间的冗余结构，我们在 SHAP 解释前后对所有候选预测因子进行了 Pearson/Spearman 相关分析、方差膨胀因子和条件数诊断。由于气象和生态水文变量之间存在不可避免的物理耦合，特别是短波/长波辐射、气温和 VPD 之间的热辐射-大气需水链条，本研究不将 SHAP 单变量贡献解释为完全独立的因果效应。相反，我们将变量归并为辐射热力、大气需水、水分供给/储存、事件属性和植被状态等机制组，并同时报告单变量和机制组尺度的 mean absolute SHAP contribution。该处理用于减少相关变量之间贡献分摊对生态解释的影响，并为后续 ALE/ICE/PDP、地理探测器和 SEM 机制路径提供一致的解释框架。")

    add_heading(doc, "七、可直接用于回复审稿人的段落", 1)
    add_para(doc, "感谢审稿人指出输入变量共线性可能影响模型解释。我们已补充 Pearson/Spearman 相关、VIF 和条件数诊断，结果显示 TMP-STRD 是最稳定的强相关组合，VPD、EVA 和 SSRD 在部分场景中也存在中等共线性。我们因此修订了 SHAP 解释方式：单变量 SHAP 排名仍用于展示模型预测贡献，但不再被解释为彼此独立的因果效应；正文新增了机制组尺度的 SHAP 汇总，将 SSRD/STRD/TMP 解释为辐射热力机制，将 VPD/WIND 解释为大气需水机制，将 PRE/EVA/SMrz 解释为水分供给和储存机制。此外，SEM 路径分析采用分层机制结构，以减少高度相关变量在同一终点方程中的并列竞争。上述修改使模型解释更符合生态物理过程，也降低了共线性对单变量归因的潜在影响。")

    doc.save(DOCX_OUT)


def build_markdown(grouped: pd.DataFrame) -> None:
    summary = pd.read_csv(COL_DIR / "all_scenarios_overall_summary.csv")
    overall_group = grouped[grouped["scope"] == "Overall"].sort_values(["metric", "group_rank"])

    def csv_block(df: pd.DataFrame) -> str:
        return "```csv\n" + df.to_csv(index=False).strip() + "\n```"

    lines = [
        "# SHAP 输入特征共线性处理说明",
        "",
        "## 核心判断",
        "",
        "当前 SHAP 输入变量确实存在共线性，最稳定的问题是 TMP-STRD 强相关。这个问题不会使模型预测或 SHAP 图本身失效，但会影响单变量贡献在相关变量之间的分配。因此，后续写作应避免把 SHAP 排名解释为完全独立的因果效应，而应采用机制组解释。",
        "",
        "## 共线性诊断摘要",
        "",
        csv_block(summary),
        "",
        "## 整体机制组 SHAP 贡献",
        "",
        csv_block(overall_group[["metric", "mechanism_group", "mean_abs_shap", "percent", "group_rank"]]),
        "",
        "## 最小修改方案",
        "",
        "1. 保留原始 SHAP 图和单变量排序。",
        "2. 新增 VIF、相关矩阵和条件数作为补充材料。",
        "3. 正文使用机制组解释：辐射热力组、大气需水组、水分供给/储存组、事件属性组。",
        "4. SEM 采用分层路径，避免所有高相关变量平行直连恢复时间。",
        "5. 如需更强稳健性，可补 reduced-feature 或 residualized-feature 敏感性模型。",
        "",
    ]
    MD_OUT.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    _, grouped = build_grouped_shap()
    plot_grouped_shap(grouped)
    build_docx(grouped)
    build_markdown(grouped)
    print(f"Wrote {DOCX_OUT}")
    print(f"Wrote {MD_OUT}")
    print(f"Wrote {OUT_DIR / 'grouped_shap_importance_by_metric_scope.csv'}")


if __name__ == "__main__":
    main()
