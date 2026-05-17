#!/usr/bin/env python3
"""Enhance the flash drought definition validation docx with detailed metrics."""

from __future__ import annotations

from pathlib import Path
import shutil

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


HERE = Path("/home/xulc/flash_drought/process/result_analysis/result_weighted/conclusion/flash_drought_definition_validation")
TARGET = HERE / "01_flash_drought_definition_validation.docx"
BACKUP = HERE / "01_flash_drought_definition_validation.before_metric_detail.docx"


def set_cn(run, font: str = "SimSun", size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def style_para(paragraph, heading: bool = False) -> None:
    for run in paragraph.runs:
        set_cn(run, "SimHei" if heading else "SimSun", 12 if heading else 10, heading)
    if not heading:
        paragraph.paragraph_format.first_line_indent = Pt(22)
        paragraph.paragraph_format.line_spacing = 1.18


def insert_paragraph_after(paragraph, text: str, heading: bool = False):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    run = new_para.add_run(text)
    set_cn(run, "SimHei" if heading else "SimSun", 12 if heading else 10, heading)
    if not heading:
        new_para.paragraph_format.first_line_indent = Pt(22)
        new_para.paragraph_format.line_spacing = 1.18
    return new_para


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def insert_table_after(doc: Document, paragraph, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(header)
        set_cn(r, "SimHei", 8, True)
        shade_cell(cell, "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(value))
            set_cn(r, "SimSun", 8)
    paragraph._p.addnext(table._tbl)
    return table


def find_para(doc: Document, startswith: str):
    for p in doc.paragraphs:
        if p.text.strip().startswith(startswith):
            return p
    raise ValueError(f"paragraph not found: {startswith}")


def enhance() -> Path:
    if not BACKUP.exists():
        shutil.copy2(TARGET, BACKUP)
    doc = Document(str(TARGET))
    for p in doc.paragraphs:
        for r in p.runs:
            if r.text:
                set_cn(r, "SimHei" if p.text.strip().startswith(tuple("123456789")) and len(p.text.strip()) < 40 else "SimSun")

    p3 = find_para(doc, "3 GLEAM与ERA5结果的空间一致性验证")
    cursor = p3
    inserts = [
        ("3.1 验证指标与计算口径", True),
        (
            "本研究对 GLEAM 与 ERA5 的一致性验证并不只依赖图面上的热点位置是否相似，而是同时使用频次数值、正频次空间范围和强度分布三个层面的指标进行判定。"
            "其中，频次指标回答“两套资料是否在相同区域识别出更多或更少的骤旱事件”；正频次格点重叠度回答“两套资料是否把相似区域识别为至少发生过骤旱”；"
            "强度指标则回答“在已经发生骤旱的区域，两套资料对事件强弱的量值估计是否一致”。这种分层验证可以避免只看平均频次或只看空间图导致的片面判断。",
            False,
        ),
        (
            "所有空间一致性指标均在共同有效格点上计算。共同有效格点要求 GLEAM 与 ERA5 对应土层的事件频次均为有限值，并且 event_count 不为缺失或无效值。"
            "这样做可以保证相关系数、重叠度和均值统计都来自同一空间样本集合，而不是由于两套资料掩膜范围不同导致虚假的一致或差异。",
            False,
        ),
    ]
    for text, heading in inserts:
        cursor = insert_paragraph_after(cursor, text, heading)

    insert_table_after(
        doc,
        cursor,
        ["指标", "计算对象", "计算方式", "解释含义"],
        [
            ["平均频次", "event_count 网格", "在共同有效陆地格点上计算 event_count 的空间平均值", "衡量某资料整体识别出的骤旱事件数量水平"],
            ["频次空间相关", "GLEAM 与 ERA5 的 event_count", "共同有效格点上将两套 event_count 展开为向量并计算 Pearson 相关系数", "衡量频次高低的空间格局是否一致，值越接近 1 表示热点梯度越相似"],
            ["正频次格点 Jaccard 重叠度", "event_count > 0 的格点集合", "J = |A ∩ B| / |A ∪ B|，A 和 B 分别为 GLEAM 与 ERA5 正频次格点集合", "衡量两套资料识别出的发生范围是否重合，避免只看频次数值大小"],
            ["平均强度", "事件强度或平均强度网格", "在有效格点上计算事件平均强度的空间平均", "衡量资料对骤旱亏缺强弱的量值估计"],
            ["强度空间相关", "GLEAM 与 ERA5 的平均强度场", "在共同有效格点上对两套强度场计算 Pearson 相关系数", "衡量同一位置强度高低排序是否一致"],
        ],
    )
    # Table insertion does not produce a paragraph cursor, use p3's next material by reinserting after the original cursor.
    cursor = insert_paragraph_after(
        cursor,
        "其中，频次空间相关与 Jaccard 重叠度的含义不同。频次空间相关使用所有共同有效格点上的 event_count 数值，关注的是“哪里频次更高、哪里频次更低”的连续空间梯度；"
        "Jaccard 重叠度只关心 event_count 是否大于 0，关注的是“哪些格点被识别为至少发生过骤旱”的空间范围。"
        "因此，一个资料可能识别出更多事件，但只要其正频次格点与另一资料高度重合，Jaccard 值仍然可以较高；这说明两套资料对发生范围判断一致，只是在事件数量幅值上不同。",
        False,
    )

    cursor = insert_paragraph_after(cursor, "3.2 频次空间相关与正频次格点 Jaccard 重叠度", True)
    cursor = insert_paragraph_after(
        cursor,
        "频次验证结果表明，GLEAM 与 ERA5 对全球骤旱发生位置的识别具有较高一致性。根层土壤湿度情景下，GLEAM 的 1980—2024 年平均频次为 7.10 次，"
        "ERA5 为 3.25 次，说明 GLEAM 在根区识别出的事件数量更丰富；但两者在共同有效格点上的频次空间相关仍达到 0.709，正频次格点 Jaccard 重叠度达到 0.823。"
        "这意味着即使 ERA5 的根区事件数量较少，两套资料仍然在大尺度上把相似区域识别为骤旱活跃区。",
        False,
    )
    cursor = insert_paragraph_after(
        cursor,
        "表层土壤湿度情景下，一致性更强。GLEAM 与 ERA5 的平均频次分别为 10.13 次和 11.86 次，二者幅值更接近；频次空间相关达到 0.851，"
        "正频次格点 Jaccard 重叠度达到 0.902。0.902 的 Jaccard 值说明，在至少发生过一次骤旱的表层格点中，两套资料识别出的发生范围高度重叠，"
        "表层骤旱热点的空间范围具有很强的跨资料源稳定性。",
        False,
    )
    insert_table_after(
        doc,
        cursor,
        ["土层", "资料", "平均频次", "平均强度", "频次空间相关", "正频次 Jaccard", "交集格点", "并集格点", "共同有效格点"],
        [
            ["SMrz", "GLEAM", "7.10", "2.17", "0.709", "0.823", "171838", "208690", "900387"],
            ["SMrz", "ERA5", "3.25", "3.24", "0.709", "0.823", "171838", "208690", "900387"],
            ["SMs", "GLEAM", "10.13", "1.58", "0.851", "0.902", "189253", "209858", "900387"],
            ["SMs", "ERA5", "11.86", "1.31", "0.851", "0.902", "189253", "209858", "900387"],
        ],
    )
    cursor = insert_paragraph_after(
        cursor,
        "需要注意，Jaccard 重叠度不是简单的“共同为正格点占全部陆地格点比例”，而是正频次格点交集除以正频次格点并集。以 SMrz 为例，"
        "GLEAM 与 ERA5 的正频次交集为 171838 个格点，并集为 208690 个格点，因此 Jaccard = 171838 / 208690 = 0.823。"
        "这一指标能够过滤掉大量两套资料都没有识别事件的无信息格点，更直接评估两套资料对骤旱发生范围的共同判断。",
        False,
    )

    # Add intensity detail after the paragraph that starts with "强度分布进一步支持".
    p_int = find_para(doc, "强度分布进一步支持")
    cursor = p_int
    for text, heading in [
        ("3.3 强度场一致性及其解释边界", True),
        (
            "与频次场相比，强度场的一致性较弱，这一点需要在文档中明确区分。频次场主要由事件是否发生以及发生次数决定，因而更稳定地反映气候带、植被带和水分限制区的空间结构；"
            "强度场则进一步依赖土壤湿度产品的绝对振幅、土层深度、土壤水容量约束、同化策略和降水输入误差。不同产品即便在相同区域都识别出骤旱，也可能对亏缺强弱给出不同量值。",
            False,
        ),
        (
            "本项目依据 GLEAM 与 ERA5 的平均强度场计算发现，根层平均强度分别为 2.17 和 3.24，表层平均强度分别为 1.58 和 1.31；"
            "但两套产品在共同有效格点上的强度空间相关仅约为 0.121 和 0.127，显著低于频次空间相关。"
            "因此，本研究将强度结果更多用于说明不同资料源对事件亏缺量级的敏感性，而不是要求 GLEAM 与 ERA5 在强度量值上完全一致。",
            False,
        ),
        (
            "这种“频次格局一致、强度量值差异较大”的结果是合理的。它说明骤旱识别方法在判断高风险区域方面具有跨资料源稳健性，但事件强弱的绝对量值仍受资料产品物理结构影响。"
            "因此，后续生态响应分析更适合以 GLEAM 作为主事件库，并将 ERA5 作为独立对照资料，而不是把两者的强度场视作可直接互换的同一量。",
            False,
        ),
    ]:
        cursor = insert_paragraph_after(cursor, text, heading)

    # Add rationale before section 4.
    p4 = find_para(doc, "4 结果合理性的综合判断")
    cursor = p4
    # insert after p4 means appears after heading; that's OK as part of section 4.
    for text, heading in [
        ("4.1 从多指标验证到事件库选择", True),
        (
            "综合频次空间相关、正频次格点 Jaccard 重叠度和强度场对比，可以得到一个更细致的判断：GLEAM 与 ERA5 并非在空间分布上相互矛盾，"
            "而是在事件数量、根区敏感性和强度幅值上存在资料源差异。根区中 GLEAM 平均频次更高，意味着其提供了更丰富的事件样本；表层中两套资料频次和发生范围高度一致，"
            "说明骤旱定义本身具有良好的跨资料源稳定性。",
            False,
        ),
        (
            "从后续碳通量响应研究的角度看，采用 GLEAM 作为主线并不意味着 ERA5 不可靠，而是因为 GLEAM 在根区和恢复阶段能够提供更充分的事件样本和更连续的生态恢复信号。"
            "ERA5 的作用则是作为独立土壤湿度产品验证关键空间格局和定义合理性。这样的设计比单一资料源分析更稳健，也比简单平均两套资料更易解释。",
            False,
        ),
    ]:
        cursor = insert_paragraph_after(cursor, text, heading)

    # Strengthen conclusion at the end.
    last = doc.paragraphs[-1]
    insert_paragraph_after(
        last,
        "更具体地说，频次空间相关说明 GLEAM 与 ERA5 对频次高低梯度的判断较一致，正频次格点 Jaccard 重叠度说明两套资料对“哪里至少发生过骤旱”的空间范围判断高度重合；"
        "强度空间相关较低则提醒我们，骤旱强度是更依赖产品物理结构的量值，不应被过度解释为两套资料必须完全一致。"
        "因此，本研究的验证逻辑是：用 GLEAM 与 ERA5 的高频次相关和高 Jaccard 重叠度证明骤旱定义的空间稳健性，用强度差异说明资料源不确定性，"
        "再结合后续生态响应需求选择 GLEAM 作为主事件库、ERA5 作为独立对照。",
        False,
    )

    doc.save(TARGET)
    return TARGET


if __name__ == "__main__":
    print(enhance())
