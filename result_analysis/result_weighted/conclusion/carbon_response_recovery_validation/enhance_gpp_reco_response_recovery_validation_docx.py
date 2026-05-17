#!/usr/bin/env python3
from pathlib import Path
import shutil

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


BASE = Path(__file__).resolve().parent
DOCX = BASE / "gpp_reco_response_recovery_validation_report.docx"
BACKUP = BASE / "gpp_reco_response_recovery_validation_report.before_gpp_method_detail.docx"


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "SimSun"
    run.font.size = Pt(9)
    if bold:
        run.font.color.rgb = RGBColor(31, 78, 121)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_paragraph_before(target, text, style="Normal"):
    p = target.insert_paragraph_before(text)
    p.style = style
    for run in p.runs:
        run.font.name = "SimSun"
        run.font.size = Pt(10.5)
    return p


def add_heading_before(target, text):
    p = target.insert_paragraph_before(text)
    p.style = "Heading 2"
    for run in p.runs:
        run.font.name = "SimHei"
        run.font.size = Pt(12)
        run.bold = True
    return p


def move_table_before(doc, table, target_paragraph):
    tbl = table._tbl
    target_paragraph._p.addprevious(tbl)


def main():
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)

    doc = Document(DOCX)

    # Replace the short original section-2 explanatory paragraphs, keeping the existing figure and caption.
    old_text_starts = [
        "0401版本对应的是“growing-season recovery gsdays”分析框架。",
        "在事件级指标上，0401版本同时保存了响应起点、峰值时刻和恢复时长等一组紧凑",
        "基线水平的定义采用事件前参考窗口的平均值",
        "在全局统计层面，0401版本给出的结果表明",
    ]
    for para in list(doc.paragraphs):
        text = para.text.strip()
        if any(text.startswith(prefix) for prefix in old_text_starts):
            delete_paragraph(para)

    # First non-caption/image paragraph after section 2 is the existing figure paragraph.
    # It is an empty paragraph that contains the inline image; insert the expanded methods before it.
    image_anchor = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == "图1  BESS 0401版本下GPP与RECO在SMrz和SMs骤旱情景中的全球平均恢复时间空间分布。":
            image_anchor = doc.paragraphs[i - 1]
            break
    if image_anchor is None:
        raise RuntimeError("Could not find figure-1 anchor paragraph.")

    add_heading_before(image_anchor, "2.1 生长季筛选与 growing-season days 的含义")
    add_paragraph_before(
        image_anchor,
        "0401版本对应的是“growing-season recovery gsdays”分析框架，其核心并不是把全部事件按日历日直接统计，而是先判断骤旱是否真正发生在植被活跃期，再用生长季有效日度量恢复过程。具体实现上，程序使用ERA5 2 m气温逐像元、逐年构建生长季掩膜：某年内首次连续5天高于278.15 K（约5℃）的首日被识别为生长季起点；起点之后首次连续5天低于278.15 K之前的最后一日被识别为生长季终点。只有当骤旱持续期内超过一半天数落在该年生长季掩膜内时，事件才会进入GPP/RECO响应与恢复统计。",
    )
    add_paragraph_before(
        image_anchor,
        "这一处理与Lu等（2025）将研究期限定在中国区域4—10月生长季的思路一致，目的都是避免把休眠季、低温限制期或生产力本底极低时期的波动误写成植被对骤旱的恢复差异。不同之处在于，本研究是全球尺度分析，因此不采用固定月份，而采用温度阈值在每个像元上动态确定生长季；这样可以同时兼容北半球、南半球以及不同气候带的物候差异。",
    )

    add_heading_before(image_anchor, "2.2 GPP响应检测：从骤旱事件到生产力损伤")
    add_paragraph_before(
        image_anchor,
        "骤旱事件本身仍来自GLEAM识别的SMrz或SMs flash drought事件库，BESS和FluxSat只作为碳通量响应的数据源。对GPP而言，骤旱响应被定义为负向生产力扰动，即GPP相对其正常状态出现下降。0401版本沿用rel0_abspeak_absrec_c30x095_w420_decline30_d5的主规则：响应判定采用相对异常阈值z <= 0并要求连续5天满足条件，同时要求干旱开始后30天内存在连续5天的下降过程，以排除没有明确生产力衰退过程的事件。程序同时记录从onset start和drought start出发的响应时间，因此输出中会出现t_response_onset_start、t_response_drought_start和t_response_drought_start_from_onset等变量。",
    )
    add_paragraph_before(
        image_anchor,
        "这里的“响应”不是恢复时间的起点，而是说明骤旱事件确实触发了可识别的碳通量异常。换言之，response_detected=1表示该事件在碳通量序列中有足够清晰的负向响应；若没有持续负异常或没有干旱后的持续下降过程，该事件不会被作为有效GPP响应事件用于恢复时间解释。",
    )

    add_heading_before(image_anchor, "2.3 恢复时间从哪里开始算起")
    add_paragraph_before(
        image_anchor,
        "恢复时间的起点不是骤旱onset，也不是drought start，而是碳通量损伤达到峰值的时刻。对GPP而言，程序在响应被确认后继续搜索5日平滑后的绝对GPP序列最低点，并将其记录为t_peak_abs；这代表干旱影响下GPP损失最强、生态系统从“受损加深”转入“恢复尝试”的分界点。最终用于SEM和后续分析的t_recover_to_baseline_abs_peak，就是从这个绝对损伤峰值开始累计到恢复终点的时间。",
    )
    add_paragraph_before(
        image_anchor,
        "这一点与Lu等（2025）的恢复定义高度一致。Lu等将恢复时间定义为GPP达到最大损失到恢复至骤旱前水平之间的时间，并在图示中明确说明恢复阶段始于植被生产损失最大时刻，终于去趋势后的生产力指数重新高于0。本研究没有直接照搬其五日合成异常指数，而是在日尺度BESS/FluxSat框架下使用绝对GPP最低点和事件前基线来实现同一生态含义：恢复时间衡量的是“从最大生产力损伤到恢复至正常水平”的持续时间，而不是整个骤旱事件长度。",
    )

    add_heading_before(image_anchor, "2.4 恢复终点与基线定义")
    add_paragraph_before(
        image_anchor,
        "恢复终点采用事件前基线，而不是事件内部的局部回升。0401版本在事件前窗口中提取灾前GPP状态，其中响应序列保留onset前60天用于构建相对异常和质量控制；绝对恢复判定则使用灾前30天GPP均值并乘以0.95作为有效恢复基线，同时在5日平滑后的绝对GPP序列上要求连续5天达到或超过该基线。相应变量中，gpp_baseline_abs记录GPP的灾前绝对基线，t_recover_to_baseline记录从相对峰值到基线恢复的时间，而t_recover_to_baseline_abs_peak记录从绝对损伤峰值到基线恢复的时间。",
    )
    add_paragraph_before(
        image_anchor,
        "采用0.95倍灾前30天均值并要求连续5天恢复，是为了避免把单日噪声、短暂天气反弹或遥感反演误差误判为真正恢复。0401版本还取消了此前100天恢复上限（norecmax），允许恢复日期跨越季节；但恢复时长本身只累计峰值之后到恢复日期之间落在生长季掩膜内的有效日。因此，一个事件可以在下一段生长季才恢复，但休眠期不会被简单计入GPP生态恢复时长。这对于GPP尤其重要，因为GPP的季节振幅远大于日际扰动，若按日历日硬算，冬季会人为拉长恢复时间。",
    )

    add_heading_before(image_anchor, "2.5 RECO与GPP在同一框架下的可比性")
    add_paragraph_before(
        image_anchor,
        "RECO在0401版本中使用与GPP一致的事件筛选、生长季掩膜、响应—峰值—恢复变量结构和恢复日累计方式，因此GPP与RECO的差异主要来自碳吸收与生态系统呼吸过程本身，而不是来自事件库、时间窗口或恢复定义的差异。需要注意的是，RECO的生态含义与GPP不同：GPP恢复代表植被光合吸收能力回到灾前水平，RECO恢复则反映呼吸通量在干旱扰动后回到灾前状态。二者采用统一框架，有利于比较“碳吸收恢复”和“碳释放恢复”是否同步，以及不同土壤水分层事件对两类碳过程的影响是否一致。",
    )

    add_heading_before(image_anchor, "2.6 与Lu等（2025）方法的对应关系")
    add_paragraph_before(
        image_anchor,
        "Lu等（2025）使用FluxSat GPP五日尺度数据评估中国生态系统骤旱后的GPP恢复时间，核心做法是先计算GPP异常，再把最大负异常到异常恢复为正值之间的时间定义为恢复时间。该论文还剔除了三类不适合解释恢复的情形：负异常期间没有发生骤旱、骤旱开始前GPP异常已经为负、以及负异常只持续一个五日段。本研究的0401版本在数据尺度和空间范围上不同，但方法逻辑是一致的：只保留具有明确骤旱—GPP损伤对应关系的事件，把最大损伤作为恢复起点，把回到灾前正常水平作为恢复终点。",
    )

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    move_table_before(doc, table, image_anchor)
    headers = ["方法要素", "Lu等（2025）", "本研究0401版本", "解释意义"]
    for j, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[j], header, bold=True)
    rows = [
        (
            "生长季限定",
            "中国区域4—10月生长季。",
            "全球逐像元逐年温度阈值生长季：连续5天>5℃开始，之后连续5天<5℃前结束；骤旱持续期>50%落入生长季才保留。",
            "避免把非生长季低生产力波动解释为GPP生态恢复。",
        ),
        (
            "GPP响应对象",
            "FluxSat GPP五日尺度异常。",
            "BESS和FluxSat日尺度GPP，使用GLEAM SMrz/SMs骤旱事件触发统计。",
            "保证事件定义一致，差异主要来自GPP产品或碳过程响应。",
        ),
        (
            "响应判定",
            "GPP异常进入负异常阶段。",
            "相对异常z <= 0连续5天，且干旱开始后30天内有连续5天下降。",
            "确保该事件确有可识别的生产力下降，而不是随机噪声。",
        ),
        (
            "恢复起点",
            "GPP达到最大损失，即负异常最低点。",
            "响应后5日平滑绝对GPP达到最低点，记录为t_peak_abs。",
            "恢复时间衡量从最大损伤后的恢复过程，而非整个骤旱历时。",
        ),
        (
            "恢复终点",
            "GPP异常回到正值，表示达到或超过灾前水平。",
            "5日平滑绝对GPP连续5天达到0.95×灾前30天均值；记录t_recover_to_baseline_abs_peak。",
            "用灾前基线定义真正恢复，减少单日反弹造成的误判。",
        ),
        (
            "恢复天数",
            "五日尺度恢复持续时间。",
            "恢复可跨季节，但只累计峰值到恢复之间的生长季有效日。",
            "避免休眠季日历时间把GPP恢复时长人为拉长。",
        ),
    ]
    for row in rows:
        cells = table.add_row().cells
        for j, value in enumerate(row):
            set_cell_text(cells[j], value)

    add_paragraph_before(
        image_anchor,
        "表1  Lu等（2025）与本研究0401版本GPP恢复时间定义的对应关系。两者的共同核心是：恢复从最大GPP损失开始，到生产力回到灾前正常水平结束；本研究在全球尺度上进一步引入逐像元生长季掩膜和生长季有效日累计。",
    )

    add_paragraph_before(
        image_anchor,
        "在全局统计层面，0401版本给出的结果表明，GPP与RECO在两类土壤层情景下都具有稳定的响应与恢复尺度。以BESS为例，GPP在SMrz与SMs两种骤旱情景下的全球平均恢复时间分别为42.07天和41.50天，长期趋势分别为0.58 d/10a和1.00 d/10a；RECO在SMrz和SMs情景下的平均恢复时间分别为43.61天和41.82天，长期趋势分别为0.71 d/10a和0.89 d/10a。上述量级说明，0401版本既能保留GPP受干旱抑制后的恢复过程，也能在同一事件框架下刻画RECO回到灾前状态的时间尺度。",
    )

    # Add a short concluding clarification to section 4 without replacing existing conclusions.
    section4_anchor = None
    for para in doc.paragraphs:
        if para.text.strip().startswith("因此，本研究认为，0401版本不仅适用于"):
            section4_anchor = para
            break
    if section4_anchor is not None:
        add_paragraph_before(
            section4_anchor,
            "特别需要强调的是，GPP恢复时间的计算起点位于最大生产力损伤，而不是骤旱开始日期；恢复终点位于连续回到灾前基线，而不是曲线出现短暂反弹。这一设置使恢复时间更接近生态韧性的含义，即生态系统在受到骤旱冲击并达到最大功能损失之后，需要多少个生长季有效日才能重新接近事件前生产力水平。BESS与FluxSat在恢复时间量级和长期延长趋势上的一致性，说明该定义不依赖单一GPP产品，具有较好的稳健性。",
        )

    doc.save(DOCX)
    print(f"updated: {DOCX}")
    print(f"backup:  {BACKUP}")


if __name__ == "__main__":
    main()
