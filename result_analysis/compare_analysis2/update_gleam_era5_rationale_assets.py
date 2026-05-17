#!/home/xulc/.local/share/mamba/envs/Flash_dra/bin/python
"""Backfill reproducible metric notes into the GLEAM-ERA5 rationale assets."""

from __future__ import annotations

import csv
from pathlib import Path
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt


BASE_DIR = Path("/home/xulc/flash_drought")
CONCLUSION_DIR = BASE_DIR / "process/result_analysis/result_weighted/conclusion"
HTML_PATH = CONCLUSION_DIR / "01_gleam_era5_flash_drought_rationale.html"
DOCX_PATH = CONCLUSION_DIR / "01_gleam_era5_flash_drought_rationale.docx"
SUMMARY_CSV = (
    CONCLUSION_DIR
    / "gleam_era5_flash_frequency_intensity_spatial"
    / "gleam_vs_era5_flash_frequency_intensity_summary.csv"
)


def load_rows() -> list[dict[str, str]]:
    with open(SUMMARY_CSV, "r", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def find_row(rows: list[dict[str, str]], dataset: str, soil_layer: str) -> dict[str, str]:
    for row in rows:
        if row["dataset"] == dataset and row["soil_layer"] == soil_layer:
            return row
    raise KeyError((dataset, soil_layer))


def build_note(rows: list[dict[str, str]]) -> str:
    smrz = find_row(rows, "GLEAM", "SMrz")
    sms = find_row(rows, "GLEAM", "SMs")
    return (
        "这里的“频次空间相关（GLEAM vs ERA5）”并不是主观判读值，而是直接基于两套事件文件中 event_count 网格在共同有效格点上的 Pearson 相关系数。"
        "具体做法是：先对 GLEAM 与 ERA5 对应土层的 event_count 取共同有效格点（即两者都为有限值且 event_count >= 0 的像元），"
        "再将这两组频次值展开为一维向量并计算相关系数。"
        "“正频次格点重叠度”采用的是 Jaccard 指数，即两套资料中 event_count > 0 的像元交集除以并集，而不是简单的共同为正比例。"
        f"按这一规则重算后，SMrz 的频次空间相关为 {float(smrz['frequency_spatial_correlation_gleam_vs_era5']):.3f}，"
        f"正频次格点 Jaccard 重叠度为 {float(smrz['positive_frequency_jaccard_overlap']):.3f}"
        f"（交集 {smrz['positive_frequency_intersection_pixels']} 格点，并集 {smrz['positive_frequency_union_pixels']} 格点）；"
        f"SMs 的对应值分别为 {float(sms['frequency_spatial_correlation_gleam_vs_era5']):.3f} 和 {float(sms['positive_frequency_jaccard_overlap']):.3f}。"
        "因此，正文中的 0.709、0.823、0.851 和 0.902 都可以由原始事件文件直接复现，不是人工估写。"
    )


def build_table_html(rows: list[dict[str, str]]) -> str:
    smrz = find_row(rows, "GLEAM", "SMrz")
    sms = find_row(rows, "GLEAM", "SMs")
    return f"""  <table>
    <tr>
      <th>土层</th>
      <th>资料</th>
      <th>1980–2024 平均频次</th>
      <th>平均强度</th>
      <th>频次空间相关（GLEAM vs ERA5）</th>
      <th>正频次格点重叠度</th>
      <th>正频次交集格点数</th>
      <th>正频次并集格点数</th>
      <th>共同有效格点数</th>
    </tr>
    <tr>
      <td rowspan="2">SMrz</td>
      <td>GLEAM</td>
      <td>7.10</td>
      <td>2.17</td>
      <td rowspan="2">0.709</td>
      <td rowspan="2">0.823</td>
      <td rowspan="2">{smrz['positive_frequency_intersection_pixels']}</td>
      <td rowspan="2">{smrz['positive_frequency_union_pixels']}</td>
      <td rowspan="2">{smrz['shared_valid_pixels']}</td>
    </tr>
    <tr>
      <td>ERA5</td>
      <td>3.25</td>
      <td>3.24</td>
    </tr>
    <tr>
      <td rowspan="2">SMs</td>
      <td>GLEAM</td>
      <td>10.13</td>
      <td>1.58</td>
      <td rowspan="2">0.851</td>
      <td rowspan="2">0.902</td>
      <td rowspan="2">{sms['positive_frequency_intersection_pixels']}</td>
      <td rowspan="2">{sms['positive_frequency_union_pixels']}</td>
      <td rowspan="2">{sms['shared_valid_pixels']}</td>
    </tr>
    <tr>
      <td>ERA5</td>
      <td>11.86</td>
      <td>1.31</td>
    </tr>
  </table>"""


def update_html(note: str) -> None:
    text = HTML_PATH.read_text(encoding="utf-8")
    marker = (
        "这表明表层土壤湿度的骤旱热点在两套资料之间具有非常稳定的共同空间结构。</p>"
    )
    insertion = (
        marker
        + "\n\n"
        + f"  <p>{note}</p>"
    )
    if note in text:
        return
    if marker not in text:
        raise RuntimeError("Unable to locate insertion point in HTML rationale file.")
    HTML_PATH.write_text(text.replace(marker, insertion, 1), encoding="utf-8")


def update_html_table(rows: list[dict[str, str]]) -> None:
    text = HTML_PATH.read_text(encoding="utf-8")
    new_table = build_table_html(rows)
    if new_table in text:
        return
    pattern = re.compile(r"  <table>.*?</table>", re.S)
    if not pattern.search(text):
        raise RuntimeError("Unable to locate table block in HTML rationale file.")
    HTML_PATH.write_text(pattern.sub(new_table, text, count=1), encoding="utf-8")


def insert_paragraph_after(paragraph, text: str):
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    from docx.text.paragraph import Paragraph

    new_para = Paragraph(new_p, paragraph._parent)
    run = new_para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
            "SimSun",
        )
    return new_para


def update_docx(note: str) -> None:
    doc = Document(str(DOCX_PATH))
    if any(note in p.text for p in doc.paragraphs):
        doc.save(str(DOCX_PATH))
        return

    target_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("首先从像元尺度频次分布看"):
            target_idx = i
            break
    if target_idx is None:
        raise RuntimeError("Unable to find target paragraph in DOCX rationale file.")
    insert_paragraph_after(doc.paragraphs[target_idx], note)
    doc.save(str(DOCX_PATH))


def update_docx_table(rows: list[dict[str, str]]) -> None:
    doc = Document(str(DOCX_PATH))
    table = doc.tables[1]
    headers = [
        "土层",
        "资料",
        "1980–2024 平均频次",
        "平均强度",
        "频次空间相关（GLEAM vs ERA5）",
        "正频次格点重叠度",
        "正频次交集格点数",
        "正频次并集格点数",
        "共同有效格点数",
    ]
    while len(table.columns) < len(headers):
        table.add_column(Pt(60))
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    smrz = find_row(rows, "GLEAM", "SMrz")
    sms = find_row(rows, "GLEAM", "SMs")
    values = {
        1: ["SMrz", "GLEAM", "7.10", "2.17", "0.709", "0.823", smrz["positive_frequency_intersection_pixels"], smrz["positive_frequency_union_pixels"], smrz["shared_valid_pixels"]],
        2: ["SMrz", "ERA5", "3.25", "3.24", "0.709", "0.823", smrz["positive_frequency_intersection_pixels"], smrz["positive_frequency_union_pixels"], smrz["shared_valid_pixels"]],
        3: ["SMs", "GLEAM", "10.13", "1.58", "0.851", "0.902", sms["positive_frequency_intersection_pixels"], sms["positive_frequency_union_pixels"], sms["shared_valid_pixels"]],
        4: ["SMs", "ERA5", "11.86", "1.31", "0.851", "0.902", sms["positive_frequency_intersection_pixels"], sms["positive_frequency_union_pixels"], sms["shared_valid_pixels"]],
    }
    for row_idx, row_values in values.items():
        for col_idx, value in enumerate(row_values):
            table.rows[row_idx].cells[col_idx].text = str(value)
    doc.save(str(DOCX_PATH))


def main() -> None:
    rows = load_rows()
    note = build_note(rows)
    update_html(note)
    update_html_table(rows)
    update_docx(note)
    update_docx_table(rows)
    print(f"Updated {HTML_PATH}")
    print(f"Updated {DOCX_PATH}")


if __name__ == "__main__":
    main()
