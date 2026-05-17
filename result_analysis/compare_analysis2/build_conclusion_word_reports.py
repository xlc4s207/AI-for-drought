#!/home/xulc/.local/share/mamba/envs/Flash_dra/bin/python
"""Build two conclusion Word reports with embedded figures."""

from __future__ import annotations

import csv
import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE_DIR = Path("/home/xulc/flash_drought")
CONCLUSION_DIR = BASE_DIR / "process/result_analysis/result_weighted/conclusion"

DOC1_DIR = CONCLUSION_DIR / "flash_drought_definition_validation"
DOC2_DIR = CONCLUSION_DIR / "carbon_response_recovery_validation"

DOC1_PATH = DOC1_DIR / "01_flash_drought_definition_validation.docx"
DOC2_PATH = DOC2_DIR / "02_gpp_reco_response_recovery_validation.docx"

GLEAM_ERA5_DIR = CONCLUSION_DIR / "gleam_era5_flash_frequency_intensity_spatial"
BESS_FLUXSAT_DIR = CONCLUSION_DIR / "BESS_Fluxsat_valid"

SUMMARY_0401 = (
    BASE_DIR
    / "process/result_analysis/result_weighted/compare_analysis2/v20260401_growingseason_recovery_gsdays/"
    "summary_table_v20260401_growingseason_recovery_gsdays.csv"
)
SUMMARY_FLUXSAT = (
    BASE_DIR
    / "process/result_analysis/result_weighted/fluxsat_compare_analysis2/fluxsat_0401_sensitivity_compare/"
    "fluxsat_0401_sensitivity_summary.csv"
)
SUMMARY_DROUGHT = GLEAM_ERA5_DIR / "gleam_vs_era5_flash_frequency_intensity_summary.csv"


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    for style_name, size in [("Title", 18), ("Heading 1", 14), ("Heading 2", 12)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.bold = True
        style.font.size = Pt(size)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
    return doc


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Title"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def add_figure(doc: Document, image_path: Path, caption: str, width_in: float = 6.3) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(image_path), width=Inches(width_in))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(8)
    run = cp.add_run(caption)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def read_csv(path: Path) -> list[dict[str, str]]:
    with open(path, "r", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def find_row(rows: list[dict[str, str]], **filters: str) -> dict[str, str]:
    for row in rows:
        if all(row.get(k) == v for k, v in filters.items()):
            return row
    raise KeyError(filters)


def f2(value: str) -> str:
    return f"{float(value):.2f}"


def build_doc1() -> None:
    DOC1_DIR.mkdir(parents=True, exist_ok=True)
    rows = read_csv(SUMMARY_DROUGHT)
    gleam_smrz = find_row(rows, dataset="GLEAM", soil_layer="SMrz")
    era5_smrz = find_row(rows, dataset="ERA5", soil_layer="SMrz")
    gleam_sms = find_row(rows, dataset="GLEAM", soil_layer="SMs")
    era5_sms = find_row(rows, dataset="ERA5", soil_layer="SMs")

    doc = setup_document()
    add_title(doc, "骤旱识别与多源验证技术说明")

    add_heading(doc, "1 研究目的与数据基础", level=1)
    add_paragraph(
        doc,
        "本说明的目标是对本研究所采用的骤旱识别方案进行方法性说明，并对基于GLEAM与ERA5两套土壤湿度产品得到的空间分布结果进行一致性验证。为增强结论的稳健性，本研究分别使用GLEAM根系土壤湿度与表层土壤湿度产品，以及ERA5-Land根层土壤湿度与表层土壤湿度产品，在统一的0.25°空间分辨率、统一的陆地掩膜条件下识别骤旱事件，并据此比较两套驱动资料在全球尺度上的频次与强度分布特征。"
    )
    add_paragraph(
        doc,
        "从变量对应关系上看，SMrz用于表征根层土壤湿度变化，SMs用于表征表层土壤湿度变化；GLEAM与ERA5分别代表不同来源的土壤湿度资料体系。由于两套资料的物理结构、资料同化方式和垂向土层定义并不完全一致，因此本研究并不要求其在像元级事件数量上绝对相等，而是重点考察它们在全球高风险区域识别、主要空间梯度以及大陆尺度热点位置上的一致性。"
    )

    add_heading(doc, "2 骤旱定义原则与事件计算方法", level=1)
    add_paragraph(
        doc,
        "本研究使用的是基于土壤湿度百分位阈值的两阶段骤旱识别框架，即Two-step Method v5.4。百分位基线采用1981—2010年多年参考期，并使用5天滑动时间窗构建同历日百分位阈值。首先，当土壤湿度下降到P20以下时，像元进入干旱状态；其次，若土壤湿度从相对正常状态快速下降并在20天以内完成骤降至干旱阈值以下，则该过程被识别为骤旱事件。实现上，本研究将1—4天快速型事件与5—20天闪旱型事件合并为统一的flash_lt20事件集，从而保证骤旱定义在不同产品间的一致性。"
    )
    add_paragraph(
        doc,
        "在单个事件尺度上，事件文件同时记录了onset_days、duration、days_below_p20、onset_drop、onset_rate和intensity等关键特征。其中，onset_days表示从快速衰减开始到进入干旱阈值所经历的时间；duration表示干旱持续时间；days_below_p20表示事件期间低于P20阈值的总天数；onset_drop表示爆发阶段的土壤湿度下降幅度；onset_rate为下降幅度与爆发时长之比；intensity则表示事件期间土壤湿度低于P20阈值的累积亏缺，是刻画事件强度的核心指标。需要强调的是，onset_rate用于事件刻画而非额外的硬判定阈值，因此最终的骤旱识别仍以土壤湿度百分位快速跌落与低于P20阈值为核心。"
    )
    add_paragraph(
        doc,
        "为避免非研究区域对结果的干扰，本研究在事件文件层面统一剔除了冰雪覆盖区和荒漠/裸地区像元，仅保留具有生态学意义的陆地像元进行分析。因此，后续的全球频次图和强度图实际上反映的是可分析陆地范围内的骤旱空间分布，而不是未经筛选的全球所有网格。"
    )

    add_heading(doc, "3 GLEAM与ERA5结果的空间一致性验证", level=1)
    add_paragraph(
        doc,
        "空间分布图显示，GLEAM与ERA5虽然在事件数量幅值上存在差别，但在骤旱高发区和高强度区的空间位置上具有明显一致性。以发生频次为例，两套资料均将北半球中高纬大陆带、欧亚大陆中部至东部、东亚季风区、北美东部和西部过渡带、南美东南部、南部非洲以及澳大利亚东南部识别为骤旱较为活跃的区域。对于根层土壤湿度情景，GLEAM给出的全球平均频次约为"
        f"{f2(gleam_smrz['mean_frequency_1980_2024'])}次，而ERA5约为{f2(era5_smrz['mean_frequency_1980_2024'])}次；对于表层土壤湿度情景，GLEAM约为{f2(gleam_sms['mean_frequency_1980_2024'])}次，ERA5约为{f2(era5_sms['mean_frequency_1980_2024'])}次。这样的差异反映的是资料源与土层定义差别所导致的幅值敏感性，而非空间结构上的根本矛盾。"
    )
    add_figure(
        doc,
        GLEAM_ERA5_DIR / "gleam_vs_era5_flash_frequency_global.png",
        "图1  GLEAM与ERA5驱动下全球骤旱发生频次空间分布比较。海洋与无效像元已被遮罩，仅显示陆地区域。",
    )
    add_paragraph(
        doc,
        "强度分布进一步支持了这种一致性。无论是GLEAM还是ERA5，强骤旱区均主要出现在水热条件波动较大、土壤水分亏缺易快速累积的过渡带和半干旱—半湿润过渡区，例如欧亚大陆中纬带、南美东南部、南部非洲及部分季风边缘区。虽然不同资料对局地强度峰值的估计有所不同，但高强度区并非随机散布，而是沿着相似的气候—植被梯度集中出现，这说明本研究所采用的骤旱定义能够在不同土壤湿度产品中稳定提取出相近的生态水文风险带。"
    )
    add_figure(
        doc,
        GLEAM_ERA5_DIR / "gleam_vs_era5_flash_intensity_global.png",
        "图2  GLEAM与ERA5驱动下全球骤旱平均强度空间分布比较。高强度区在不同资料源间表现出相近的大陆尺度热点结构。",
    )
    add_paragraph(
        doc,
        "总体而言，GLEAM与ERA5之间的差异主要体现在事件幅值和局部细节上，而在大陆尺度的空间范围、主要热点位置以及高频高强度带的组织方式上保持了良好的可比性。考虑到两套资料在资料来源、垂向结构和同化约束上的天然差异，这种“空间格局一致、绝对数量有别”的结果本身正是方法合理性的体现。它表明，本研究提出的骤旱识别原则不是依赖某一特定资料源才成立，而是能够在独立的土壤湿度产品上重复得到相似的风险空间结构。"
    )

    add_heading(doc, "4 结果合理性的综合判断", level=1)
    add_paragraph(
        doc,
        "从方法论角度看，本研究的骤旱识别体系具备三方面合理性。首先，识别原则建立在同历日百分位阈值和快速跌落过程之上，能够同时约束季节背景与发展速度，避免把普通缓慢转干过程误判为骤旱。其次，GLEAM与ERA5两套独立土壤湿度资料在根层和表层两个土层上均可得到结构相近的全球分布图，说明方法具有跨资料源稳健性。再次，频次与强度图所呈现的高风险区与既有干旱易发带和水分限制生态区高度吻合，进一步证明本研究得到的骤旱事件能够反映真实且有生态意义的全球水分胁迫格局。"
    )
    add_paragraph(
        doc,
        "因此，本研究认为，基于GLEAM和ERA5分别构建的根层与表层骤旱事件库可以作为后续生态系统碳通量响应分析的可信事件基础。其中，GLEAM与ERA5的并行计算不仅起到了结果验证作用，也提高了后续GPP与RECO响应结论的可解释性与外部稳健性。"
    )
    doc.save(str(DOC1_PATH))


def build_doc2() -> None:
    DOC2_DIR.mkdir(parents=True, exist_ok=True)
    rows_0401 = read_csv(SUMMARY_0401)
    rows_flux = read_csv(SUMMARY_FLUXSAT)

    gpp_smrz = find_row(rows_0401, variable="GPP", code="code1", soil_layer="SMrz")
    gpp_sms = find_row(rows_0401, variable="GPP", code="code2", soil_layer="SMs")
    reco_smrz = find_row(rows_0401, variable="RECO", code="code1", soil_layer="SMrz")
    reco_sms = find_row(rows_0401, variable="RECO", code="code2", soil_layer="SMs")
    flux_smrz = find_row(rows_flux, dataset="FluxSat 0401 rec100cap", code="code1", soil_layer="SMrz")
    flux_sms = find_row(rows_flux, dataset="FluxSat 0401 rec100cap", code="code2", soil_layer="SMs")
    bess_smrz = find_row(rows_flux, dataset="BESS 0401", code="code1", soil_layer="SMrz")
    bess_sms = find_row(rows_flux, dataset="BESS 0401", code="code2", soil_layer="SMs")

    doc = setup_document()
    add_title(doc, "GPP与RECO骤旱响应恢复计算及BESS-FluxSat验证说明")

    add_heading(doc, "1 分析目的与总体框架", level=1)
    add_paragraph(
        doc,
        "本说明围绕两个问题展开。其一，说明0401版本下GPP与RECO在骤旱事件后的响应与恢复是如何计算的；其二，利用BESS与FluxSat两套独立GPP资料在长期恢复趋势上的对比结果，验证本研究对“骤旱后GPP恢复时间”的度量是否具有合理性。这里的核心思想是：先用统一定义的GLEAM骤旱事件驱动碳通量分析，再通过多资料对照考察恢复时间的数量级、空间分布和长期趋势是否稳定，从而检验指标构建的可靠性。"
    )

    add_heading(doc, "2 0401版本中GPP与RECO响应/恢复时间的计算原则", level=1)
    add_paragraph(
        doc,
        "0401版本对应的是“growing-season recovery gsdays”分析框架。该版本保留与年生长季具有明确重叠关系的骤旱事件，用于避免在植被非活跃期统计到不具生态响应意义的碳通量波动；同时，恢复过程允许跨季节延续，但恢复时长只累计生长季有效天数，以保证恢复时间更接近生态系统在有效生长条件下完成修复所需要的实际时长。"
    )
    add_paragraph(
        doc,
        "在事件级指标上，0401版本同时保存了响应起点、峰值时刻和恢复时长等一组紧凑而彼此兼容的时间变量。具体而言，response_detected用于标记某次骤旱后是否出现明确的碳通量响应；t_response_onset_start和t_response_drought_start分别表示相对于骤旱发展起点和干旱开始时刻的响应时间；t_peak与t_peak_drought_start用于刻画达到峰值损伤的时间位置；t_recover_to_baseline则表示峰值损伤后恢复到基线水平所需的时间，是本文讨论恢复过程的核心指标。"
    )
    add_paragraph(
        doc,
        "基线水平的定义采用事件前参考窗口的平均值，这一点在GPP与RECO的绝对量变量中都可以得到印证，例如gpp_baseline_abs和reco_baseline_abs分别记录了骤旱前的碳吸收或生态系统呼吸基线。换言之，0401版本的恢复并不是简单地回到事件期间的局部波动水平，而是要求碳通量从峰值损伤状态重新回到骤旱发生前的正常背景，这使得t_recover_to_baseline具有清晰的生态学含义。"
    )
    add_paragraph(
        doc,
        "在全局统计层面，0401版本给出的结果表明，GPP与RECO在两类土壤层情景下都具有稳定的响应与恢复尺度。以BESS为例，GPP在SMrz与SMs两种骤旱情景下的全球平均恢复时间分别为"
        f"{f2(gpp_smrz['recovery_mean_days'])}天和{f2(gpp_sms['recovery_mean_days'])}天，长期趋势分别为{f2(gpp_smrz['recovery_mean_slope_days_per_decade'])} d/10a和{f2(gpp_sms['recovery_mean_slope_days_per_decade'])} d/10a；RECO对应的全球平均恢复时间分别为{f2(reco_smrz['recovery_mean_days'])}天和{f2(reco_sms['recovery_mean_days'])}天，趋势分别为{f2(reco_smrz['recovery_mean_slope_days_per_decade'])} d/10a和{f2(reco_sms['recovery_mean_slope_days_per_decade'])} d/10a。上述结果说明，0401版本能够在不同碳通量变量之间产生结构一致、数量级合理的恢复时长估计。"
    )
    add_figure(
        doc,
        BESS_FLUXSAT_DIR / "bess_0401_gpp_reco_recovery_mean_global.png",
        "图1  BESS 0401版本下GPP与RECO在SMrz和SMs骤旱情景中的全球平均恢复时间空间分布。",
    )

    add_heading(doc, "3 BESS与FluxSat长期恢复趋势对比及其解释", level=1)
    add_paragraph(
        doc,
        "为了进一步检验GPP恢复时间指标的合理性，本研究使用与BESS独立的数据产品FluxSat进行交叉验证。两套资料在事件驱动上保持一致，即均由同一套GLEAM骤旱事件库触发统计，因此它们之间的差异主要反映GPP产品本身的观测/反演特性，而不是事件定义改变带来的混杂。对比结果显示，BESS与FluxSat在长期趋势方向上具有较好一致性：在SMrz情景中，BESS的全球平均恢复时间趋势为"
        f"{f2(bess_smrz['recovery_mean_slope_days_per_decade'])} d/10a，FluxSat为{f2(flux_smrz['recovery_mean_slope_days_per_decade'])} d/10a；在SMs情景中，BESS为{f2(bess_sms['recovery_mean_slope_days_per_decade'])} d/10a，FluxSat为{f2(flux_sms['recovery_mean_slope_days_per_decade'])} d/10a。也就是说，两套独立GPP资料都支持“骤旱后恢复时间总体呈延长趋势”这一判断。"
    )
    add_paragraph(
        doc,
        "在绝对恢复时长上，FluxSat略低于BESS，但数量级保持接近。SMrz情景下，BESS与FluxSat的全球平均恢复时间分别为"
        f"{f2(bess_smrz['recovery_mean_days'])}天和{f2(flux_smrz['recovery_mean_days'])}天；SMs情景下分别为{f2(bess_sms['recovery_mean_days'])}天和{f2(flux_sms['recovery_mean_days'])}天。两者之间约2—4天的差距说明不同GPP产品对碳恢复过程的振幅与平滑程度存在一定敏感性，但这种差距并未改变恢复时长的总体数量级，也未改变长期延长这一方向性结论。"
    )
    add_figure(
        doc,
        BESS_FLUXSAT_DIR / "fluxsat_0401_sensitivity_recovery_trend.png",
        "图2  BESS与FluxSat在SMrz和SMs骤旱情景下的全球平均GPP恢复时间年度变化趋势比较。两套资料均呈现恢复时间延长趋势。",
    )
    add_paragraph(
        doc,
        "空间对比结果也支持这一判断。BESS与FluxSat在全球恢复时间高值区和低值区的分布上具有明显对应关系，尤其是在热带—亚热带季节性水分限制区、部分南半球草地/灌丛区以及欧亚大陆中纬带，较长恢复时间的主要空间位置在两套资料之间基本一致。虽然FluxSat的空间分布整体更平滑、局部极值略弱，但其并未改变恢复时间高值区的主要地理组织方式，这说明本研究的恢复时间指标并不是BESS单一数据产品特有的产物，而是可以在独立GPP资料中被重复识别。"
    )
    add_figure(
        doc,
        BESS_FLUXSAT_DIR / "smrz_recovery_mean_bess0401_vs_fluxsat_fixlon.png",
        "图3  SMrz骤旱情景下BESS与FluxSat平均恢复时间空间分布对比。",
    )
    add_figure(
        doc,
        BESS_FLUXSAT_DIR / "sms_recovery_mean_bess0401_vs_fluxsat_fixlon.png",
        "图4  SMs骤旱情景下BESS与FluxSat平均恢复时间空间分布对比。",
    )

    add_heading(doc, "4 对GPP恢复时间计算合理性的综合判断", level=1)
    add_paragraph(
        doc,
        "综合0401版本的事件定义、恢复指标结构、BESS与RECO结果的一致数量级以及FluxSat对长期趋势方向的独立支持，可以认为本研究关于“骤旱后GPP恢复时间”的计算是合理且具有生态解释力的。首先，恢复定义以骤旱前基线为参照，而非以事件内部任意波动为参照，因此t_recover_to_baseline能够反映生态系统恢复到正常生产力水平所需的真实时间。其次，0401版本通过生长季有效天数统计避免了非活跃季节对恢复时长的人为拉长，使恢复时间更贴近植被有效修复过程。再次，FluxSat与BESS在长期变化方向上的一致性表明，恢复时间延长并不是某一数据集独有的算法假象，而是可以在不同GPP产品中重复观察到的共同信号。"
    )
    add_paragraph(
        doc,
        "因此，本研究认为，0401版本不仅适用于描述骤旱后GPP与RECO响应—恢复的一般规律，而且为解释长期碳恢复时间变化趋势提供了较可靠的方法基础。基于该框架得到的结果，可以用于后续论文中关于生态系统碳功能脆弱性、恢复拖延和区域差异的讨论。"
    )
    doc.save(str(DOC2_PATH))


def main() -> None:
    build_doc1()
    build_doc2()
    print(f"Wrote {DOC1_PATH}")
    print(f"Wrote {DOC2_PATH}")


if __name__ == "__main__":
    main()
