#!/home/xulc/.local/share/mamba/envs/Flash_dra/bin/python
"""Build RECO 0401 recovery Word report."""

from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE_DIR = Path("/home/xulc/flash_drought")
OUT_DIR = BASE_DIR / "process/result_analysis/result_weighted/conclusion/RECO_recovery_valid"
DOCX_PATH = OUT_DIR / "03_reco_0401_recovery_validation.docx"
SUMMARY_CSV = OUT_DIR / "reco_0401_recovery_summary.csv"
SPATIAL_PNG = OUT_DIR / "reco_0401_recovery_mean_global.png"
TREND_PNG = OUT_DIR / "reco_0401_recovery_trend.png"


def read_csv(path: Path) -> list[dict[str, str]]:
    with open(path, "r", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def find_row(rows: list[dict[str, str]], soil_layer: str) -> dict[str, str]:
    for row in rows:
        if row["soil_layer"] == soil_layer:
            return row
    raise KeyError(soil_layer)


def fmt(v: str) -> str:
    return f"{float(v):.2f}"


def setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    for style_name, size in [("Title", 18), ("Heading 1", 14)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.bold = True
        style.font.size = Pt(size)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
    return doc


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Title"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text).bold = True


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def add_figure(doc: Document, image_path: Path, caption: str, width: float = 6.2) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cp.add_run(caption)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def main() -> None:
    rows = read_csv(SUMMARY_CSV)
    smrz = find_row(rows, "SMrz")
    sms = find_row(rows, "SMs")

    doc = setup_doc()
    add_title(doc, "RECO骤旱恢复时间空间分布与趋势变化说明")

    doc.add_heading("1 分析目的", level=1)
    add_para(
        doc,
        "本说明针对RECO在0401版本下受到骤旱影响后的恢复过程进行单独总结，重点讨论两方面内容，即根层土壤湿度驱动的SMrz骤旱与表层土壤湿度驱动的SMs骤旱所对应的全球恢复时间空间分布，以及1982—2021年期间全球平均恢复时间的长期变化趋势。该说明与前面针对GPP的分析口径保持一致，但这里不再引入FluxSat对比，而是聚焦BESS RECO结果本身的空间结构与时间演变。"
    )

    doc.add_heading("2 空间分布特征", level=1)
    add_para(
        doc,
        "RECO恢复时间的全球空间分布表明，SMrz与SMs两种骤旱情景在大尺度空间结构上具有较高一致性。两种情景下，较长恢复时间主要集中于低纬到副热带水分受限区域、季风边缘区以及部分半干旱—半湿润过渡带，而高纬湿润区和部分寒冷区域的恢复时间相对较短。就总体水平而言，SMrz情景的全球平均恢复时间为"
        f"{fmt(smrz['global_mean_days'])}天，略高于SMs情景的{fmt(sms['global_mean_days'])}天；两者的全球中位数均约为14天，说明全球大部分事件的典型恢复过程处于相近量级，但在高值区和尾部分布上SMrz更容易出现更长的恢复时长。"
    )
    add_para(
        doc,
        "从空间统计量看，SMrz情景的空间平均恢复时间约为"
        f"{fmt(smrz['spatial_mean_days'])}天，95百分位约为{fmt(smrz['spatial_p95_days'])}天；SMs情景的空间平均恢复时间约为{fmt(sms['spatial_mean_days'])}天，95百分位约为{fmt(sms['spatial_p95_days'])}天。该结果说明，虽然两类事件在全球平均水平上差距不大，但SMrz在局地高值区更容易出现恢复拖延，从而表现出更高的空间尾部值。"
    )
    add_figure(
        doc,
        SPATIAL_PNG,
        "图1  RECO在SMrz与SMs骤旱情景下的全球平均恢复时间空间分布。",
    )

    doc.add_heading("3 长期趋势变化", level=1)
    add_para(
        doc,
        "RECO的全球平均恢复时间在两类骤旱情景下均表现为缓慢上升。SMrz情景的趋势斜率为"
        f"{fmt(smrz['trend_days_per_decade'])} d/10a，SMs情景的趋势斜率为{fmt(sms['trend_days_per_decade'])} d/10a，后者略高于前者。这说明从长期变化上看，RECO在表层土壤湿度驱动的骤旱事件下，恢复时间延长的幅度略更明显。"
    )
    add_para(
        doc,
        "从首尾年份对比看，SMrz情景的全球平均恢复时间由"
        f"{fmt(smrz['first_year_mean'])}天增加到{fmt(smrz['last_year_mean'])}天；SMs情景则由{fmt(sms['first_year_mean'])}天变化到{fmt(sms['last_year_mean'])}天。尽管年际波动仍然存在，尤其在若干年份中会出现阶段性回落，但线性拟合结果显示两条序列总体都沿着上升方向变化，说明RECO恢复过程在长期上存在一定程度的拖延。"
    )
    add_figure(
        doc,
        TREND_PNG,
        "图2  RECO在SMrz与SMs骤旱情景下的全球平均恢复时间年度变化及线性趋势。",
    )

    doc.add_heading("4 结果解释", level=1)
    add_para(
        doc,
        "RECO恢复时间整体偏长，且在长期上呈现缓慢延长趋势，说明骤旱不仅会影响生态系统瞬时呼吸过程，还可能通过改变土壤水热环境、底物供给和植被生理状态，使生态系统呼吸在事件结束后仍经历较长的调整阶段。SMrz与SMs结果方向一致而幅度略有差异，进一步表明该恢复信号并非依赖单一土层定义，而是在不同土壤湿度层次上均可被识别。"
    )
    add_para(
        doc,
        "因此，0401版本下RECO的恢复时间分布与趋势结果可以作为后续讨论生态系统呼吸脆弱性和恢复拖延的重要依据。若将其与GPP恢复过程联合解释，则有助于从碳吸收与碳释放两个维度共同理解骤旱对生态系统碳循环的持续影响。"
    )

    doc.save(str(DOCX_PATH))
    print(f"Wrote {DOCX_PATH}")


if __name__ == "__main__":
    main()
