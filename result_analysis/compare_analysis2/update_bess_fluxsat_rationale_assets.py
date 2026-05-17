#!/home/xulc/.local/share/mamba/envs/Flash_dra/bin/python
"""Backfill hotspot-overlap notes into the BESS-FluxSat rationale assets."""

from __future__ import annotations

import csv
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt
import netCDF4 as nc
import numpy as np


BASE_DIR = Path("/home/xulc/flash_drought")
CONCLUSION_DIR = BASE_DIR / "process/result_analysis/result_weighted/conclusion"
HTML_PATH = CONCLUSION_DIR / "02_bess_fluxsat_gpp_reco_rationale.html"
DOCX_PATH = CONCLUSION_DIR / "02_bess_fluxsat_gpp_reco_rationale.docx"
SUMMARY_CSV = (
    BASE_DIR
    / "process/result_analysis/result_weighted/fluxsat_compare_analysis2/"
    / "bess0401_vs_fluxsat_fixlon_spatial_compare/bess_fluxsat_spatial_consistency_summary.csv"
)
START_YEAR = 2001
END_YEAR = 2018
ANNUAL_BESS_CSV = (
    BASE_DIR
    / "process/result_analysis/result_weighted/compare_analysis2/"
    / "v20260401_growingseason_recovery_gsdays/"
    / "annual_response_recovery_trends_v20260401_growingseason_recovery_gsdays.csv"
)
ANNUAL_FLUXSAT_CSV = (
    BASE_DIR
    / "process/result_analysis/result_weighted/fluxsat_compare_analysis2/"
    / "fluxsat_0401_sensitivity_compare/fluxsat_0401_sensitivity_annual.csv"
)
BESS_FLUXSAT_41_FILES = {
    ("BESS", "SMrz"): BASE_DIR
    / "process/GPP-draught-analysis/code1/results/"
    / "gpp_response_SMrz_events_global_v20260401_growingseason_recovery_gsdays_rel0_abspeak_absrec_c30x095_w420_decline30_d5_norecmax.nc",
    ("FluxSat 0401 rec100cap", "SMrz"): BASE_DIR
    / "process/fluxsat-draught-analysis/code1/results/"
    / "fluxsat_gpp_response_SMrz_events_global_v20260401_growingseason_recovery_gsdays_rel0_abspeak_absrec_c30x095_w420_decline30_d5_rec100cap_fixlon_v20260426.nc",
    ("BESS", "SMs"): BASE_DIR
    / "process/GPP-draught-analysis/code2_SMs/results/"
    / "gpp_response_SMs_events_global_v20260401_growingseason_recovery_gsdays_rel0_abspeak_absrec_c30x095_w420_decline30_d5_norecmax.nc",
    ("FluxSat 0401 rec100cap", "SMs"): BASE_DIR
    / "process/fluxsat-draught-analysis/code2_SMs/results/"
    / "fluxsat_gpp_response_SMs_events_global_v20260401_growingseason_recovery_gsdays_rel0_abspeak_absrec_c30x095_w420_decline30_d5_rec100cap_fixlon_v20260426.nc",
}


def load_rows() -> list[dict[str, str]]:
    with open(SUMMARY_CSV, "r", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def find_row(rows: list[dict[str, str]], scenario: str) -> dict[str, str]:
    for row in rows:
        if row["scenario"] == scenario:
            return row
    raise KeyError(scenario)


def to_numpy(var) -> np.ndarray:
    arr = var[:]
    if np.ma.isMaskedArray(arr):
        arr = arr.filled(np.nan)
    arr = np.asarray(arr)
    if np.issubdtype(arr.dtype, np.integer):
        arr = arr.astype(np.float64)
    fill_value = getattr(var, "_FillValue", None)
    if fill_value is not None:
        arr = arr.astype(np.float64, copy=False)
        arr[np.isclose(arr, float(fill_value), equal_nan=False)] = np.nan
    return arr


def load_41_rows() -> list[dict[str, object]]:
    annual_bess_rows = list(csv.DictReader(open(ANNUAL_BESS_CSV, "r", encoding="utf-8")))
    annual_fluxsat_rows = list(csv.DictReader(open(ANNUAL_FLUXSAT_CSV, "r", encoding="utf-8")))
    out = []
    for dataset, soil_layer in [
        ("BESS", "SMrz"),
        ("FluxSat 0401 rec100cap", "SMrz"),
        ("BESS", "SMs"),
        ("FluxSat 0401 rec100cap", "SMs"),
    ]:
        code = "code1" if soil_layer == "SMrz" else "code2"
        with nc.Dataset(str(BESS_FLUXSAT_41_FILES[(dataset, soil_layer)]), "r") as ds:
            onset_year = to_numpy(ds.variables["onset_year"])
            if "t_response_onset_start" in ds.variables:
                t_response = to_numpy(ds.variables["t_response_onset_start"])
            else:
                t_response = to_numpy(ds.variables["t_response"])
            t_recover = to_numpy(ds.variables["t_recover_to_baseline"])
        mask = np.isfinite(onset_year) & (onset_year >= START_YEAR) & (onset_year <= END_YEAR)
        resp_valid = np.isfinite(t_response[mask]) & (t_response[mask] >= 0)
        rec_valid = np.isfinite(t_recover[mask]) & (t_recover[mask] >= 0)
        if dataset == "BESS":
            annual_rows = annual_bess_rows
            years = np.array(
                [
                    float(r["year"])
                    for r in annual_rows
                    if r["variable"] == "GPP"
                    and r["code"] == code
                    and r["soil_layer"] == soil_layer
                    and START_YEAR <= int(r["year"]) <= END_YEAR
                ],
                dtype=np.float64,
            )
            resp_vals = np.array(
                [
                    float(r["response_mean"])
                    for r in annual_rows
                    if r["variable"] == "GPP"
                    and r["code"] == code
                    and r["soil_layer"] == soil_layer
                    and START_YEAR <= int(r["year"]) <= END_YEAR
                ],
                dtype=np.float64,
            )
            rec_vals = np.array(
                [
                    float(r["recovery_mean"])
                    for r in annual_rows
                    if r["variable"] == "GPP"
                    and r["code"] == code
                    and r["soil_layer"] == soil_layer
                    and START_YEAR <= int(r["year"]) <= END_YEAR
                ],
                dtype=np.float64,
            )
        else:
            annual_rows = annual_fluxsat_rows
            years = np.array(
                [
                    float(r["year"])
                    for r in annual_rows
                    if r["dataset"] == dataset
                    and r["code"] == code
                    and r["soil_layer"] == soil_layer
                    and START_YEAR <= int(r["year"]) <= END_YEAR
                ],
                dtype=np.float64,
            )
            resp_vals = np.array(
                [
                    float(r["response_mean"])
                    for r in annual_rows
                    if r["dataset"] == dataset
                    and r["code"] == code
                    and r["soil_layer"] == soil_layer
                    and START_YEAR <= int(r["year"]) <= END_YEAR
                ],
                dtype=np.float64,
            )
            rec_vals = np.array(
                [
                    float(r["recovery_mean"])
                    for r in annual_rows
                    if r["dataset"] == dataset
                    and r["code"] == code
                    and r["soil_layer"] == soil_layer
                    and START_YEAR <= int(r["year"]) <= END_YEAR
                ],
                dtype=np.float64,
            )
        response_slope = float(np.polyfit(years, resp_vals, 1)[0] * 10.0)
        recovery_slope = float(np.polyfit(years, rec_vals, 1)[0] * 10.0)
        out.append(
            {
                "dataset": dataset,
                "soil_layer": soil_layer,
                "event_total": int(mask.sum()),
                "response_valid_pct": float(np.mean(resp_valid) * 100.0),
                "response_mean_days": float(np.nanmean(np.where(resp_valid, t_response[mask], np.nan))),
                "recovery_valid_pct": float(np.mean(rec_valid) * 100.0),
                "recovery_mean_days": float(np.nanmean(np.where(rec_valid, t_recover[mask], np.nan))),
                "response_mean_slope_days_per_decade": response_slope,
                "recovery_mean_slope_days_per_decade": recovery_slope,
            }
        )
    return out


def find_41_row(rows: list[dict[str, object]], dataset: str, soil_layer: str) -> dict[str, object]:
    for row in rows:
        if row["dataset"] == dataset and row["soil_layer"] == soil_layer:
            return row
    raise KeyError((dataset, soil_layer))


def build_41_paragraph(rows: list[dict[str, object]]) -> str:
    f_smrz = find_41_row(rows, "FluxSat 0401 rec100cap", "SMrz")
    f_sms = find_41_row(rows, "FluxSat 0401 rec100cap", "SMs")
    b_smrz = find_41_row(rows, "BESS", "SMrz")
    b_sms = find_41_row(rows, "BESS", "SMs")
    return (
        f"在 {START_YEAR}–{END_YEAR} 的共同覆盖时段内，FluxSat 0401 的 100 天恢复上限版本为检验 BESS 0401 的独立合理性提供了更可比的参照。"
        f"在这一统一口径下，FluxSat 在根区和表层情景下的响应有效率分别为 {f_smrz['response_valid_pct']:.2f}% 和 {f_sms['response_valid_pct']:.2f}%，"
        f"平均响应时间分别为 {f_smrz['response_mean_days']:.2f} 天和 {f_sms['response_mean_days']:.2f} 天；"
        f"对应的 BESS 值分别为 {b_smrz['response_valid_pct']:.2f}%、{b_smrz['response_mean_days']:.2f} 天，以及 {b_sms['response_valid_pct']:.2f}%、{b_sms['response_mean_days']:.2f} 天。"
        f"更重要的是，两套产品在年际响应变慢趋势上保持同向：SMrz 情景下 BESS 与 FluxSat 分别为 {b_smrz['response_mean_slope_days_per_decade']:.2f} d/10a 和 {f_smrz['response_mean_slope_days_per_decade']:.2f} d/10a，"
        f"SMs 情景下分别为 {b_sms['response_mean_slope_days_per_decade']:.2f} d/10a 和 {f_sms['response_mean_slope_days_per_decade']:.2f} d/10a。"
        "这说明在采用一致时间窗口和一致恢复上限之后，两个独立产品仍对“骤旱响应随时间推移而变慢”这一方向性结论给出了稳定支持。"
    )


def build_41_followup_paragraph() -> str:
    return (
        "不过，FluxSat 与 BESS 在恢复阶段的绝对幅度、有效恢复事件保留比例以及空间覆盖完整性上仍存在差异。"
        "FluxSat 的时间覆盖较短，且需要经历月文件拼接、重采样、经度修正和 100 天恢复上限筛选；这些步骤会共同影响可进入恢复统计的事件集合。"
        "因此，本项目在总体量级对照之外，还进一步引入共同期空间对照与热点重叠分析，用于判断两者在全球格局层面是否保持一致。"
    )


def build_41_table_html(rows: list[dict[str, object]]) -> str:
    ordered = [
        find_41_row(rows, "BESS", "SMrz"),
        find_41_row(rows, "FluxSat 0401 rec100cap", "SMrz"),
        find_41_row(rows, "BESS", "SMs"),
        find_41_row(rows, "FluxSat 0401 rec100cap", "SMs"),
    ]
    labels = ["BESS", "FluxSat 0401 rec100cap", "BESS", "FluxSat 0401 rec100cap"]
    scenarios = ["SMrz", "SMrz", "SMs", "SMs"]
    lines = [
        "  <table>",
        "    <tr>",
        "      <th>数据源</th>",
        "      <th>情景</th>",
        "      <th>时间窗口</th>",
        "      <th>事件数</th>",
        "      <th>响应有效率（%）</th>",
        "      <th>平均响应时间（d）</th>",
        "      <th>恢复有效率（%）</th>",
        "      <th>平均恢复时间（d）</th>",
        "      <th>响应趋势（d/10a）</th>",
        "      <th>恢复趋势（d/10a）</th>",
        "    </tr>",
    ]
    for label, scenario, row in zip(labels, scenarios, ordered):
        lines.extend(
            [
                "    <tr>",
                f"      <td>{label}</td>",
                f"      <td>{scenario}</td>",
                f"      <td>{START_YEAR}-{END_YEAR}</td>",
                f"      <td>{row['event_total']}</td>",
                f"      <td>{row['response_valid_pct']:.2f}</td>",
                f"      <td>{row['response_mean_days']:.2f}</td>",
                f"      <td>{row['recovery_valid_pct']:.2f}</td>",
                f"      <td>{row['recovery_mean_days']:.2f}</td>",
                f"      <td>{row['response_mean_slope_days_per_decade']:.2f}</td>",
                f"      <td>{row['recovery_mean_slope_days_per_decade']:.2f}</td>",
                "    </tr>",
            ]
        )
    lines.append("  </table>")
    return "\n".join(lines)


def build_note(rows: list[dict[str, str]]) -> str:
    smrz = find_row(rows, "SMrz")
    sms = find_row(rows, "SMs")
    start_year = smrz["start_year"]
    end_year = smrz["end_year"]
    return (
        "由于 BESS 与 FluxSat 0401 100 天恢复上限版本在这里比较的是连续型恢复时间场，而不是像 GLEAM 与 ERA5 那样比较“是否发生事件”的频次网格，"
        "因此不适合直接定义“正频次格点重叠度”。为补充空间结构的一致性证据，这里额外定义高恢复时间热点区："
        f"在 {start_year}–{end_year} 共同时间窗口内的共同有效格点上，分别取各自产生的恢复时间场前 20% 高值格点作为热点区，再计算两者热点集合的 Jaccard 重叠度。"
        f"按这一规则，SMrz 情景下 BESS 热点阈值为 {float(smrz['bess_hotspot_threshold_days']):.2f} d，"
        f"FluxSat 热点阈值为 {float(smrz['fluxsat_hotspot_threshold_days']):.2f} d，热点区 Jaccard 重叠度为 {float(smrz['hotspot_jaccard_overlap']):.3f}"
        f"（交集 {smrz['hotspot_intersection_pixels']} 格点，并集 {smrz['hotspot_union_pixels']} 格点）；"
        f"SMs 情景下对应阈值分别为 {float(sms['bess_hotspot_threshold_days']):.2f} d 和 {float(sms['fluxsat_hotspot_threshold_days']):.2f} d，"
        f"热点区 Jaccard 重叠度为 {float(sms['hotspot_jaccard_overlap']):.3f}。"
        "这一指标的含义不是逐像元数值完全相等，而是检验两个独立 GPP 产品是否会把“恢复最慢的区域”落在相近的全球位置上。"
    )


def build_spatial_paragraph(rows: list[dict[str, str]]) -> str:
    smrz = find_row(rows, "SMrz")
    sms = find_row(rows, "SMs")
    start_year = smrz["start_year"]
    end_year = smrz["end_year"]
    return (
        f"为了避免单纯依赖全球均值而忽略区域结构，本项目进一步将 BESS 0401 与 FluxSat 0401 的 100 天恢复上限版本都严格限制在 {start_year}–{end_year} 的共同时间窗口内，"
        "并在统一网格上对恢复时间格点均值进行了逐像元比较。"
        f"在共同有效区域上，根区情景的空间相关系数达到 {float(smrz['spatial_correlation']):.3f}，"
        f"BESS 与 FluxSat 的面积加权平均恢复时间分别为 {float(smrz['bess_area_weighted_recovery_mean_days']):.2f} 天和 {float(smrz['fluxsat_area_weighted_recovery_mean_days']):.2f} 天；"
        f"表层情景的空间相关系数达到 {float(sms['spatial_correlation']):.3f}，对应面积加权平均恢复时间分别为 {float(sms['bess_area_weighted_recovery_mean_days']):.2f} 天和 {float(sms['fluxsat_area_weighted_recovery_mean_days']):.2f} 天。"
        "换言之，虽然两个产品在绝对恢复时长上仍存在系统差异，但在全球尺度的大区域分布格局上，两者表现出中等偏强的一致性。"
    )


def build_table_html(rows: list[dict[str, str]]) -> str:
    smrz = find_row(rows, "SMrz")
    sms = find_row(rows, "SMs")
    return f"""  <table>
    <tr>
      <th>情景</th>
      <th>时间窗口</th>
      <th>共同有效格点数</th>
      <th>BESS 面积加权恢复均值（d）</th>
      <th>FluxSat 面积加权恢复均值（d）</th>
      <th>空间相关系数</th>
      <th>平均差值（BESS − FluxSat，d）</th>
      <th>热点区定义</th>
      <th>热点区 Jaccard 重叠度</th>
      <th>热点交集格点数</th>
      <th>热点并集格点数</th>
    </tr>
    <tr>
      <td>SMrz</td>
      <td>{smrz['start_year']}-{smrz['end_year']}</td>
      <td>{smrz['shared_valid_pixels']}</td>
      <td>{float(smrz['bess_area_weighted_recovery_mean_days']):.2f}</td>
      <td>{float(smrz['fluxsat_area_weighted_recovery_mean_days']):.2f}</td>
      <td>{float(smrz['spatial_correlation']):.3f}</td>
      <td>{float(smrz['mean_difference_bess_minus_fluxsat_days']):.2f}</td>
      <td>共同有效格点内各自前20%高值区</td>
      <td>{float(smrz['hotspot_jaccard_overlap']):.3f}</td>
      <td>{smrz['hotspot_intersection_pixels']}</td>
      <td>{smrz['hotspot_union_pixels']}</td>
    </tr>
    <tr>
      <td>SMs</td>
      <td>{sms['start_year']}-{sms['end_year']}</td>
      <td>{sms['shared_valid_pixels']}</td>
      <td>{float(sms['bess_area_weighted_recovery_mean_days']):.2f}</td>
      <td>{float(sms['fluxsat_area_weighted_recovery_mean_days']):.2f}</td>
      <td>{float(sms['spatial_correlation']):.3f}</td>
      <td>{float(sms['mean_difference_bess_minus_fluxsat_days']):.2f}</td>
      <td>共同有效格点内各自前20%高值区</td>
      <td>{float(sms['hotspot_jaccard_overlap']):.3f}</td>
      <td>{sms['hotspot_intersection_pixels']}</td>
      <td>{sms['hotspot_union_pixels']}</td>
    </tr>
  </table>"""


def insert_paragraph_after(paragraph, text: str):
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    from docx.text.paragraph import Paragraph

    new_para = Paragraph(new_p, paragraph._parent)
    run = new_para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
            "SimSun",
        )
    return new_para


def update_html(note: str, rows: list[dict[str, str]], rows_41: list[dict[str, object]]) -> None:
    text = HTML_PATH.read_text(encoding="utf-8")
    paragraph_41 = build_41_paragraph(rows_41)
    followup_41 = build_41_followup_paragraph()
    table_41 = build_41_table_html(rows_41)
    p41_pattern = re.compile(r"  <p>在 .*?核心方向性结论给出了同向支持。</p>", re.S)
    if not p41_pattern.search(text):
        raise RuntimeError("Unable to locate 4.1 summary paragraph in BESS-FluxSat HTML.")
    text = p41_pattern.sub(f"  <p>{paragraph_41}</p>", text, count=1)
    p41b_pattern = re.compile(r"  <p>不过，仅用主运行的绝对恢复时间.*?单一全局平均表上。</p>", re.S)
    if not p41b_pattern.search(text):
        raise RuntimeError("Unable to locate 4.1 followup paragraph in BESS-FluxSat HTML.")
    text = p41b_pattern.sub(f"  <p>{followup_41}</p>", text, count=1)
    table41_pattern = re.compile(r"  <table>\n    <tr>\n      <th>数据源</th>.*?</table>", re.S)
    if not table41_pattern.search(text):
        raise RuntimeError("Unable to locate 4.1 table in BESS-FluxSat HTML.")
    text = table41_pattern.sub(table_41, text, count=1)
    spatial_paragraph = build_spatial_paragraph(rows)
    para_pattern = re.compile(
        r"  <p>为了避免单纯依赖全球均值而忽略区域结构.*?</p>",
        re.S,
    )
    if not para_pattern.search(text):
        raise RuntimeError("Unable to locate spatial paragraph in BESS-FluxSat HTML.")
    text = para_pattern.sub(f"  <p>{spatial_paragraph}</p>", text, count=1)
    note_pattern = re.compile(
        r"\n\n  <p>由于 BESS 与 FluxSat 在这里比较的是连续型恢复时间场.*?落在相近的全球位置上。</p>",
        re.S,
    )
    text = note_pattern.sub("", text)
    marker = f"  <p>{spatial_paragraph}</p>"
    text = text.replace(marker, marker + "\n\n" + f"  <p>{note}</p>", 1)
    pattern = re.compile(r"  <table>\n    <tr>\n      <th>情景</th>.*?</table>", re.S)
    new_table = build_table_html(rows)
    if not pattern.search(text):
        raise RuntimeError("Unable to locate spatial consistency table in BESS-FluxSat HTML.")
    text = pattern.sub(new_table, text, count=1)
    HTML_PATH.write_text(text, encoding="utf-8")


def update_docx(note: str, rows: list[dict[str, str]], rows_41: list[dict[str, object]]) -> None:
    doc = Document(str(DOCX_PATH))
    paragraph_41 = build_41_paragraph(rows_41)
    followup_41 = build_41_followup_paragraph()
    for p in doc.paragraphs:
        if p.text.strip().startswith("在 2000–2019") or p.text.strip().startswith(f"在 {START_YEAR}–{END_YEAR} 的共同覆盖时段内"):
            p.text = paragraph_41
            break
    for p in doc.paragraphs:
        if p.text.strip().startswith("不过，仅用主运行的绝对恢复时间") or p.text.strip().startswith("不过，FluxSat 与 BESS 在恢复阶段的绝对幅度"):
            p.text = followup_41
            break
    spatial_paragraph = build_spatial_paragraph(rows)
    for p in doc.paragraphs:
        if p.text.strip().startswith("为了避免单纯依赖全球均值而忽略区域结构"):
            p.text = spatial_paragraph
            break
    note_prefix = "由于 BESS 与 FluxSat 0401 100 天恢复上限版本在这里比较的是连续型恢复时间场"
    note_paragraphs = [p for p in doc.paragraphs if p.text.strip().startswith(note_prefix)]
    for old_p in note_paragraphs[1:]:
        old_p._element.getparent().remove(old_p._element)
    if note_paragraphs:
        note_paragraphs[0].text = note
    else:
        for p in doc.paragraphs:
            if p.text.strip().startswith("为了避免单纯依赖全球均值而忽略区域结构"):
                insert_paragraph_after(p, note)
                break
    table = None
    for t in doc.tables:
        if t.rows and t.rows[0].cells and t.rows[0].cells[0].text == "情景":
            table = t
            break
    if table is None:
        raise RuntimeError("Unable to locate spatial consistency table in BESS-FluxSat DOCX.")
    headers = [
        "情景",
        "时间窗口",
        "共同有效格点数",
        "BESS 面积加权恢复均值（d）",
        "FluxSat 面积加权恢复均值（d）",
        "空间相关系数",
        "平均差值（BESS − FluxSat，d）",
        "热点区定义",
        "热点区 Jaccard 重叠度",
        "热点交集格点数",
        "热点并集格点数",
    ]
    while len(table.columns) < len(headers):
        table.add_column(Pt(60))
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    smrz = find_row(rows, "SMrz")
    sms = find_row(rows, "SMs")
    values = [
        ["SMrz", f"{smrz['start_year']}-{smrz['end_year']}", smrz["shared_valid_pixels"], f"{float(smrz['bess_area_weighted_recovery_mean_days']):.2f}", f"{float(smrz['fluxsat_area_weighted_recovery_mean_days']):.2f}", f"{float(smrz['spatial_correlation']):.3f}", f"{float(smrz['mean_difference_bess_minus_fluxsat_days']):.2f}", "共同有效格点内各自前20%高值区", f"{float(smrz['hotspot_jaccard_overlap']):.3f}", smrz["hotspot_intersection_pixels"], smrz["hotspot_union_pixels"]],
        ["SMs", f"{sms['start_year']}-{sms['end_year']}", sms["shared_valid_pixels"], f"{float(sms['bess_area_weighted_recovery_mean_days']):.2f}", f"{float(sms['fluxsat_area_weighted_recovery_mean_days']):.2f}", f"{float(sms['spatial_correlation']):.3f}", f"{float(sms['mean_difference_bess_minus_fluxsat_days']):.2f}", "共同有效格点内各自前20%高值区", f"{float(sms['hotspot_jaccard_overlap']):.3f}", sms["hotspot_intersection_pixels"], sms["hotspot_union_pixels"]],
    ]
    for row_idx, row_values in enumerate(values, start=1):
        for col_idx, value in enumerate(row_values):
            table.rows[row_idx].cells[col_idx].text = str(value)
    table41 = None
    for t in doc.tables:
        if t.rows and t.rows[0].cells and t.rows[0].cells[0].text == "数据源":
            table41 = t
            break
    if table41 is None:
        raise RuntimeError("Unable to locate 4.1 comparison table in BESS-FluxSat DOCX.")
    headers41 = [
        "数据源",
        "情景",
        "时间窗口",
        "事件数",
        "响应有效率（%）",
        "平均响应时间（d）",
        "恢复有效率（%）",
        "平均恢复时间（d）",
        "响应趋势（d/10a）",
        "恢复趋势（d/10a）",
    ]
    while len(table41.columns) < len(headers41):
        table41.add_column(Pt(60))
    for i, h in enumerate(headers41):
        table41.rows[0].cells[i].text = h
    rows41_ordered = [
        ("BESS", "SMrz"),
        ("FluxSat 0401 rec100cap", "SMrz"),
        ("BESS", "SMs"),
        ("FluxSat 0401 rec100cap", "SMs"),
    ]
    for row_idx, (dataset, soil_layer) in enumerate(rows41_ordered, start=1):
        row = find_41_row(rows_41, dataset, soil_layer)
        vals = [
            dataset,
            soil_layer,
            f"{START_YEAR}-{END_YEAR}",
            row["event_total"],
            f"{row['response_valid_pct']:.2f}",
            f"{row['response_mean_days']:.2f}",
            f"{row['recovery_valid_pct']:.2f}",
            f"{row['recovery_mean_days']:.2f}",
            f"{row['response_mean_slope_days_per_decade']:.2f}",
            f"{row['recovery_mean_slope_days_per_decade']:.2f}",
        ]
        for col_idx, value in enumerate(vals):
            table41.rows[row_idx].cells[col_idx].text = str(value)
    doc.save(str(DOCX_PATH))


def main() -> None:
    rows = load_rows()
    rows_41 = load_41_rows()
    note = build_note(rows)
    update_html(note, rows, rows_41)
    update_docx(note, rows, rows_41)
    print(f"Updated {HTML_PATH}")
    print(f"Updated {DOCX_PATH}")


if __name__ == "__main__":
    main()
